#!/usr/bin/env python3
"""Comprehensive analysis of DOCX structure"""

from docx import Document
import re

docx_path = "5°_2° - Proyecto de Ingeniería Mecatrónica.docx"
doc = Document(docx_path)

print("🔍 ANÁLISIS COMPLETO DEL DOCUMENTO")
print("=" * 70)

# 1. Buscar menciones de "Ingeniería" o "carrera" en TODO el documento
print("\n1️⃣ BÚSQUEDA DE PALABRAS CLAVE EN EL DOCUMENTO:")
print("-" * 70)

keywords = ['ingeniería', 'carrera', 'plan de estudio', 'ciclo', 'especialidad']
found_keywords = {}

for para in doc.paragraphs:
    text_lower = para.text.lower()
    for keyword in keywords:
        if keyword in text_lower:
            if keyword not in found_keywords:
                found_keywords[keyword] = []
            found_keywords[keyword].append(para.text.strip()[:100])

for keyword, occurrences in found_keywords.items():
    print(f"\n'{keyword.upper()}' encontrado {len(occurrences)} veces:")
    for idx, occurrence in enumerate(occurrences[:3], 1):  # Show first 3
        print(f"  {idx}. {occurrence}...")

# 2. Analizar todas las tablas
print("\n\n2️⃣ ANÁLISIS DE TODAS LAS TABLAS:")
print("-" * 70)

for table_idx, table in enumerate(doc.tables):
    # Get table dimensions
    rows = len(table.rows)
    cols = len(table.columns) if table.columns else 0
    
    # Get all text from table
    table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
    
    print(f"\nTabla {table_idx}: {rows} filas × {cols} columnas")
    print(f"  Contenido resumido: {table_text[:150]}...")
    
    # Si tiene muchas celdas, mostrar la estructura
    if rows <= 5:
        print(f"  Estructura:")
        for row_idx, row in enumerate(table.rows):
            row_text = " | ".join([cell.text.strip()[:30] for cell in row.cells])
            print(f"    Row {row_idx}: {row_text}")

# 3. Información de documento
print("\n\n3️⃣ PROPIEDADES DEL DOCUMENTO:")
print("-" * 70)
print(f"Total de párrafos: {len(doc.paragraphs)}")
print(f"Total de tablas: {len(doc.tables)}")
print(f"Total de secciones: {len(doc.sections)}")

# Core properties
try:
    print("\nMetadatos (si están disponibles):")
    print(f"  Title: {doc.core_properties.title or 'No definido'}")
    print(f"  Subject: {doc.core_properties.subject or 'No definido'}")
    print(f"  Author: {doc.core_properties.author or 'No definido'}")
except Exception as e:
    print(f"  Error accediendo a metadatos: {e}")

print("\n✅ Análisis completado")
