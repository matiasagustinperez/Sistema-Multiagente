#!/usr/bin/env python
"""Debug de extracción de prácticos"""
from docx import Document

file_path = r'c:\TesisMCD\backend\data\uploads\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(file_path)

print("Buscando tablas con prácticos...")
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if 'práctico' in cell.text.lower() or 'practico' in cell.text.lower():
                print(f"\nTABLA {table_idx}, FILA {row_idx}:")
                for c_idx, c in enumerate(row.cells):
                    print(f"  [{c_idx}] {c.text[:80]}")

print("\n" + "=" * 80)
print("Analizando tabla 7 (conocida con prácticos)...")
if len(doc.tables) > 7:
    table = doc.tables[7]
    print(f"Tabla 7: {len(table.rows)} filas x {len(table.columns)} columnas")
    for row_idx, row in enumerate(table.rows[:10]):
        print(f"\nFila {row_idx}:")
        for col_idx, cell in enumerate(row.cells):
            text = cell.text[:120]
            print(f"  [{col_idx}] {repr(text)}")
