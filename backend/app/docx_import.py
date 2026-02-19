"""
Módulo mejorado para importar datos desde archivos DOCX.
Estrategia: Búsqueda por secciones numeradas (1., 2., 3., etc.)
"""
from __future__ import annotations

import re
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from docx.table import Table


def extract_text_from_table_cell(cell) -> str:
    """Extrae todo el texto de una celda, preservando párrafos."""
    return '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])


def normalize_docx_text(value: str) -> str:
    """Normaliza texto para comparaciones robustas (sin tildes, minúsculas)."""
    if not value:
        return ""
    replacements = str.maketrans({
        'á': 'a', 'à': 'a', 'ä': 'a', 'â': 'a',
        'é': 'e', 'è': 'e', 'ë': 'e', 'ê': 'e',
        'í': 'i', 'ì': 'i', 'ï': 'i', 'î': 'i',
        'ó': 'o', 'ò': 'o', 'ö': 'o', 'ô': 'o',
        'ú': 'u', 'ù': 'u', 'ü': 'u', 'û': 'u',
        'ñ': 'n',
    })
    normalized = value.strip().lower().translate(replacements)
    return re.sub(r'\s+', ' ', normalized)


def find_section_paragraphs(doc: Document) -> Dict[str, Tuple[int, int]]:
    """
    Encuentra índices de párrafos con secciones (con o sin números).
    Retorna: {nombre_sección: (inicio_idx, fin_idx)}
    
    Secciones esperadas (sin números en DOCX real):
    - CONTENIDOS MÍNIMOS:
    - FUNDAMENTOS:
    - OBJETIVOS:
    - CONTENIDOS DE LA ASIGNATURA: (UNIDADES)
    - PROGRAMA DE TRABAJOS PRÁCTICOS:
    - METODOLOGÍA:
    - EVALUACIÓN:
    - BIBLIOGRAFÍA:
    - OBSERVACIONES:
    """
    sections = {}
    section_starts = {}
    section_order = []
    
    # Palabras clave para identificar secciones
    section_keywords = {
        'CONTENIDOS MÍNIMOS': 0,
        'CONTENIDOS MINIMOS': 0,
        'FUNDAMENTOS': 1,
        'OBJETIVOS': 2,
        'CONTENIDOS DE LA ASIGNATURA': 3,
        'CONTENIDOS DE LA ASIGNATUR': 3,  # Por si está cortado
        'PROGRAMA DE TRABAJOS PRÁCTICOS': 4,
        'PROGRAMA DE TRABAJOS PRACTICOS': 4,
        'PROGRAMA DE TRABAJO PRACTICO': 4,
        'PROGRAMA DE TRABAJO PRÁCTICO': 4,
        'METODOLOGÍA': 5,
        'METODOLOGIA': 5,
        'EVALUACIÓN': 6,
        'EVALUACION': 6,
        'BIBLIOGRAFÍA': 7,
        'BIBLIOGRAFIA': 7,
        'OBSERVACIONES': 8,
    }
    
    # Buscar inicio de cada sección
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        
        # Buscar coincidencia con palabras clave
        for keyword, order in section_keywords.items():
            if text.startswith(keyword):
                # Limpiar el nombre de la sección
                section_name = para.text.strip().rstrip(':')
                section_starts[order] = (idx, section_name)
                section_order.append((order, idx, section_name))
    
    # Crear rangos: desde cada sección hasta la siguiente
    section_order.sort(key=lambda x: x[0])  # Ordenar por número de sección
    for i, (order_num, start_idx, section_name) in enumerate(section_order):
        # Buscar índice final (inicio de siguiente sección)
        if i < len(section_order) - 1:
            end_idx = section_order[i + 1][1]
        else:
            end_idx = len(doc.paragraphs)
        
        sections[section_name] = (start_idx, end_idx)
    
    return sections


def extract_section_content_from_tables(doc: Document, keywords: List[str]) -> str:
    """
    Extrae contenido de TABLAS que coincidan con los keywords dados.
    Usado cuando los párrafos entre secciones están vacíos y el contenido
    está dentro de tablas.
    """
    for table in doc.tables:
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
        
        # Buscar tabla que contenga todos los keywords
        if all(keyword.lower() in table_text for keyword in keywords):
            content_parts = []
            for row in table.rows[1:]:  # Saltar encabezado
                row_text = '\n'.join([extract_text_from_table_cell(cell) for cell in row.cells])
                if row_text.strip():
                    content_parts.append(row_text)
            return '\n---\n'.join(content_parts)
    
    return ''


def extract_labeled_table_content(doc: Document, label_keywords: List[str]) -> str:
    """Extrae contenido desde tablas buscando celdas con una etiqueta."""
    if not label_keywords:
        return ''

    for table in doc.tables:
        for row in table.rows:
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_norm = normalize_docx_text(cell_text)
                if any(keyword in cell_norm for keyword in label_keywords):
                    if len(row.cells) > 1:
                        other_texts = [
                            extract_text_from_table_cell(c)
                            for idx, c in enumerate(row.cells)
                            if idx != cell_idx
                        ]
                        content = '\n'.join([t for t in other_texts if t.strip()]).strip()
                        if content:
                            return content

                    content = extract_text_from_table_cell(cell)
                    label_pattern = r'(?i)(' + '|'.join([re.escape(k) for k in label_keywords]) + r')\s*:?' 
                    content = re.sub(label_pattern, '', content, count=1).strip()
                    return content

    return ''


def strip_observations_footer(text: str) -> str:
    """Elimina el cierre estándar que no es una observación real."""
    if not text:
        return ''

    split_patterns = [
        r'(?i)\bchilecito\s*:',
        r'(?i)\belevo\s+el\s+presente\b',
        r'(?i)\bprofesor/a\b',
        r'(?i)\bprofesor\b',
    ]
    for pattern in split_patterns:
        match = re.search(pattern, text)
        if match:
            text = text[:match.start()].strip()
            break

    lines = [line.strip() for line in text.splitlines()]
    footer_starts = [
        'chilecito',
        'elevo el presente',
        'profesor/a',
        'profesor',
    ]
    cleaned_lines = []
    for line in lines:
        normalized = normalize_docx_text(line)
        if any(normalized.startswith(fs) for fs in footer_starts):
            break
        if line:
            cleaned_lines.append(line)

    return '\n'.join(cleaned_lines).strip()


def parse_global_bibliography(text: str) -> Tuple[str, str, str]:
    """Separa bibliografía global en básica y complementaria."""
    if not text:
        return "", "", ""

    normalized = normalize_docx_text(text)
    basic_label = r'bibliograf[ií]a\s+b[áa]sica'
    comp_label = r'bibliograf[ií]a\s+complementaria'

    if re.search(basic_label, normalized) or re.search(comp_label, normalized):
        basic_text = ""
        comp_text = ""

        parts = re.split(r'(?i)' + basic_label + r'\s*\(.*?\)?\s*:?', text, maxsplit=1)
        after_basic = parts[1] if len(parts) > 1 else text

        comp_parts = re.split(r'(?i)' + comp_label + r'\s*:?', after_basic, maxsplit=1)
        if len(comp_parts) > 1:
            basic_text = comp_parts[0].strip()
            comp_text = comp_parts[1].strip()
        else:
            basic_text = after_basic.strip()

        combined = basic_text
        if comp_text:
            combined = (basic_text + "\n\nBibliografía complementaria:\n" + comp_text).strip()

        return combined, basic_text, comp_text

    return text.strip(), "", ""


def extract_section_content(doc: Document, section_start_idx: int, section_end_idx: int) -> str:
    """
    Extrae contenido de texto entre dos índices de párrafo.
    Como muchas secciones usan tablas, esto busca el contenido textual.
    """
    paragraphs = []
    
    # Extraer párrafos (aunque muchos estén vacíos)
    for idx in range(section_start_idx + 1, section_end_idx):  # +1 para saltar título
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if text and text not in [':', ''] and len(text) > 2:
            paragraphs.append(text)
    
    return '\n'.join(paragraphs)


def extract_section_tables_content(doc: Document, section_start_idx: int, section_end_idx: int) -> str:
    """Extrae texto de tablas ubicadas entre dos párrafos de sección."""
    body_elements = []
    para_idx = 0
    table_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            body_elements.append(('p', para_idx))
            para_idx += 1
        elif child.tag.endswith('}tbl'):
            body_elements.append(('tbl', table_idx))
            table_idx += 1

    para_positions = {}
    for pos, (kind, idx) in enumerate(body_elements):
        if kind == 'p':
            para_positions[idx] = pos

    if section_start_idx not in para_positions:
        return ''

    start_pos = para_positions.get(section_start_idx, 0)
    end_pos = para_positions.get(section_end_idx, len(body_elements))
    if end_pos <= start_pos:
        end_pos = len(body_elements)

    table_indices = []
    for pos in range(start_pos + 1, min(end_pos, len(body_elements))):
        kind, idx = body_elements[pos]
        if kind == 'tbl':
            table_indices.append(idx)

    if not table_indices:
        return ''

    parts = []
    for idx in table_indices:
        if idx >= len(doc.tables):
            continue
        table = doc.tables[idx]
        for row in table.rows:
            row_text = '\n'.join([extract_text_from_table_cell(cell) for cell in row.cells])
            if row_text.strip():
                parts.append(row_text.strip())

    return '\n'.join(parts).strip()


def extract_section_tables(doc: Document, section_start_idx: int, section_end_idx: int) -> List[Table]:
    """Extrae todas las tablas dentro de un rango de párrafos."""
    tables_in_range = []
    
    # Obtener posiciones de tablas
    para_idx = 0
    for table_idx, table in enumerate(doc.tables):
        # Buscar en qué posición está cada tabla (aproximado)
        # Nota: python-docx no proporciona índice directo, así que usamos heurística
        # Contamos párrafos y tablas hasta encontrar las tablas en el rango
        pass
    
    # Alternativa: buscar tablas por contenido
    for table in doc.tables:
        table_text = extract_text_from_table_cell(table.rows[0].cells[0]) if table.rows else ""
        # Heurística: si la tabla está en este rango, incluirla
        # (Esto es aproximado; mejor sería mapeo directo)
        tables_in_range.append(table)
    
    return tables_in_range


def extract_generic_competencies(text: str) -> List[Dict[str, str]]:
    """
    Extrae Competencias Genéricas del texto usando un enfoque basado en bloques.
    Captura TODAS las variantes de competencias genéricas:
    - CGT (Competencias Genéricas Transversales)  
    - CGS (Competencias Genéricas Específicas)
    - Cualquier CG[LETRA][NÚMERO]
    
    Soporta formatos:
    - Línea separada: "- CGT1 - Descripción - Nivel"
    - Línea única: "CGT1 - Descripción - Nivel - CGT2 - Descripción - Nivel - ..."
    
    Retorna: [{'code': 'CGT1', 'description': 'Descripción', 'level': 'Nivel'}, ...]
    """
    competencies = []
    
    # Paso 1: Encontrar todos los códigos CG en el texto
    # Patrón: CG + [cualquier letra: T, S, etc.] + [números]
    code_pattern = r'([Cc][Gg][A-Za-z]\d+)'
    codes = []
    for match in re.finditer(code_pattern, text):
        codes.append((match.group(1), match.start(), match.end()))
    
    if not codes:
        return competencies
    
    # Paso 2: Para cada código CG, extraer description y nivel
    for idx, (code, code_start, code_end) in enumerate(codes):
        # El texto para este código va desde después del código hasta antes del siguiente código
        if idx < len(codes) - 1:
            next_code_start = codes[idx + 1][1]
            block_text = text[code_end:next_code_start]
        else:
            block_text = text[code_end:]
        
        # Extraer la descripción (todo hasta el nivel)
        level_pattern = r'-\s*(Alto|Medio|Bajo)(?=\s*(?:$|-\s*[Cc][Gg][A-Za-z]|\s*\)))'
        level_match = re.search(level_pattern, block_text, re.IGNORECASE)
        
        if level_match:
            # La descripción es todo lo que hay entre el "-" después del código y donde comienza el "- Nivel"
            desc_text = block_text[:level_match.start()].strip()
            # Limpiar el guión inicial si existe
            desc_text = desc_text.lstrip('-').strip()
            # Normalizar espacios (reemplazar múltiples espacios y newlines)
            description = ' '.join(desc_text.split())
            level = level_match.group(1).capitalize()
            
            if description and len(description) > 2:
                competencies.append({
                    'code': code.upper(),
                    'description': description,
                    'level': level
                })
    
    return competencies


def extract_specific_competencies(text: str) -> List[Dict[str, str]]:
    """
    Extrae Competencias Específicas del texto usando un enfoque basado en bloques.
    Soporta formatos:
    - Línea separada: "- CE1 - Descripción - Nivel"
    - Línea única: "CE1 - Descripción - Nivel - CE2 - Descripción - Nivel - ..."
    Retorna: [{'code': 'CE1', 'description': 'Descripción', 'level': 'Nivel'}, ...]
    
    ESTRATEGIA MEJORADA (similar a RAs): Buscar todos los códigos primero, 
    luego extraer los datos entre ellos para evitar capturar a través de límites.
    """
    competencies = []
    
    # Paso 1: Encontrar todos los códigos CE en el texto
    code_pattern = r'([Cc][Ee]\d+)'
    codes = []
    for match in re.finditer(code_pattern, text):
        codes.append((match.group(1), match.start(), match.end()))
    
    if not codes:
        return competencies
    
    # Paso 2: Para cada código CE, extraer description y nivel
    for idx, (code, code_start, code_end) in enumerate(codes):
        # El texto para este código va desde después del código hasta antes del siguiente código
        if idx < len(codes) - 1:
            next_code_start = codes[idx + 1][1]
            block_text = text[code_end:next_code_start]
        else:
            block_text = text[code_end:]
        
        # Extraer la descripción (todo hasta el nivel)
        level_pattern = r'-\s*(Alto|Medio|Bajo)(?=\s*(?:$|-\s*[Cc][Ee]\d+|\s*\)))'
        level_match = re.search(level_pattern, block_text, re.IGNORECASE)
        
        if level_match:
            # La descripción es todo lo que hay entre el "-" después del código y donde comienza el "- Nivel"
            desc_text = block_text[:level_match.start()].strip()
            # Limpiar el guión inicial si existe
            desc_text = desc_text.lstrip('-').strip()
            # Normalizar espacios (reemplazar múltiples espacios y newlines)
            description = ' '.join(desc_text.split())
            level = level_match.group(1).capitalize()
            
            if description and len(description) > 2:
                competencies.append({
                    'code': code.upper(),
                    'description': description,
                    'level': level
                })
    
    return competencies


def extract_learning_outcomes_parsed(text: str) -> List[Dict[str, str]]:
    """
    Extrae Resultados de Aprendizaje (RAx) del texto con descripción completa.
    Formatos soportados:
    - "- RA1 - Descripción" (con guión al inicio)
    - "- RA1. Descripción"
    - "- RA1: Descripción"
    - "RA 1: Descripción" (sin guión, con espacio entre RA y número)
    - "RA1: Descripción" (sin guión, sin espacio)
    - Múltiples líneas de descripción
    
    Retorna: [{'code': 'RA1', 'description': 'Descripción'}, ...]
    """
    outcomes = []
    
    # Patrón más flexible: busca RA en cualquier posición (inicio de línea o después de espacio)
    # Ahora captura:
    # (?:^|\n|-\s|) - Inicio de línea, salto, o después de guión
    # \s*(?:[-•]\s*)? - Espacios opcionales y guión/bullet opcional
    # ([Rr][Aa]\s*\d+) - RA con espacio opcional (CAPTURA 1)
    # \s*[-:.]?\s* - Espacios y separadores opcionales
    # ((?:[^\n]|(?:\n(?!\s*[Rr][Aa]\s*\d+)))*?) - Descripción que puede incluir saltos de línea (CAPTURA 2)
    
    # Una alternativa: buscar cada RA y luego extraer su descripción hasta el siguiente
    ra_pattern = r'(?:^|\n)\s*(?:[-•]\s*)?([Rr][Aa]\s*\d+)\s*[-:.]?\s*'
    ra_matches = list(re.finditer(ra_pattern, text, re.MULTILINE))
    
    if not ra_matches:
        # Si no hay match, intentar patrón más simple (sin línea nueva)
        ra_pattern = r'(\bRA\s*\d+)\s*[-:.]?\s*'
        ra_matches = list(re.finditer(ra_pattern, text))
    
    seen_codes = set()
    for idx, match in enumerate(ra_matches):
        code = match.group(1).upper().replace(' ', '')  # RA 1 -> RA1
        
        # Extraer descripción desde el final del código hasta:
        # - El siguiente RA, o
        # - El final del texto, o
        # - Hasta 1000 caracteres (límite razonable)
        start_desc = match.end()
        
        if idx + 1 < len(ra_matches):
            end_desc = ra_matches[idx + 1].start()
        else:
            end_desc = len(text)
        
        description = text[start_desc:end_desc].strip()
        
        # Limpiar descripción: convertir saltos de línea en espacios y normalizar espacios
        description = ' '.join(description.split())
        
        # Limpiar descripción de separadores leading
        description = re.sub(r'^[-:.\s]+', '', description).strip()
        
        # Limitar a 1000 caracteres si es muy larga
        if len(description) > 1000:
            description = description[:1000].rsplit(' ', 1)[0] + '...'
        
        # Evitar duplicados y descripciones muy cortas
        if code not in seen_codes and description and len(description) > 2:
            seen_codes.add(code)
            outcomes.append({
                'code': code,
                'description': description
            })
    
    # Renumerar RAs para que sean consecutivos (RA1, RA2, RA3, ...)
    outcomes = normalize_learning_outcomes(outcomes)
    return outcomes


def normalize_learning_outcomes(outcomes: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """
    Normaliza los RAs para que sean consecutivos (RA1, RA2, RA3, ...).
    Si hay RAs con números no consecutivos (ej: RA1, RA3, RA5), 
    se renumeran como RA1, RA2, RA3.
    
    Args:
        outcomes: Lista de RAs extraídos
        
    Returns:
        Lista de RAs con códigos consecutivos
    """
    if not outcomes:
        return []
    
    # Renumerar cada RA con su posición en la lista
    normalized = []
    for idx, outcome in enumerate(outcomes, start=1):
        normalized.append({
            'code': f'RA{idx}',
            'description': outcome.get('description', '')
        })
    
    return normalized


def extract_competencies_from_table(doc: Document, table_idx: int) -> Tuple[List[Dict[str, str]], List[Dict[str, str]]]:
    """
    Extrae competencias genéricas y específicas desde una tabla específica.
    La tabla generalmente está entre 'OBJETIVOS' y 'CONTENIDOS DE LA ASIGNATURA'.
    Captura TODAS las variantes genéricas: CGT, CGS, etc.
    Usa enfoque basado en bloques (similar a RAs) para evitar capturar entre límites.
    """
    gen_comp = []
    spec_comp = []
    
    if table_idx >= len(doc.tables):
        return gen_comp, spec_comp
    
    table = doc.tables[table_idx]
    
    # Buscar en todas las celdas de la tabla
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                continue
            
            # ===  COMPETENCIAS GENÉRICAS (CGT, CGS, y cualquier CG[LETRA]) ===
            # Paso 1: Encontrar todos los códigos CG (CGT, CGS, CGA, etc.)
            cg_codes = []
            for match in re.finditer(r'([Cc][Gg][A-Za-z]\d+)', text):
                cg_codes.append((match.group(1), match.start(), match.end()))
            
            # Paso 2: Extraer para cada código
            for idx, (code, code_start, code_end) in enumerate(cg_codes):
                # El bloque va desde después del código hasta antes del siguiente código
                if idx < len(cg_codes) - 1:
                    next_code_start = cg_codes[idx + 1][1]
                    block_text = text[code_end:next_code_start]
                else:
                    block_text = text[code_end:]
                
                # Buscar el nivel en este bloque
                level_match = re.search(r'-\s*(Alto|Medio|Bajo)(?=\s*(?:-\s*[Cc][Gg]|$|\n))', block_text, re.IGNORECASE)
                if level_match:
                    # La descripción es todo antes del "- Nivel"
                    desc_text = block_text[:level_match.start()].strip()
                    desc_text = desc_text.lstrip('-').strip()
                    description = ' '.join(desc_text.split())
                    level = level_match.group(1).capitalize()
                    
                    if description and len(description) > 2 and code.upper() not in [c['code'] for c in gen_comp]:
                        gen_comp.append({
                            'code': code.upper(),
                            'description': description,
                            'level': level
                        })
            
            # === COMPETENCIAS ESPECÍFICAS ===
            # Paso 1: Encontrar todos los códigos CE
            ce_codes = []
            for match in re.finditer(r'([Cc][Ee]\d+)', text):
                ce_codes.append((match.group(1), match.start(), match.end()))
            
            # Paso 2: Extraer para cada código
            for idx, (code, code_start, code_end) in enumerate(ce_codes):
                # El bloque va desde después del código hasta antes del siguiente código
                if idx < len(ce_codes) - 1:
                    next_code_start = ce_codes[idx + 1][1]
                    block_text = text[code_end:next_code_start]
                else:
                    block_text = text[code_end:]
                
                # Buscar el nivel en este bloque
                level_match = re.search(r'-\s*(Alto|Medio|Bajo)(?=\s*(?:-\s*[Cc][Ee]|$|\n))', block_text, re.IGNORECASE)
                if level_match:
                    # La descripción es todo antes del "- Nivel"
                    desc_text = block_text[:level_match.start()].strip()
                    desc_text = desc_text.lstrip('-').strip()
                    description = ' '.join(desc_text.split())
                    level = level_match.group(1).capitalize()
                    
                    if description and len(description) > 2 and code.upper() not in [c['code'] for c in spec_comp]:
                        spec_comp.append({
                            'code': code.upper(),
                            'description': description,
                            'level': level
                        })
    
    return gen_comp, spec_comp


def extract_fundamentals_from_table(doc: Document, table_idx: int) -> Tuple[str, str]:
    """
    Extrae secciones de 'Importancia' y 'Perfil Profesional' desde una tabla.
    Típicamente la Tabla 3 en documentos con tablas de fundamentos.
    """
    importance = ""
    professional_profile = ""
    
    if table_idx >= len(doc.tables):
        return importance, professional_profile
    
    table = doc.tables[table_idx]
    
    # Buscar en las celdas de la tabla
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            
            # Buscar "Importancia"
            if "importancia" in text.lower() and "plan" in text.lower():
                # Extraer todo lo que viene después de los dos puntos
                match = re.search(r'importancia[^:]*:\s*(.+?)(?=relación|$)', text, re.IGNORECASE | re.DOTALL)
                if match:
                    importance = match.group(1).strip()
            
            # Buscar "Perfil Profesional"
            if "relación" in text.lower() and "perfil" in text.lower():
                # Extraer todo lo que viene después
                match = re.search(r'relación[^:]*:\s*(.+?)$', text, re.IGNORECASE | re.DOTALL)
                if match:
                    professional_profile = match.group(1).strip()
    
    return importance, professional_profile


def extract_programa_analitico(doc: Document) -> Dict[str, str]:
    """
    Extrae TODOS los campos de la tabla "Programa Analítico":
    Carácter, Régimen, Carga Horaria Tot, Hs Teóricas, Hs Prácticas, Hs Sem
    """
    program_data = {
        'character': '',
        'regime': '',
        'total_hours': '',
        'theoretical_hours': '',
        'practical_hours': '',
        'weekly_hours': '',
    }
    
    for table in doc.tables:
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        # Buscar tabla con "régimen" o "programa"
        if ('régimen' in table_text or 'regimen' in table_text or 'carácter' in table_text):
            if len(table.rows) >= 2:
                header_row = table.rows[0]
                data_row = table.rows[1]
                
                # Mapear encabezados a índices
                headers = {}
                for col_idx, cell in enumerate(header_row.cells):
                    header_lower = cell.text.lower().strip()
                    headers[header_lower] = col_idx
                
                # Extraer valores basados en encabezados
                for header_key, col_idx in headers.items():
                    if col_idx < len(data_row.cells):
                        value = data_row.cells[col_idx].text.strip()
                        
                        if 'carácter' in header_key or 'caracter' in header_key:
                            program_data['character'] = value
                        elif 'régimen' in header_key or 'regimen' in header_key:
                            program_data['regime'] = value
                        elif 'carga horaria tot' in header_key or 'carga horaria total' in header_key:
                            program_data['total_hours'] = value
                        elif 'hs teórica' in header_key or 'hs teorica' in header_key:
                            program_data['theoretical_hours'] = value
                        elif 'hs prácti' in header_key or 'hs practi' in header_key:
                            program_data['practical_hours'] = value
                        elif 'hs sem' in header_key:
                            program_data['weekly_hours'] = value
    
    return program_data


def extract_equipo_docente(doc: Document) -> List[Dict[str, str]]:
    """Extrae tabla de Equipo Docente con profesores, categoría, correo."""
    teaching_team = []
    
    for table in doc.tables:
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        # Buscar tabla de equipo docente
        if ('equipo' in table_text and 'docente' in table_text) or \
           ('profesor' in table_text and 'categoría' in table_text and 'correo' in table_text):
            
            # Encontrar fila de encabezado
            header_row_idx = 0
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.lower() for cell in row.cells])
                if 'profesor' in row_text or 'categoría' in row_text:
                    header_row_idx = row_idx
                    break
            
            # Mapear columnas desde encabezado
            header_row = table.rows[header_row_idx]
            name_idx = 0
            category_idx = 1
            email_idx = 2
            
            for col_idx, cell in enumerate(header_row.cells):
                cell_lower = cell.text.lower().strip()
                if 'profesor' in cell_lower or 'docente' in cell_lower or 'nombre' in cell_lower:
                    name_idx = col_idx
                elif 'categoría' in cell_lower or 'categoria' in cell_lower:
                    category_idx = col_idx
                elif 'correo' in cell_lower or 'email' in cell_lower:
                    email_idx = col_idx
            
            # Extraer filas de datos
            for row_idx in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[row_idx]
                
                if len(row.cells) > max(name_idx, category_idx, email_idx):
                    name = row.cells[name_idx].text.strip() if name_idx < len(row.cells) else ''
                    category = row.cells[category_idx].text.strip() if category_idx < len(row.cells) else ''
                    email = row.cells[email_idx].text.strip() if email_idx < len(row.cells) else ''
                    
                    # Solo agregar si tiene nombre válido
                    if name and name not in ['', '{{doc1}}', '{{doc2}}', '{{doc3}}']:
                        teaching_team.append({
                            'name': name,
                            'category': category,
                            'email': email,
                        })
    
    return teaching_team


def extract_units_from_docx(doc: Document, section_start_idx: int, section_end_idx: int) -> List[Dict[str, str]]:
    """
    Extrae unidades de secciones numeradas (Unidad N°: 1, Unidad N°: 2, etc.).
    Maneja múltiples formatos:
    1. Tablas con filas alternadas: [Unidad N°], [Contenidos]
    2. Párrafos con patrones de unidad
    3. Tablas simples con estructura fila=unidad, columnas con datos
    """
    def parse_unit_block(text: str) -> Tuple[str, str, str]:
        """Parsea el bloque de una unidad en contenido y bibliografias."""
        if not text:
            return "", "", ""

        normalized = normalize_docx_text(text)
        # Separar contenido y bibliografia especifica
        split_keyword = "bibliografia especifica de la unidad"
        if split_keyword in normalized:
            parts = re.split(r'(?i)bibliograf[ií]a\s+espec[ií]fica\s+de\s+la\s+unidad\s*:?', text, maxsplit=1)
            content_raw = parts[0]
            biblio_raw = parts[1] if len(parts) > 1 else ""
        else:
            content_raw = text
            biblio_raw = ""

        # Limpiar prefijo de contenidos
        content = re.sub(r'(?i)contenidos\s*:?', '', content_raw).strip()

        # Separar bibliografia basica y complementaria
        bib_basic = ""
        bib_comp = ""
        if biblio_raw:
            basic_parts = re.split(r'(?i)bibliograf[ií]a\s+basica\s*\(.*?\)\s*:?', biblio_raw, maxsplit=1)
            if len(basic_parts) > 1:
                biblio_after_basic = basic_parts[1]
            else:
                biblio_after_basic = biblio_raw

            comp_parts = re.split(r'(?i)bibliograf[ií]a\s+complementaria\s*:?', biblio_after_basic, maxsplit=1)
            if len(comp_parts) > 1:
                bib_basic = comp_parts[0].strip()
                bib_comp = comp_parts[1].strip()
            else:
                bib_basic = biblio_after_basic.strip()

        return content, bib_basic, bib_comp

    units = []
    
    # Buscar tablas que tengan unidades
    for table_idx, table in enumerate(doc.tables):
        # Detectar si esta tabla tiene unidades
        table_has_units = False
        for row in table.rows:
            for cell in row.cells:
                if re.search(r'unidad\s*n[°º]?', cell.text, re.IGNORECASE):
                    table_has_units = True
                    break
            if table_has_units:
                break
        
        if not table_has_units:
            continue
        
        # Procesar tabla buscando pares de [Unidad/Nombre] + [Contenidos]
        i = 0
        while i < len(table.rows):
            row = table.rows[i]
            row_cells = [cell.text.strip() for cell in row.cells]
            row_text = '\n'.join(row_cells).strip()
            
            # Detectar línea de unidad
            unit_match = re.search(r'unidad\s*n[°º]?\s*:?\s*(\d+)\s*(.*?)(?:\n|$)', row_text, re.IGNORECASE)
            if unit_match:
                unit_number = unit_match.group(1).strip()
                unit_name = unit_match.group(2).strip() if unit_match.group(2) else ''
                
                # Si no hay nombre en el mismo patrón, intentar obtener del siguiente texto
                if not unit_name and len(row_cells) > 1:
                    unit_name = row_cells[1].strip()
                
                content = ""
                bib_basic = ""
                bib_comp = ""
                
                # Ver si el contenido está en la siguiente fila (estructura alternada)
                if i + 1 < len(table.rows):
                    next_row = table.rows[i + 1]
                    next_row_cells = [cell.text.strip() for cell in next_row.cells]
                    next_row_text = '\n'.join(next_row_cells).strip()
                    
                    # Si la siguiente fila tiene "Contenidos", extraer de ahí
                    if 'contenidos' in next_row_text.lower():
                        content, bib_basic, bib_comp = parse_unit_block(next_row_text)
                        i += 2  # Saltar la fila de contenidos
                    else:
                        i += 1
                else:
                    i += 1
                
                # Si hay contenido, agregar la unidad
                if content or unit_name:
                    units.append({
                        'number': unit_number or str(len(units) + 1),
                        'name': unit_name,
                        'contenidos': content,
                        'bib_basic': bib_basic,
                        'bib_comp': bib_comp,
                    })
            else:
                i += 1
    
    # Convertir formato interno al formato esperado
    result = []
    for unit in units:
        result.append({
            'name': unit['name'],
            'contenidos': unit['contenidos'],
            'bib_basica': unit['bib_basic'],
            'bib_complementaria': unit['bib_comp'],
        })
    
    return result


def extract_practicals_from_docx(doc: Document) -> List[Dict[str, str]]:
    """
    Extrae practicales desde sección 5 (PROGRAMA DE TRABAJOS PRÁCTICOS).
    Busca tablas con "Práctico" y "Objetivo".
    """
    def parse_practical_block(text: str) -> Dict[str, Any]:
        if not text:
            return {
                'objective': '',
                'activities': '',
                'materials': '',
                'scope': '',
                'ra_codes': [],
            }

        label_patterns = [
            ('objective', re.compile(r'(?i)objetivo(?:\s*\(.*?\))?\s*:?')),
            ('activities', re.compile(r'(?i)actividades?\s+a\s+desarrollar(?:\s*\(.*?\))?\s*:?')),
            ('materials', re.compile(r'(?i)materiales?\s*:?')),
            ('scope', re.compile(r'(?i)[áa]mbito(?:\s+de\s+pr[áa]ctica)?\s*:?')),
        ]

        def strip_trailing_labels(segment: str) -> str:
            if not segment:
                return ''
            earliest = None
            for _, pattern in label_patterns:
                match = pattern.search(segment)
                if match:
                    earliest = match.start() if earliest is None else min(earliest, match.start())
            if earliest is None:
                return segment.strip()
            return segment[:earliest].strip()

        matches = []
        for label, pattern in label_patterns:
            match = pattern.search(text)
            if match:
                matches.append((match.start(), match.end(), label))

        matches.sort(key=lambda x: x[0])
        segments: Dict[str, str] = {}

        for idx, (start, end, label) in enumerate(matches):
            next_start = matches[idx + 1][0] if idx + 1 < len(matches) else len(text)
            segment = text[end:next_start].strip()
            # Do NOT clean - the strip_trailing_labels was too aggressive and cutting off content
            segments[label] = segment  # strip_trailing_labels(segment)

        if not segments.get('scope') and segments.get('materials'):
            scope_match = label_patterns[3][1].search(segments['materials'])
            if scope_match:
                before = segments['materials'][:scope_match.start()].strip()
                after = segments['materials'][scope_match.end():].strip()
                segments['materials'] = before
                segments['scope'] = after

        objective_raw = segments.get('objective', '')
        ra_text = ''
        line_sections = None

        if not segments.get('activities') and not segments.get('materials') and not segments.get('scope'):
            line_sections = {
                'objective': [],
                'ra': [],
                'activities': [],
                'materials': [],
                'scope': [],
            }
            current = None

            for raw_line in text.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                norm_line = normalize_docx_text(line)

                if re.search(r'\bobjetivo\b', norm_line):
                    current = 'objective'
                    continue
                if re.search(r'resultados?\s+de\s+aprendizaje', norm_line):
                    current = 'ra'
                    continue
                if 'actividades' in norm_line and 'desarrollar' in norm_line:
                    current = 'activities'
                    continue
                if re.search(r'\bmateriales?\b', norm_line):
                    current = 'materials'
                    continue
                if 'ambito' in norm_line:
                    current = 'scope'
                    continue

                if current:
                    line_sections[current].append(line)

            if any(line_sections.values()):
                if line_sections['objective']:
                    objective_raw = '\n'.join(line_sections['objective']).strip()
                if line_sections['ra']:
                    ra_text = '\n'.join(line_sections['ra']).strip()
                if line_sections['activities']:
                    segments['activities'] = '\n'.join(line_sections['activities']).strip()
                if line_sections['materials']:
                    segments['materials'] = '\n'.join(line_sections['materials']).strip()
                if line_sections['scope']:
                    segments['scope'] = '\n'.join(line_sections['scope']).strip()

        ra_codes = []
        ra_source = ra_text or objective_raw
        # Busca RAs en formatos: "RA 3", "RA3", "RA 3.", "RA 3:", etc.
        for match in re.finditer(r'RA\s*(\d+)', ra_source, re.IGNORECASE):
            code = f"RA{match.group(1)}"
            if code not in ra_codes:
                ra_codes.append(code)

        objective_clean = re.split(r'(?i)resultados?\s+de\s+aprendizaje.*', objective_raw, maxsplit=1)[0].strip()
        if not objective_clean and objective_raw:
            cleaned_lines = []
            for line in objective_raw.splitlines():
                if re.search(r'(?i)resultados?\s+de\s+aprendizaje', line):
                    continue
                if re.search(r'\bRA\s*\d+\b', line, re.IGNORECASE):
                    continue
                if line.strip():
                    cleaned_lines.append(line.strip())
            objective_clean = ' '.join(cleaned_lines).strip()
        if not objective_clean and objective_raw:
            objective_clean = objective_raw.strip()

        def clean_scope(value: str) -> str:
            if not value:
                return ''
            cleaned = re.sub(r'(?i)[áa]mbito(?:\s+de\s+pr[áa]ctica)?\s*:?','', value).strip()
            earliest = None
            for _, pattern in label_patterns:
                match = pattern.search(cleaned)
                if match:
                    earliest = match.start() if earliest is None else min(earliest, match.start())
            if earliest is not None:
                cleaned = cleaned[:earliest].strip()
            for line in cleaned.splitlines():
                line = line.strip()
                if line:
                    return line
            return cleaned.strip()

        return {
            'objective': objective_clean,
            'activities': segments.get('activities', '').strip(),
            'materials': segments.get('materials', '').strip(),
            'scope': clean_scope(segments.get('scope', '')),
            'ra_codes': ra_codes,
        }

    practicals = []

    header_pattern = re.compile(r'pr[áa]ctico\s*n[°º]?\s*:?\s*(\d+)\s*(.*)', re.IGNORECASE)

    for table in doc.tables:
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
        table_text_lower = normalize_docx_text(table_text)

        if 'practico' not in table_text_lower or 'objetivo' not in table_text_lower:
            continue

        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]
            header_match = None
            header_cell_idx = None

            for cell_idx, cell_text in enumerate(row_cells):
                match = header_pattern.search(cell_text)
                if match:
                    header_match = match
                    header_cell_idx = cell_idx
                    break

            if not header_match:
                continue

            tp_number = header_match.group(1).strip() if header_match.group(1) else ''
            tp_name = header_match.group(2).strip() if header_match.group(2) else ''
            if not tp_name and header_cell_idx is not None and header_cell_idx + 1 < len(row_cells):
                tp_name = row_cells[header_cell_idx + 1].strip()

            practical = {
                'number': tp_number or str(len(practicals) + 1),
                'name': tp_name,
                'objective': '',
                'activities': '',
                'materials': '',
                'scope': '',
            }

            if row_idx + 1 < len(table.rows):
                block_cells = table.rows[row_idx + 1].cells
                block_text = '\n'.join([extract_text_from_table_cell(cell) for cell in block_cells]).strip()
                parsed = parse_practical_block(block_text)
                practical['objective'] = parsed.get('objective', '')
                practical['activities'] = parsed.get('activities', '')
                practical['materials'] = parsed.get('materials', '')
                practical['scope'] = parsed.get('scope', '')
                if parsed.get('ra_codes'):
                    practical['ra_codes'] = parsed.get('ra_codes')

            practicals.append(practical)

    return practicals


def extract_header_fields_improved(doc: Document, filename: str = "") -> Dict[str, str]:
    """
    Extrae campos del encabezado mejorado.
    """
    fields = {
        'career': '',
        'subject': '',
        'academic_year': '',
        'year_of_career': '',
        'quarter': '',
        'study_plan': '',
    }

    def _collapse_row_cells(row) -> List[str]:
        texts = []
        last = None
        for cell in row.cells:
            text = extract_text_from_table_cell(cell).strip()
            if not text:
                continue
            if text == last:
                continue
            texts.append(text)
            last = text
        return texts

    def _normalize_year_of_career(value: str) -> str:
        if not value:
            return ""
        match = re.search(r'(\d+)', value)
        if match:
            return match.group(1)
        return value.strip()

    def _extract_from_header_paragraphs() -> None:
        for section in doc.sections:
            for header_attr in ('header', 'first_page_header', 'even_page_header'):
                header = getattr(section, header_attr, None)
                if header is None:
                    continue
                for para in header.paragraphs:
                    text = para.text.strip()
                    if not text or ':' not in text:
                        continue

                    normalized = normalize_docx_text(text)
                    value = text.split(':', 1)[1].strip()
                    if not value:
                        continue

                    if 'carrera' in normalized and not fields['career']:
                        fields['career'] = value
                    elif 'asignatura' in normalized and not fields['subject']:
                        fields['subject'] = value
                    elif 'plan' in normalized and not fields['study_plan']:
                        fields['study_plan'] = value
                    elif normalized.startswith('ciclo') and not fields['academic_year']:
                        fields['academic_year'] = value
                    elif 'ano de carrera' in normalized and not fields['year_of_career']:
                        fields['year_of_career'] = _normalize_year_of_career(value)
                    elif normalized.startswith('ano') and not fields['year_of_career']:
                        fields['year_of_career'] = _normalize_year_of_career(value)
                    elif 'cuatrimestre' in normalized and not fields['quarter']:
                        fields['quarter'] = value

    def _extract_from_header_table() -> None:
        tables = list(doc.tables)
        for section in doc.sections:
            for header_attr in ('header', 'first_page_header', 'even_page_header'):
                header = getattr(section, header_attr, None)
                if header is None:
                    continue
                tables.extend(list(getattr(header, 'tables', [])))

        for table in tables:
            table_text = normalize_docx_text(' '.join(
                cell.text for row in table.rows for cell in row.cells
            ))
            if 'escuela de ingenier' not in table_text:
                continue

            for row in table.rows:
                row_cells = _collapse_row_cells(row)
                if not row_cells:
                    continue

                normalized_cells = [normalize_docx_text(text) for text in row_cells]

                for cell_idx, cell_text in enumerate(row_cells):
                    label_raw = None
                    value_raw = None

                    if ':' in cell_text:
                        label_raw, value_raw = cell_text.split(':', 1)
                    else:
                        parts = [p.strip() for p in cell_text.splitlines() if p.strip()]
                        if len(parts) >= 2:
                            label_raw, value_raw = parts[0], ' '.join(parts[1:])

                    if label_raw is None:
                        continue

                    label = normalize_docx_text(label_raw)
                    value = (value_raw or '').strip()

                    # Si el valor esta en otra celda (label y valor separados)
                    if not value:
                        if cell_idx + 1 < len(row_cells):
                            value = row_cells[cell_idx + 1].strip()
                        elif len(row_cells) > 1:
                            value = next((c.strip() for i, c in enumerate(row_cells) if i != cell_idx and c.strip()), '')

                    if not value:
                        continue

                    if 'carrera' in label and not fields['career']:
                        fields['career'] = value
                    elif 'asignatura' in label and not fields['subject']:
                        fields['subject'] = value
                    elif 'plan' in label and not fields['study_plan']:
                        fields['study_plan'] = value
                    elif label.startswith('ciclo') and not fields['academic_year']:
                        fields['academic_year'] = value
                    elif 'ano' in label and 'carrera' in label and not fields['year_of_career']:
                        fields['year_of_career'] = _normalize_year_of_career(value)
                    elif label.startswith('ano') and not fields['year_of_career']:
                        fields['year_of_career'] = _normalize_year_of_career(value)
                    elif 'cuatrimestre' in label and not fields['quarter']:
                        fields['quarter'] = value

                # Row with Escuela de Ingenieria: [Escuela..., Carrera, Plan, Ciclo]
                if any('escuela de ingenier' in text for text in normalized_cells):
                    if not fields['career'] and len(row_cells) >= 2:
                        fields['career'] = row_cells[1]
                    if not fields['study_plan'] and len(row_cells) >= 3:
                        fields['study_plan'] = row_cells[2]
                    if not fields['academic_year'] and len(row_cells) >= 4:
                        fields['academic_year'] = row_cells[3]

            break
    
    # Priorizar encabezado (parrafos y tablas) por sobre otras fuentes.
    _extract_from_header_paragraphs()
    _extract_from_header_table()
    
    # Buscar campos en párrafos iniciales
    for para in doc.paragraphs[:80]:
        text = para.text.strip()
        if not text or ':' not in text:
            continue

        normalized = normalize_docx_text(text)
        value = text.split(':', 1)[1].strip()

        if 'carrera' in normalized and not fields['career']:
            fields['career'] = value
        elif 'asignatura' in normalized and not fields['subject']:
            fields['subject'] = value
        elif 'plan' in normalized and not fields['study_plan']:
            fields['study_plan'] = value
        elif normalized.startswith('ciclo') and not fields['academic_year']:
            fields['academic_year'] = value
        elif 'ano de carrera' in normalized and not fields['year_of_career']:
            fields['year_of_career'] = _normalize_year_of_career(value)
        elif normalized.startswith('ano') and not fields['year_of_career']:
            fields['year_of_career'] = _normalize_year_of_career(value)
        elif 'cuatrimestre' in normalized and not fields['quarter']:
            fields['quarter'] = value

    
    return fields


def import_proposal_from_docx(file_path: str, filename: str = "") -> Dict[str, Any]:
    """
    Importa una propuesta completa desde un DOCX.
    Estrategia: Buscar secciones en párrafos TAMBIÉN en tablas.
    """
    doc = Document(file_path)
    
    # Extraer metadata del encabezado
    header_fields = extract_header_fields_improved(doc, filename)
    
    # Extraer tabla de Programa Analítico (TODOS los campos)
    programa_analitico = extract_programa_analitico(doc)
    
    # Extraer equipo docente
    teaching_team = extract_equipo_docente(doc)
    if teaching_team:
        category_priority = {
            'titular': 0,
            'asociado': 1,
            'adjunto': 2,
            'jtp': 3,
            'ayudante 1o': 4,
            'ayudante 1º': 4,
            'ayudante 1': 4,
            'ayudante': 4,
        }

        def teacher_rank(entry: Dict[str, str], original_idx: int) -> Tuple[int, int]:
            category_raw = entry.get('category') or ''
            normalized = normalize_docx_text(category_raw)
            normalized = normalized.replace('1°', '1o').replace('1º', '1o')
            rank = category_priority.get(normalized, 99)
            return (rank, original_idx)

        teaching_team = [
            entry
            for _, entry in sorted(
                enumerate(teaching_team),
                key=lambda pair: teacher_rank(pair[1], pair[0])
            )
        ]
    teaching_team_str = '; '.join([
        f"{t['name']} ({t['category']})"
        for t in teaching_team if t.get('name')
    ])
    
    # Buscar secciones numeradas en párrafos
    sections = find_section_paragraphs(doc)
    
    # Extraer contenidos mínimos (sección 1)
    contenidos_minimos = ""
    for section_key in sections.keys():
        if 'CONTENIDOS MÍNIMOS' in section_key or 'CONTENIDOS MINIMOS' in section_key:
            start_idx, end_idx = sections[section_key]
            contenidos_minimos = extract_section_content(doc, start_idx, end_idx)
            break
    
    # Si no se encontró en párrafos, buscar en tablas (Tabla 2)
    if not contenidos_minimos and len(doc.tables) > 2:
        table = doc.tables[2]
        contenidos_minimos = extract_text_from_table_cell(table.rows[0].cells[0]) if table.rows else ""
    
    # Extraer fundamentos (sección 2)
    fundamentals = ""
    importance = ""
    professional_profile = ""
    
    # Primero intentar desde tabla (más probable)  
    if len(doc.tables) > 3:
        importance, professional_profile = extract_fundamentals_from_table(doc, 3)
    
    # Si no se encontró, buscar en párrafos
    if not importance:
        for section_key in sections.keys():
            if 'FUNDAMENTOS' in section_key:
                start_idx, end_idx = sections[section_key]
                fundamentals_text = extract_section_content(doc, start_idx, end_idx)
                
                # Separar subsecciones con regex
                import_match = re.search(r'importancia\s+en\s+el\s+plan\s+de\s+estudio\s*:?\s*(.+?)(?=relación\s+con\s+el\s+perfil|$)', 
                                        fundamentals_text, re.IGNORECASE | re.DOTALL)
                if import_match:
                    importance = import_match.group(1).strip()
                
                profile_match = re.search(r'relación\s+con\s+el\s+perfil\s+profesional\s+esperado\s*:?\s*(.+?)$', 
                                         fundamentals_text, re.IGNORECASE | re.DOTALL)
                if profile_match:
                    professional_profile = profile_match.group(1).strip()
                
                fundamentals = fundamentals_text
                break
    
    # Extraer objetivos (sección 3)
    generic_competencies_list = []
    specific_competencies_list = []
    learning_outcomes = []
    
    # Primero intentar desde tabla de competencias (Tabla 4)
    if len(doc.tables) > 4:
        gen_comp_from_table, spec_comp_from_table = extract_competencies_from_table(doc, 4)
        if gen_comp_from_table or spec_comp_from_table:
            generic_competencies_list = gen_comp_from_table
            specific_competencies_list = spec_comp_from_table
    
    # Si no se encontró en tablas, buscar en párrafos
    if not generic_competencies_list:
        for section_key in sections.keys():
            if 'OBJETIVOS' in section_key:
                start_idx, end_idx = sections[section_key]
                objectives_text = extract_section_content(doc, start_idx, end_idx)
                
                generic_competencies_list = extract_generic_competencies(objectives_text)
                specific_competencies_list = extract_specific_competencies(objectives_text)
                learning_outcomes = extract_learning_outcomes_parsed(objectives_text)
                
                break
    
    # Si no se encontraron RAs en párrafos, intentar desde TABLAS
    if not learning_outcomes:
        for table_idx, table in enumerate(doc.tables):
            # Buscar tablas que contengan "RA"
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if 'RA1' in cell_text or 'RA2' in cell_text or 'RA3' in cell_text or 'RA 1' in cell_text:
                        # Encontrada tabla con RAs - usar función mejorada
                        learning_outcomes = extract_learning_outcomes_parsed(cell_text)
                        if learning_outcomes:
                            break
                if learning_outcomes:
                    break
            if learning_outcomes:
                break
            
            # Si ya encontramos RAs, salir del bucle externo
            if learning_outcomes:
                break
    
    # Extraer unidades (sección 4)
    units = extract_units_from_docx(doc, 0, len(doc.paragraphs))
    
    # Extraer prácticos (sección 5)
    practicals = extract_practicals_from_docx(doc)
    
    # Extraer metodología (sección 6)
    methodology = ""
    for section_key in sections.keys():
        if 'METODOLOGÍA' in section_key or 'METODOLOGIA' in section_key:
            start_idx, end_idx = sections[section_key]
            methodology = extract_section_content(doc, start_idx, end_idx)
            table_methodology = extract_section_tables_content(doc, start_idx, end_idx)
            if table_methodology:
                methodology = table_methodology
            break
    if not methodology:
        methodology = extract_labeled_table_content(doc, ['metodologia', 'metodología'])
    
    # Extraer evaluación (sección 7)
    evaluation = ""
    for section_key in sections.keys():
        if 'EVALUACIÓN' in section_key or 'EVALUACION' in section_key:
            start_idx, end_idx = sections[section_key]
            evaluation = extract_section_content(doc, start_idx, end_idx)
            table_evaluation = extract_section_tables_content(doc, start_idx, end_idx)
            if table_evaluation:
                evaluation = table_evaluation
            break
    if not evaluation:
        evaluation = extract_labeled_table_content(doc, ['evaluacion', 'evaluación'])
    
    # Extraer bibliografía (sección 8)
    bibliography = ""
    bibliography_basic = ""
    bibliography_complementary = ""
    for section_key in sections.keys():
        if 'BIBLIOGRAFÍA' in section_key or 'BIBLIOGRAFIA' in section_key:
            start_idx, end_idx = sections[section_key]
            bibliography = extract_section_content(doc, start_idx, end_idx)
            table_bibliography = extract_section_tables_content(doc, start_idx, end_idx)
            if table_bibliography:
                bibliography = table_bibliography
            break
    if not bibliography:
        bibliography = extract_labeled_table_content(doc, ['bibliografia', 'bibliografía'])
    bibliography, bibliography_basic, bibliography_complementary = parse_global_bibliography(bibliography)
    
    # Extraer observaciones (sección 9)
    observations = ""
    for section_key in sections.keys():
        if 'OBSERVACIONES' in section_key:
            start_idx, end_idx = sections[section_key]
            observations = extract_section_content(doc, start_idx, end_idx)
            table_observations = extract_section_tables_content(doc, start_idx, end_idx)
            if table_observations:
                observations = table_observations
            break
    if not observations:
        observations = extract_labeled_table_content(doc, ['observaciones', 'observación', 'observacion'])
    observations = strip_observations_footer(observations)

    quarter_raw = header_fields.get('quarter', '')
    quarter_norm = normalize_docx_text(quarter_raw)
    regime_value = programa_analitico.get('regime', '')
    if quarter_norm:
        if 'anual' in quarter_norm or quarter_norm.strip() == 'a':
            regime_value = 'Anual'
        elif '1' in quarter_norm or '2' in quarter_norm:
            regime_value = 'Cuatrimestral'
    
    # Compilar respuesta final
    data = {
        # Encabezado
        'career': header_fields.get('career', ''),
        'subject': header_fields.get('subject', ''),
        'study_plan': header_fields.get('study_plan', ''),
        'academic_year': header_fields.get('academic_year', ''),
        'year_of_career': header_fields.get('year_of_career', ''),
        'quarter': header_fields.get('quarter', ''),
        
        # Programa Analítico
        'character': programa_analitico.get('character', ''),
        'regime': regime_value,
        'total_hours': programa_analitico.get('total_hours', ''),
        'theoretical_hours': programa_analitico.get('theoretical_hours', ''),
        'practical_hours': programa_analitico.get('practical_hours', ''),
        'weekly_hours': programa_analitico.get('weekly_hours', ''),
        
        # Equipo docente
        'teachers': teaching_team_str,
        'teaching_team': teaching_team,
        
        # Secciones del documento
        'minimum_content': contenidos_minimos,
        'importance': importance.strip() if importance else "",
        'professional_profile': professional_profile.strip() if professional_profile else "",
        'fundamentals': fundamentals,
        'generic_competencies': generic_competencies_list,
        'specific_competencies': specific_competencies_list,
        'learning_outcomes': learning_outcomes,
        
        # Contenido
        'units': units,
        'practicals': practicals,
        'methodology': methodology,
        'evaluation': evaluation,
        'bibliography': bibliography,
        'bibliography_basic': bibliography_basic,
        'bibliography_complementary': bibliography_complementary,
        'observations': observations,
    }
    
    return data
