import os
import glob
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, BackgroundTasks, Body, Request
from fastapi.responses import FileResponse, Response, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func, or_
from . import models, schemas
from .database import SessionLocal, init_db
from .auth import hash_password, verify_password, create_access_token, decode_access_token
from agents import extract as extract_agent
from .docx_import import import_proposal_from_docx
from openai import OpenAI
import shutil
import tempfile
import subprocess
import unicodedata
import re
import hashlib
import json
import ast
import mimetypes
import xml.etree.ElementTree as ET
import textwrap
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from io import BytesIO
from urllib.parse import quote
import requests

app = FastAPI(title="TesisMCD API")

# Load environment variables from backend/.env
backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
env_file = os.path.join(backend_dir, '.env')
load_dotenv(dotenv_path=env_file, override=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:8011",
        "http://127.0.0.1:8011",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

# --- Google Docs import logic and endpoint ---
def get_docx_from_gdoc_url(gdoc_url: str) -> bytes:
    """
    Given a public Google Docs URL, convert to export DOCX and download the file content as bytes.
    """
    # Accept both https://docs.google.com/document/d/ID/edit and https://drive.google.com/open?id=ID
    doc_id = None
    # Try to extract from /d/{id}/
    m = re.search(r"/d/([\w-]+)", gdoc_url)
    if m:
        doc_id = m.group(1)
    else:
        # Try to extract from ?id={id}
        m = re.search(r"[?&]id=([\w-]+)", gdoc_url)
        if m:
            doc_id = m.group(1)
    if not doc_id:
        raise HTTPException(status_code=400, detail="No se pudo extraer el ID del documento de Google Docs.")
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    resp = requests.get(export_url)
    if resp.status_code != 200 or not resp.content or resp.headers.get('Content-Type','').find('application/vnd.openxmlformats-officedocument.wordprocessingml.document') == -1:
        raise HTTPException(status_code=400, detail="No se pudo descargar el DOCX desde Google Docs. ¿El documento es público?")
    return resp.content

@app.post("/proposals/import-gdoc-url")
async def import_proposal_gdoc_url(
    url: str = Body(..., embed=True, description="URL pública de Google Docs")
):
    """
    Importa una propuesta desde un Google Docs público (URL), descargando como DOCX y extrayendo los datos.
    Retorna JSON para previsualización.
    """
    try:
        docx_bytes = get_docx_from_gdoc_url(url)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name
        try:
            extracted_data = import_proposal_from_docx(tmp_path, "imported_from_gdoc.docx")
            extracted_data["gdoc_url"] = url
            return {
                "success": True,
                "data": extracted_data,
                "gdoc_url": url,
                "preview": {
                    "career": extracted_data.get('career', ''),
                    "subject": extracted_data.get('subject', ''),
                    "teachers": extracted_data.get('teachers', ''),
                    "year": extracted_data.get('year_of_career', ''),
                    "quarter": extracted_data.get('quarter', ''),
                    "total_hours": extracted_data.get('total_hours', ''),
                    "theoretical_hours": extracted_data.get('theoretical_hours', ''),
                    "practical_hours": extracted_data.get('practical_hours', ''),
                    "weekly_hours": extracted_data.get('weekly_hours', ''),
                    "regime": extracted_data.get('regime', ''),
                    "units_count": len(extracted_data.get('units', [])),
                    "practicals_count": len(extracted_data.get('practicals', [])),
                    "ra_count": len(extracted_data.get('learning_outcomes', [])),
                    "generic_comp_count": len(extracted_data.get('generic_competencies', [])),
                    "specific_comp_count": len(extracted_data.get('specific_competencies', [])),
                }
            }
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error importando desde Google Docs: {str(e)}")

# Load environment variables from backend/.env (if not already loaded)

UPLOAD_FOLDER = os.getenv("LOCAL_UPLOAD_PATH", "./data/uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
INSTRUMENTS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(UPLOAD_FOLDER)) if not os.path.isabs(UPLOAD_FOLDER) else os.path.dirname(UPLOAD_FOLDER), "instruments")
TEMPLATE_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Template Propuestas.docx"))


class AiPrompt(BaseModel):
    prompt: str


class LoginRequest(BaseModel):
    username: str
    password: str


class ResetPasswordsRequest(BaseModel):
    user_ids: list[int]
    new_password: str = ""
    use_email_as_password: bool = False



class GdocStatusRequest(BaseModel):
    ids: list[int]

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def normalize_subject_name(value: str | None) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).lower()


def normalize_title_name(value: str | None) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).lower()


def split_preview_lines(text: str, max_lines: int = 20, max_chars: int = 3000) -> list[str]:
    if not text:
        return []
    clean = str(text).replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(clean) > max_chars:
        clean = clean[:max_chars]
    lines = [line.strip() for line in clean.split("\n") if line.strip()]
    if not lines:
        return []
    return lines[:max_lines]


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return ""
    try:
        document = Document(BytesIO(content))
        chunks: list[str] = []
        chunks.extend([p.text for p in document.paragraphs if str(p.text or "").strip()])
        for table in document.tables:
            for row in table.rows:
                row_cells = [str(cell.text or "").strip() for cell in row.cells if str(cell.text or "").strip()]
                if row_cells:
                    chunks.append(" | ".join(row_cells))
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _extract_xlsx_text(content: bytes) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception:
        return ""
    try:
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        chunks: list[str] = []
        for ws in wb.worksheets:
            chunks.append(f"[Hoja] {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if values:
                    chunks.append(" | ".join(values))
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _extract_pdf_text(content: bytes, max_pages: int | None = None) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            return ""
    try:
        reader = PdfReader(BytesIO(content))
        chunks: list[str] = []
        for page_index, page in enumerate(reader.pages):
            if max_pages is not None and page_index >= max_pages:
                break
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text)
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _extract_text_from_image_ocr(content: bytes) -> str:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except Exception:
        return ""
    try:
        tesseract_cmd = resolve_tesseract_executable()
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        image = Image.open(BytesIO(content))
        return (pytesseract.image_to_string(image, lang="spa+eng") or "").strip()
    except Exception:
        return ""


def resolve_tesseract_executable() -> str | None:
    env_candidates = [
        os.getenv("TESSERACT_CMD"),
        os.getenv("PYTESSERACT_TESSERACT_CMD"),
    ]
    for candidate in env_candidates:
        if candidate and os.path.isfile(candidate):
            return candidate

    in_path = shutil.which("tesseract")
    if in_path:
        return in_path

    common_paths = [
        r"C:/Program Files/Tesseract-OCR/tesseract.exe",
        r"C:/Program Files (x86)/Tesseract-OCR/tesseract.exe",
    ]
    for candidate in common_paths:
        if os.path.isfile(candidate):
            return candidate
    return None


def resolve_poppler_bin_dir() -> str | None:
    env_candidate = os.getenv("POPPLER_PATH")
    if env_candidate:
        if os.path.isdir(env_candidate):
            return env_candidate
        if os.path.isfile(env_candidate):
            return os.path.dirname(env_candidate)

    pdftoppm_in_path = shutil.which("pdftoppm")
    if pdftoppm_in_path:
        return os.path.dirname(pdftoppm_in_path)

    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        winget_glob = os.path.join(
            local_app_data,
            "Microsoft",
            "WinGet",
            "Packages",
            "oschwartz10612.Poppler*",
            "poppler-*",
            "Library",
            "bin",
        )
        matches = sorted(glob.glob(winget_glob), reverse=True)
        for candidate in matches:
            if os.path.isdir(candidate) and os.path.isfile(os.path.join(candidate, "pdftoppm.exe")):
                return candidate

    common_dirs = [
        r"C:/Program Files/poppler/Library/bin",
        r"C:/Program Files (x86)/poppler/Library/bin",
    ]
    for candidate in common_dirs:
        if os.path.isdir(candidate) and os.path.isfile(os.path.join(candidate, "pdftoppm.exe")):
            return candidate

    return None


def _extract_pdf_text_with_ocr(content: bytes, max_pages: int | None = None) -> str:
    try:
        import pytesseract  # type: ignore
    except Exception:
        return ""

    # 1) Try pdf2image + Poppler (when available)
    try:
        tesseract_cmd = resolve_tesseract_executable()
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        poppler_bin_dir = resolve_poppler_bin_dir()
        kwargs = {"dpi": 200}
        if poppler_bin_dir:
            kwargs["poppler_path"] = poppler_bin_dir

        from pdf2image import convert_from_bytes  # type: ignore
        images = convert_from_bytes(content, **kwargs)
        chunks: list[str] = []
        for page_index, image in enumerate(images):
            if max_pages is not None and page_index >= max_pages:
                break
            txt = (pytesseract.image_to_string(image, lang="spa+eng") or "").strip()
            if txt:
                chunks.append(txt)
        merged = "\n".join(chunks).strip()
        if merged:
            return merged
    except Exception:
        pass

    # 2) Fallback: pypdfium2 render (no Poppler required)
    try:
        import pypdfium2 as pdfium  # type: ignore
        pdf = pdfium.PdfDocument(content)
        chunks: list[str] = []
        for page_index in range(len(pdf)):
            if max_pages is not None and page_index >= max_pages:
                break
            page = pdf[page_index]
            bitmap = page.render(scale=2.0)
            pil_image = bitmap.to_pil()
            txt = (pytesseract.image_to_string(pil_image, lang="spa+eng") or "").strip()
            if txt:
                chunks.append(txt)
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def extract_document_preview(filename: str, content: bytes, preview_mode: bool = False) -> dict:
    name = str(filename or "").strip().lower()
    guessed_mime, _ = mimetypes.guess_type(name)

    text = ""
    method = "none"
    ocr_applied = False
    pdf_text_pages = 1 if preview_mode else None
    pdf_ocr_pages = 1 if preview_mode else None

    if name.endswith(".docx"):
        text = _extract_docx_text(content)
        method = "docx"
    elif name.endswith(".xlsx") or name.endswith(".xlsm") or name.endswith(".xls"):
        text = _extract_xlsx_text(content)
        method = "xlsx"
    elif name.endswith(".pdf"):
        text = _extract_pdf_text(content, max_pages=pdf_text_pages)
        method = "pdf-text"
        if len(text.strip()) < 40:
            has_tesseract = bool(resolve_tesseract_executable())
            if has_tesseract:
                ocr_text = _extract_pdf_text_with_ocr(content, max_pages=pdf_ocr_pages)
                if ocr_text:
                    text = ocr_text
                    method = "pdf-ocr"
                    ocr_applied = True
                else:
                    method = "pdf-ocr-empty"
            else:
                method = "pdf-ocr-unavailable"
    elif guessed_mime and guessed_mime.startswith("image/"):
        text = _extract_text_from_image_ocr(content)
        method = "image-ocr"
        ocr_applied = bool(text)
    else:
        # Fallback for txt/csv/json and unknown text-like files
        try:
            text = content.decode("utf-8")
            method = "text-utf8"
        except Exception:
            try:
                text = content.decode("latin-1")
                method = "text-latin1"
            except Exception:
                text = ""
                method = "binary-unsupported"

    preview_lines = split_preview_lines(text)
    return {
        "extraction_method": method,
        "ocr_applied": ocr_applied,
        "extracted_char_count": len(text or ""),
        "preview_lines": preview_lines,
    }


def normalize_hash_value(value):
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value.strip())
    if isinstance(value, list):
        return [normalize_hash_value(item) for item in value]
    if isinstance(value, dict):
        return {key: normalize_hash_value(value[key]) for key in sorted(value.keys())}
    return value


def compute_payload_hash(payload: dict) -> str:
    normalized = normalize_hash_value(payload)
    encoded = json.dumps(normalized, sort_keys=True, separators=(",", ":"), ensure_ascii=True)
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()


def extract_gdoc_payload(gdoc_url: str) -> tuple[str, str, str, dict]:
    docx_bytes = get_docx_from_gdoc_url(gdoc_url)
    extracted_subject = ""
    extracted_title = ""
    extracted_payload: dict = {}
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name
    try:
        extracted_payload = import_proposal_from_docx(tmp_path, "imported_from_gdoc.docx")
        extracted_subject = extracted_payload.get("subject") or ""
        extracted_title = extracted_payload.get("title") or ""
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    doc_hash = compute_payload_hash(build_extracted_snapshot(extracted_payload))
    return extracted_subject, extracted_title, doc_hash, extracted_payload


def extract_drive_folder_id(url_or_id: str | None) -> str | None:
    raw = str(url_or_id or "").strip()
    if not raw:
        return None
    folder_match = re.search(r"/folders/([\w-]+)", raw)
    if folder_match:
        return folder_match.group(1)
    query_match = re.search(r"[?&]id=([\w-]+)", raw)
    if query_match:
        return query_match.group(1)
    if re.fullmatch(r"[\w-]+", raw):
        return raw
    return None


def extract_drive_file_id(url_or_id: str | None) -> str | None:
    raw = str(url_or_id or "").strip()
    if not raw:
        return None
    file_match = re.search(r"/file/d/([\w-]+)", raw)
    if file_match:
        return file_match.group(1)
    open_match = re.search(r"/open\?id=([\w-]+)", raw)
    if open_match:
        return open_match.group(1)
    query_match = re.search(r"[?&]id=([\w-]+)", raw)
    if query_match:
        return query_match.group(1)
    if re.fullmatch(r"[\w-]+", raw):
        return raw
    return None


def load_gcp_service_account_info() -> dict:
    raw = (os.getenv("GCP_SERVICE_ACCOUNT_JSON") or "").strip()
    if not raw:
        raise HTTPException(status_code=500, detail="Falta configurar GCP_SERVICE_ACCOUNT_JSON en backend/.env")

    if os.path.exists(raw):
        try:
            with open(raw, "r", encoding="utf-8") as handler:
                return json.load(handler)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"No se pudo leer el archivo de service account: {exc}")

    try:
        return json.loads(raw)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"GCP_SERVICE_ACCOUNT_JSON inválido: {exc}")


def load_google_oauth_user_credentials(required: bool = False):
    client_id = (os.getenv("GOOGLE_OAUTH_CLIENT_ID") or "").strip()
    client_secret = (os.getenv("GOOGLE_OAUTH_CLIENT_SECRET") or "").strip()
    refresh_token = (os.getenv("GOOGLE_OAUTH_REFRESH_TOKEN") or "").strip()
    access_token = (os.getenv("GOOGLE_OAUTH_ACCESS_TOKEN") or "").strip() or None
    token_uri = (os.getenv("GOOGLE_OAUTH_TOKEN_URI") or "https://oauth2.googleapis.com/token").strip()

    has_any = any([client_id, client_secret, refresh_token, access_token])
    has_required = bool(client_id and client_secret and refresh_token)

    if not has_required:
        if required:
            raise HTTPException(
                status_code=500,
                detail=(
                    "Modo OAuth de usuario activo, pero faltan variables: "
                    "GOOGLE_OAUTH_CLIENT_ID, GOOGLE_OAUTH_CLIENT_SECRET, GOOGLE_OAUTH_REFRESH_TOKEN"
                ),
            )
        if has_any:
            raise HTTPException(
                status_code=500,
                detail=(
                    "Configuración OAuth de usuario incompleta. Debes definir: "
                    "GOOGLE_OAUTH_CLIENT_ID, GOOGLE_OAUTH_CLIENT_SECRET, GOOGLE_OAUTH_REFRESH_TOKEN"
                ),
            )
        return None

    try:
        from google.oauth2.credentials import Credentials
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Faltan dependencias para OAuth de Google. Instala google-auth y google-api-python-client",
        )

    return Credentials(
        token=access_token,
        refresh_token=refresh_token,
        token_uri=token_uri,
        client_id=client_id,
        client_secret=client_secret,
        scopes=["https://www.googleapis.com/auth/drive"],
    )


def get_google_drive_service():
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Faltan dependencias de Google Drive. Instala google-api-python-client y google-auth",
        )

    auth_mode = (os.getenv("GDRIVE_AUTH_MODE") or "auto").strip().lower()
    oauth_aliases = {"oauth", "oauth-user", "user"}
    service_aliases = {"service-account", "service", "sa"}

    def build_service(credentials):
        return build("drive", "v3", credentials=credentials)

    if auth_mode in oauth_aliases:
        oauth_credentials = load_google_oauth_user_credentials(required=True)
        try:
            return build_service(oauth_credentials)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"No se pudo inicializar Google Drive API (OAuth usuario): {exc}")

    if auth_mode in service_aliases:
        service_account_info = load_gcp_service_account_info()
        try:
            credentials = service_account.Credentials.from_service_account_info(
                service_account_info,
                scopes=["https://www.googleapis.com/auth/drive"],
            )
            return build_service(credentials)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"No se pudo inicializar Google Drive API (Service Account): {exc}")

    oauth_error = None
    service_error = None

    try:
        oauth_credentials = load_google_oauth_user_credentials(required=False)
        if oauth_credentials is not None:
            return build_service(oauth_credentials)
    except HTTPException as exc:
        oauth_error = exc.detail
    except Exception as exc:
        oauth_error = str(exc)

    try:
        service_account_info = load_gcp_service_account_info()
        credentials = service_account.Credentials.from_service_account_info(
            service_account_info,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
        return build_service(credentials)
    except HTTPException as exc:
        service_error = exc.detail
    except Exception as exc:
        service_error = str(exc)

    raise HTTPException(
        status_code=500,
        detail=(
            "No se pudo inicializar Google Drive API. "
            f"OAuth usuario: {oauth_error or 'no configurado'}. "
            f"Service Account: {service_error or 'no configurado'}."
        ),
    )


@app.get("/drive-auth-debug")
def drive_auth_debug():
    service = get_google_drive_service()
    try:
        about = service.about().get(
            fields="user(displayName,emailAddress),storageQuota(limit,usage,usageInDrive,usageInDriveTrash)"
        ).execute()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"No se pudo consultar datos de Drive: {exc}")

    user = about.get("user") or {}
    quota = about.get("storageQuota") or {}

    def to_int(value):
        try:
            return int(value)
        except Exception:
            return None

    return {
        "auth_mode": (os.getenv("GDRIVE_AUTH_MODE") or "auto").strip().lower(),
        "user": {
            "display_name": user.get("displayName"),
            "email": user.get("emailAddress"),
        },
        "storage_quota": {
            "limit": to_int(quota.get("limit")),
            "usage": to_int(quota.get("usage")),
            "usage_in_drive": to_int(quota.get("usageInDrive")),
            "usage_in_drive_trash": to_int(quota.get("usageInDriveTrash")),
        },
        "raw": about,
    }


# ── OAuth Google Re-Auth ──────────────────────────────────────────────────────

_OAUTH_SCOPES = ["https://www.googleapis.com/auth/drive"]
_OAUTH_REDIRECT_URI = "http://127.0.0.1:8011/auth/google/callback"


def _build_oauth_flow(redirect_uri: str = _OAUTH_REDIRECT_URI):
    """Crea un Flow de OAuth usando oauth-client.json o variables de entorno."""
    from google_auth_oauthlib.flow import Flow

    client_secrets_path = os.path.join(backend_dir, "secrets", "oauth-client.json")
    if os.path.exists(client_secrets_path):
        return Flow.from_client_secrets_file(
            client_secrets_path,
            scopes=_OAUTH_SCOPES,
            redirect_uri=redirect_uri,
        )
    # Fallback a variables de entorno
    client_id = os.getenv("GOOGLE_OAUTH_CLIENT_ID", "").strip()
    client_secret = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET", "").strip()
    if not client_id or not client_secret:
        raise HTTPException(status_code=500, detail="No se encontraron credenciales OAuth de Google (oauth-client.json o variables de entorno).")
    client_config = {
        "installed": {
            "client_id": client_id,
            "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [redirect_uri],
        }
    }
    return Flow.from_client_config(client_config, scopes=_OAUTH_SCOPES, redirect_uri=redirect_uri)


@app.get("/auth/google/authorize")
def google_auth_authorize():
    """Genera la URL de autorización OAuth para Google Drive y la devuelve al frontend."""
    try:
        flow = _build_oauth_flow()
    except ImportError:
        raise HTTPException(status_code=500, detail="google-auth-oauthlib no está instalado.")
    auth_url, _ = flow.authorization_url(
        access_type="offline",
        prompt="consent",
        include_granted_scopes="true",
    )
    return {"auth_url": auth_url}


@app.get("/auth/google/callback", response_class=HTMLResponse)
def google_auth_callback(code: str = None, error: str = None):
    """
    Callback OAuth de Google. Intercambia el código por tokens,
    actualiza GOOGLE_OAUTH_REFRESH_TOKEN en backend/.env y recarga las variables.
    """
    _SUCCESS_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>Autorización exitosa – TesisMCD</title>
  <style>
    body {{ font-family: system-ui, sans-serif; display: flex; align-items: center;
           justify-content: center; min-height: 100vh; margin: 0; background: #f0fdf4; }}
    .card {{ background: white; border-radius: 12px; padding: 40px; text-align: center;
             box-shadow: 0 4px 24px rgba(0,0,0,.08); max-width: 400px; }}
    h1 {{ color: #15803d; margin-top: 0; }}
    p {{ color: #555; line-height: 1.6; }}
    .icon {{ font-size: 48px; }}
  </style>
</head>
<body>
  <div class="card">
    <div class="icon">✅</div>
    <h1>Autorización exitosa</h1>
    <p>La aplicación ya tiene acceso a Google Drive.<br>
       Podés cerrar esta ventana y reintentar la operación.</p>
    <script>setTimeout(() => window.close(), 3000);</script>
  </div>
</body>
</html>"""

    _ERROR_HTML_TPL = """<!DOCTYPE html>
<html lang="es">
<head><meta charset="utf-8"><title>Error de autorización</title>
<style>body{{font-family:system-ui,sans-serif;display:flex;align-items:center;
justify-content:center;min-height:100vh;margin:0;background:#fff5f5;}}
.card{{background:white;border-radius:12px;padding:40px;text-align:center;
box-shadow:0 4px 24px rgba(0,0,0,.08);max-width:400px;}}
h1{{color:#c62828;margin-top:0;}}p{{color:#555;}}</style></head>
<body><div class="card"><div style="font-size:48px">❌</div>
<h1>Error de autorización</h1><p>{msg}</p></div></body></html>"""

    if error:
        return HTMLResponse(_ERROR_HTML_TPL.format(msg=f"Google reportó: {error}"), status_code=400)
    if not code:
        return HTMLResponse(_ERROR_HTML_TPL.format(msg="No se recibió el código de autorización."), status_code=400)

    try:
        flow = _build_oauth_flow()
    except ImportError:
        return HTMLResponse(_ERROR_HTML_TPL.format(msg="google-auth-oauthlib no está instalado."), status_code=500)
    except HTTPException as exc:
        return HTMLResponse(_ERROR_HTML_TPL.format(msg=exc.detail), status_code=500)

    try:
        import os as _os
        _os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"  # permite http en localhost
        flow.fetch_token(code=code)
    except Exception as exc:
        return HTMLResponse(_ERROR_HTML_TPL.format(msg=f"No se pudo intercambiar el código: {exc}"), status_code=500)

    credentials = flow.credentials
    new_refresh_token = credentials.refresh_token

    if new_refresh_token:
        # Actualizar GOOGLE_OAUTH_REFRESH_TOKEN en el archivo .env
        try:
            env_content = ""
            if os.path.exists(env_file):
                with open(env_file, "r", encoding="utf-8") as f:
                    env_content = f.read()

            import re as _re
            token_line = f"GOOGLE_OAUTH_REFRESH_TOKEN={new_refresh_token}"
            if _re.search(r"^GOOGLE_OAUTH_REFRESH_TOKEN\s*=", env_content, _re.MULTILINE):
                env_content = _re.sub(
                    r"^GOOGLE_OAUTH_REFRESH_TOKEN\s*=.*$",
                    token_line,
                    env_content,
                    flags=_re.MULTILINE,
                )
            else:
                env_content += f"\n{token_line}\n"

            with open(env_file, "w", encoding="utf-8") as f:
                f.write(env_content)

            load_dotenv(dotenv_path=env_file, override=True)
        except Exception as exc:
            return HTMLResponse(_ERROR_HTML_TPL.format(msg=f"Token obtenido pero no se pudo guardar en .env: {exc}"), status_code=500)

    return HTMLResponse(_SUCCESS_HTML)


def resolve_drive_settings_for_proposal(db: Session, proposal: models.Proposal) -> models.DriveSettings | None:
    if not proposal.career:
        return None

    study_plan = (proposal.study_plan or "").strip()
    if study_plan:
        scoped = db.query(models.DriveSettings).filter(
            models.DriveSettings.career == proposal.career,
            models.DriveSettings.plan_name == study_plan,
        ).first()
        if scoped:
            return scoped

    return db.query(models.DriveSettings).filter(
        models.DriveSettings.career == proposal.career,
        models.DriveSettings.plan_name.is_(None),
    ).first()


def build_proposal_docx_filename(proposal: models.Proposal) -> str:
    year = proposal.year_of_career or "0"
    quarter_raw = proposal.quarter or "0"
    subject = proposal.subject or proposal.title or "Sin_Asignatura"

    quarter_lower = str(quarter_raw).lower()
    if "anual" in quarter_lower or quarter_lower.strip() == "a":
        quarter = "A"
    elif "1" in quarter_lower or "primer" in quarter_lower:
        quarter = "1°"
    elif "2" in quarter_lower or "segundo" in quarter_lower:
        quarter = "2°"
    else:
        quarter = quarter_raw

    subject_clean = re.sub(r'[<>:"/\\|?*]', '', subject)
    year_display = f"{year}°" if str(year).strip() else "0°"
    return f"{year_display}_{quarter} - {subject_clean}.docx"


def build_download_headers(filename: str) -> dict:
    ascii_fallback = unicodedata.normalize("NFKD", filename).encode("ascii", "ignore").decode("ascii")
    ascii_fallback = ascii_fallback or "export"
    encoded = quote(filename)
    return {
        "Content-Disposition": f"attachment; filename=\"{ascii_fallback}\"; filename*=utf-8''{encoded}"
    }


def build_proposal_export_basename(proposal: models.Proposal) -> str:
    return os.path.splitext(build_proposal_docx_filename(proposal))[0]


def sanitize_xml_tag(tag: str) -> str:
    candidate = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(tag or "item"))
    if not candidate:
        candidate = "item"
    if not re.match(r"^[a-zA-Z_]", candidate):
        candidate = f"n_{candidate}"
    return candidate


def append_xml_value(parent: ET.Element, key: str, value) -> None:
    tag = sanitize_xml_tag(key)

    if isinstance(value, dict):
        node = ET.SubElement(parent, tag)
        for child_key, child_value in value.items():
            append_xml_value(node, child_key, child_value)
        return

    if isinstance(value, list):
        node = ET.SubElement(parent, tag)
        for item in value:
            append_xml_value(node, "item", item)
        return

    node = ET.SubElement(parent, tag)
    node.text = "" if value is None else str(value)


def build_proposal_xml_bytes(payload: dict) -> bytes:
    root = ET.Element("proposal")
    for key, value in payload.items():
        append_xml_value(root, key, value)

    xml_buffer = BytesIO()
    ET.ElementTree(root).write(xml_buffer, encoding="utf-8", xml_declaration=True)
    return xml_buffer.getvalue()


def resolve_soffice_executable() -> str | None:
    env_candidate = os.getenv("SOFFICE_PATH")
    if env_candidate and os.path.isfile(env_candidate):
        return env_candidate

    for cmd in ("soffice", "soffice.exe"):
        in_path = shutil.which(cmd)
        if in_path:
            return in_path

    common_paths = [
        r"C:/Program Files/LibreOffice/program/soffice.exe",
        r"C:/Program Files (x86)/LibreOffice/program/soffice.exe",
    ]
    for candidate in common_paths:
        if os.path.isfile(candidate):
            return candidate

    return None


def convert_docx_to_pdf_with_docx2pdf(docx_path: str, pdf_path: str) -> tuple[str | None, str | None]:
    try:
        from docx2pdf import convert as docx2pdf_convert  # type: ignore
    except Exception as exc:
        return None, str(exc)

    try:
        docx2pdf_convert(docx_path, pdf_path)
    except Exception as exc:
        return None, str(exc)

    if os.path.exists(pdf_path):
        return pdf_path, None
    return None, "docx2pdf finalizó pero no generó archivo PDF"


def convert_docx_to_pdf(docx_path: str) -> tuple[str, str]:
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    converted_by_word, word_error = convert_docx_to_pdf_with_docx2pdf(docx_path, pdf_path)
    if converted_by_word:
        return converted_by_word, "word"

    soffice = resolve_soffice_executable()
    if not soffice:
        raise HTTPException(
            status_code=500,
            detail=(
                "No se pudo exportar a PDF manteniendo el formato DOCX. "
                "Instalá LibreOffice (soffice en PATH) o Microsoft Word con docx2pdf. "
                f"Detalle docx2pdf: {word_error or 'no disponible'}"
            ),
        )

    output_dir = os.path.dirname(docx_path)
    command = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        output_dir,
        docx_path,
    ]

    try:
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
    except subprocess.TimeoutExpired:
        raise HTTPException(
            status_code=500,
            detail="No se pudo convertir DOCX a PDF: tiempo de espera agotado en LibreOffice.",
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"No se pudo convertir DOCX a PDF: {str(exc)}",
        )

    if result.returncode != 0 or not os.path.exists(pdf_path):
        error = (result.stderr or result.stdout or "").strip()
        raise HTTPException(
            status_code=500,
            detail=f"No se pudo convertir DOCX a PDF. {error or 'Sin detalle adicional.'}",
        )

    return pdf_path, "libreoffice"


def convert_docx_to_pdf_with_libreoffice(docx_path: str) -> str:
    soffice = resolve_soffice_executable()
    if not soffice:
        raise HTTPException(
            status_code=500,
            detail="No se pudo exportar a PDF: LibreOffice (soffice) no está disponible.",
        )

    output_dir = os.path.dirname(docx_path)
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    command = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        output_dir,
        docx_path,
    ]

    try:
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
    except subprocess.TimeoutExpired:
        raise HTTPException(
            status_code=500,
            detail="No se pudo convertir DOCX a PDF: tiempo de espera agotado en LibreOffice.",
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"No se pudo convertir DOCX a PDF con LibreOffice: {str(exc)}",
        )

    if result.returncode != 0 or not os.path.exists(pdf_path):
        error = (result.stderr or result.stdout or "").strip()
        raise HTTPException(
            status_code=500,
            detail=f"No se pudo convertir DOCX a PDF con LibreOffice. {error or 'Sin detalle adicional.'}",
        )

    return pdf_path


def make_json_compatible(value):
    if isinstance(value, datetime):
        return value.isoformat()

    if isinstance(value, dict):
        return {str(key): make_json_compatible(item) for key, item in value.items()}

    if isinstance(value, list):
        return [make_json_compatible(item) for item in value]

    if isinstance(value, tuple):
        return [make_json_compatible(item) for item in value]

    return value


def _format_export_field_label(key: str) -> str:
    cleaned = str(key or "").strip().replace("_", " ")
    return cleaned.capitalize() if cleaned else "Campo"


def _flatten_export_payload_lines(value, level: int = 0, key_name: str | None = None) -> list[str]:
    indent = "  " * max(level, 0)
    lines: list[str] = []

    if isinstance(value, dict):
        if key_name is not None:
            lines.append(f"{indent}{_format_export_field_label(key_name)}:")
        next_level = level + 1 if key_name is not None else level
        for child_key, child_value in value.items():
            lines.extend(_flatten_export_payload_lines(child_value, next_level, str(child_key)))
        return lines

    if isinstance(value, list):
        if key_name is not None:
            lines.append(f"{indent}{_format_export_field_label(key_name)}:")
        next_level = level + 1 if key_name is not None else level
        if not value:
            lines.append(f"{'  ' * next_level}-")
            return lines
        for index, item in enumerate(value, start=1):
            if isinstance(item, (dict, list)):
                lines.append(f"{'  ' * next_level}- Item {index}")
                lines.extend(_flatten_export_payload_lines(item, next_level + 1))
            else:
                text = str(item or "").strip()
                lines.append(f"{'  ' * next_level}- {text}")
        return lines

    text_value = str(value or "").strip()
    if key_name is not None:
        lines.append(f"{indent}{_format_export_field_label(key_name)}: {text_value}")
    else:
        lines.append(f"{indent}{text_value}")
    return lines


def generate_pdf_from_payload(payload: dict, output_path: str) -> str:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="No se pudo exportar a PDF: falta la dependencia reportlab. Instalá requirements.txt.",
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    title = str(payload.get("subject") or payload.get("title") or "Propuesta")
    c = canvas.Canvas(output_path, pagesize=A4)
    page_width, page_height = A4
    margin_x = 40
    margin_top = 48
    margin_bottom = 40
    line_height = 13

    def new_page() -> tuple[object, float]:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin_x, page_height - margin_top, "Propuesta Académica")
        c.setFont("Helvetica", 11)
        c.drawString(margin_x, page_height - margin_top - 16, title)
        text_obj = c.beginText(margin_x, page_height - margin_top - 36)
        text_obj.setFont("Helvetica", 9)
        return text_obj, page_height - margin_top - 36

    text_obj, current_y = new_page()

    lines = _flatten_export_payload_lines(payload)
    for raw_line in lines:
        line_chunks = str(raw_line).splitlines() or [""]
        for chunk in line_chunks:
            wrapped = textwrap.wrap(chunk, width=120, break_long_words=False, break_on_hyphens=False)
            if not wrapped:
                wrapped = [""]
            for part in wrapped:
                if current_y <= margin_bottom:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj, current_y = new_page()
                text_obj.textLine(part)
                current_y -= line_height

    c.drawText(text_obj)
    c.save()
    return output_path


@app.post("/proposals/{proposal_id}/validate-gdoc")
def validate_proposal_gdoc(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if not proposal.gdoc_url:
        proposal.gdoc_status = "missing"
        proposal.gdoc_last_checked = datetime.utcnow()
        db.add(proposal)
        db.commit()
        return {"status": "missing", "message": "La propuesta no tiene enlace de Google Docs."}

    try:
        extracted_subject, extracted_title, doc_hash, _ = extract_gdoc_payload(proposal.gdoc_url)
    except HTTPException:
        proposal.gdoc_url = None
        proposal.gdoc_status = "lost"
        proposal.gdoc_last_checked = datetime.utcnow()
        db.add(proposal)
        db.commit()
        return {
            "status": "missing",
            "message": "No se pudo acceder al documento de Google Docs. Se eliminó el enlace.",
            "gdoc_url": None,
        }

    extracted_name = extracted_subject or extracted_title
    if not extracted_name:
        proposal.gdoc_url = None
        proposal.gdoc_status = "lost"
        proposal.gdoc_last_checked = datetime.utcnow()
        db.add(proposal)
        db.commit()
        return {
            "status": "mismatch",
            "message": "No se pudo identificar la asignatura en el documento. Se eliminó el enlace.",
            "gdoc_url": None,
        }

    if proposal.subject and extracted_name:
        if normalize_subject_name(proposal.subject) != normalize_subject_name(extracted_name):
            proposal.gdoc_url = None
            proposal.gdoc_status = "lost"
            proposal.gdoc_last_checked = datetime.utcnow()
            db.add(proposal)
            db.commit()
            return {
                "status": "mismatch",
                "message": "El documento no coincide con la propuesta. Se eliminó el enlace.",
                "gdoc_url": None,
                "extracted_subject": extracted_subject,
                "extracted_title": extracted_title,
            }

    if proposal.title and (extracted_title or extracted_subject):
        compare_value = extracted_title or extracted_subject
        if normalize_title_name(proposal.title) != normalize_title_name(compare_value):
            proposal.gdoc_url = None
            proposal.gdoc_status = "lost"
            proposal.gdoc_last_checked = datetime.utcnow()
            db.add(proposal)
            db.commit()
            return {
                "status": "mismatch",
                "message": "El documento no coincide con el título de la propuesta. Se eliminó el enlace.",
                "gdoc_url": None,
                "extracted_subject": extracted_subject,
                "extracted_title": extracted_title,
            }

    proposal.gdoc_last_checked = datetime.utcnow()

    if not proposal.gdoc_hash or not proposal.gdoc_last_synced:
        proposal.gdoc_hash = doc_hash
        if not proposal.gdoc_last_synced:
            proposal.gdoc_last_synced = proposal.gdoc_last_checked
        proposal.gdoc_status = "ok"
        db.add(proposal)
        db.commit()
        return {
            "status": "ok",
            "message": "Enlace válido.",
            "gdoc_url": proposal.gdoc_url,
            "extracted_subject": extracted_subject or None,
            "extracted_title": extracted_title or None,
            "gdoc_hash": proposal.gdoc_hash,
            "gdoc_last_checked": proposal.gdoc_last_checked,
            "gdoc_last_synced": proposal.gdoc_last_synced,
        }

    if proposal.gdoc_hash != doc_hash:
        proposal.gdoc_status = "updated"
        db.add(proposal)
        db.commit()
        return {
            "status": "updated",
            "message": "El documento fue actualizado en Google Docs.",
            "gdoc_url": proposal.gdoc_url,
            "extracted_subject": extracted_subject or None,
            "extracted_title": extracted_title or None,
            "new_hash": doc_hash,
            "gdoc_last_checked": proposal.gdoc_last_checked,
            "gdoc_last_synced": proposal.gdoc_last_synced,
        }

    proposal.gdoc_status = "ok"
    db.add(proposal)
    db.commit()
    return {
        "status": "ok",
        "message": "Enlace válido.",
        "gdoc_url": proposal.gdoc_url,
        "extracted_subject": extracted_subject or None,
        "extracted_title": extracted_title or None,
        "gdoc_hash": proposal.gdoc_hash,
        "gdoc_last_checked": proposal.gdoc_last_checked,
        "gdoc_last_synced": proposal.gdoc_last_synced,
    }


@app.post("/proposals/{proposal_id}/link-gdoc")
def link_proposal_gdoc(
    proposal_id: int,
    url: str = Body(..., embed=True, description="URL pública de Google Docs"),
    db: Session = Depends(get_db),
):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if not url or not str(url).strip():
        raise HTTPException(status_code=400, detail="URL de Google Docs requerida")

    extracted_subject, extracted_title, doc_hash, _ = extract_gdoc_payload(url)
    extracted_name = extracted_subject or extracted_title
    if not extracted_name:
        raise HTTPException(status_code=400, detail="No se pudo identificar la asignatura en el documento")

    if proposal.subject and normalize_subject_name(proposal.subject) != normalize_subject_name(extracted_name):
        raise HTTPException(status_code=400, detail="El documento no coincide con la propuesta")

    if proposal.title and (extracted_title or extracted_subject):
        compare_value = extracted_title or extracted_subject
        if normalize_title_name(proposal.title) != normalize_title_name(compare_value):
            raise HTTPException(status_code=400, detail="El documento no coincide con el título de la propuesta")

    proposal.gdoc_url = url
    proposal.source_type = "gdoc"
    proposal.gdoc_hash = doc_hash
    proposal.gdoc_last_checked = datetime.utcnow()
    proposal.gdoc_last_synced = proposal.gdoc_last_checked
    proposal.gdoc_status = "ok"
    db.add(proposal)
    db.commit()
    return {
        "status": "linked",
        "gdoc_url": proposal.gdoc_url,
        "extracted_subject": extracted_subject or None,
        "extracted_title": extracted_title or None,
    }


@app.post("/proposals/{proposal_id}/unlink-gdoc")
def unlink_proposal_gdoc(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    proposal.gdoc_url = None
    proposal.gdoc_status = "missing"
    proposal.gdoc_last_checked = datetime.utcnow()
    db.add(proposal)
    db.commit()
    return {"status": "unlinked", "gdoc_url": None}


def try_create_gdoc_for_proposal(db: Session, proposal: models.Proposal) -> tuple[bool, str | None]:
    """
    Intenta crear un GDoc para una propuesta existente.
    Retorna True si se crea exitosamente, False si falla (no lanza excepción).
    """
    if not proposal:
        return False, "proposal-not-found"
    if proposal.gdoc_url:
        return True, None
    
    if not os.path.exists(TEMPLATE_PATH):
        return False, "template-not-found"

    settings = resolve_drive_settings_for_proposal(db, proposal)
    if not settings or not settings.root_folder_url or not settings.pdf_folder_url:
        return False, "drive-settings-missing"

    folder_id = extract_drive_folder_id(settings.root_folder_url)
    if not folder_id:
        return False, "drive-folder-invalid"

    try:
        from googleapiclient.http import MediaFileUpload
        from googleapiclient.errors import HttpError
        from .docx_export import generate_proposal_docx
    except ImportError:
        return False, "drive-dependencies-missing"

    try:
        drive_service = get_google_drive_service()
        output_path = generate_proposal_docx(proposal, TEMPLATE_PATH)
        output_dir = os.path.dirname(output_path)
        
        try:
            filename = build_proposal_docx_filename(proposal)
            file_title = filename[:-5] if filename.lower().endswith(".docx") else filename
            media = MediaFileUpload(
                output_path,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                resumable=False,
            )

            created = drive_service.files().create(
                body={
                    "name": file_title,
                    "mimeType": "application/vnd.google-apps.document",
                    "parents": [folder_id]
                },
                media_body=media,
                fields="id, webViewLink, parents",
                supportsAllDrives=True,
            ).execute()

            gdoc_id = created.get("id")
            if gdoc_id:
                gdoc_url = f"https://docs.google.com/document/d/{gdoc_id}/edit"
                # Extract new hash
                _, _, new_hash, _ = extract_gdoc_payload(gdoc_url)
                # Update proposal with GDoc info
                proposal.gdoc_url = gdoc_url
                proposal.gdoc_hash = new_hash
                proposal.gdoc_last_synced = datetime.utcnow()
                proposal.gdoc_last_checked = datetime.utcnow()
                proposal.gdoc_status = "ok"
                proposal.source_type = "gdoc"
                db.add(proposal)
                db.flush()
                return True, None
        finally:
            shutil.rmtree(output_dir, ignore_errors=True)
    except Exception as exc:
        return False, str(exc)

    return False, "drive-create-failed"


@app.post("/proposals/{proposal_id}/create-gdoc", response_model=schemas.Proposal)
def create_proposal_gdoc(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    if proposal.gdoc_url:
        return build_proposal_response(db, proposal)

    if not os.path.exists(TEMPLATE_PATH):
        raise HTTPException(status_code=500, detail="Template Propuestas.docx not found")

    settings = resolve_drive_settings_for_proposal(db, proposal)
    if not settings or not settings.root_folder_url or not settings.pdf_folder_url:
        raise HTTPException(
            status_code=400,
            detail="Debes configurar Carpeta Raíz y Carpeta PDF de Drive para esta carrera/plan",
        )

    folder_id = extract_drive_folder_id(settings.root_folder_url)
    if not folder_id:
        raise HTTPException(status_code=400, detail="La URL de Carpeta Raíz (Drive) es inválida")

    try:
        from googleapiclient.http import MediaFileUpload
        from .docx_export import generate_proposal_docx
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Faltan dependencias para generar/subir DOCX. Verifica python-docx y google-api-python-client",
        )

    drive_service = get_google_drive_service()
    output_path = generate_proposal_docx(proposal, TEMPLATE_PATH)
    output_dir = os.path.dirname(output_path)
    try:
        from googleapiclient.errors import HttpError

        filename = build_proposal_docx_filename(proposal)
        file_title = filename[:-5] if filename.lower().endswith(".docx") else filename
        media = MediaFileUpload(
            output_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=False,
        )

        def extract_google_error(exc: Exception) -> tuple[int | None, list[str], str]:
            status_code = None
            reasons: list[str] = []
            message = str(exc)
            if isinstance(exc, HttpError):
                status_code = getattr(exc.resp, "status", None)
                content = getattr(exc, "content", None)
                if content:
                    try:
                        payload = json.loads(content.decode("utf-8") if isinstance(content, (bytes, bytearray)) else str(content))
                        err = payload.get("error") if isinstance(payload, dict) else None
                        if isinstance(err, dict):
                            parsed_message = err.get("message")
                            if parsed_message:
                                message = str(parsed_message)
                            parsed_reasons = err.get("errors")
                            if isinstance(parsed_reasons, list):
                                reasons = [str(item.get("reason")) for item in parsed_reasons if isinstance(item, dict) and item.get("reason")]
                    except Exception:
                        pass
            return status_code, reasons, message

        def create_google_doc(target_folder_id: str | None):
            body = {
                "name": file_title,
                "mimeType": "application/vnd.google-apps.document",
            }
            if target_folder_id:
                body["parents"] = [target_folder_id]
            return drive_service.files().create(
                body=body,
                media_body=media,
                fields="id, webViewLink, parents",
                supportsAllDrives=True,
            ).execute()

        try:
            created = create_google_doc(folder_id)
        except Exception as first_exc:
            status_code, reasons, message = extract_google_error(first_exc)
            lower_message = message.lower()
            has_quota_error = "storagequotaexceeded" in lower_message or "drive storage quota" in lower_message or "storageQuotaExceeded" in reasons

            if has_quota_error:
                raise HTTPException(
                    status_code=507,
                    detail="No se pudo crear el documento: la cuota de almacenamiento de la cuenta autenticada en Google Drive está excedida.",
                )

            if "invalid_grant" in lower_message or "token has been expired" in lower_message or "token has been revoked" in lower_message:
                raise HTTPException(
                    status_code=401,
                    detail="GOOGLE_AUTH_EXPIRED: El token de Google Drive expiró o fue revocado. Necesitás reautorizar la aplicación desde el modal de la propuesta.",
                )

            # Fallback probado: crear primero en raíz y luego mover a carpeta destino.
            try:
                created = create_google_doc(None)
                doc_id_for_move = created.get("id")
                if doc_id_for_move and folder_id:
                    try:
                        current_parents = created.get("parents") or []
                        remove_parents = ",".join([parent for parent in current_parents if parent and parent != folder_id])
                        move_kwargs = {
                            "fileId": doc_id_for_move,
                            "addParents": folder_id,
                            "fields": "id, webViewLink, parents",
                            "supportsAllDrives": True,
                        }
                        if remove_parents:
                            move_kwargs["removeParents"] = remove_parents
                        created = drive_service.files().update(**move_kwargs).execute()
                    except Exception:
                        # Si no se puede mover, igual dejamos el documento creado y vinculado.
                        pass
            except Exception as fallback_exc:
                fb_status, fb_reasons, fb_message = extract_google_error(fallback_exc)
                fb_lower = fb_message.lower()
                if "storagequotaexceeded" in fb_lower or "drive storage quota" in fb_lower or "storageQuotaExceeded" in fb_reasons:
                    raise HTTPException(
                        status_code=507,
                        detail="No se pudo crear el documento: la cuota de almacenamiento de la cuenta autenticada en Google Drive está excedida.",
                    )
                if "invalid_grant" in fb_lower or "token has been expired" in fb_lower or "token has been revoked" in fb_lower:
                    raise HTTPException(
                        status_code=401,
                        detail="GOOGLE_AUTH_EXPIRED: El token de Google Drive expiró o fue revocado. Necesitás reautorizar la aplicación desde el modal de la propuesta.",
                    )
                status_text = f"{fb_status}" if fb_status is not None else "unknown"
                reason_text = ", ".join(fb_reasons) if fb_reasons else "none"
                raise HTTPException(
                    status_code=502,
                    detail=f"Error al crear el documento en Google Drive (status={status_text}, reasons={reason_text}): {fb_message}",
                )
    except HTTPException:
        raise
    except Exception as exc:
        raw_message = str(exc)
        raise HTTPException(
            status_code=502,
            detail=f"Error al crear el documento en Google Drive: {raw_message}",
        )

    try:

        doc_id = created.get("id")
        if not doc_id:
            raise HTTPException(status_code=500, detail="Google Drive no devolvió el ID del documento creado")

        gdoc_url = created.get("webViewLink") or f"https://docs.google.com/document/d/{doc_id}/edit"
        now = datetime.utcnow()
        proposal.gdoc_url = gdoc_url
        proposal.source_type = "gdoc"
        proposal.gdoc_last_checked = now
        proposal.gdoc_last_synced = now
        proposal.gdoc_status = "ok"
        try:
            _, _, extracted_hash, _ = extract_gdoc_payload(gdoc_url)
            proposal.gdoc_hash = extracted_hash
        except Exception:
            proposal.gdoc_hash = compute_payload_hash(build_extracted_snapshot(build_proposal_snapshot(db, proposal)))
        db.add(proposal)
        db.commit()
        db.refresh(proposal)
        return build_proposal_response(db, proposal)
    finally:
        shutil.rmtree(output_dir, ignore_errors=True)


@app.post("/proposals/gdoc-status")
def get_gdoc_statuses(request: GdocStatusRequest, db: Session = Depends(get_db)):
    ids = request.ids or []
    if not ids:
        return {"statuses": {}}

    proposals = db.query(models.Proposal).filter(models.Proposal.id.in_(ids)).all()
    statuses: dict[int, dict] = {}
    now = datetime.utcnow()
    has_updates = False

    for proposal in proposals:
        status = "missing"
        if proposal.gdoc_url:
            try:
                _, _, doc_hash, _ = extract_gdoc_payload(proposal.gdoc_url)
            except HTTPException:
                status = "lost"
            else:
                if not proposal.gdoc_hash:
                    proposal.gdoc_hash = doc_hash
                    proposal.gdoc_last_synced = proposal.gdoc_last_synced or now
                    has_updates = True
                elif proposal.gdoc_hash != doc_hash:
                    status = "updated"
                else:
                    status = "ok"
                proposal.gdoc_last_checked = now
                has_updates = True
        else:
            if proposal.source_type == "gdoc":
                status = "lost"
            else:
                status = "missing"

        proposal.gdoc_status = status
        proposal.gdoc_last_checked = proposal.gdoc_last_checked or now
        has_updates = True

        statuses[proposal.id] = {"status": status}

    if has_updates:
        db.add_all(proposals)
        db.commit()

    return {"statuses": statuses}


@app.post("/proposals/{proposal_id}/gdoc-accept-latest")
def accept_latest_gdoc(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if not proposal.gdoc_url:
        raise HTTPException(status_code=400, detail="La propuesta no tiene enlace de Google Docs")

    _, _, doc_hash, _ = extract_gdoc_payload(proposal.gdoc_url)
    proposal.gdoc_hash = doc_hash
    proposal.gdoc_last_checked = datetime.utcnow()
    proposal.gdoc_last_synced = proposal.gdoc_last_checked
    proposal.gdoc_status = "ok"
    db.add(proposal)
    db.commit()
    return {"status": "ok", "gdoc_hash": proposal.gdoc_hash}


def apply_extracted_payload_to_proposal(db: Session, proposal: models.Proposal, payload: dict) -> None:
    def has_text(value) -> bool:
        return isinstance(value, str) and value.strip() != ""

    def has_items(value) -> bool:
        return isinstance(value, list) and len(value) > 0

    def normalize_learning_outcomes_for_storage(items) -> list[dict]:
        normalized: list[dict] = []
        for idx, item in enumerate(items or []):
            if isinstance(item, dict):
                description = str(item.get("description") or item.get("descripcion") or "").strip()
                observable_verb = str(item.get("observable_verb") or item.get("verbo_observable") or "").strip()
            else:
                description = str(item or "").strip()
                observable_verb = ""
            if not description:
                continue
            normalized.append({
                "id": idx + 1,
                "description": description,
                "observable_verb": observable_verb,
            })
        return normalized

    def normalize_units_for_storage(items) -> list[dict]:
        normalized: list[dict] = []
        for idx, item in enumerate(items or []):
            if not isinstance(item, dict):
                continue
            name = str(
                item.get("name")
                or item.get("nombre")
                or item.get("title")
                or item.get("titulo")
                or ""
            ).strip()
            content = str(item.get("content") or item.get("contenidos") or item.get("contents") or "").strip()
            bibliography_basic = str(
                item.get("bibliography_basic")
                or item.get("bib_basica")
                or item.get("bib_basic")
                or item.get("bibliografia_basica")
                or ""
            ).strip()
            bibliography_complementary = str(
                item.get("bibliography_complementary")
                or item.get("bib_complementaria")
                or item.get("bib_comp")
                or item.get("bibliografia_complementaria")
                or ""
            ).strip()
            if not (name or content or bibliography_basic or bibliography_complementary):
                continue
            normalized.append(
                {
                    "id": idx + 1,
                    "name": name,
                    "content": content,
                    "bibliography_basic": bibliography_basic,
                    "bibliography_complementary": bibliography_complementary,
                }
            )
        return normalized

    def normalize_practicals_for_storage(items) -> list[dict]:
        normalized: list[dict] = []
        for idx, item in enumerate(items or []):
            if not isinstance(item, dict):
                continue
            number = str(item.get("number") or item.get("numero") or idx + 1).strip()
            name = str(item.get("name") or item.get("nombre") or "").strip()
            objective = str(item.get("objective") or item.get("objetivo") or "").strip()
            activities = str(item.get("activities") or item.get("actividades") or "").strip()
            materials = str(item.get("materials") or item.get("materiales") or "").strip()
            scope = str(item.get("scope") or item.get("ambito") or "").strip()
            if not (name or objective or activities or materials or scope):
                continue
            normalized.append(
                {
                    "id": idx + 1,
                    "number": number,
                    "name": name,
                    "objective": objective,
                    "activities": activities,
                    "materials": materials,
                    "scope": scope,
                }
            )
        return normalized

    proposal.career = payload.get("career") or proposal.career
    proposal.subject = payload.get("subject") or proposal.subject
    proposal.study_plan = payload.get("study_plan") or proposal.study_plan
    proposal.academic_year = payload.get("academic_year") or proposal.academic_year
    proposal.year_of_career = payload.get("year_of_career") or proposal.year_of_career
    payload_quarter = payload.get("quarter")
    proposal.quarter = normalize_term_name(payload_quarter) if payload_quarter else proposal.quarter
    proposal.character = payload.get("character") or proposal.character
    proposal.regime = payload.get("regime") or proposal.regime
    proposal.total_hours = parse_int(payload.get("total_hours"), proposal.total_hours)
    proposal.theoretical_hours = parse_int(payload.get("theoretical_hours"), proposal.theoretical_hours)
    proposal.practical_hours = parse_int(payload.get("practical_hours"), proposal.practical_hours)
    proposal.weekly_hours = parse_int(payload.get("weekly_hours"), proposal.weekly_hours)
    minimum_content = payload.get("minimum_content")
    if has_text(minimum_content):
        proposal.minimum_content = minimum_content

    fundamentals_part1 = payload.get("importance")
    if has_text(fundamentals_part1):
        proposal.fundamentals_part1 = fundamentals_part1

    fundamentals_part2 = payload.get("professional_profile")
    if has_text(fundamentals_part2):
        proposal.fundamentals_part2 = fundamentals_part2

    methodology = payload.get("methodology")
    if has_text(methodology):
        proposal.methodology = methodology

    evaluation = payload.get("evaluation")
    if has_text(evaluation):
        proposal.evaluation = evaluation

    bibliography = payload.get("bibliography")
    if has_text(bibliography):
        proposal.bibliography = bibliography

    observations = payload.get("observations")
    if has_text(observations):
        proposal.observations = observations

    learning_outcomes = normalize_learning_outcomes_for_storage(payload.get("learning_outcomes"))
    if has_items(learning_outcomes):
        proposal.learning_outcomes = learning_outcomes

    units = normalize_units_for_storage(payload.get("units"))
    if has_items(units):
        has_rich_unit_data = any(
            bool((unit.get("content") or "").strip())
            or bool((unit.get("bibliography_basic") or "").strip())
            or bool((unit.get("bibliography_complementary") or "").strip())
            for unit in units
        )
        existing_units = proposal.units or []
        if has_rich_unit_data or not existing_units:
            proposal.units = units

    practicals = normalize_practicals_for_storage(payload.get("practicals"))
    if has_items(practicals):
        proposal.practicals = practicals

    teaching_team = payload.get("teaching_team")
    if isinstance(teaching_team, list):
        teacher_objs = []
        teacher_ids = []
        for entry in teaching_team:
            if not isinstance(entry, dict):
                continue
            teacher = upsert_teacher(db, entry)
            if teacher:
                db.flush()
                ensure_teacher_career(db, teacher.id, proposal.career)
                teacher_objs.append(teacher)
                teacher_ids.append(teacher.id)
        if teacher_ids:
            replace_proposal_teachers(db, proposal.id, teacher_ids)
        proposal.teaching_team = build_teaching_team_payload(teacher_objs)

    generic_items = normalize_competency_items(payload.get("generic_competencies") or [])
    specific_items = normalize_competency_items(payload.get("specific_competencies") or [])
    if generic_items:
        proposal.generic_competencies = build_competencies_text(generic_items)
        ensure_competency_catalog(db, proposal.career, generic_items, "generic", proposal.study_plan)
        replace_proposal_competencies(db, proposal.id, generic_items, "generic")
    if specific_items:
        proposal.specific_competencies = build_competencies_text(specific_items)
        ensure_competency_catalog(db, proposal.career, specific_items, "specific", proposal.study_plan)
        replace_proposal_competencies(db, proposal.id, specific_items, "specific")

    sync_subject_from_proposal(db, proposal)


def build_proposal_snapshot(db: Session, proposal: models.Proposal) -> dict:
    competencies = get_proposal_competencies(db, proposal.id)
    return {
        "minimum_content": proposal.minimum_content or "",
        "importance": proposal.fundamentals_part1 or "",
        "professional_profile": proposal.fundamentals_part2 or "",
        "learning_outcomes": normalize_learning_outcomes(proposal.learning_outcomes or []),
        "units": normalize_unit_items(proposal.units or []),
        "practicals": normalize_practical_items(proposal.practicals or []),
        "teaching_team": normalize_teaching_team_items(proposal.teaching_team or []),
        "methodology": proposal.methodology or "",
        "evaluation": proposal.evaluation or "",
        "generic_competencies": normalize_competency_items(competencies.get("generic") or []),
        "specific_competencies": normalize_competency_items(competencies.get("specific") or []),
    }


def build_extracted_snapshot(payload: dict) -> dict:
    return {
        "minimum_content": payload.get("minimum_content") or "",
        "importance": payload.get("importance") or "",
        "professional_profile": payload.get("professional_profile") or "",
        "learning_outcomes": normalize_learning_outcomes(payload.get("learning_outcomes") or []),
        "units": normalize_unit_items(payload.get("units") or []),
        "practicals": normalize_practical_items(payload.get("practicals") or []),
        "teaching_team": normalize_teaching_team_items(payload.get("teaching_team") or []),
        "methodology": payload.get("methodology") or "",
        "evaluation": payload.get("evaluation") or "",
        "generic_competencies": normalize_competency_items(payload.get("generic_competencies") or []),
        "specific_competencies": normalize_competency_items(payload.get("specific_competencies") or []),
    }


def format_diff_value(value) -> str:
    if isinstance(value, list) and all(isinstance(item, str) for item in value):
        return "\n".join([item for item in value if item])
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value or "")


def normalize_list_items(items: list, drop_keys: set[str] | None = None) -> list:
    drop_keys = drop_keys or set()
    normalized = []
    for item in items or []:
        if isinstance(item, dict):
            cleaned = {key: item.get(key) for key in item.keys() if key not in drop_keys}
            normalized.append(cleaned)
        else:
            normalized.append(item)
    return normalized


def normalize_snapshot_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def normalize_unit_items(items: list) -> list[dict]:
    normalized: list[dict] = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "name": normalize_snapshot_text(item.get("name")),
                "content": normalize_snapshot_text(item.get("content") or item.get("contenidos")),
                "bibliography_basic": normalize_snapshot_text(
                    item.get("bibliography_basic") or item.get("bib_basica") or item.get("bib_basic")
                ),
                "bibliography_complementary": normalize_snapshot_text(
                    item.get("bibliography_complementary") or item.get("bib_complementaria") or item.get("bib_comp")
                ),
            }
        )
    return normalized


def normalize_practical_items(items: list) -> list[dict]:
    normalized: list[dict] = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "number": normalize_snapshot_text(item.get("number") or item.get("numero")),
                "name": normalize_snapshot_text(item.get("name") or item.get("nombre")),
                "objective": normalize_snapshot_text(item.get("objective") or item.get("objetivo")),
                "activities": normalize_snapshot_text(item.get("activities") or item.get("actividades")),
                "materials": normalize_snapshot_text(item.get("materials") or item.get("materiales")),
                "scope": normalize_snapshot_text(item.get("scope") or item.get("ambito")),
            }
        )
    return normalized


def normalize_teaching_team_items(items: list) -> list[dict]:
    normalized: list[dict] = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "name": re.sub(r"\s+", " ", str(item.get("name") or item.get("nombre") or "").strip()),
                "email": re.sub(r"\s+", " ", str(item.get("email") or item.get("correo") or "").strip().lower()),
                "category": re.sub(r"\s+", " ", str(item.get("category") or item.get("categoria") or "").strip()),
            }
        )
    normalized.sort(key=lambda value: (value.get("name") or "", value.get("email") or "", value.get("category") or ""))
    return normalized


def normalize_learning_outcomes(items: list) -> list[str]:
    normalized = []
    for item in items or []:
        if isinstance(item, dict):
            value = item.get("description") or item.get("descripcion") or ""
        else:
            value = str(item)
        normalized.append(re.sub(r"\s+", " ", value).strip())
    return normalized


@app.get("/proposals/{proposal_id}/gdoc-diff")
def get_gdoc_diff(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if not proposal.gdoc_url:
        raise HTTPException(status_code=400, detail="La propuesta no tiene enlace de Google Docs")

    extracted_subject, extracted_title, doc_hash, extracted_payload = extract_gdoc_payload(proposal.gdoc_url)
    current_snapshot = build_proposal_snapshot(db, proposal)
    latest_snapshot = build_extracted_snapshot(extracted_payload)

    labels = {
        "teaching_team": "Equipo docente",
        "minimum_content": "Contenidos mínimos",
        "importance": "Importancia",
        "professional_profile": "Perfil profesional",
        "learning_outcomes": "Resultados de aprendizaje",
        "units": "Unidades",
        "practicals": "Trabajos prácticos",
        "generic_competencies": "Competencias genéricas",
        "specific_competencies": "Competencias específicas",
        "methodology": "Metodología",
        "evaluation": "Evaluación",
    }

    review_required_keys = {"minimum_content", "teaching_team"}

    changes = {}
    for key, label in labels.items():
        current_val = current_snapshot.get(key)
        latest_val = latest_snapshot.get(key)
        if compute_payload_hash({"value": current_val}) != compute_payload_hash({"value": latest_val}):
            changes[key] = {
                "label": label,
                "current": current_val,
                "latest": latest_val,
                "current_display": format_diff_value(current_val),
                "latest_display": format_diff_value(latest_val),
                "review_required": key in review_required_keys,
            }

    return {
        "current": current_snapshot,
        "latest": latest_snapshot,
        "changes": changes,
        "extracted_subject": extracted_subject or None,
        "extracted_title": extracted_title or None,
        "gdoc_hash": doc_hash,
    }


@app.post("/proposals/{proposal_id}/sync-gdoc", response_model=schemas.Proposal)
def sync_proposal_gdoc(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if not proposal.gdoc_url:
        raise HTTPException(status_code=400, detail="La propuesta no tiene enlace de Google Docs")

    extracted_subject, extracted_title, doc_hash, extracted_payload = extract_gdoc_payload(proposal.gdoc_url)
    extracted_name = extracted_subject or extracted_title
    if not extracted_name:
        raise HTTPException(status_code=400, detail="No se pudo identificar la asignatura en el documento")

    if proposal.subject and normalize_subject_name(proposal.subject) != normalize_subject_name(extracted_name):
        raise HTTPException(status_code=400, detail="El documento no coincide con la propuesta")

    if proposal.title and (extracted_title or extracted_subject):
        compare_value = extracted_title or extracted_subject
        if normalize_title_name(proposal.title) != normalize_title_name(compare_value):
            raise HTTPException(status_code=400, detail="El documento no coincide con el título de la propuesta")

    apply_extracted_payload_to_proposal(db, proposal, extracted_payload)
    proposal.gdoc_hash = doc_hash
    proposal.gdoc_last_checked = datetime.utcnow()
    proposal.gdoc_last_synced = proposal.gdoc_last_checked
    proposal.gdoc_status = "ok"
    db.add(proposal)
    db.commit()
    db.refresh(proposal)
    return build_proposal_response(db, proposal)


@app.get("/proposals/{proposal_id}/local-diff")
def get_local_diff(proposal_id: int, db: Session = Depends(get_db)):
    """
    Detecta cambios locales comparando el snapshot actual con el GDoc actual.
    Similar a gdoc-diff, pero muestra qué cambió EN LOCAL para enviarlo a GDoc.
    """
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    
    if not proposal.gdoc_url:
        raise HTTPException(status_code=400, detail="La propuesta no tiene enlace de Google Docs configurado")

    extracted_subject, extracted_title, doc_hash, extracted_payload = extract_gdoc_payload(proposal.gdoc_url)
    
    # Snapshot actual en local
    local_snapshot = build_proposal_snapshot(db, proposal)
    # Snapshot del GDoc actual
    gdoc_snapshot = build_extracted_snapshot(extracted_payload)

    labels = {
        "teaching_team": "Equipo docente",
        "minimum_content": "Contenidos mínimos",
        "importance": "Importancia",
        "professional_profile": "Perfil profesional",
        "learning_outcomes": "Resultados de aprendizaje",
        "units": "Unidades",
        "practicals": "Trabajos prácticos",
        "generic_competencies": "Competencias genéricas",
        "specific_competencies": "Competencias específicas",
        "methodology": "Metodología",
        "evaluation": "Evaluación",
    }

    review_required_keys = {"minimum_content", "teaching_team"}
    
    changes = {}
    for key, label in labels.items():
        local_val = local_snapshot.get(key)
        gdoc_val = gdoc_snapshot.get(key)
        
        # Solo mostrar como cambio si LOCAL es diferente de GDoc
        if compute_payload_hash({"value": local_val}) != compute_payload_hash({"value": gdoc_val}):
            changes[key] = {
                "label": label,
                "local": local_val,
                "gdoc": gdoc_val,
                "local_display": format_diff_value(local_val),
                "gdoc_display": format_diff_value(gdoc_val),
                "review_required": key in review_required_keys,
            }

    return {
        "local": local_snapshot,
        "gdoc": gdoc_snapshot,
        "changes": changes,
        "gdoc_hash": doc_hash,
        "message": f"Se encontraron {len(changes)} diferencia(s) entre local y GDoc. Selecciona cuáles enviar a GDoc.",
    }


@app.post("/proposals/{proposal_id}/push-to-gdoc-direct")
def push_proposal_to_gdoc_direct(
    proposal_id: int,
    changes_to_apply: dict = Body(..., embed=True, description="Dict con los cambios a aplicar: {field: True/False}"),
    db: Session = Depends(get_db),
):
    """
    Aplica cambios locales directamente al documento de Google Docs.
    Genera DOCX actualizado y reemplaza el contenido en GDoc preservando la estructura.
    Actualiza: GDoc + gdoc_hash en BD.
    """
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    
    if not proposal.gdoc_url:
        raise HTTPException(status_code=400, detail="La propuesta no tiene enlace de Google Docs")

    if not changes_to_apply or not isinstance(changes_to_apply, dict):
        raise HTTPException(status_code=400, detail="changes_to_apply es requerido")

    # Verificar que tengo cambios que aplicar
    if not any(changes_to_apply.values()):
        return {
            "status": "ok",
            "message": "No hay cambios seleccionados para aplicar",
            "gdoc_url": proposal.gdoc_url,
            "updated_fields": []
        }

    # Validar que el template existe
    if not os.path.exists(TEMPLATE_PATH):
        raise HTTPException(status_code=500, detail="Template Propuestas.docx no encontrado")

    try:
        from .docx_export import generate_proposal_docx
        from googleapiclient.http import MediaFileUpload
        from googleapiclient.errors import HttpError
    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Faltan dependencias: {str(e)}",
        )

    # Extraer ID del documento de GDoc
    doc_id_match = re.search(r"/d/([\w-]+)", proposal.gdoc_url)
    if not doc_id_match:
        raise HTTPException(status_code=400, detail="URL de Google Docs inválida")
    
    doc_id = doc_id_match.group(1)
    drive_service = get_google_drive_service()

    # Registrar qué campos se actualizarán
    updated_fields = [field for field, apply in changes_to_apply.items() if apply]
    
    try:
        # Generar DOCX actualizado
        output_path = generate_proposal_docx(proposal, TEMPLATE_PATH)
        output_dir = os.path.dirname(output_path)

        try:
            # Leer el DOCX actualizado como binario
            with open(output_path, 'rb') as f:
                docx_content = f.read()

            # Actualizar el contenido del documento en Google Drive
            # Esto preserva la estructura básica del documento pero reemplaza el contenido
            media = MediaFileUpload(
                output_path,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                resumable=False,
            )

            try:
                drive_service.files().update(
                    fileId=doc_id,
                    media_body=media,
                    fields="id, webViewLink"
                ).execute()
            except HttpError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Error al actualizar Google Docs: {e.reason if hasattr(e, 'reason') else str(e)}"
                )

            # Extraer nuevo hash del documento actualizado
            _, _, new_hash, _ = extract_gdoc_payload(proposal.gdoc_url)
            
            # Actualizar BD con el nuevo hash
            proposal.gdoc_hash = new_hash
            proposal.gdoc_last_checked = datetime.utcnow()
            proposal.gdoc_last_synced = proposal.gdoc_last_checked
            proposal.gdoc_status = "ok"
            db.add(proposal)
            db.commit()

            return {
                "status": "success",
                "message": f"Se aplicaron {len(updated_fields)} cambios en Google Docs exitosamente",
                "gdoc_url": proposal.gdoc_url,
                "updated_fields": updated_fields,
                "gdoc_hash": proposal.gdoc_hash,
            }

        finally:
            shutil.rmtree(output_dir, ignore_errors=True)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al aplicar cambios: {str(e)}"
        )





def sync_teachers_from_existing_proposals() -> None:
    db = SessionLocal()
    try:
        proposals = db.query(models.Proposal).all()
        for proposal in proposals:
            if not proposal.teaching_team:
                continue
            teacher_payload = proposal.teaching_team
            if not isinstance(teacher_payload, list):
                continue
            teacher_objs = []
            teacher_ids = []
            for entry in teacher_payload:
                if not isinstance(entry, dict):
                    continue
                teacher = upsert_teacher(db, entry)
                if teacher:  # Only process if teacher was created
                    db.flush()
                    ensure_teacher_career(db, teacher.id, proposal.career)
                    teacher_objs.append(teacher)
                    teacher_ids.append(teacher.id)
            if teacher_ids:
                replace_proposal_teachers(db, proposal.id, teacher_ids)
                proposal.teaching_team = build_teaching_team_payload(teacher_objs)
                db.add(proposal)
        db.commit()
    finally:
        db.close()


def get_openai_client():
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("OPENAI_API_KEY not set")
    return OpenAI(api_key=key)


def model_uses_max_completion_tokens(model: str | None) -> bool:
    name = str(model or "").strip().lower()
    if not name:
        return False
    normalized = name.replace("_", "-")
    tokens = [token for token in re.split(r"[^a-z0-9]+", normalized) if token]
    for token in tokens:
        if token.startswith(("o1", "o3", "o4", "gpt-5")):
            return True
    return normalized.startswith(("o1", "o3", "o4", "gpt-5"))


def model_restricts_temperature_to_default(model: str | None) -> bool:
    name = str(model or "").strip().lower()
    if not name:
        return False
    normalized = name.replace("_", "-")
    tokens = [token for token in re.split(r"[^a-z0-9]+", normalized) if token]
    for token in tokens:
        if token.startswith(("o1", "o3", "o4", "gpt-5")):
            return True
    return normalized.startswith(("o1", "o3", "o4", "gpt-5"))


def create_chat_completion_compatible(
    client: OpenAI,
    *,
    model: str,
    messages: list[dict],
    temperature: float | int | None = None,
    max_tokens: int | None = None,
):
    kwargs = {
        "model": model,
        "messages": messages,
    }
    if temperature is not None:
        if model_restricts_temperature_to_default(model):
            try:
                temp_value = float(temperature)
            except Exception:
                temp_value = None
            if temp_value is not None and abs(temp_value - 1.0) < 1e-9:
                kwargs["temperature"] = 1
        else:
            kwargs["temperature"] = temperature
    if max_tokens is not None:
        if model_uses_max_completion_tokens(model):
            kwargs["max_completion_tokens"] = int(max_tokens)
        else:
            kwargs["max_tokens"] = int(max_tokens)

    try:
        return client.chat.completions.create(**kwargs)
    except Exception as exc:
        message = str(exc)
        lower_message = message.lower()
        retriable = False
        token_param_unsupported = (
            "unsupported parameter" in lower_message
            and ("max_tokens" in lower_message or "max_completion_tokens" in lower_message)
        )
        if token_param_unsupported and "max_tokens" in kwargs:
            kwargs["max_completion_tokens"] = kwargs.pop("max_tokens")
            retriable = True
        elif token_param_unsupported and "max_completion_tokens" in kwargs:
            kwargs["max_tokens"] = kwargs.pop("max_completion_tokens")
            retriable = True
        temperature_unsupported = "unsupported parameter: 'temperature'" in lower_message
        temperature_value_unsupported = (
            "unsupported value" in lower_message
            and "temperature" in lower_message
            and "default (1)" in lower_message
        )
        if (temperature_unsupported or temperature_value_unsupported) and "temperature" in kwargs:
            kwargs.pop("temperature", None)
            retriable = True
        if retriable:
            return client.chat.completions.create(**kwargs)
        raise


INTELLIGENT_TOPIC_ALIASES = {
    "equipo docente": "teaching_team",
    "equipo_docente": "teaching_team",
    "docentes": "teaching_team",
    "fundamentacion": "fundamentals",
    "fundamentación": "fundamentals",
    "bibliografia": "bibliography",
    "bibliografía": "bibliography",
    "contenidos minimos": "minimum_content",
    "contenidos mínimos": "minimum_content",
    "contenidos": "minimum_content",
    "resultados de aprendizaje": "learning_outcomes",
    "objetivos": "learning_outcomes",
    "unidades": "units",
    "trabajos practicos": "practicals",
    "trabajos prácticos": "practicals",
    "metodologia": "methodology",
    "metodología": "methodology",
    "evaluacion": "evaluation",
    "evaluación": "evaluation",
}

INTELLIGENT_AVAILABLE_MODELS = [
    "gpt-5.2",
    "gpt-5.2-pro",
    "gpt-5.1",
    "gpt-5-mini",
    "gpt-4o",
    "o3",
    "o3-pro",
    "o4-mini",
    "gpt-4.1",
    "gpt-4.1-mini",
]


def default_intelligent_mode_settings() -> dict:
    return {
        "guepardo": {
            "model": os.getenv("OPENAI_MODEL_FAST", os.getenv("OPENAI_MODEL", "gpt-4o-mini")),
            "temperature": 0.15,
            "max_tokens": 420,
        },
        "delfin": {
            "model": os.getenv("OPENAI_MODEL_BALANCED", os.getenv("OPENAI_MODEL", "gpt-4o-mini")),
            "temperature": 0.1,
            "max_tokens": 500,
        },
        "ballena": {
            "model": os.getenv("OPENAI_MODEL_PRECISE", os.getenv("OPENAI_MODEL", "gpt-4o")),
            "temperature": 0.1,
            "max_tokens": 700,
        },
    }


def sanitize_mode_temperature(value, default: float) -> float:
    try:
        temp = float(value)
    except Exception:
        return float(default)
    if temp < 0:
        return 0.0
    if temp > 2:
        return 2.0
    return temp


def sanitize_mode_max_tokens(value, default: int) -> int:
    try:
        tokens = int(value)
    except Exception:
        return int(default)
    if tokens < 100:
        return 100
    if tokens > 4000:
        return 4000
    return tokens


def build_effective_intelligent_mode_settings(settings: models.IntelligentControlSettings | None) -> dict:
    defaults = default_intelligent_mode_settings()
    if not settings:
        return defaults

    result = {}
    for mode in ("guepardo", "delfin", "ballena"):
        default_mode = defaults[mode]
        model = getattr(settings, f"{mode}_model", None) or default_mode["model"]
        temperature = sanitize_mode_temperature(getattr(settings, f"{mode}_temperature", None), default_mode["temperature"])
        max_tokens = sanitize_mode_max_tokens(getattr(settings, f"{mode}_max_tokens", None), default_mode["max_tokens"])
        result[mode] = {
            "model": model,
            "temperature": temperature,
            "max_tokens": max_tokens,
        }
    return result


def serialize_intelligent_settings(settings: models.IntelligentControlSettings) -> dict:
    mode_settings = build_effective_intelligent_mode_settings(settings)
    return {
        "director_last_mode": normalize_intelligent_mode(settings.director_last_mode, default="delfin"),
        "docente_mode": normalize_intelligent_mode(settings.docente_mode, default="guepardo"),
        "guepardo": mode_settings["guepardo"],
        "delfin": mode_settings["delfin"],
        "ballena": mode_settings["ballena"],
        "available_models": INTELLIGENT_AVAILABLE_MODELS,
        "updated_at": settings.updated_at,
    }


def normalize_intelligent_mode(value: str | None, default: str = "delfin") -> str:
    mode = normalize_header(value or default)
    if mode not in {"guepardo", "delfin", "ballena"}:
        return default
    return mode


def get_or_create_intelligent_settings(db: Session) -> models.IntelligentControlSettings:
    settings = db.query(models.IntelligentControlSettings).order_by(models.IntelligentControlSettings.id.asc()).first()
    defaults = default_intelligent_mode_settings()
    if settings:
        changed = False
        if not settings.director_last_mode:
            settings.director_last_mode = "delfin"
            changed = True
        if not settings.docente_mode:
            settings.docente_mode = "guepardo"
            changed = True
        for mode in ("guepardo", "delfin", "ballena"):
            model_key = f"{mode}_model"
            temp_key = f"{mode}_temperature"
            tokens_key = f"{mode}_max_tokens"
            if not getattr(settings, model_key, None):
                setattr(settings, model_key, defaults[mode]["model"])
                changed = True
            if getattr(settings, temp_key, None) is None:
                setattr(settings, temp_key, defaults[mode]["temperature"])
                changed = True
            if getattr(settings, tokens_key, None) is None:
                setattr(settings, tokens_key, defaults[mode]["max_tokens"])
                changed = True
        if changed:
            db.add(settings)
            db.commit()
            db.refresh(settings)
        return settings

    settings = models.IntelligentControlSettings(
        director_last_mode="delfin",
        docente_mode="guepardo",
        guepardo_model=defaults["guepardo"]["model"],
        guepardo_temperature=defaults["guepardo"]["temperature"],
        guepardo_max_tokens=defaults["guepardo"]["max_tokens"],
        delfin_model=defaults["delfin"]["model"],
        delfin_temperature=defaults["delfin"]["temperature"],
        delfin_max_tokens=defaults["delfin"]["max_tokens"],
        ballena_model=defaults["ballena"]["model"],
        ballena_temperature=defaults["ballena"]["temperature"],
        ballena_max_tokens=defaults["ballena"]["max_tokens"],
    )
    db.add(settings)
    db.commit()
    db.refresh(settings)
    return settings


def normalize_intelligent_topic(topic: str | None) -> str:
    raw = normalize_header(topic or "")
    if not raw:
        return "minimum_content"
    return INTELLIGENT_TOPIC_ALIASES.get(raw, raw.replace(" ", "_"))


def normalize_associated_topics(topics: list[str] | None, main_topic: str | None = None) -> list[str]:
    if not isinstance(topics, list):
        return []
    normalized_main = normalize_intelligent_topic(main_topic or "") if main_topic else None
    result: list[str] = []
    seen: set[str] = set()
    for item in topics:
        topic = normalize_intelligent_topic(str(item or "").strip())
        if not topic:
            continue
        if normalized_main and topic == normalized_main:
            continue
        if topic in seen:
            continue
        seen.add(topic)
        result.append(topic)
    return result


def build_topic_payload(proposal: models.Proposal, topic: str) -> dict:
    normalized = normalize_intelligent_topic(topic)
    if normalized == "teaching_team":
        team = proposal.teaching_team or []
        def normalize_teacher_category(value: str | None) -> str:
            raw = normalize_header(value or "")
            if not raw:
                return "SIN INFORMAR"
            mapping = {
                "titular": "TITULAR",
                "asociado": "ASOCIADO",
                "adjunto": "ADJUNTO",
                "jtp": "JTP",
                "ayudante 1": "AYUDANTE 1º",
                "ayudante 1o": "AYUDANTE 1º",
                "ayudante 1º": "AYUDANTE 1º",
            }
            return mapping.get(raw, str(value or "").strip().upper() or "SIN INFORMAR")

        normalized_team = []
        categories_count: dict[str, int] = {}
        senior_categories = {"ADJUNTO", "ASOCIADO", "TITULAR"}
        senior_teachers = []
        for member in team:
            category = normalize_teacher_category((member or {}).get("category"))
            name = str((member or {}).get("name") or "").strip()
            normalized_member = {
                "id": (member or {}).get("id"),
                "name": name,
                "category": category,
                "email": (member or {}).get("email"),
            }
            normalized_team.append(normalized_member)
            categories_count[category] = categories_count.get(category, 0) + 1
            if category in senior_categories:
                senior_teachers.append({"name": name, "category": category})

        return {
            "topic": normalized,
            "content": {
                "team": normalized_team,
                "summary": {
                    "total_teachers": len(normalized_team),
                    "categories_count": categories_count,
                    "senior_categories": sorted(list(senior_categories)),
                    "has_senior_teacher": len(senior_teachers) > 0,
                    "senior_teachers": senior_teachers,
                },
            },
            "has_content": bool(team),
        }
    if normalized == "fundamentals":
        payload = {
            "importancia": proposal.fundamentals_part1 or "",
            "perfil_profesional": proposal.fundamentals_part2 or "",
        }
        return {
            "topic": normalized,
            "content": payload,
            "has_content": bool(payload["importancia"].strip() or payload["perfil_profesional"].strip()),
        }
    if normalized == "minimum_content":
        text = proposal.minimum_content or ""
        return {"topic": normalized, "content": text, "has_content": bool(text.strip())}
    if normalized == "learning_outcomes":
        outcomes = proposal.learning_outcomes or []
        return {"topic": normalized, "content": outcomes, "has_content": bool(outcomes)}
    if normalized == "units":
        units = proposal.units or []
        return {"topic": normalized, "content": units, "has_content": bool(units)}
    if normalized == "practicals":
        practicals = proposal.practicals or []
        return {"topic": normalized, "content": practicals, "has_content": bool(practicals)}
    if normalized == "methodology":
        text = proposal.methodology or ""
        return {"topic": normalized, "content": text, "has_content": bool(text.strip())}
    if normalized == "evaluation":
        text = proposal.evaluation or ""
        return {"topic": normalized, "content": text, "has_content": bool(text.strip())}
    if normalized == "bibliography":
        text = proposal.bibliography or ""
        return {"topic": normalized, "content": text, "has_content": bool(text.strip())}

    text = proposal.observations or ""
    return {"topic": normalized, "content": text, "has_content": bool(text.strip())}


def build_associated_topic_payloads(proposal: models.Proposal, associated_topics: list[str] | None) -> list[dict]:
    topics = normalize_associated_topics(associated_topics)
    payloads: list[dict] = []
    for topic in topics:
        topic_payload = build_topic_payload(proposal, topic)
        payloads.append({
            "topic": topic,
            "has_content": bool(topic_payload.get("has_content")),
            "content": topic_payload.get("content"),
        })
    return payloads


def parse_llm_json_response(content: str) -> dict:
    raw = (content or "").strip()
    raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"```$", "", raw).strip()

    def try_parse(candidate: str):
        text = str(candidate or "").strip()
        if not text:
            return None
        for parser in (
            lambda v: json.loads(v),
            lambda v: json.loads(re.sub(r",\s*([}\]])", r"\1", v)),
        ):
            try:
                parsed = parser(text)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                continue
        normalized_quotes = (
            text.replace("“", '"')
            .replace("”", '"')
            .replace("’", "'")
            .replace("‘", "'")
        )
        try:
            parsed = ast.literal_eval(normalized_quotes)
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            return None
        return None

    parsed = try_parse(raw)
    if parsed is not None:
        return parsed

    match = re.search(r"\{[\s\S]*\}", raw)
    if match:
        parsed = try_parse(match.group(0))
        if parsed is not None:
            return parsed

    if not raw:
        return {
            "pass": False,
            "what_failed": "No se pudo interpretar respuesta JSON del modelo.",
            "why_failed": "El modelo devolvió una respuesta vacía.",
            "suggestion": "Reintentar con más tokens o con otro modo/modelo.",
            "proposed_text": "",
            "summary": "Respuesta vacía del modelo",
        }

    return {
        "pass": False,
        "what_failed": "No se pudo interpretar respuesta JSON del modelo.",
        "why_failed": raw[:1500],
        "suggestion": "Revisar manualmente este tópico y volver a ejecutar el control.",
        "proposed_text": "",
        "summary": "Respuesta no estructurada del modelo",
    }


ENGLISH_HINT_WORDS = {
    "the", "and", "or", "with", "without", "requires", "require", "should", "must", "only",
    "because", "reason", "suggestion", "add", "include", "team", "teacher", "teachers", "senior",
    "presence", "meets", "does", "not", "fail", "failed", "rule", "rules", "compliance",
}

TERMINOLOGY_POLICY_PROMPT = (
    "Convención terminológica obligatoria: usa SIEMPRE 'Asignatura' (nunca 'Materia'), "
    "'Docente' (nunca 'Profesor') y 'Estudiante' (nunca 'Alumno'). "
    "Si el texto de entrada trae esos términos prohibidos, normalízalos respetando el sentido original."
)


def _match_case(source: str, target: str) -> str:
    src = str(source or "")
    if not src:
        return target
    if src.isupper():
        return target.upper()
    if src[:1].isupper():
        return target.capitalize()
    return target


def normalize_terminology_text(text: str) -> str:
    value = str(text or "")
    if not value:
        return value
    replacements = [
        (r"\bmateria(s)?\b", "asignatura"),
        (r"\bprofesor(?:a|es|as)?\b", "docente"),
        (r"\balumno(?:s|as)?\b", "estudiante"),
    ]
    result = value
    for pattern, replacement in replacements:
        result = re.sub(pattern, lambda m: _match_case(m.group(0), replacement), result, flags=re.IGNORECASE)
    return result


def text_looks_english(text: str) -> bool:
    value = str(text or "").strip()
    if not value:
        return False
    words = re.findall(r"[A-Za-z']+", value.lower())
    if not words:
        return False
    english_hits = sum(1 for token in words if token in ENGLISH_HINT_WORDS)
    return english_hits >= 2


def force_intelligent_feedback_spanish(data: dict, client: OpenAI) -> dict:
    if not isinstance(data, dict):
        return data

    fields = ["what_failed", "why_failed", "suggestion", "proposed_text", "summary"]
    payload = {field: str(data.get(field) or "").strip() for field in fields}
    if not any(text_looks_english(payload[field]) for field in fields):
        return data

    system_prompt = (
        "Eres traductor técnico académico. Convierte al español claro y natural el contenido recibido. "
        "Debes devolver SOLO JSON válido con las mismas claves: what_failed, why_failed, suggestion, proposed_text, summary."
    )
    user_prompt = (
        "Traduce al español (sin cambiar sentido) este JSON:\n"
        f"{json.dumps(payload, ensure_ascii=False)}"
    )
    try:
        translation = create_chat_completion_compatible(
            client,
            model=os.getenv("OPENAI_MODEL_FAST", os.getenv("OPENAI_MODEL", "gpt-4o-mini")),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0,
            max_tokens=450,
        )
        translated_raw = (translation.choices[0].message.content or "").strip()
        translated = parse_llm_json_response(translated_raw)
        for field in fields:
            translated_value = str(translated.get(field) or "").strip()
            if translated_value:
                data[field] = translated_value
    except Exception:
        return data
    for field in fields:
        data[field] = normalize_terminology_text(str(data.get(field) or ""))
    return data


def evaluate_control_with_llm(
    control: models.IntelligentControl,
    proposal: models.Proposal,
    topic_payload: dict,
    mode: str = "delfin",
    associated_context: list[dict] | None = None,
) -> dict:
    selected_mode = normalize_header(mode or "delfin")
    if selected_mode not in {"guepardo", "delfin", "ballena"}:
        selected_mode = "delfin"

    db = SessionLocal()
    try:
        persisted_settings = get_or_create_intelligent_settings(db)
        mode_settings = build_effective_intelligent_mode_settings(persisted_settings)
    finally:
        db.close()
    effective_mode = selected_mode
    topic_normalized = normalize_intelligent_topic(control.topic)
    instruction_text = str(control.instruction or "")
    senior_category_rule = bool(re.search(r"adjunt|asociad|titular", normalize_header(instruction_text)))
    if topic_normalized == "teaching_team" and senior_category_rule and selected_mode in {"guepardo", "delfin"}:
        effective_mode = "ballena"

    config = mode_settings[effective_mode]

    system_prompt = (
        "Eres un evaluador académico estricto de programas analíticos universitarios. "
        "Debes responder SOLO JSON válido con claves: "
        "pass (boolean), what_failed (string), why_failed (string), suggestion (string), proposed_text (string), summary (string). "
        "En proposed_text devuelve texto reformulado sugerido SOLO si aplica (p.ej. evaluación, metodología, contenidos, fundamentación). "
        "Si no aplica, devuelve cadena vacía. "
        "Debes escribir SIEMPRE en español (nunca en inglés). "
        "No inventes datos ni contradigas evidencia explícita del JSON. "
        "Si el JSON trae un resumen estructurado, úsalo como fuente principal para decidir. "
        f"{TERMINOLOGY_POLICY_PROMPT}"
    )
    user_prompt = (
        f"Control: {control.name}\n"
        f"Tópico: {control.topic}\n"
        f"Regla de control:\n{control.instruction}\n\n"
        f"Datos de la propuesta (JSON):\n{json.dumps(topic_payload.get('content'), ensure_ascii=False, indent=2)}\n\n"
        + (
            f"Contexto adicional asociado (JSON por tópico):\n{json.dumps(associated_context or [], ensure_ascii=False, indent=2)}\n\n"
            if associated_context else ""
        )
        + "Evalúa si cumple estrictamente la regla."
    )
    client = get_openai_client()
    response = create_chat_completion_compatible(
        client,
        model=config["model"],
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=config["temperature"],
        max_tokens=config["max_tokens"],
    )
    response_choice = response.choices[0]
    finish_reason = str(getattr(response_choice, "finish_reason", "") or "").lower()
    content = (response_choice.message.content or "").strip()
    retried_for_length = False
    used_max_tokens = int(config["max_tokens"])

    if not content and finish_reason == "length":
        retry_tokens = min(max(used_max_tokens + 300, used_max_tokens * 2), 4000)
        if retry_tokens > used_max_tokens:
            retry_response = create_chat_completion_compatible(
                client,
                model=config["model"],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=config["temperature"],
                max_tokens=retry_tokens,
            )
            response = retry_response
            response_choice = retry_response.choices[0]
            finish_reason = str(getattr(response_choice, "finish_reason", "") or "").lower()
            content = (response_choice.message.content or "").strip()
            retried_for_length = True
            used_max_tokens = retry_tokens

    data = parse_llm_json_response(content)
    for field in ["what_failed", "why_failed", "suggestion", "proposed_text", "summary"]:
        if field in data:
            data[field] = normalize_terminology_text(str(data.get(field) or ""))
    data = force_intelligent_feedback_spanish(data, client)
    passed = bool(data.get("pass") or data.get("passed") or data.get("ok"))
    return {
        "pass": passed,
        "what_failed": str(data.get("what_failed") or "").strip(),
        "why_failed": str(data.get("why_failed") or "").strip(),
        "suggestion": str(data.get("suggestion") or "").strip(),
        "proposed_text": str(data.get("proposed_text") or "").strip(),
        "summary": str(data.get("summary") or "").strip(),
        "raw_response": {
            **(data if isinstance(data, dict) else {"data": data}),
            "requested_mode": selected_mode,
            "effective_mode": effective_mode,
            "finish_reason": finish_reason,
            "retried_for_length": retried_for_length,
            "used_max_tokens": used_max_tokens,
            "model": str(config.get("model") or ""),
            "raw_content": content[:3000],
            "associated_topics": normalize_associated_topics(getattr(control, "associated_topics", None), control.topic),
        },
    }


def compute_intelligent_status(active_controls: list[models.IntelligentControl], results_by_control_id: dict[int, models.ProposalIntelligentControlResult]) -> str:
    if not active_controls:
        return "Sin ejecutar"
    if not results_by_control_id:
        return "Sin ejecutar"
    for control in active_controls:
        result = results_by_control_id.get(control.id)
        if not result:
            return "Con sugerencias"
        if not result.passed:
            return "Con sugerencias"
    return "Validada"


def build_intelligent_summary(db: Session, proposal: models.Proposal) -> dict:
    active_controls = db.query(models.IntelligentControl).filter(models.IntelligentControl.is_active == True).order_by(
        models.IntelligentControl.topic.asc(),
        models.IntelligentControl.sort_order.asc().nulls_last(),
        models.IntelligentControl.id.asc(),
    ).all()

    control_ids = [control.id for control in active_controls]
    results = []
    if control_ids:
        results = db.query(models.ProposalIntelligentControlResult).filter(
            models.ProposalIntelligentControlResult.proposal_id == proposal.id,
            models.ProposalIntelligentControlResult.control_id.in_(control_ids),
        ).all()

    results_by_control_id = {row.control_id: row for row in results}
    status = compute_intelligent_status(active_controls, results_by_control_id)
    proposal.intelligent_status = status
    db.add(proposal)
    db.commit()
    db.refresh(proposal)

    result_items = []
    passed_controls = 0
    failed_controls = 0
    last_updated = None

    controls_by_id = {control.id: control for control in active_controls}
    for control in active_controls:
        row = results_by_control_id.get(control.id)
        if not row:
            continue
        if row.passed:
            passed_controls += 1
        else:
            failed_controls += 1
        if row.checked_at and (last_updated is None or row.checked_at > last_updated):
            last_updated = row.checked_at
        result_items.append({
            "id": row.id,
            "proposal_id": row.proposal_id,
            "control_id": row.control_id,
            "control_topic": controls_by_id[row.control_id].topic,
            "control_name": controls_by_id[row.control_id].name,
            "passed": row.passed,
            "what_failed": row.what_failed,
            "why_failed": row.why_failed,
            "suggestion": row.suggestion,
            "proposed_text": row.proposed_text,
            "summary": row.summary,
            "checked_at": row.checked_at,
        })

    return {
        "proposal_id": proposal.id,
        "intelligent_status": status,
        "total_controls": len(active_controls),
        "passed_controls": passed_controls,
        "failed_controls": failed_controls,
        "results": result_items,
        "updated_at": last_updated,
    }


LEVEL_LABELS = {
    0: "Nulo",
    1: "Bajo",
    2: "Medio",
    3: "Alto",
}

LEVEL_VALUES = {
    "nulo": 0,
    "0": 0,
    "bajo": 1,
    "1": 1,
    "medio": 2,
    "2": 2,
    "alto": 3,
    "3": 3,
}

CATEGORY_ORDER = {
    "TITULAR": 5,
    "ASOCIADO": 4,
    "ADJUNTO": 3,
    "JTP": 2,
    "AYUDANTE 1º": 1,
}

CATEGORY_NORMALIZATION = {
    "titular": "TITULAR",
    "asociado": "ASOCIADO",
    "adjunto": "ADJUNTO",
    "jtp": "JTP",
    "ayudante 1": "AYUDANTE 1º",
    "ayudante 1o": "AYUDANTE 1º",
    "ayudante 1º": "AYUDANTE 1º",
}


def normalize_term_name(value: str) -> str:
    if not value:
        return ""
    normalized = str(value).strip().lower()
    if "anual" in normalized or normalized == "a":
        return "Anual"
    if "1" in normalized or "primer" in normalized:
        return "1er Cuatrimestre"
    if "2" in normalized or "segundo" in normalized:
        return "2do Cuatrimestre"
    return str(value).strip()


def parse_int(value, default=None):
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return int(value)
    text = str(value).strip()
    if not text:
        return default
    match = re.search(r"\d+", text)
    return int(match.group(0)) if match else default


def normalize_blocks(value: str) -> list[str]:
    if not value:
        return []
    raw = str(value)
    parts = [part.strip() for part in re.split(r"\s*-\s*", raw) if part.strip()]
    return parts


def get_or_create_study_plan(db: Session, career: str, name: str) -> models.StudyPlan:
    plan = db.query(models.StudyPlan).filter(
        models.StudyPlan.career == career,
        models.StudyPlan.name == name,
    ).first()
    if plan:
        return plan
    plan = models.StudyPlan(career=career, name=name)
    db.add(plan)
    db.flush()
    return plan


def get_or_create_study_year(db: Session, plan_id: int, year_number: int, label: str | None = None) -> models.StudyYear:
    year = db.query(models.StudyYear).filter(
        models.StudyYear.plan_id == plan_id,
        models.StudyYear.year_number == year_number,
    ).first()
    if year:
        return year
    year = models.StudyYear(
        plan_id=plan_id,
        year_number=year_number,
        label=label,
        sort_order=year_number,
    )
    db.add(year)
    db.flush()
    return year


def get_or_create_study_term(db: Session, year_id: int, name: str) -> models.StudyTerm:
    term = db.query(models.StudyTerm).filter(
        models.StudyTerm.year_id == year_id,
        models.StudyTerm.name == name,
    ).first()
    if term:
        return term
    sort_order = 3
    if "1" in name:
        sort_order = 1
    elif "2" in name:
        sort_order = 2
    term = models.StudyTerm(year_id=year_id, name=name, sort_order=sort_order)
    db.add(term)
    db.flush()
    return term


def find_subject_by_plan_name(db: Session, plan_id: int, name: str) -> models.StudySubject | None:
    """
    Find subject by name within a plan, normalizing accents and case.
    This allows "Programacion I" to match "Programación I"
    """
    if not name:
        return None
    
    # Get all subjects in the plan
    subjects = db.query(models.StudySubject).join(models.StudyTerm).join(models.StudyYear).filter(
        models.StudyYear.plan_id == plan_id
    ).all()
    
    # Normalize the search name
    search_normalized = normalize_header(name)
    
    # Find first matching subject (normalized comparison)
    for subject in subjects:
        if normalize_header(subject.name) == search_normalized:
            return subject
    
    return None


def get_or_create_study_subject(db: Session, plan_id: int, term_id: int, name: str) -> models.StudySubject:
    subject = find_subject_by_plan_name(db, plan_id, name)
    if subject:
        if subject.term_id != term_id:
            subject.term_id = term_id
            db.add(subject)
        return subject
    subject = models.StudySubject(term_id=term_id, name=name)
    db.add(subject)
    db.flush()
    return subject


def replace_subject_prerequisites(db: Session, subject_id: int, prerequisite_ids: list[int]) -> None:
    db.query(models.StudySubjectPrerequisite).filter(
        models.StudySubjectPrerequisite.subject_id == subject_id
    ).delete()
    for prereq_id in prerequisite_ids:
        db.add(models.StudySubjectPrerequisite(subject_id=subject_id, prerequisite_id=prereq_id))


def build_practice_scope_from_proposal(proposal: models.Proposal) -> str:
    scopes = []
    for tp in proposal.practicals or []:
        scope = ""
        if isinstance(tp, dict):
            scope = tp.get("scope") or ""
        if scope:
            scopes.append(scope.strip())
    unique = [s for s in dict.fromkeys(scopes) if s]
    return "\n".join(unique)


def sync_subject_from_proposal(db: Session, proposal: models.Proposal) -> None:
    if not proposal.career or not proposal.subject:
        return
    plan_name = proposal.study_plan or "Plan"
    plan = get_or_create_study_plan(db, proposal.career, plan_name)
    year_number = parse_int(proposal.year_of_career, default=0)
    term_name = normalize_term_name(proposal.quarter)
    year = get_or_create_study_year(db, plan.id, year_number, label=str(proposal.year_of_career or ""))
    term = get_or_create_study_term(db, year.id, term_name or "Sin Cuatrimestre")
    subject = get_or_create_study_subject(db, plan.id, term.id, proposal.subject)
    subject.character = proposal.character
    subject.regime = proposal.regime
    subject.theoretical_hours = proposal.theoretical_hours
    subject.practical_hours = proposal.practical_hours
    subject.total_hours = proposal.total_hours
    subject.weekly_hours = proposal.weekly_hours
    subject.minimum_content = proposal.minimum_content
    subject.generic_competencies = proposal.generic_competencies
    subject.specific_competencies = proposal.specific_competencies
    practice_scope = build_practice_scope_from_proposal(proposal)
    subject.practice_scope = practice_scope or subject.practice_scope
    db.add(subject)
    proposal.study_subject_id = subject.id
    db.add(proposal)


def normalize_header(value: str) -> str:
    if not value:
        return ""
    normalized = unicodedata.normalize("NFKD", str(value))
    normalized = "".join([c for c in normalized if not unicodedata.combining(c)])
    return normalized.strip().lower()


def build_subject_out(db: Session, subject: models.StudySubject) -> dict:
    prereq_ids = [
        row.prerequisite_id
        for row in db.query(models.StudySubjectPrerequisite)
        .filter(models.StudySubjectPrerequisite.subject_id == subject.id)
        .all()
    ]
    
    # Get associated proposals
    proposals = db.query(models.Proposal).filter(
        models.Proposal.study_subject_id == subject.id
    ).all()
    associated_proposals = [
        {
            "id": p.id,
            "title": p.title,
            "status": p.status,
            "gdoc_url": p.gdoc_url,
            "gdoc_status": p.gdoc_status,
        }
        for p in proposals
    ]
    
    return {
        "id": subject.id,
        "term_id": subject.term_id,
        "code": subject.code,
        "name": subject.name,
        "character": subject.character,
        "regime": subject.regime,
        "theoretical_hours": subject.theoretical_hours,
        "practical_hours": subject.practical_hours,
        "total_hours": subject.total_hours,
        "weekly_hours": subject.weekly_hours,
        "practice_scope": subject.practice_scope,
        "minimum_content": subject.minimum_content,
        "generic_competencies": subject.generic_competencies,
        "specific_competencies": subject.specific_competencies,
        "blocks": subject.blocks or [],
        "prerequisite_ids": prereq_ids,
        "associated_proposals": associated_proposals,
        "created_at": subject.created_at,
        "updated_at": subject.updated_at,
    }

DEDICATION_OPTIONS = {
    "simple": "Simple",
    "parcial": "Parcial",
    "parcial + simple": "Parcial + Simple",
    "exclusivo": "Exclusivo",
    "sin informar": "Sin Informar",
}


def normalize_competency_level(value) -> int:
    if value is None:
        return 0
    if isinstance(value, bool):
        return 0
    if isinstance(value, (int, float)):
        try:
            level = int(value)
            return level if level in (0, 1, 2, 3) else 0
        except Exception:
            return 0
    text = strip_accents(str(value)).strip().lower()
    if not text:
        return 0
    if text.isdigit():
        level = int(text)
        return level if level in (0, 1, 2, 3) else 0
    if text in LEVEL_VALUES:
        return LEVEL_VALUES[text]
    return 0


@app.get("/study-plans", response_model=list[schemas.StudyPlanOut])
def list_study_plans(career: str = "", db: Session = Depends(get_db)):
    query = db.query(models.StudyPlan)
    if career:
        query = query.filter(models.StudyPlan.career == career)
    return query.order_by(models.StudyPlan.name.asc()).all()


@app.post("/study-plans", response_model=schemas.StudyPlanOut)
def create_study_plan(payload: schemas.StudyPlanCreate, db: Session = Depends(get_db)):
    existing = db.query(models.StudyPlan).filter(
        models.StudyPlan.career == payload.career,
        models.StudyPlan.name == payload.name,
    ).first()
    if existing:
        return existing
    plan = models.StudyPlan(career=payload.career, name=payload.name)
    db.add(plan)
    db.commit()
    db.refresh(plan)
    return plan


@app.get("/study-plans-storage", response_model=list[schemas.StudyPlanStorageOut])
def list_study_plans_storage(career: str = "", db: Session = Depends(get_db)):
    query = db.query(models.StudyPlan)
    if career:
        query = query.filter(models.StudyPlan.career == career)
    return query.order_by(models.StudyPlan.name.asc()).all()


@app.post("/study-plans-storage", response_model=schemas.StudyPlanStorageOut)
def upsert_study_plan_storage(payload: schemas.StudyPlanStorageIn, db: Session = Depends(get_db)):
    plan = None
    if payload.id:
        plan = db.query(models.StudyPlan).filter(models.StudyPlan.id == payload.id).first()
    if not plan:
        plan = db.query(models.StudyPlan).filter(
            models.StudyPlan.career == payload.career,
            models.StudyPlan.name == payload.name,
        ).first()
    if not plan:
        plan = models.StudyPlan(career=payload.career, name=payload.name)
        db.add(plan)
        db.flush()

    plan.career = payload.career
    plan.name = payload.name
    if payload.is_active is not None:
        plan.is_active = bool(payload.is_active)
    if payload.payload is not None:
        plan.payload = payload.payload

    if plan.is_active:
        db.query(models.StudyPlan).filter(
            models.StudyPlan.career == payload.career,
            models.StudyPlan.id != plan.id,
        ).update({models.StudyPlan.is_active: False})

    db.commit()
    db.refresh(plan)
    return plan


@app.post("/study-plans-storage/{plan_id}/activate", response_model=schemas.StudyPlanStorageOut)
def activate_study_plan_storage(plan_id: int, db: Session = Depends(get_db)):
    plan = db.query(models.StudyPlan).filter(models.StudyPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Study plan not found")
    db.query(models.StudyPlan).filter(
        models.StudyPlan.career == plan.career,
        models.StudyPlan.id != plan.id,
    ).update({models.StudyPlan.is_active: False})
    plan.is_active = True
    db.add(plan)
    db.commit()
    db.refresh(plan)
    return plan


@app.delete("/study-plans-storage/{plan_id}")
def delete_study_plan_storage(plan_id: int, db: Session = Depends(get_db)):
    plan = db.query(models.StudyPlan).filter(models.StudyPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Study plan not found")
    db.delete(plan)
    db.commit()
    return {"status": "deleted", "id": plan_id}


@app.get("/drive-settings", response_model=schemas.DriveSettingsOut | None)
def get_drive_settings(career: str, plan_name: str | None = None, db: Session = Depends(get_db)):
    if not career:
        raise HTTPException(status_code=400, detail="Career is required")
    query = db.query(models.DriveSettings).filter(models.DriveSettings.career == career)
    if plan_name:
        query = query.filter(models.DriveSettings.plan_name == plan_name)
    else:
        query = query.filter(models.DriveSettings.plan_name.is_(None))
    settings = query.first()
    if not settings:
        return None
    return settings


@app.put("/drive-settings", response_model=schemas.DriveSettingsOut)
def upsert_drive_settings(payload: schemas.DriveSettingsCreate, db: Session = Depends(get_db)):
    career = (payload.career or "").strip()
    plan_name = (payload.plan_name or "").strip() or None
    root_folder_url = (payload.root_folder_url or "").strip() or None
    pdf_folder_url = (payload.pdf_folder_url or "").strip() or None
    if not career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not root_folder_url and not pdf_folder_url:
        raise HTTPException(status_code=400, detail="At least one Drive URL is required")
    query = db.query(models.DriveSettings).filter(models.DriveSettings.career == career)
    if plan_name:
        query = query.filter(models.DriveSettings.plan_name == plan_name)
    else:
        query = query.filter(models.DriveSettings.plan_name.is_(None))
    settings = query.first()
    if not settings:
        settings = models.DriveSettings(
            career=career,
            plan_name=plan_name,
            root_folder_url=root_folder_url,
            pdf_folder_url=pdf_folder_url,
        )
    else:
        settings.plan_name = plan_name
        settings.root_folder_url = root_folder_url
        settings.pdf_folder_url = pdf_folder_url
    db.add(settings)
    db.commit()
    db.refresh(settings)
    return settings


def serialize_accreditation_evidence(item: models.AccreditationEvidenceRegistry) -> dict:
    return {
        "id": item.id,
        "career": item.career,
        "title": item.title,
        "evidence_type": item.evidence_type,
        "source_kind": item.source_kind,
        "source_reference": item.source_reference,
        "source_file_id": item.source_file_id,
        "source_filename": item.source_filename,
        "normalized_filename": item.normalized_filename,
        "destination_folder_url": item.destination_folder_url,
        "destination_file_url": item.destination_file_url,
        "destination_file_id": item.destination_file_id,
        "checksum_sha256": item.checksum_sha256,
        "version_number": item.version_number,
        "status": item.status,
        "ocr_applied": bool(item.ocr_applied),
        "access_error": item.access_error,
        "metadata": item.metadata_json,
        "created_by": item.created_by,
        "created_at": item.created_at,
        "updated_at": item.updated_at,
    }


def normalize_evidence_source_kind(value: str | None) -> str:
    normalized = normalize_header(value or "")
    if normalized in {"local", "drive-url", "drive_url", "drive folder", "drive-folder", "drive_folder"}:
        if "folder" in normalized:
            return "drive-folder"
        if "url" in normalized:
            return "drive-url"
        return "local"
    return "local"


def infer_evidence_source_kind(source_reference: str, source_kind: str | None = None) -> str:
    explicit = normalize_evidence_source_kind(source_kind)
    if source_kind:
        return explicit
    reference = str(source_reference or "").strip().lower()
    if "/folders/" in reference:
        return "drive-folder"
    if "drive.google.com" in reference:
        return "drive-url"
    return "local"


def is_http_reference(source_reference: str | None) -> bool:
    value = str(source_reference or "").strip().lower()
    return value.startswith("http://") or value.startswith("https://")


def upload_binary_to_drive_folder(
    drive_service,
    *,
    destination_folder_id: str,
    filename: str,
    content: bytes,
    mime_type: str | None = None,
) -> tuple[str | None, str | None]:
    from googleapiclient.http import MediaInMemoryUpload

    guessed_mime, _ = mimetypes.guess_type(filename)
    upload_mime = mime_type or guessed_mime or "application/octet-stream"
    media = MediaInMemoryUpload(content or b"", mimetype=upload_mime, resumable=False)
    created = drive_service.files().create(
        body={"name": filename, "parents": [destination_folder_id]},
        media_body=media,
        fields="id, webViewLink",
        supportsAllDrives=True,
    ).execute()
    file_id = created.get("id")
    file_url = created.get("webViewLink") or (f"https://drive.google.com/file/d/{file_id}/view" if file_id else None)
    return file_id, file_url


def transfer_drive_file_to_destination(
    drive_service,
    *,
    source_file_id: str,
    destination_folder_id: str,
    move_source_file: bool,
    source_folder_id: str | None = None,
) -> tuple[str | None, str | None]:
    if move_source_file:
        current = drive_service.files().get(
            fileId=source_file_id,
            fields="id, parents, webViewLink",
            supportsAllDrives=True,
        ).execute()
        parents = current.get("parents") or []
        if source_folder_id and source_folder_id in parents:
            remove_parents = source_folder_id
        else:
            remove_candidates = [parent for parent in parents if parent and parent != destination_folder_id]
            remove_parents = ",".join(remove_candidates) if remove_candidates else None

        move_kwargs = {
            "fileId": source_file_id,
            "addParents": destination_folder_id,
            "fields": "id, webViewLink",
            "supportsAllDrives": True,
        }
        if remove_parents:
            move_kwargs["removeParents"] = remove_parents
        moved = drive_service.files().update(**move_kwargs).execute()
        moved_id = moved.get("id") or source_file_id
        moved_url = moved.get("webViewLink") or f"https://drive.google.com/file/d/{moved_id}/view"
        return moved_id, moved_url

    copied = drive_service.files().copy(
        fileId=source_file_id,
        body={"parents": [destination_folder_id]},
        fields="id, webViewLink",
        supportsAllDrives=True,
    ).execute()
    copied_id = copied.get("id")
    copied_url = copied.get("webViewLink") or (f"https://drive.google.com/file/d/{copied_id}/view" if copied_id else None)
    return copied_id, copied_url


def normalize_simple_list(values, fallback: list[str]) -> list[str]:
    if not isinstance(values, list):
        return fallback
    cleaned: list[str] = []
    seen: set[str] = set()
    for item in values:
        text = re.sub(r"\s+", " ", str(item or "").strip())
        if not text:
            continue
        key = normalize_header(text)
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(text)
    return cleaned


def normalize_accreditation_actors(values) -> list[dict]:
    if not isinstance(values, list):
        return []
    result: list[dict] = []
    seen: set[str] = set()
    for item in values:
        if not isinstance(item, dict):
            continue
        name = re.sub(r"\s+", " ", str(item.get("name") or "").strip())
        role = re.sub(r"\s+", " ", str(item.get("role") or "").strip())
        actor_type = re.sub(r"\s+", " ", str(item.get("type") or "manual").strip().lower())
        teacher_id = item.get("teacher_id")
        if not name:
            continue
        if actor_type not in {"teacher", "manual"}:
            actor_type = "manual"
        key = f"{normalize_header(name)}::{normalize_header(role)}::{actor_type}::{teacher_id or ''}"
        if key in seen:
            continue
        seen.add(key)
        result.append({
            "name": name,
            "role": role or None,
            "type": actor_type,
            "teacher_id": teacher_id,
        })
    return result


def normalize_evidence_filename(source_filename: str | None, source_reference: str | None) -> str | None:
    if source_filename and str(source_filename).strip():
        candidate = str(source_filename).strip()
    elif source_reference and str(source_reference).strip():
        raw_reference = str(source_reference).strip().rstrip("/")
        parts = [part for part in raw_reference.split("/") if part]
        candidate = parts[-1] if parts else raw_reference
    else:
        candidate = ""

    if not candidate:
        return None
    cleaned = re.sub(r"\s+", " ", candidate).strip()
    cleaned = re.sub(r"[<>:\"/\\|?*]", "", cleaned)
    return cleaned or None


def list_drive_folder_files(drive_service, folder_id: str, recursive: bool = True) -> list[dict]:
    collected: list[dict] = []

    def traverse(current_folder_id: str):
        page_token = None
        while True:
            response = drive_service.files().list(
                q=f"'{current_folder_id}' in parents and trashed = false",
                fields="nextPageToken, files(id, name, mimeType, webViewLink)",
                pageToken=page_token,
                pageSize=200,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            ).execute()
            files = response.get("files") or []
            for file_item in files:
                mime_type = str(file_item.get("mimeType") or "")
                is_folder = mime_type == "application/vnd.google-apps.folder"
                if is_folder and recursive:
                    traverse(file_item.get("id"))
                elif not is_folder:
                    collected.append(file_item)
            page_token = response.get("nextPageToken")
            if not page_token:
                break

    traverse(folder_id)
    return collected


def download_drive_file_bytes(drive_service, file_id: str) -> tuple[bytes, str | None, str | None, str | None]:
    try:
        metadata = drive_service.files().get(
            fileId=file_id,
            fields="id, name, mimeType, webViewLink",
            supportsAllDrives=True,
        ).execute()
    except Exception as exc:
        raise RuntimeError(f"No se pudo consultar metadatos del archivo de Drive: {exc}")

    mime_type = str(metadata.get("mimeType") or "")
    file_name = metadata.get("name")
    web_link = metadata.get("webViewLink")

    try:
        from googleapiclient.http import MediaIoBaseDownload  # type: ignore
    except Exception:
        raise RuntimeError("Faltan dependencias de Google Drive (google-api-python-client)")

    request = None
    if mime_type == "application/vnd.google-apps.document":
        request = drive_service.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        if file_name and not str(file_name).lower().endswith(".docx"):
            file_name = f"{file_name}.docx"
    elif mime_type == "application/vnd.google-apps.spreadsheet":
        request = drive_service.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        if file_name and not str(file_name).lower().endswith(".xlsx"):
            file_name = f"{file_name}.xlsx"
    elif mime_type == "application/vnd.google-apps.presentation":
        request = drive_service.files().export_media(
            fileId=file_id,
            mimeType="application/pdf",
        )
        if file_name and not str(file_name).lower().endswith(".pdf"):
            file_name = f"{file_name}.pdf"
    elif mime_type == "application/vnd.google-apps.folder":
        raise RuntimeError("La referencia corresponde a una carpeta, no a un archivo")
    else:
        request = drive_service.files().get_media(fileId=file_id, supportsAllDrives=True)

    output = BytesIO()
    downloader = MediaIoBaseDownload(output, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return output.getvalue(), file_name, mime_type, web_link


@app.get("/accreditation/evidences", response_model=list[schemas.AccreditationEvidenceOut])
def list_accreditation_evidences(career: str | None = None, db: Session = Depends(get_db)):
    query = db.query(models.AccreditationEvidenceRegistry)
    if career:
        query = query.filter(models.AccreditationEvidenceRegistry.career == career)
    rows = query.order_by(models.AccreditationEvidenceRegistry.created_at.desc()).all()
    return [serialize_accreditation_evidence(row) for row in rows]


@app.get("/accreditation/settings", response_model=schemas.AccreditationSettingsOut | None)
def get_accreditation_settings(career: str, study_plan: str | None = None, db: Session = Depends(get_db)):
    normalized_career = (career or "").strip()
    normalized_study_plan = (study_plan or "").strip() or None
    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not normalized_study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")
    settings = db.query(models.AccreditationSettings).filter(
        models.AccreditationSettings.career == normalized_career,
        models.AccreditationSettings.study_plan == normalized_study_plan,
    ).first()
    if not settings:
        return None
    settings.evidence_types = normalize_simple_list(settings.evidence_types, ["General"])
    settings.actor_roles = normalize_simple_list(settings.actor_roles, [])
    settings.actors = normalize_accreditation_actors(settings.actors)
    db.add(settings)
    db.commit()
    db.refresh(settings)
    return settings


@app.put("/accreditation/settings", response_model=schemas.AccreditationSettingsOut)
def upsert_accreditation_settings(payload: schemas.AccreditationSettingsCreate, db: Session = Depends(get_db)):
    normalized_career = (payload.career or "").strip()
    normalized_study_plan = (payload.study_plan or "").strip() or None
    source_folder_url = (payload.source_folder_url or "").strip() or None
    destination_folder_url = (payload.destination_folder_url or "").strip() or None
    process_mode = (payload.process_mode or "move").strip().lower() or "move"
    recursive_scan = bool(payload.recursive_scan)
    evidence_types = normalize_simple_list(payload.evidence_types, ["General"])
    actor_roles = normalize_simple_list(payload.actor_roles, [])
    actors = normalize_accreditation_actors(payload.actors)

    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not normalized_study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")
    if process_mode not in {"move", "copy"}:
        raise HTTPException(status_code=400, detail="process_mode must be 'move' or 'copy'")
    if not source_folder_url and not destination_folder_url:
        raise HTTPException(status_code=400, detail="At least one folder URL is required")

    settings = db.query(models.AccreditationSettings).filter(
        models.AccreditationSettings.career == normalized_career,
        models.AccreditationSettings.study_plan == normalized_study_plan,
    ).first()
    if not settings:
        settings = models.AccreditationSettings(
            career=normalized_career,
            study_plan=normalized_study_plan,
            source_folder_url=source_folder_url,
            destination_folder_url=destination_folder_url,
            process_mode=process_mode,
            recursive_scan=recursive_scan,
            evidence_types=evidence_types,
            actor_roles=actor_roles,
            actors=actors,
        )
    else:
        settings.study_plan = normalized_study_plan
        settings.source_folder_url = source_folder_url
        settings.destination_folder_url = destination_folder_url
        settings.process_mode = process_mode
        settings.recursive_scan = recursive_scan
        settings.evidence_types = evidence_types
        settings.actor_roles = actor_roles
        settings.actors = actors

    db.add(settings)
    db.commit()
    db.refresh(settings)
    return settings


@app.post("/accreditation/evidences", response_model=schemas.AccreditationEvidenceOut)
def create_accreditation_evidence(payload: schemas.AccreditationEvidenceCreate, db: Session = Depends(get_db)):
    created_by = (payload.created_by or "").strip()
    if not created_by:
        raise HTTPException(status_code=400, detail="Actor es obligatorio para registrar la evidencia")

    item = models.AccreditationEvidenceRegistry(
        career=payload.career,
        title=payload.title,
        evidence_type=payload.evidence_type,
        source_kind=payload.source_kind,
        source_reference=payload.source_reference,
        source_file_id=payload.source_file_id,
        source_filename=payload.source_filename,
        normalized_filename=payload.normalized_filename,
        destination_folder_url=payload.destination_folder_url,
        destination_file_url=payload.destination_file_url,
        destination_file_id=payload.destination_file_id,
        checksum_sha256=payload.checksum_sha256,
        version_number=1,
        status=payload.status,
        ocr_applied=bool(payload.ocr_applied),
        access_error=payload.access_error,
        metadata_json=payload.metadata,
        created_by=created_by,
    )
    db.add(item)
    db.flush()

    version = models.AccreditationEvidenceVersion(
        evidence_id=item.id,
        version_number=1,
        source_reference=item.source_reference,
        source_file_id=item.source_file_id,
        source_filename=item.source_filename,
        destination_file_url=item.destination_file_url,
        destination_file_id=item.destination_file_id,
        checksum_sha256=item.checksum_sha256,
        status=item.status,
        note=payload.version_note or "Registro inicial",
        created_by=created_by,
    )
    db.add(version)

    audit = models.AccreditationEvidenceAuditLog(
        evidence_id=item.id,
        action="create",
        changed_fields={"created": True},
        note=payload.version_note or "Alta de evidencia",
        actor=created_by,
    )
    db.add(audit)

    db.commit()
    db.refresh(item)
    return serialize_accreditation_evidence(item)


@app.post("/accreditation/ingest", response_model=schemas.AccreditationIngestResult)
def ingest_accreditation_evidences(payload: schemas.AccreditationIngestRequest, db: Session = Depends(get_db)):
    career = (payload.career or "").strip()
    study_plan = (payload.study_plan or "").strip() or None
    if not career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")
    if not payload.items:
        raise HTTPException(status_code=400, detail="At least one item is required")

    settings = db.query(models.AccreditationSettings).filter(
        models.AccreditationSettings.career == career,
        models.AccreditationSettings.study_plan == study_plan,
    ).first()
    if not settings:
        raise HTTPException(status_code=400, detail="No hay configuración de acreditación para la carrera")

    actor = (payload.actor or "").strip()
    if not actor:
        raise HTTPException(status_code=400, detail="Actor es obligatorio para la ingesta de evidencias")
    source_folder_url = settings.source_folder_url
    destination_folder_url = settings.destination_folder_url
    process_mode = (settings.process_mode or "move").strip().lower()
    source_folder_id = extract_drive_folder_id(source_folder_url)
    destination_folder_id = extract_drive_folder_id(destination_folder_url)
    if not destination_folder_id:
        raise HTTPException(status_code=400, detail="Configurá una carpeta DESTINO válida en Acreditación > Configuración")

    drive_service = None
    drive_service_error = None

    def get_drive_service_for_ingest():
        nonlocal drive_service, drive_service_error
        if drive_service is not None:
            return drive_service
        if drive_service_error:
            raise RuntimeError(drive_service_error)
        try:
            drive_service = get_google_drive_service()
            return drive_service
        except Exception as exc:
            drive_service_error = str(getattr(exc, "detail", "") or exc)
            raise RuntimeError(drive_service_error)

    def create_incident(
        *,
        source_kind: str,
        source_reference: str,
        source_file_id: str | None,
        source_filename: str | None,
        title: str | None,
        evidence_type: str | None,
        metadata_value,
        error_message: str,
    ):
        normalized_filename = normalize_evidence_filename(source_filename, source_reference)
        incident = models.AccreditationEvidenceRegistry(
            career=career,
            title=title,
            evidence_type=evidence_type,
            source_kind=source_kind,
            source_reference=source_reference,
            source_file_id=source_file_id,
            source_filename=source_filename,
            normalized_filename=normalized_filename,
            destination_folder_url=destination_folder_url,
            destination_file_url=None,
            destination_file_id=None,
            checksum_sha256=None,
            version_number=1,
            status="error",
            ocr_applied=False,
            access_error=error_message,
            metadata_json=metadata_value,
            created_by=actor,
        )
        db.add(incident)
        db.flush()
        db.add(models.AccreditationEvidenceVersion(
            evidence_id=incident.id,
            version_number=1,
            source_reference=source_reference,
            source_file_id=source_file_id,
            source_filename=source_filename,
            destination_file_url=None,
            destination_file_id=None,
            checksum_sha256=None,
            status="error",
            note="Incidencia de acceso/permisos durante ingesta",
            created_by=actor,
        ))
        db.add(models.AccreditationEvidenceAuditLog(
            evidence_id=incident.id,
            action="create",
            changed_fields={"access_error": error_message},
            note="Incidencia registrada en ingesta",
            actor=actor,
        ))
        return {
            "evidence_id": incident.id,
            "version_number": 1,
            "action": "incident",
            "source_kind": source_kind,
            "source_reference": source_reference,
            "source_file_id": source_file_id,
            "normalized_filename": normalized_filename,
            "status": "error",
            "access_error": error_message,
        }

    expanded_items: list[dict] = []
    for item_payload in payload.items:
        source_reference = (item_payload.source_reference or "").strip()
        if not source_reference:
            continue
        source_kind = infer_evidence_source_kind(source_reference, item_payload.source_kind)

        if source_kind == "drive-folder":
            folder_id = extract_drive_folder_id(source_reference)
            if not folder_id:
                expanded_items.append({
                    "kind": "incident",
                    "source_kind": source_kind,
                    "source_reference": source_reference,
                    "source_file_id": item_payload.source_file_id,
                    "source_filename": item_payload.source_filename,
                    "title": item_payload.title,
                    "evidence_type": item_payload.evidence_type,
                    "metadata": item_payload.metadata,
                    "error": "No se pudo extraer ID de carpeta de Drive",
                })
                continue
            try:
                service = get_drive_service_for_ingest()
                folder_files = list_drive_folder_files(service, folder_id, recursive=bool(settings.recursive_scan))
                for folder_item in folder_files:
                    file_id = folder_item.get("id")
                    file_name = folder_item.get("name")
                    web_url = folder_item.get("webViewLink") or (f"https://drive.google.com/file/d/{file_id}/view" if file_id else source_reference)
                    expanded_items.append({
                        "kind": "item",
                        "source_kind": "drive-url",
                        "source_reference": web_url,
                        "source_file_id": file_id,
                        "source_filename": file_name,
                        "title": item_payload.title,
                        "evidence_type": item_payload.evidence_type,
                        "metadata": item_payload.metadata,
                    })
                if not folder_files:
                    expanded_items.append({
                        "kind": "incident",
                        "source_kind": source_kind,
                        "source_reference": source_reference,
                        "source_file_id": None,
                        "source_filename": None,
                        "title": item_payload.title,
                        "evidence_type": item_payload.evidence_type,
                        "metadata": item_payload.metadata,
                        "error": "No se encontraron archivos en la carpeta indicada",
                    })
            except Exception as exc:
                expanded_items.append({
                    "kind": "incident",
                    "source_kind": source_kind,
                    "source_reference": source_reference,
                    "source_file_id": None,
                    "source_filename": None,
                    "title": item_payload.title,
                    "evidence_type": item_payload.evidence_type,
                    "metadata": item_payload.metadata,
                    "error": f"No se pudo listar carpeta de Drive: {exc}",
                })
            continue

        expanded_items.append({
            "kind": "item",
            "source_kind": source_kind,
            "source_reference": source_reference,
            "source_file_id": item_payload.source_file_id,
            "source_filename": item_payload.source_filename,
            "title": item_payload.title,
            "evidence_type": item_payload.evidence_type,
            "metadata": item_payload.metadata,
        })

    result_items: list[dict] = []
    created = 0
    versioned = 0
    skipped = 0

    for item_payload in expanded_items:
        source_reference = (item_payload.get("source_reference") or "").strip()
        if not source_reference:
            skipped += 1
            continue

        if item_payload.get("kind") == "incident":
            incident_item = create_incident(
                source_kind=item_payload.get("source_kind") or "local",
                source_reference=source_reference,
                source_file_id=item_payload.get("source_file_id"),
                source_filename=item_payload.get("source_filename"),
                title=item_payload.get("title"),
                evidence_type=item_payload.get("evidence_type"),
                metadata_value=item_payload.get("metadata"),
                error_message=item_payload.get("error") or "Incidencia desconocida",
            )
            created += 1
            result_items.append(incident_item)
            continue

        source_kind = infer_evidence_source_kind(source_reference, item_payload.get("source_kind"))
        source_file_id = item_payload.get("source_file_id")
        source_filename = item_payload.get("source_filename")
        destination_file_id = None
        destination_file_url = None
        transfer_error = None

        if source_kind == "drive-url":
            file_id = source_file_id or extract_drive_file_id(source_reference)
            if file_id:
                source_file_id = file_id
                try:
                    service = get_drive_service_for_ingest()
                    drive_file = service.files().get(
                        fileId=file_id,
                        fields="id, name, webViewLink, parents",
                        supportsAllDrives=True,
                    ).execute()
                    source_filename = source_filename or drive_file.get("name")
                    source_reference = drive_file.get("webViewLink") or source_reference
                    parents = drive_file.get("parents") or []
                    if destination_folder_id:
                        move_source_file = bool(process_mode == "move" and source_folder_id and source_folder_id in parents)
                        try:
                            destination_file_id, destination_file_url = transfer_drive_file_to_destination(
                                service,
                                source_file_id=source_file_id,
                                destination_folder_id=destination_folder_id,
                                move_source_file=move_source_file,
                                source_folder_id=source_folder_id,
                            )
                        except Exception as transfer_exc:
                            transfer_error = f"No se pudo transferir a carpeta destino: {transfer_exc}"
                except Exception as exc:
                    incident_item = create_incident(
                        source_kind=source_kind,
                        source_reference=source_reference,
                        source_file_id=source_file_id,
                        source_filename=source_filename,
                        title=item_payload.get("title"),
                        evidence_type=item_payload.get("evidence_type"),
                        metadata_value=item_payload.get("metadata"),
                        error_message=f"Sin acceso de lectura al recurso de Drive: {exc}",
                    )
                    created += 1
                    result_items.append(incident_item)
                    continue
        elif destination_folder_id and is_http_reference(source_reference):
            try:
                response = requests.get(source_reference, timeout=45)
                response.raise_for_status()
                content = response.content or b""
                parsed_name = source_filename or os.path.basename(source_reference.split("?", 1)[0]) or "archivo_remoto"
                guessed_mime = response.headers.get("Content-Type") or None
                service = get_drive_service_for_ingest()
                destination_file_id, destination_file_url = upload_binary_to_drive_folder(
                    service,
                    destination_folder_id=destination_folder_id,
                    filename=parsed_name,
                    content=content,
                    mime_type=guessed_mime,
                )
            except Exception as transfer_exc:
                transfer_error = f"No se pudo descargar/subir URL remota a destino: {transfer_exc}"

        normalized_filename = normalize_evidence_filename(source_filename, source_reference)

        existing_query = db.query(models.AccreditationEvidenceRegistry).filter(
            models.AccreditationEvidenceRegistry.career == career,
        )
        if normalized_filename:
            existing_query = existing_query.filter(or_(
                models.AccreditationEvidenceRegistry.source_reference == source_reference,
                models.AccreditationEvidenceRegistry.normalized_filename == normalized_filename,
            ))
        else:
            existing_query = existing_query.filter(
                models.AccreditationEvidenceRegistry.source_reference == source_reference,
            )
        existing = existing_query.first()

        if existing:
            next_version = int(existing.version_number or 1) + 1
            existing.version_number = next_version
            existing.source_kind = source_kind
            existing.source_reference = source_reference
            existing.source_filename = source_filename or existing.source_filename
            existing.normalized_filename = normalized_filename or existing.normalized_filename
            existing.source_file_id = source_file_id or existing.source_file_id
            existing.destination_folder_url = destination_folder_url or existing.destination_folder_url
            existing.destination_file_id = destination_file_id or existing.destination_file_id
            existing.destination_file_url = destination_file_url or existing.destination_file_url
            existing.status = "versioned"
            existing.access_error = transfer_error
            if item_payload.get("title"):
                existing.title = item_payload.get("title")
            if item_payload.get("evidence_type"):
                existing.evidence_type = item_payload.get("evidence_type")
            if item_payload.get("metadata") is not None:
                existing.metadata_json = item_payload.get("metadata")
            db.add(existing)

            db.add(models.AccreditationEvidenceVersion(
                evidence_id=existing.id,
                version_number=next_version,
                source_reference=source_reference,
                source_file_id=source_file_id or existing.source_file_id,
                source_filename=source_filename or existing.source_filename,
                destination_file_url=destination_file_url or existing.destination_file_url,
                destination_file_id=destination_file_id or existing.destination_file_id,
                checksum_sha256=existing.checksum_sha256,
                status="error" if transfer_error else "versioned",
                note=payload.version_note or "Nueva versión por ingesta",
                created_by=actor,
            ))
            db.add(models.AccreditationEvidenceAuditLog(
                evidence_id=existing.id,
                action="version",
                changed_fields={
                    "version_number": {"from": next_version - 1, "to": next_version},
                    "source_reference": source_reference,
                },
                note=payload.version_note or "Versionado automático",
                actor=actor,
            ))
            versioned += 1
            result_items.append({
                "evidence_id": existing.id,
                "version_number": next_version,
                "action": "versioned",
                "source_kind": source_kind,
                "source_reference": source_reference,
                "source_file_id": source_file_id,
                "normalized_filename": normalized_filename,
                "status": "error" if transfer_error else "versioned",
                "access_error": transfer_error,
            })
            continue

        record = models.AccreditationEvidenceRegistry(
            career=career,
            title=item_payload.get("title"),
            evidence_type=item_payload.get("evidence_type"),
            source_kind=source_kind,
            source_reference=source_reference,
            source_file_id=source_file_id,
            source_filename=source_filename,
            normalized_filename=normalized_filename,
            destination_folder_url=destination_folder_url,
            destination_file_url=destination_file_url,
            destination_file_id=destination_file_id,
            checksum_sha256=None,
            version_number=1,
            status="error" if transfer_error else "registered",
            ocr_applied=False,
            access_error=transfer_error,
            metadata_json=item_payload.get("metadata"),
            created_by=actor,
        )
        db.add(record)
        db.flush()

        db.add(models.AccreditationEvidenceVersion(
            evidence_id=record.id,
            version_number=1,
            source_reference=source_reference,
            source_file_id=source_file_id,
            source_filename=source_filename,
            destination_file_url=destination_file_url,
            destination_file_id=destination_file_id,
            checksum_sha256=None,
            status="error" if transfer_error else "registered",
            note=payload.version_note or "Registro inicial por ingesta",
            created_by=actor,
        ))
        db.add(models.AccreditationEvidenceAuditLog(
            evidence_id=record.id,
            action="create",
            changed_fields={"ingested": True},
            note=payload.version_note or "Alta automática por ingesta",
            actor=actor,
        ))
        created += 1
        result_items.append({
            "evidence_id": record.id,
            "version_number": 1,
            "action": "created",
            "source_kind": source_kind,
            "source_reference": source_reference,
            "source_file_id": source_file_id,
            "normalized_filename": normalized_filename,
            "status": "error" if transfer_error else "registered",
            "access_error": transfer_error,
        })

    db.commit()
    return {
        "processed": len(expanded_items),
        "created": created,
        "versioned": versioned,
        "skipped": skipped,
        "items": result_items,
    }


@app.post("/accreditation/ingest-preview", response_model=schemas.AccreditationPreviewResult)
def preview_accreditation_evidences(payload: schemas.AccreditationPreviewRequest, db: Session = Depends(get_db)):
    career = (payload.career or "").strip()
    study_plan = (payload.study_plan or "").strip() or None
    if not career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")
    if not payload.items:
        raise HTTPException(status_code=400, detail="At least one item is required")

    settings = db.query(models.AccreditationSettings).filter(
        models.AccreditationSettings.career == career,
        models.AccreditationSettings.study_plan == study_plan,
    ).first()
    if not settings:
        raise HTTPException(status_code=400, detail="No hay configuración de acreditación para la carrera")

    drive_service = None
    drive_service_error = None

    def get_drive_service_for_preview():
        nonlocal drive_service, drive_service_error
        if drive_service is not None:
            return drive_service
        if drive_service_error:
            raise RuntimeError(drive_service_error)
        try:
            drive_service = get_google_drive_service()
            return drive_service
        except Exception as exc:
            drive_service_error = str(getattr(exc, "detail", "") or exc)
            raise RuntimeError(drive_service_error)

    expanded_items: list[dict] = []
    for item_payload in payload.items:
        source_reference = (item_payload.source_reference or "").strip()
        if not source_reference:
            continue
        source_kind = infer_evidence_source_kind(source_reference, item_payload.source_kind)

        if source_kind == "drive-folder":
            folder_id = extract_drive_folder_id(source_reference)
            if not folder_id:
                expanded_items.append({
                    "kind": "incident",
                    "source_kind": source_kind,
                    "source_reference": source_reference,
                    "source_file_id": None,
                    "source_filename": None,
                    "error": "No se pudo extraer ID de carpeta de Drive",
                })
                continue
            try:
                service = get_drive_service_for_preview()
                folder_files = list_drive_folder_files(service, folder_id, recursive=bool(settings.recursive_scan))
                for folder_item in folder_files:
                    file_id = folder_item.get("id")
                    file_name = folder_item.get("name")
                    web_url = folder_item.get("webViewLink") or (f"https://drive.google.com/file/d/{file_id}/view" if file_id else source_reference)
                    expanded_items.append({
                        "kind": "item",
                        "source_kind": "drive-url",
                        "source_reference": web_url,
                        "source_file_id": file_id,
                        "source_filename": file_name,
                    })
                if not folder_files:
                    expanded_items.append({
                        "kind": "incident",
                        "source_kind": source_kind,
                        "source_reference": source_reference,
                        "source_file_id": None,
                        "source_filename": None,
                        "error": "No se encontraron archivos en la carpeta indicada",
                    })
            except Exception as exc:
                expanded_items.append({
                    "kind": "incident",
                    "source_kind": source_kind,
                    "source_reference": source_reference,
                    "source_file_id": None,
                    "source_filename": None,
                    "error": f"No se pudo listar carpeta de Drive: {exc}",
                })
            continue

        expanded_items.append({
            "kind": "item",
            "source_kind": source_kind,
            "source_reference": source_reference,
            "source_file_id": item_payload.source_file_id,
            "source_filename": item_payload.source_filename,
        })

    skipped = 0
    result_items: list[dict] = []

    for item_payload in expanded_items:
        source_reference = (item_payload.get("source_reference") or "").strip()
        if not source_reference:
            skipped += 1
            continue

        if item_payload.get("kind") == "incident":
            result_items.append({
                "source_kind": item_payload.get("source_kind") or "drive-url",
                "source_reference": source_reference,
                "source_file_id": item_payload.get("source_file_id"),
                "source_filename": item_payload.get("source_filename"),
                "normalized_filename": normalize_evidence_filename(item_payload.get("source_filename"), source_reference),
                "status": "error",
                "access_error": item_payload.get("error") or "Incidencia desconocida",
                "extraction_method": None,
                "extracted_char_count": 0,
                "ocr_applied": False,
                "preview_lines": [],
            })
            continue

        source_kind = infer_evidence_source_kind(source_reference, item_payload.get("source_kind"))
        if source_kind != "drive-url":
            result_items.append({
                "source_kind": source_kind,
                "source_reference": source_reference,
                "source_file_id": item_payload.get("source_file_id"),
                "source_filename": item_payload.get("source_filename"),
                "normalized_filename": normalize_evidence_filename(item_payload.get("source_filename"), source_reference),
                "status": "error",
                "access_error": "Solo se admite preview para referencias de Drive en esta operación",
                "extraction_method": None,
                "extracted_char_count": 0,
                "ocr_applied": False,
                "preview_lines": [],
            })
            continue

        file_id = item_payload.get("source_file_id") or extract_drive_file_id(source_reference)
        if not file_id:
            result_items.append({
                "source_kind": source_kind,
                "source_reference": source_reference,
                "source_file_id": None,
                "source_filename": item_payload.get("source_filename"),
                "normalized_filename": normalize_evidence_filename(item_payload.get("source_filename"), source_reference),
                "status": "error",
                "access_error": "No se pudo extraer ID de archivo de Drive",
                "extraction_method": None,
                "extracted_char_count": 0,
                "ocr_applied": False,
                "preview_lines": [],
            })
            continue

        try:
            service = get_drive_service_for_preview()
            content, resolved_name, _, resolved_ref = download_drive_file_bytes(service, file_id)
            source_filename = resolved_name or item_payload.get("source_filename")
            final_reference = resolved_ref or source_reference
            extraction_preview = extract_document_preview(source_filename or source_reference, content, preview_mode=True)
            result_items.append({
                "source_kind": "drive-url",
                "source_reference": final_reference,
                "source_file_id": file_id,
                "source_filename": source_filename,
                "normalized_filename": normalize_evidence_filename(source_filename, final_reference),
                "status": "previewed",
                "access_error": None,
                "extraction_method": extraction_preview.get("extraction_method"),
                "extracted_char_count": extraction_preview.get("extracted_char_count"),
                "ocr_applied": bool(extraction_preview.get("ocr_applied")),
                "preview_lines": extraction_preview.get("preview_lines") or [],
            })
        except Exception as exc:
            result_items.append({
                "source_kind": "drive-url",
                "source_reference": source_reference,
                "source_file_id": file_id,
                "source_filename": item_payload.get("source_filename"),
                "normalized_filename": normalize_evidence_filename(item_payload.get("source_filename"), source_reference),
                "status": "error",
                "access_error": f"No se pudo previsualizar archivo de Drive: {exc}",
                "extraction_method": None,
                "extracted_char_count": 0,
                "ocr_applied": False,
                "preview_lines": [],
            })

    return {
        "processed": len(expanded_items),
        "skipped": skipped,
        "items": result_items,
    }


@app.post("/accreditation/ingest-local", response_model=schemas.AccreditationIngestResult)
async def ingest_accreditation_local_files(
    career: str = Form(...),
    actor: str = Form(""),
    evidence_type: str = Form("General"),
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    normalized_career = (career or "").strip()
    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not files:
        raise HTTPException(status_code=400, detail="At least one local file is required")

    settings = db.query(models.AccreditationSettings).filter(models.AccreditationSettings.career == normalized_career).first()
    if not settings:
        raise HTTPException(status_code=400, detail="No hay configuración de acreditación para la carrera")

    actor_value = (actor or "").strip()
    if not actor_value:
        raise HTTPException(status_code=400, detail="Actor es obligatorio para la carga local de evidencias")
    evidence_type_value = (evidence_type or "").strip() or "General"
    destination_folder_id = extract_drive_folder_id(settings.destination_folder_url)
    if not destination_folder_id:
        raise HTTPException(status_code=400, detail="Configurá una carpeta DESTINO válida en Acreditación > Configuración")
    accreditation_local_dir = os.path.join(UPLOAD_FOLDER, "accreditation")
    os.makedirs(accreditation_local_dir, exist_ok=True)

    drive_service = None
    drive_service_error = None

    def get_drive_service_for_local_ingest():
        nonlocal drive_service, drive_service_error
        if drive_service is not None:
            return drive_service
        if drive_service_error:
            raise RuntimeError(drive_service_error)
        try:
            drive_service = get_google_drive_service()
            return drive_service
        except Exception as exc:
            drive_service_error = str(getattr(exc, "detail", "") or exc)
            raise RuntimeError(drive_service_error)

    created = 0
    versioned = 0
    skipped = 0
    result_items: list[dict] = []

    for upload in files:
        original_name = str(upload.filename or "").strip()
        if not original_name:
            skipped += 1
            continue

        normalized_filename = normalize_evidence_filename(original_name, original_name)
        timestamp_prefix = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        safe_name = normalized_filename or f"archivo_{timestamp_prefix}"
        target_name = f"{timestamp_prefix}_{safe_name}"
        target_path = os.path.join(accreditation_local_dir, target_name)

        content = await upload.read()
        extraction_preview = extract_document_preview(original_name, content, preview_mode=True)
        with open(target_path, "wb") as handler:
            handler.write(content)

        destination_file_id = None
        destination_file_url = None
        transfer_error = None
        if destination_folder_id:
            try:
                service = get_drive_service_for_local_ingest()
                destination_file_id, destination_file_url = upload_binary_to_drive_folder(
                    service,
                    destination_folder_id=destination_folder_id,
                    filename=original_name,
                    content=content,
                    mime_type=upload.content_type,
                )
            except Exception as transfer_exc:
                transfer_error = f"No se pudo subir archivo local a carpeta destino: {transfer_exc}"

        source_reference = target_path.replace("\\", "/")
        checksum = hashlib.sha256(content).hexdigest() if content else None

        existing = db.query(models.AccreditationEvidenceRegistry).filter(
            models.AccreditationEvidenceRegistry.career == normalized_career,
            models.AccreditationEvidenceRegistry.normalized_filename == normalized_filename,
        ).first()

        if existing:
            next_version = int(existing.version_number or 1) + 1
            existing.version_number = next_version
            existing.source_kind = "local"
            existing.source_reference = source_reference
            existing.source_filename = original_name
            existing.destination_folder_url = settings.destination_folder_url or existing.destination_folder_url
            existing.destination_file_id = destination_file_id or existing.destination_file_id
            existing.destination_file_url = destination_file_url or existing.destination_file_url
            existing.status = "error" if transfer_error else "versioned"
            existing.evidence_type = evidence_type_value
            existing.checksum_sha256 = checksum
            existing.access_error = transfer_error
            existing.ocr_applied = bool(extraction_preview.get("ocr_applied"))
            db.add(existing)

            db.add(models.AccreditationEvidenceVersion(
                evidence_id=existing.id,
                version_number=next_version,
                source_reference=source_reference,
                source_file_id=None,
                source_filename=original_name,
                destination_file_url=destination_file_url or existing.destination_file_url,
                destination_file_id=destination_file_id or existing.destination_file_id,
                checksum_sha256=checksum,
                status="error" if transfer_error else "versioned",
                note="Nueva versión por carga local",
                created_by=actor_value,
            ))
            db.add(models.AccreditationEvidenceAuditLog(
                evidence_id=existing.id,
                action="version",
                changed_fields={"version_number": {"from": next_version - 1, "to": next_version}},
                note="Carga local versionada",
                actor=actor_value,
            ))
            versioned += 1
            result_items.append({
                "evidence_id": existing.id,
                "version_number": next_version,
                "action": "versioned",
                "source_kind": "local",
                "source_reference": source_reference,
                "source_file_id": None,
                "normalized_filename": normalized_filename,
                "status": "error" if transfer_error else "versioned",
                "access_error": transfer_error,
                "extraction_method": extraction_preview.get("extraction_method"),
                "ocr_applied": bool(extraction_preview.get("ocr_applied")),
                "extracted_char_count": extraction_preview.get("extracted_char_count"),
                "preview_lines": extraction_preview.get("preview_lines") or [],
            })
            continue

        record = models.AccreditationEvidenceRegistry(
            career=normalized_career,
            title=original_name,
            evidence_type=evidence_type_value,
            source_kind="local",
            source_reference=source_reference,
            source_file_id=None,
            source_filename=original_name,
            normalized_filename=normalized_filename,
            destination_folder_url=settings.destination_folder_url,
            destination_file_url=destination_file_url,
            destination_file_id=destination_file_id,
            checksum_sha256=checksum,
            version_number=1,
            status="error" if transfer_error else "registered",
            ocr_applied=bool(extraction_preview.get("ocr_applied")),
            access_error=transfer_error,
            metadata_json=None,
            created_by=actor_value,
        )
        db.add(record)
        db.flush()
        db.add(models.AccreditationEvidenceVersion(
            evidence_id=record.id,
            version_number=1,
            source_reference=source_reference,
            source_file_id=None,
            source_filename=original_name,
            destination_file_url=destination_file_url,
            destination_file_id=destination_file_id,
            checksum_sha256=checksum,
            status="error" if transfer_error else "registered",
            note="Registro inicial por carga local",
            created_by=actor_value,
        ))
        db.add(models.AccreditationEvidenceAuditLog(
            evidence_id=record.id,
            action="create",
            changed_fields={"local_upload": True},
            note="Alta por carga local",
            actor=actor_value,
        ))
        created += 1
        result_items.append({
            "evidence_id": record.id,
            "version_number": 1,
            "action": "created",
            "source_kind": "local",
            "source_reference": source_reference,
            "source_file_id": None,
            "normalized_filename": normalized_filename,
            "status": "error" if transfer_error else "registered",
            "access_error": transfer_error,
            "extraction_method": extraction_preview.get("extraction_method"),
            "ocr_applied": bool(extraction_preview.get("ocr_applied")),
            "extracted_char_count": extraction_preview.get("extracted_char_count"),
            "preview_lines": extraction_preview.get("preview_lines") or [],
        })

    db.commit()
    return {
        "processed": len(files),
        "created": created,
        "versioned": versioned,
        "skipped": skipped,
        "items": result_items,
    }


@app.post("/accreditation/ingest-local-preview", response_model=schemas.AccreditationPreviewResult)
async def preview_accreditation_local_files(
    career: str = Form(...),
    files: list[UploadFile] = File(...),
):
    normalized_career = (career or "").strip()
    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not files:
        raise HTTPException(status_code=400, detail="At least one local file is required")

    skipped = 0
    result_items: list[dict] = []
    for upload in files:
        original_name = str(upload.filename or "").strip()
        if not original_name:
            skipped += 1
            continue
        try:
            content = await upload.read()
            extraction_preview = extract_document_preview(original_name, content, preview_mode=True)
            result_items.append({
                "source_kind": "local",
                "source_reference": original_name,
                "source_file_id": None,
                "source_filename": original_name,
                "normalized_filename": normalize_evidence_filename(original_name, original_name),
                "status": "previewed",
                "access_error": None,
                "extraction_method": extraction_preview.get("extraction_method"),
                "extracted_char_count": extraction_preview.get("extracted_char_count"),
                "ocr_applied": bool(extraction_preview.get("ocr_applied")),
                "preview_lines": extraction_preview.get("preview_lines") or [],
            })
        except Exception as exc:
            result_items.append({
                "source_kind": "local",
                "source_reference": original_name,
                "source_file_id": None,
                "source_filename": original_name,
                "normalized_filename": normalize_evidence_filename(original_name, original_name),
                "status": "error",
                "access_error": f"No se pudo previsualizar el archivo local: {exc}",
                "extraction_method": None,
                "extracted_char_count": 0,
                "ocr_applied": False,
                "preview_lines": [],
            })

    return {
        "processed": len(files),
        "skipped": skipped,
        "items": result_items,
    }


@app.patch("/accreditation/evidences/{evidence_id}", response_model=schemas.AccreditationEvidenceOut)
def update_accreditation_evidence(evidence_id: int, payload: schemas.AccreditationEvidenceUpdate, db: Session = Depends(get_db)):
    item = db.query(models.AccreditationEvidenceRegistry).filter(models.AccreditationEvidenceRegistry.id == evidence_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Evidencia no encontrada")

    updates = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    actor = updates.pop("actor", None)
    note = updates.pop("note", None)

    changed_fields = {}
    field_map = {
        "title": "title",
        "evidence_type": "evidence_type",
        "created_by": "created_by",
        "source_reference": "source_reference",
        "source_filename": "source_filename",
        "normalized_filename": "normalized_filename",
        "destination_folder_url": "destination_folder_url",
        "destination_file_url": "destination_file_url",
        "status": "status",
        "ocr_applied": "ocr_applied",
        "access_error": "access_error",
        "metadata": "metadata_json",
    }

    for payload_key, model_key in field_map.items():
        if payload_key not in updates:
            continue
        next_value = updates[payload_key]
        current_value = getattr(item, model_key)
        if current_value != next_value:
            changed_fields[payload_key] = {"from": current_value, "to": next_value}
            setattr(item, model_key, next_value)

    if not changed_fields:
        return serialize_accreditation_evidence(item)

    db.add(item)
    db.add(models.AccreditationEvidenceAuditLog(
        evidence_id=item.id,
        action="update",
        changed_fields=changed_fields,
        note=note,
        actor=actor,
    ))
    db.commit()
    db.refresh(item)
    return serialize_accreditation_evidence(item)


@app.delete("/accreditation/evidences/{evidence_id}")
def delete_accreditation_evidence(evidence_id: int, db: Session = Depends(get_db)):
    item = db.query(models.AccreditationEvidenceRegistry).filter(models.AccreditationEvidenceRegistry.id == evidence_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Evidencia no encontrada")

    db.query(models.AccreditationEvidenceVersion).filter(
        models.AccreditationEvidenceVersion.evidence_id == evidence_id
    ).delete(synchronize_session=False)
    db.query(models.AccreditationEvidenceAuditLog).filter(
        models.AccreditationEvidenceAuditLog.evidence_id == evidence_id
    ).delete(synchronize_session=False)
    db.delete(item)
    db.commit()
    return {"status": "deleted", "id": evidence_id}


@app.get("/accreditation/evidences/{evidence_id}/versions", response_model=list[schemas.AccreditationEvidenceVersionOut])
def list_accreditation_evidence_versions(evidence_id: int, db: Session = Depends(get_db)):
    exists = db.query(models.AccreditationEvidenceRegistry.id).filter(models.AccreditationEvidenceRegistry.id == evidence_id).first()
    if not exists:
        raise HTTPException(status_code=404, detail="Evidencia no encontrada")
    return db.query(models.AccreditationEvidenceVersion).filter(
        models.AccreditationEvidenceVersion.evidence_id == evidence_id
    ).order_by(models.AccreditationEvidenceVersion.version_number.desc()).all()


@app.get("/accreditation/evidences/{evidence_id}/audit", response_model=list[schemas.AccreditationEvidenceAuditOut])
def list_accreditation_evidence_audit(evidence_id: int, db: Session = Depends(get_db)):
    exists = db.query(models.AccreditationEvidenceRegistry.id).filter(models.AccreditationEvidenceRegistry.id == evidence_id).first()
    if not exists:
        raise HTTPException(status_code=404, detail="Evidencia no encontrada")
    return db.query(models.AccreditationEvidenceAuditLog).filter(
        models.AccreditationEvidenceAuditLog.evidence_id == evidence_id
    ).order_by(models.AccreditationEvidenceAuditLog.created_at.desc()).all()


WORKPLAN_ALLOWED_STATUS = {"pending", "started", "completed", "delayed", "cancelled"}


def normalize_workplan_status(value: str | None) -> str:
    status = str(value or "pending").strip().lower()
    if status not in WORKPLAN_ALLOWED_STATUS:
        raise HTTPException(status_code=400, detail="Estado inválido para plan de trabajo")
    return status


def compute_activity_number(stage_order: int, sub_stage_order: int, activity_order: int) -> str:
    return f"{int(stage_order)}.{int(sub_stage_order)}.{int(activity_order)}"


def to_naive_utc(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    if value.tzinfo is None:
        return value
    return value.astimezone(timezone.utc).replace(tzinfo=None)


def is_workplan_overdue(deadline: datetime | None) -> bool:
    normalized_deadline = to_naive_utc(deadline)
    if not normalized_deadline:
        return False
    return datetime.utcnow() > normalized_deadline


def serialize_workplan_activity(db: Session, row: models.AccreditationWorkPlanActivity):
    tasks = db.query(models.AccreditationWorkPlanTask).filter(
        models.AccreditationWorkPlanTask.activity_id == row.id
    ).order_by(models.AccreditationWorkPlanTask.status_date.asc(), models.AccreditationWorkPlanTask.id.asc()).all()
    return {
        "id": row.id,
        "career": row.career,
        "study_plan": row.study_plan,
        "stage": row.stage,
        "stage_order": row.stage_order,
        "sub_stage": row.sub_stage,
        "sub_stage_order": row.sub_stage_order,
        "activity": row.activity,
        "activity_order": row.activity_order,
        "activity_number": row.activity_number,
        "responsible_actor": row.responsible_actor,
        "collaborators": list(row.collaborators or []),
        "start_date": row.start_date,
        "deadline": row.deadline,
        "end_date": row.end_date,
        "status": row.status,
        "deadline_history": list(row.deadline_history or []),
        "observations": row.observations,
        "tasks": tasks,
        "created_at": row.created_at,
        "updated_at": row.updated_at,
    }


@app.get("/accreditation/workplan", response_model=list[schemas.AccreditationWorkPlanActivityOut])
def list_accreditation_workplan(career: str, study_plan: str | None = None, db: Session = Depends(get_db)):
    normalized_career = (career or "").strip()
    normalized_study_plan = (study_plan or "").strip() or None
    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not normalized_study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")
    rows = db.query(models.AccreditationWorkPlanActivity).filter(
        models.AccreditationWorkPlanActivity.career == normalized_career,
        models.AccreditationWorkPlanActivity.study_plan == normalized_study_plan,
    ).order_by(
        models.AccreditationWorkPlanActivity.stage_order.asc(),
        models.AccreditationWorkPlanActivity.sub_stage_order.asc(),
        models.AccreditationWorkPlanActivity.activity_order.asc(),
        models.AccreditationWorkPlanActivity.id.asc(),
    ).all()

    dirty = False
    for row in rows:
        if row.status not in {"completed", "cancelled"} and is_workplan_overdue(row.deadline):
            if row.status != "delayed":
                row.status = "delayed"
                row.end_date = None
                db.add(row)
                dirty = True

    if dirty:
        db.commit()
        for row in rows:
            db.refresh(row)

    return [serialize_workplan_activity(db, row) for row in rows]


@app.post("/accreditation/workplan/activities", response_model=schemas.AccreditationWorkPlanActivityOut)
def create_accreditation_workplan_activity(payload: schemas.AccreditationWorkPlanActivityCreate, db: Session = Depends(get_db)):
    normalized_career = (payload.career or "").strip()
    normalized_study_plan = (payload.study_plan or "").strip() or None
    if not normalized_career:
        raise HTTPException(status_code=400, detail="Career is required")
    if not normalized_study_plan:
        raise HTTPException(status_code=400, detail="Study plan is required")

    start_date = to_naive_utc(payload.start_date)
    deadline = to_naive_utc(payload.deadline)

    if deadline < start_date:
        raise HTTPException(status_code=400, detail="Deadline no puede ser anterior a Fecha Inicio")

    status = normalize_workplan_status(payload.status)
    if status not in {"completed", "cancelled"} and is_workplan_overdue(deadline):
        status = "delayed"
    end_date = datetime.utcnow() if status == "completed" else None
    activity_number = compute_activity_number(payload.stage_order, payload.sub_stage_order, payload.activity_order)

    row = models.AccreditationWorkPlanActivity(
        career=normalized_career,
        study_plan=normalized_study_plan,
        stage=(payload.stage or "").strip(),
        stage_order=int(payload.stage_order),
        sub_stage=(payload.sub_stage or "").strip(),
        sub_stage_order=int(payload.sub_stage_order),
        activity=(payload.activity or "").strip(),
        activity_order=int(payload.activity_order),
        activity_number=activity_number,
        responsible_actor=(payload.responsible_actor or "").strip() or None,
        collaborators=[str(item).strip() for item in (payload.collaborators or []) if str(item).strip()],
        start_date=start_date,
        deadline=deadline,
        end_date=end_date,
        status=status,
        deadline_history=[],
        observations=(payload.observations or "").strip() or None,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return serialize_workplan_activity(db, row)


@app.patch("/accreditation/workplan/activities/{activity_id}", response_model=schemas.AccreditationWorkPlanActivityOut)
def update_accreditation_workplan_activity(activity_id: int, payload: schemas.AccreditationWorkPlanActivityUpdate, db: Session = Depends(get_db)):
    row = db.query(models.AccreditationWorkPlanActivity).filter(models.AccreditationWorkPlanActivity.id == activity_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Actividad no encontrada")

    updates = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)

    next_stage_order = int(updates.get("stage_order", row.stage_order))
    next_sub_stage_order = int(updates.get("sub_stage_order", row.sub_stage_order))
    next_activity_order = int(updates.get("activity_order", row.activity_order))
    next_start_date = to_naive_utc(updates.get("start_date", row.start_date))
    next_deadline = to_naive_utc(updates.get("deadline", row.deadline))

    if next_deadline < next_start_date:
        raise HTTPException(status_code=400, detail="Deadline no puede ser anterior a Fecha Inicio")

    if "stage" in updates:
        row.stage = (updates.get("stage") or "").strip()
    if "sub_stage" in updates:
        row.sub_stage = (updates.get("sub_stage") or "").strip()
    if "activity" in updates:
        row.activity = (updates.get("activity") or "").strip()
    if "stage_order" in updates:
        row.stage_order = next_stage_order
    if "sub_stage_order" in updates:
        row.sub_stage_order = next_sub_stage_order
    if "activity_order" in updates:
        row.activity_order = next_activity_order
    if "responsible_actor" in updates:
        row.responsible_actor = (updates.get("responsible_actor") or "").strip() or None
    if "collaborators" in updates:
        row.collaborators = [str(item).strip() for item in (updates.get("collaborators") or []) if str(item).strip()]
    if "start_date" in updates:
        row.start_date = next_start_date
    if "observations" in updates:
        row.observations = (updates.get("observations") or "").strip() or None

    if "deadline" in updates and next_deadline != row.deadline:
        current_deadline = to_naive_utc(row.deadline)
        is_extension = bool(current_deadline and next_deadline and next_deadline > current_deadline)
        is_overdue_now = bool(current_deadline and datetime.utcnow() > current_deadline)

        if is_extension and is_overdue_now:
            history = list(row.deadline_history or [])
            history.append({
                "changed_at": datetime.utcnow().isoformat(),
                "previous_deadline": current_deadline.isoformat() if current_deadline else None,
                "new_deadline": next_deadline.isoformat() if next_deadline else None,
            })
            row.deadline_history = history

        row.deadline = next_deadline

    next_status = normalize_workplan_status(updates.get("status")) if "status" in updates else row.status
    effective_deadline = next_deadline if "deadline" in updates else to_naive_utc(row.deadline)
    if next_status not in {"completed", "cancelled"} and is_workplan_overdue(effective_deadline):
        next_status = "delayed"

    row.status = next_status
    if next_status == "completed":
        row.end_date = datetime.utcnow()
    elif next_status in {"pending", "started", "delayed", "cancelled"}:
        row.end_date = None

    row.activity_number = compute_activity_number(row.stage_order, row.sub_stage_order, row.activity_order)

    db.add(row)
    db.commit()
    db.refresh(row)
    return serialize_workplan_activity(db, row)


@app.post("/accreditation/workplan/activities/{activity_id}/tasks", response_model=schemas.AccreditationWorkPlanTaskOut)
def create_accreditation_workplan_task(activity_id: int, payload: schemas.AccreditationWorkPlanTaskCreate, db: Session = Depends(get_db)):
    activity = db.query(models.AccreditationWorkPlanActivity).filter(models.AccreditationWorkPlanActivity.id == activity_id).first()
    if not activity:
        raise HTTPException(status_code=404, detail="Actividad no encontrada")
    status = normalize_workplan_status(payload.status)
    task = models.AccreditationWorkPlanTask(
        activity_id=activity_id,
        name=(payload.name or "").strip(),
        status=status,
        status_date=payload.status_date,
        notes=(payload.notes or "").strip() or None,
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    return task


@app.patch("/accreditation/workplan/tasks/{task_id}", response_model=schemas.AccreditationWorkPlanTaskOut)
def update_accreditation_workplan_task(task_id: int, payload: schemas.AccreditationWorkPlanTaskUpdate, db: Session = Depends(get_db)):
    task = db.query(models.AccreditationWorkPlanTask).filter(models.AccreditationWorkPlanTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Tarea no encontrada")
    updates = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    if "name" in updates:
        task.name = (updates.get("name") or "").strip()
    if "status" in updates:
        task.status = normalize_workplan_status(updates.get("status"))
        now_ba = datetime.now(ZoneInfo("America/Argentina/Buenos_Aires"))
        task.status_date = now_ba.astimezone(timezone.utc).replace(tzinfo=None)
    elif "status_date" in updates:
        task.status_date = to_naive_utc(updates.get("status_date"))
    if "notes" in updates:
        task.notes = (updates.get("notes") or "").strip() or None
    db.add(task)
    db.commit()
    db.refresh(task)
    return task


@app.delete("/accreditation/workplan/tasks/{task_id}")
def delete_accreditation_workplan_task(task_id: int, db: Session = Depends(get_db)):
    task = db.query(models.AccreditationWorkPlanTask).filter(models.AccreditationWorkPlanTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Tarea no encontrada")
    db.delete(task)
    db.commit()
    return {"status": "deleted", "id": task_id}


@app.delete("/accreditation/workplan/activities/{activity_id}")
def delete_accreditation_workplan_activity(activity_id: int, db: Session = Depends(get_db)):
    activity = db.query(models.AccreditationWorkPlanActivity).filter(models.AccreditationWorkPlanActivity.id == activity_id).first()
    if not activity:
        raise HTTPException(status_code=404, detail="Actividad no encontrada")

    tasks = db.query(models.AccreditationWorkPlanTask).filter(
        models.AccreditationWorkPlanTask.activity_id == activity_id
    ).all()
    tasks_deleted = len(tasks)
    for task in tasks:
        db.delete(task)

    db.delete(activity)
    db.commit()
    return {"status": "deleted", "id": activity_id, "tasks_deleted": tasks_deleted}


@app.get("/study-plans/{plan_id}", response_model=schemas.StudyPlanDetail)
def get_study_plan(plan_id: int, db: Session = Depends(get_db)):
    plan = db.query(models.StudyPlan).filter(models.StudyPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Study plan not found")
    years = db.query(models.StudyYear).filter(models.StudyYear.plan_id == plan.id).order_by(
        models.StudyYear.sort_order.asc().nullslast(),
        models.StudyYear.year_number.asc(),
    ).all()
    years_payload = []
    for year in years:
        terms = db.query(models.StudyTerm).filter(models.StudyTerm.year_id == year.id).order_by(
            models.StudyTerm.sort_order.asc().nullslast(),
            models.StudyTerm.name.asc(),
        ).all()
        terms_payload = []
        for term in terms:
            subjects = db.query(models.StudySubject).filter(models.StudySubject.term_id == term.id).order_by(
                models.StudySubject.name.asc()
            ).all()
            subjects_payload = [build_subject_out(db, subject) for subject in subjects]
            terms_payload.append({
                "id": term.id,
                "year_id": year.id,
                "name": term.name,
                "sort_order": term.sort_order,
                "subjects": subjects_payload,
            })
        years_payload.append({
            "id": year.id,
            "plan_id": plan.id,
            "year_number": year.year_number,
            "label": year.label,
            "sort_order": year.sort_order,
            "terms": terms_payload,
        })
    return {
        "id": plan.id,
        "career": plan.career,
        "name": plan.name,
        "created_at": plan.created_at,
        "updated_at": plan.updated_at,
        "years": years_payload,
    }


@app.post("/study-plans/{plan_id}/years", response_model=schemas.StudyYearOut)
def add_study_year(plan_id: int, payload: schemas.StudyYearBase, db: Session = Depends(get_db)):
    plan = db.query(models.StudyPlan).filter(models.StudyPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Study plan not found")
    year = get_or_create_study_year(db, plan.id, payload.year_number, payload.label)
    if payload.sort_order is not None:
        year.sort_order = payload.sort_order
        db.add(year)
    db.commit()
    db.refresh(year)
    return year


@app.post("/study-years/{year_id}/terms", response_model=schemas.StudyTermOut)
def add_study_term(year_id: int, payload: schemas.StudyTermBase, db: Session = Depends(get_db)):
    year = db.query(models.StudyYear).filter(models.StudyYear.id == year_id).first()
    if not year:
        raise HTTPException(status_code=404, detail="Study year not found")
    term = get_or_create_study_term(db, year.id, payload.name)
    if payload.sort_order is not None:
        term.sort_order = payload.sort_order
        db.add(term)
    db.commit()
    db.refresh(term)
    return {
        "id": term.id,
        "year_id": term.year_id,
        "name": term.name,
        "sort_order": term.sort_order,
        "subjects": [],
    }


@app.post("/study-terms/{term_id}/subjects", response_model=schemas.StudySubjectOut)
def add_study_subject(term_id: int, payload: schemas.StudySubjectBase, db: Session = Depends(get_db)):
    term = db.query(models.StudyTerm).filter(models.StudyTerm.id == term_id).first()
    if not term:
        raise HTTPException(status_code=404, detail="Study term not found")
    plan_id = db.query(models.StudyYear.plan_id).filter(models.StudyYear.id == term.year_id).scalar()
    subject = get_or_create_study_subject(db, plan_id, term.id, payload.name)
    subject.code = payload.code
    subject.character = payload.character
    subject.regime = payload.regime
    subject.theoretical_hours = payload.theoretical_hours
    subject.practical_hours = payload.practical_hours
    subject.total_hours = payload.total_hours
    subject.weekly_hours = payload.weekly_hours
    subject.practice_scope = payload.practice_scope
    subject.minimum_content = payload.minimum_content
    subject.generic_competencies = payload.generic_competencies
    subject.specific_competencies = payload.specific_competencies
    subject.blocks = payload.blocks or []
    db.add(subject)
    replace_subject_prerequisites(db, subject.id, payload.prerequisite_ids or [])
    db.commit()
    db.refresh(subject)
    return build_subject_out(db, subject)


@app.patch("/study-subjects/{subject_id}", response_model=schemas.StudySubjectOut)
def update_study_subject(subject_id: int, payload: schemas.StudySubjectUpdate, db: Session = Depends(get_db)):
    subject = db.query(models.StudySubject).filter(models.StudySubject.id == subject_id).first()
    if not subject:
        raise HTTPException(status_code=404, detail="Study subject not found")
    data = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    prereq_ids = data.pop("prerequisite_ids", None)
    for key, value in data.items():
        setattr(subject, key, value)
    if prereq_ids is not None:
        replace_subject_prerequisites(db, subject.id, prereq_ids)
    db.add(subject)
    db.commit()
    db.refresh(subject)
    return build_subject_out(db, subject)


@app.delete("/study-subjects/{subject_id}")
def delete_study_subject(subject_id: int, db: Session = Depends(get_db)):
    subject = db.query(models.StudySubject).filter(models.StudySubject.id == subject_id).first()
    if not subject:
        raise HTTPException(status_code=404, detail="Study subject not found")
    db.delete(subject)
    db.commit()
    return {"status": "deleted", "id": subject_id}


@app.get("/study-subjects")
def list_study_subjects(career: str = "", plan: str = "", db: Session = Depends(get_db)):
    query = db.query(models.StudySubject, models.StudyTerm, models.StudyYear, models.StudyPlan).join(
        models.StudyTerm, models.StudySubject.term_id == models.StudyTerm.id
    ).join(
        models.StudyYear, models.StudyTerm.year_id == models.StudyYear.id
    ).join(
        models.StudyPlan, models.StudyYear.plan_id == models.StudyPlan.id
    )
    if career:
        query = query.filter(models.StudyPlan.career == career)
    if plan:
        query = query.filter(models.StudyPlan.name == plan)
    rows = query.order_by(models.StudyYear.year_number.asc(), models.StudyTerm.sort_order.asc()).all()
    return [
        {
            "id": subject.id,
            "code": subject.code,
            "name": subject.name,
            "year_number": year.year_number,
            "term_name": term.name,
            "plan_name": plan_row.name,
        }
        for subject, term, year, plan_row in rows
    ]


@app.post("/study-plans/import-xlsx")
async def import_study_plan_xlsx(
    career: str = Form(...),
    plan_name: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
        try:
            import openpyxl  # type: ignore
        except ImportError:
            raise HTTPException(status_code=500, detail="openpyxl no esta instalado")
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Archivo vacio")
        filename = file.filename or "Plan"
        plan_name = plan_name or os.path.splitext(filename)[0]
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        sheet = wb.active
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            raise HTTPException(status_code=400, detail="Archivo sin filas")
        header = rows[0]
        header_map = {normalize_header(value): idx for idx, value in enumerate(header)}
        def col(name: str) -> int | None:
            return header_map.get(normalize_header(name))
        idx_asignatura = col("Asignatura")
        if idx_asignatura is None:
            raise HTTPException(status_code=400, detail="Columna Asignatura no encontrada")
        plan = get_or_create_study_plan(db, career, plan_name)
        imported = 0
        for row in rows[1:]:
            subject_name = row[idx_asignatura] if idx_asignatura < len(row) else None
            if not subject_name:
                continue
            year_number = parse_int(row[col("Año")] if col("Año") is not None else None, default=0)
            term_name = normalize_term_name(row[col("Cuatrimestre")] if col("Cuatrimestre") is not None else "")
            year = get_or_create_study_year(db, plan.id, year_number, label=str(row[col("Año")]) if col("Año") is not None else None)
            term = get_or_create_study_term(db, year.id, term_name or "Sin Cuatrimestre")
            subject = get_or_create_study_subject(db, plan.id, term.id, str(subject_name).strip())
            blocks = normalize_blocks(row[col("Bloque / Campo")] if col("Bloque / Campo") is not None else "")
            subject.blocks = sorted({*(subject.blocks or []), *blocks}) if blocks else subject.blocks
            subject.practice_scope = row[col("Ámbito de Práctica")] if col("Ámbito de Práctica") is not None else subject.practice_scope
            subject.minimum_content = row[col("Contenido Mínimo")] if col("Contenido Mínimo") is not None else subject.minimum_content
            subject.character = row[col("Caracter")] if col("Caracter") is not None else subject.character
            subject.regime = row[col("Régimen")] if col("Régimen") is not None else subject.regime
            subject.total_hours = parse_int(row[col("Horas Totales")] if col("Horas Totales") is not None else None, subject.total_hours)
            subject.theoretical_hours = parse_int(row[col("Hs Teoría")] if col("Hs Teoría") is not None else None, subject.theoretical_hours)
            subject.practical_hours = parse_int(row[col("Hs Práctica")] if col("Hs Práctica") is not None else None, subject.practical_hours)
            subject.weekly_hours = parse_int(row[col("Hs Semanales")] if col("Hs Semanales") is not None else None, subject.weekly_hours)
            subject.generic_competencies = row[col("Competencias Genéricas")] if col("Competencias Genéricas") is not None else subject.generic_competencies
            subject.specific_competencies = row[col("Competencias Específicas")] if col("Competencias Específicas") is not None else subject.specific_competencies
            db.add(subject)
            imported += 1
        db.commit()
        return {"status": "ok", "plan_id": plan.id, "imported": imported}


def normalize_teacher_category(value: str | None) -> str | None:
    if not value:
        return None
    text = str(value).strip().lower()
    text = text.replace("1o", "1º").replace("1°", "1º")
    return CATEGORY_NORMALIZATION.get(text, str(value).strip().upper())


def normalize_teacher_dedication(value: str | None) -> str:
    if not value:
        return "Sin Informar"
    text = str(value).strip().lower()
    return DEDICATION_OPTIONS.get(text, "Sin Informar")


def strip_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFD", value)
    return "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")


def normalize_teacher_key(name: str | None) -> str:
    if not name:
        return ""
    cleaned = strip_accents(str(name)).lower()
    cleaned = cleaned.replace(",", " ")
    cleaned = "".join(ch if ch.isalnum() or ch.isspace() else " " for ch in cleaned)
    tokens = [token for token in cleaned.split() if token]
    tokens.sort()
    return " ".join(tokens)


def normalize_teacher_name(name: str | None) -> str:
    if not name:
        return ""
    return str(name).strip().upper()


def normalize_teacher_tokens(name: str | None) -> list[str]:
    key = normalize_teacher_key(name)
    return key.split() if key else []


def find_teacher_by_name_fuzzy(db: Session, name: str | None) -> models.Teacher | None:
    tokens = normalize_teacher_tokens(name)
    if len(tokens) < 2:
        return None
    token_set = set(tokens)
    best_match = None
    best_score = 0
    candidates = db.query(models.Teacher).filter(models.Teacher.normalized_key != "").all()
    for candidate in candidates:
        candidate_tokens = (candidate.normalized_key or "").split()
        if len(candidate_tokens) < 2:
            continue
        candidate_set = set(candidate_tokens)
        size_diff = abs(len(candidate_tokens) - len(tokens))
        if size_diff > 1:
            continue
        if token_set.issubset(candidate_set) or candidate_set.issubset(token_set):
            score = min(len(token_set), len(candidate_set))
            if score > best_score:
                best_match = candidate
                best_score = score
    return best_match


def resolve_teacher_category(existing: str | None, incoming: str | None) -> str | None:
    existing_norm = normalize_teacher_category(existing)
    incoming_norm = normalize_teacher_category(incoming)
    if not existing_norm:
        return incoming_norm
    if not incoming_norm:
        return existing_norm
    return incoming_norm if CATEGORY_ORDER.get(incoming_norm, 0) > CATEGORY_ORDER.get(existing_norm, 0) else existing_norm


def upsert_teacher(db: Session, teacher_data: dict) -> models.Teacher:
    if not teacher_data or not isinstance(teacher_data, dict):
        return None
    
    name = normalize_teacher_name(teacher_data.get("name"))
    email = (teacher_data.get("email") or "").strip().lower() or None
    category = normalize_teacher_category(teacher_data.get("category"))
    dedication = normalize_teacher_dedication(teacher_data.get("dedication"))
    normalized_key = normalize_teacher_key(name) if name else None

    # If no valid data, return None (don't create empty teacher)
    if not name and not email:
        return None

    teacher = None
    if email:
        teacher = db.query(models.Teacher).filter(models.Teacher.email == email).first()
    if not teacher and normalized_key:
        teacher = db.query(models.Teacher).filter(models.Teacher.normalized_key == normalized_key).first()
    if not teacher and name:
        teacher = find_teacher_by_name_fuzzy(db, name)

    if teacher:
        if name and (not teacher.name or len(name) > len(teacher.name)):
            teacher.name = name
        if email and not teacher.email:
            teacher.email = email
        # If no password set yet and we have email, seed it
        if not teacher.password_hash and (email or teacher.email):
            pw = email or teacher.email
            teacher.password_hash = hash_password(pw)
            teacher.password_reset_at = datetime.now(timezone.utc)
        teacher.category = resolve_teacher_category(teacher.category, category)
        if dedication and (teacher.dedication in (None, "", "Sin Informar")):
            teacher.dedication = dedication
        if normalized_key:
            teacher.normalized_key = normalized_key
        db.add(teacher)
        return teacher

    # Only create new teacher if we have at least a name
    if not name:
        return None
    
    teacher = models.Teacher(
        name=name,
        normalized_key=normalized_key,
        email=email,
        category=category,
        dedication=dedication,
    )
    if email:
        teacher.password_hash = hash_password(email)
        teacher.password_reset_at = datetime.now(timezone.utc)
    db.add(teacher)
    return teacher


def _canonical_career_name(db: Session, career: str) -> str:
    """Return the canonical (proper-case) career name from the Career table, falling back to the raw value."""
    if not career:
        return career
    raw_lower = career.lower().strip()
    match = db.query(models.Career).all()
    for c in match:
        if c.name.lower().strip() == raw_lower:
            return c.name
    return career.strip()


def ensure_teacher_career(db: Session, teacher_id: int, career: str | None) -> None:
    if not career:
        return
    # Normalize to canonical Career.name to avoid uppercase/case duplicates
    canonical = _canonical_career_name(db, career)
    existing = db.query(models.TeacherCareer).filter(
        models.TeacherCareer.teacher_id == teacher_id,
        models.TeacherCareer.career == canonical,
    ).first()
    if not existing:
        db.add(models.TeacherCareer(teacher_id=teacher_id, career=canonical))


def replace_proposal_teachers(db: Session, proposal_id: int, teacher_ids: list[int]) -> None:
    db.query(models.ProposalTeacher).filter(
        models.ProposalTeacher.proposal_id == proposal_id
    ).delete(synchronize_session=False)
    for teacher_id in teacher_ids:
        db.add(models.ProposalTeacher(
            proposal_id=proposal_id,
            teacher_id=teacher_id,
        ))


def build_teaching_team_payload(teachers: list[models.Teacher]) -> list[dict]:
    payload = []
    for teacher in teachers:
        payload.append({
            "id": teacher.id,
            "name": teacher.name,
            "category": teacher.category or "",
            "email": teacher.email or "",
            "dedication": teacher.dedication or "Sin Informar",
        })
    return payload


def build_competencies_text(items: list[dict]) -> str:
    if not items:
        return ""
    lines = []
    for item in items:
        code = (item.get("code") or "").strip()
        description = (item.get("description") or "").strip()
        level = normalize_competency_level(item.get("level"))
        level_label = LEVEL_LABELS.get(level, "Nulo")
        if code and description:
            lines.append(f"{code} - {description} - {level_label}")
        elif description:
            lines.append(f"{description} - {level_label}")
    return "\n".join(lines)


def normalize_competency_items(items: list) -> list[dict]:
    normalized = []
    for item in items or []:
        if isinstance(item, dict):
            code = (item.get("code") or "").strip()
            description = (item.get("description") or item.get("descripcion") or "").strip()
            level = normalize_competency_level(item.get("level") or item.get("level_label") or item.get("nivel"))
        else:
            data = item.model_dump() if hasattr(item, "model_dump") else item.dict()
            code = (data.get("code") or "").strip()
            description = (data.get("description") or data.get("descripcion") or "").strip()
            level = normalize_competency_level(data.get("level") or data.get("level_label") or data.get("nivel"))
        if not code and not description:
            continue
        normalized.append({
            "code": code,
            "description": description,
            "level": level,
        })
    return normalized


def replace_proposal_competencies(db: Session, proposal_id: int, items: list[dict], competency_type: str) -> None:
    db.query(models.ProposalCompetency).filter(
        models.ProposalCompetency.proposal_id == proposal_id,
        models.ProposalCompetency.competency_type == competency_type,
    ).delete(synchronize_session=False)
    for item in items:
        db.add(models.ProposalCompetency(
            proposal_id=proposal_id,
            competency_type=competency_type,
            code=item.get("code", ""),
            description=item.get("description", ""),
            level=item.get("level", 0),
        ))


def ensure_competency_catalog(
    db: Session,
    career: str,
    items: list[dict],
    competency_type: str,
    plan_name: str | None = None,
) -> None:
    if not career:
        return
    for item in items or []:
        code = (item.get("code") or "").strip()
        description = (item.get("description") or "").strip()
        if not code or not description:
            continue
        existing = db.query(models.CompetencyCatalog).filter(
            models.CompetencyCatalog.career == career,
            models.CompetencyCatalog.competency_type == competency_type,
            models.CompetencyCatalog.code == code,
            models.CompetencyCatalog.plan_name == (plan_name or None),
        ).first()
        if existing:
            if not existing.description and description:
                existing.description = description
                db.add(existing)
            continue
        db.add(models.CompetencyCatalog(
            career=career,
            plan_name=plan_name or None,
            competency_type=competency_type,
            code=code,
            description=description,
        ))


def get_proposal_competencies(db: Session, proposal_id: int) -> dict[str, list[dict]]:
    rows = db.query(models.ProposalCompetency).filter(
        models.ProposalCompetency.proposal_id == proposal_id
    ).all()
    generic_items = []
    specific_items = []
    for row in rows:
        item = {
            "id": row.id,
            "code": row.code,
            "description": row.description,
            "level": row.level,
            "level_label": LEVEL_LABELS.get(row.level, "Nulo"),
        }
        if row.competency_type == "specific":
            specific_items.append(item)
        else:
            generic_items.append(item)
    return {
        "generic": generic_items,
        "specific": specific_items,
    }


def build_proposal_response(db: Session, proposal: models.Proposal) -> dict:
    if hasattr(schemas.Proposal, "model_validate"):
        base = schemas.Proposal.model_validate(proposal).model_dump()
    else:
        base = schemas.Proposal.from_orm(proposal).dict()
    competencies = get_proposal_competencies(db, proposal.id)
    base["generic_competencies_items"] = competencies["generic"]
    base["specific_competencies_items"] = competencies["specific"]
    return base


# ── Auth helpers & CRUD ────────────────────────────────────────────────────

ADMIN_EMAIL = "admin"   # special username for the admin user
ADMIN_DEFAULT_PASSWORD = "admin"


def _build_user_payload(teacher: "models.Teacher", db) -> dict:
    """Build the JWT payload + public user info dict from a Teacher row."""
    # Compute role from career assignments
    career_assignments = []

    # 1) director/secretario assignments (from Career.director_id / secretario_id)
    careers = db.query(models.Career).filter(
        (models.Career.director_id == teacher.id) |
        (models.Career.secretario_id == teacher.id)
    ).all()
    assigned_career_names = set()
    for c in careers:
        if c.director_id == teacher.id:
            career_assignments.append({"role": "director", "careerId": c.id, "careerName": c.name})
            assigned_career_names.add(c.name)
        if c.secretario_id == teacher.id:
            career_assignments.append({"role": "secretario", "careerId": c.id, "careerName": c.name})
            assigned_career_names.add(c.name)

    # 2) docente career assignments (from TeacherCareer junction table)
    # Canonicalize career name against Career table to avoid case mismatches
    canonical_map = {c.name.lower().strip(): c.name for c in db.query(models.Career).all()}
    teacher_careers = db.query(models.TeacherCareer).filter(
        models.TeacherCareer.teacher_id == teacher.id
    ).all()
    for tc in teacher_careers:
        if not tc.career:
            continue
        career_name = canonical_map.get(tc.career.lower().strip(), tc.career)
        # Always add docente assignment even if the person is also director/secretario
        # in the same career — they should be able to switch to the docente view.
        # Only skip if this exact role+career combo already exists.
        already = any(a["role"] == "docente" and a["careerName"] == career_name for a in career_assignments)
        if not already:
            career_assignments.append({"role": "docente", "careerId": None, "careerName": career_name})

    if teacher.is_admin:
        top_role = "admin"
    elif any(a["role"] == "director" for a in career_assignments):
        top_role = "director"
    elif any(a["role"] == "secretario" for a in career_assignments):
        top_role = "secretario"
    else:
        top_role = "docente"

    return {
        "id": teacher.id,
        "name": teacher.name,
        "email": teacher.email,
        "is_admin": teacher.is_admin,
        "role": top_role,
        "career_assignments": career_assignments,
        "last_login": teacher.last_login.isoformat() if teacher.last_login else None,
    }


def seed_admin_user():
    """Ensure an admin user exists in the teachers table."""
    db = SessionLocal()
    try:
        admin = db.query(models.Teacher).filter(models.Teacher.is_admin == True).first()
        if not admin:
            admin = models.Teacher(
                name="Administrador",
                normalized_key="administrador",
                email=ADMIN_EMAIL,
                is_admin=True,
                password_hash=hash_password(ADMIN_DEFAULT_PASSWORD),
                category="Admin",
                dedication="",
            )
            db.add(admin)
            db.commit()
            print("[INFO] Admin user created (username=admin, password=admin)")
    finally:
        db.close()


def seed_default_passwords():
    """
    For any non-admin teacher without a password, set their password to their email.
    For teachers without an email, skip (they cannot log in yet).
    """
    db = SessionLocal()
    try:
        teachers = db.query(models.Teacher).filter(
            models.Teacher.is_admin == False,
            models.Teacher.password_hash == None,
            models.Teacher.email != None,
        ).all()
        for t in teachers:
            t.password_hash = hash_password(t.email)
            t.password_reset_at = datetime.now(timezone.utc)
        if teachers:
            db.commit()
            print(f"[INFO] Set default passwords for {len(teachers)} teachers (password = email)")
    finally:
        db.close()


@app.post("/auth/login", tags=["Auth"])
def login(payload: LoginRequest, db: Session = Depends(get_db)):
    username = payload.username.strip()
    password = payload.password

    # Admin login
    if username.lower() == ADMIN_EMAIL:
        teacher = db.query(models.Teacher).filter(models.Teacher.is_admin == True).first()
    else:
        # Regular user: look up by email
        teacher = db.query(models.Teacher).filter(
            func.lower(models.Teacher.email) == func.lower(username),
            models.Teacher.is_admin == False,
        ).first()

    if not teacher:
        raise HTTPException(status_code=401, detail="Usuario no encontrado.")
    if not teacher.password_hash:
        raise HTTPException(status_code=401, detail="El usuario no tiene contrase\u00f1a configurada. Contact\u00e1 al administrador.")
    if not verify_password(password, teacher.password_hash):
        raise HTTPException(status_code=401, detail="Contrase\u00f1a incorrecta.")

    # Update last_login
    teacher.last_login = datetime.now(timezone.utc)
    db.commit()
    db.refresh(teacher)

    user_info = _build_user_payload(teacher, db)
    token = create_access_token(user_info)
    return {"access_token": token, "token_type": "bearer", "user": user_info}


@app.get("/auth/me", tags=["Auth"])
def auth_me(request: Request, db: Session = Depends(get_db)):
    """Validate a stored JWT and return fresh user info."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="No autorizado.")
    token = auth_header[7:]
    data = decode_access_token(token)
    if not data:
        raise HTTPException(status_code=401, detail="Token inv\u00e1lido o expirado.")
    teacher = db.query(models.Teacher).filter(models.Teacher.id == data.get("id")).first()
    if not teacher:
        raise HTTPException(status_code=401, detail="Usuario no encontrado.")
    return _build_user_payload(teacher, db)


@app.post("/auth/reset-passwords", tags=["Auth"])
def reset_passwords(request: Request, payload: ResetPasswordsRequest, db: Session = Depends(get_db)):
    """Admin-only: reset password for one or more users."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=403, detail="No autorizado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data or not token_data.get("is_admin"):
        raise HTTPException(status_code=403, detail="Solo el administrador puede resetear contraseñas.")
    if not payload.user_ids:
        raise HTTPException(status_code=422, detail="Debe especificar al menos un usuario.")
    if not payload.use_email_as_password and len(payload.new_password) < 4:
        raise HTTPException(status_code=422, detail="La contraseña debe tener al menos 4 caracteres.")
    updated = []
    for uid in payload.user_ids:
        t = db.query(models.Teacher).filter(models.Teacher.id == uid).first()
        if not t:
            continue
        if payload.use_email_as_password:
            if t.email:
                t.password_hash = hash_password(t.email)
                t.password_reset_at = datetime.now(timezone.utc)
                updated.append(uid)
        else:
            t.password_hash = hash_password(payload.new_password)
            t.password_reset_at = datetime.now(timezone.utc)
            updated.append(uid)
    db.commit()
    return {"updated": updated, "count": len(updated)}


class ChangeMyPasswordRequest(BaseModel):
    current_password: str
    new_password: str


@app.post("/auth/change-my-password", tags=["Auth"])
def change_my_password(request: Request, payload: ChangeMyPasswordRequest, db: Session = Depends(get_db)):
    """Any authenticated user can change their own password (requires current password)."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=403, detail="No autenticado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data:
        raise HTTPException(status_code=403, detail="Token inválido o expirado.")
    teacher = db.query(models.Teacher).filter(models.Teacher.id == token_data["id"]).first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Usuario no encontrado.")
    if not verify_password(payload.current_password, teacher.password_hash):
        raise HTTPException(status_code=400, detail="La contraseña actual es incorrecta.")
    if len(payload.new_password) < 4:
        raise HTTPException(status_code=422, detail="La nueva contraseña debe tener al menos 4 caracteres.")
    teacher.password_hash = hash_password(payload.new_password)
    db.commit()
    return {"ok": True}


# ── Gmail Notifications ───────────────────────────────────────────────────────

_GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.send"]
_GMAIL_REDIRECT_URI = "http://127.0.0.1:8011/notifications/gmail/callback"


def _build_gmail_flow(redirect_uri: str = _GMAIL_REDIRECT_URI):
    """Build an OAuth2 flow for Gmail send scope."""
    from google_auth_oauthlib.flow import Flow
    client_secrets_path = os.path.join(backend_dir, "secrets", "oauth-client.json")
    if os.path.exists(client_secrets_path):
        return Flow.from_client_secrets_file(client_secrets_path, scopes=_GMAIL_SCOPES, redirect_uri=redirect_uri)
    client_id = os.getenv("GOOGLE_OAUTH_CLIENT_ID", "").strip()
    client_secret = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET", "").strip()
    if not client_id or not client_secret:
        raise HTTPException(status_code=500, detail="No se encontraron credenciales OAuth de Google.")
    client_config = {
        "installed": {
            "client_id": client_id,
            "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [redirect_uri],
        }
    }
    return Flow.from_client_config(client_config, scopes=_GMAIL_SCOPES, redirect_uri=redirect_uri)


@app.get("/notifications/gmail/status", tags=["Notifications"])
def gmail_status(request: Request, db: Session = Depends(get_db)):
    """Returns whether the current user has Gmail connected."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="No autorizado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data:
        raise HTTPException(status_code=401, detail="Token inválido.")
    teacher = db.query(models.Teacher).filter(models.Teacher.id == token_data["id"]).first()
    if not teacher:
        raise HTTPException(status_code=404)
    return {"connected": bool(teacher.gmail_refresh_token), "email": teacher.email}


@app.get("/notifications/gmail/authorize", tags=["Notifications"])
def gmail_authorize(request: Request, db: Session = Depends(get_db)):
    """Generate a Gmail OAuth2 URL for the current user."""
    import base64 as _b64, json as _json
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="No autorizado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data:
        raise HTTPException(status_code=401, detail="Token inválido.")
    state = _b64.urlsafe_b64encode(_json.dumps({"tid": token_data["id"]}).encode()).decode()
    flow = _build_gmail_flow()
    auth_url, _ = flow.authorization_url(access_type="offline", prompt="consent", state=state)
    return {"auth_url": auth_url}


@app.get("/notifications/gmail/callback", response_class=HTMLResponse, tags=["Notifications"])
def gmail_callback(code: str = None, state: str = None, error: str = None, db: Session = Depends(get_db)):
    """OAuth callback: exchange code for tokens and persist the refresh token."""
    _OK_HTML = """<!DOCTYPE html><html lang="es"><head><meta charset="utf-8">
<title>Gmail conectado</title>
<style>body{font-family:system-ui,sans-serif;display:flex;align-items:center;justify-content:center;
min-height:100vh;margin:0;background:#f0fdf4;}
.card{background:#fff;border-radius:12px;padding:40px;text-align:center;
box-shadow:0 4px 24px rgba(0,0,0,.08);max-width:400px;}
h1{color:#15803d;margin-top:0;}p{color:#555;line-height:1.6;}.icon{font-size:48px;}</style>
</head><body><div class="card"><div class="icon">✅</div>
<h1>Gmail conectado</h1>
<p>Tu cuenta de Gmail quedó vinculada.<br>Podés cerrar esta ventana.</p>
<script>setTimeout(()=>window.close(),2500);</script>
</div></body></html>"""

    _ERR = lambda msg: HTMLResponse(
        f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8"><title>Error</title>
<style>body{{font-family:system-ui,sans-serif;display:flex;align-items:center;justify-content:center;
min-height:100vh;margin:0;background:#fff5f5;}}
.card{{background:#fff;border-radius:12px;padding:40px;text-align:center;
box-shadow:0 4px 24px rgba(0,0,0,.08);max-width:400px;}}
h1{{color:#c62828;margin-top:0;}}p{{color:#555;}}</style></head>
<body><div class="card"><div style="font-size:48px">❌</div>
<h1>Error</h1><p>{msg}</p></div></body></html>""", status_code=400)

    if error:
        return _ERR(f"Google reportó: {error}")
    if not code or not state:
        return _ERR("Parámetros inválidos.")

    import base64 as _b64, json as _json, os as _os
    try:
        payload = _json.loads(_b64.urlsafe_b64decode(state + "==").decode())
        teacher_id = int(payload["tid"])
    except Exception:
        return _ERR("Estado de sesión inválido.")

    teacher = db.query(models.Teacher).filter(models.Teacher.id == teacher_id).first()
    if not teacher:
        return _ERR("Usuario no encontrado.")

    _os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
    try:
        flow = _build_gmail_flow()
        flow.fetch_token(code=code)
    except Exception as exc:
        return _ERR(f"No se pudo obtener el token: {exc}")

    credentials = flow.credentials
    if credentials.refresh_token:
        teacher.gmail_refresh_token = credentials.refresh_token
        db.commit()

    return HTMLResponse(_OK_HTML)


@app.post("/notifications/gmail/disconnect", tags=["Notifications"])
def gmail_disconnect(request: Request, db: Session = Depends(get_db)):
    """Remove stored Gmail token for the current user."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="No autorizado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data:
        raise HTTPException(status_code=401, detail="Token inválido.")
    teacher = db.query(models.Teacher).filter(models.Teacher.id == token_data["id"]).first()
    if teacher:
        teacher.gmail_refresh_token = None
        db.commit()
    return {"ok": True}


class GmailSendRequest(BaseModel):
    teacher_ids: list[int]
    subject: str
    body: str


@app.post("/notifications/gmail/send", tags=["Notifications"])
async def gmail_send(
    request: Request,
    teacher_ids: str = Form(...),
    subject: str = Form(...),
    body: str = Form(...),
    sender_name: str = Form(""),
    extra_recipients: str = Form(default="[]"),
    use_template: str = Form(default="true"),
    personal_name: str = Form(default=""),
    include_personal_name: str = Form(default="true"),
    attachments: list[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    """Send a Gmail notification to one or more teachers."""
    import json as _json, base64 as _b64, os as _os
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email.mime.image import MIMEImage
    from email.utils import formataddr as _formataddr
    from email import encoders as _encoders

    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="No autorizado.")
    token_data = decode_access_token(auth_header[7:])
    if not token_data:
        raise HTTPException(status_code=401, detail="Token inválido.")

    # Only director or secretario can send notifications
    role = token_data.get("active_role") or token_data.get("role") or ""
    if role not in ("director", "secretario", "admin"):
        raise HTTPException(status_code=403, detail="Solo el Director o Secretario pueden enviar notificaciones.")

    sender = db.query(models.Teacher).filter(models.Teacher.id == token_data["id"]).first()
    if not sender:
        raise HTTPException(status_code=404, detail="Remitente no encontrado.")
    if not sender.gmail_refresh_token:
        raise HTTPException(status_code=400, detail="Primero conectá tu cuenta de Gmail.")

    # Build Gmail credentials
    _os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
    client_id = os.getenv("GOOGLE_OAUTH_CLIENT_ID", "")
    client_secret = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET", "")
    client_secrets_path = os.path.join(backend_dir, "secrets", "oauth-client.json")
    if os.path.exists(client_secrets_path):
        try:
            import json as _j
            with open(client_secrets_path) as f:
                sec = _j.load(f)
            entry = sec.get("installed") or sec.get("web") or {}
            client_id = entry.get("client_id", client_id)
            client_secret = entry.get("client_secret", client_secret)
        except Exception:
            pass

    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build as _gbuild
        creds = Credentials(
            token=None,
            refresh_token=sender.gmail_refresh_token,
            token_uri="https://oauth2.googleapis.com/token",
            client_id=client_id,
            client_secret=client_secret,
            scopes=_GMAIL_SCOPES,
        )
        service = _gbuild("gmail", "v1", credentials=creds)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"No se pudo conectar con Gmail: {exc}")

    # Parse recipient list
    try:
        ids = _json.loads(teacher_ids)
    except Exception:
        ids = [int(x) for x in teacher_ids.split(",") if x.strip()]

    # Parse extra email recipients
    try:
        extra_emails: list[str] = _json.loads(extra_recipients)
        if not isinstance(extra_emails, list):
            extra_emails = []
    except Exception:
        extra_emails = []

    _do_template = (use_template.strip().lower() not in ("false", "0", "no"))
    _do_personal = (include_personal_name.strip().lower() not in ("false", "0", "no")) and bool((personal_name or "").strip())

    recipients = db.query(models.Teacher).filter(
        models.Teacher.id.in_(ids),
        models.Teacher.email != None,
    ).all()
    if not recipients and not extra_emails:
        raise HTTPException(status_code=422, detail="Ningún destinatario tiene email registrado.")

    # Read attachment bytes
    attachment_data = []
    for f in attachments:
        content = await f.read()
        attachment_data.append((f.filename, f.content_type or "application/octet-stream", content))

    sent = []
    errors = []

    # Logo for inline CID attachment
    _logo_path = os.path.join(backend_dir, "..", "frontend", "Logo MACAU.png")
    _logo_bytes = None
    try:
        with open(_logo_path, "rb") as _lf:
            _logo_bytes = _lf.read()
    except Exception:
        pass

    def _wrap_html(raw_body: str, _subject: str = "", _personal: str = "") -> str:
        _img_tag = (
            '<img src="cid:macau_logo" alt="MACAU"'
            ' style="height:52px;margin-bottom:12px;display:block;margin-left:auto;margin-right:auto;"/>'
        ) if _logo_bytes else ''
        _personal_line = (
            f'<div style="color:#555;font-size:12px;margin-top:8px;">Enviado por: <strong>{_personal}</strong></div>'
        ) if _personal else ''
        return (
            '<!DOCTYPE html>'
            '<html lang="es"><head><meta charset="utf-8">'
            '<style>'
            'body{margin:0;padding:0;background:#f0f4f8;font-family:Arial,Helvetica,sans-serif;}'
            'ul,ol{padding-left:24px;margin:8px 0;}'
            'li{margin:2px 0;}'
            '</style></head>'
            '<body style="margin:0;padding:0;background:#f0f4f8;">'
            '<table width="100%" cellpadding="0" cellspacing="0" style="background:#f0f4f8;padding:32px 0;">'
            '<tr><td align="center">'
            '<table width="83%" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:14px;overflow:hidden;box-shadow:0 4px 24px rgba(0,0,0,.10);">'
            '<tr><td style="background:#1a237e;padding:22px 36px;">'
            '<div style="color:#fff;font-size:20px;font-weight:700;letter-spacing:0.5px;">'
            + (_subject or 'Notificaci&oacute;n') +
            '</div>'
            '</td></tr>'
            '<tr><td style="padding:32px 36px;color:#222;font-size:14px;line-height:1.75;">'
            + raw_body +
            '</td></tr>'
            '<tr><td style="background:#f5f7fa;border-top:2px solid #e8eaf6;padding:24px 36px;text-align:center;">'
            + _img_tag +
            '<div style="color:#555;font-size:13px;margin-top:4px;">Este correo ha sido enviado desde el <strong>Sistema MACAU</strong></div>'
            '<div style="color:#aaa;font-size:11px;margin-top:4px;">Multiagente para la Acreditaci&oacute;n ante CONEAU</div>'
            + _personal_line +
            '</td></tr>'
            '</table></td></tr></table>'
            '</body></html>'
        )

    for recipient in recipients:
        try:
            outer = MIMEMultipart("mixed")
            outer["to"] = recipient.email
            _from_name = (sender_name or "").strip()
            outer["from"] = _formataddr((_from_name, sender.email)) if _from_name else (sender.email or "me")
            outer["subject"] = subject
            html_body = _wrap_html(body, subject, personal_name if _do_personal else "") if _do_template else body
            related = MIMEMultipart("related")
            related.attach(MIMEText(html_body, "html", "utf-8"))
            if _do_template and _logo_bytes:
                logo_part = MIMEImage(_logo_bytes, _subtype="png")
                logo_part.add_header("Content-ID", "<macau_logo>")
                logo_part.add_header("Content-Disposition", "inline", filename="macau_logo.png")
                related.attach(logo_part)
            outer.attach(related)
            for fname, ftype, fcontent in attachment_data:
                part = MIMEBase(*ftype.split("/", 1) if "/" in ftype else ("application", "octet-stream"))
                part.set_payload(fcontent)
                _encoders.encode_base64(part)
                part.add_header("Content-Disposition", "attachment", filename=fname)
                outer.attach(part)
            raw = _b64.urlsafe_b64encode(outer.as_bytes()).decode()
            service.users().messages().send(userId="me", body={"raw": raw}).execute()
            sent.append(recipient.email)
        except Exception as exc:
            errors.append({"email": recipient.email, "error": str(exc)})

    # Send to extra (manual) email recipients
    for em in extra_emails:
        try:
            outer = MIMEMultipart("mixed")
            outer["to"] = em
            _from_name = (sender_name or "").strip()
            outer["from"] = _formataddr((_from_name, sender.email)) if _from_name else (sender.email or "me")
            outer["subject"] = subject
            html_body = _wrap_html(body, subject, personal_name if _do_personal else "") if _do_template else body
            related = MIMEMultipart("related")
            related.attach(MIMEText(html_body, "html", "utf-8"))
            if _do_template and _logo_bytes:
                logo_part = MIMEImage(_logo_bytes, _subtype="png")
                logo_part.add_header("Content-ID", "<macau_logo>")
                logo_part.add_header("Content-Disposition", "inline", filename="macau_logo.png")
                related.attach(logo_part)
            outer.attach(related)
            for fname, ftype, fcontent in attachment_data:
                part = MIMEBase(*ftype.split("/", 1) if "/" in ftype else ("application", "octet-stream"))
                part.set_payload(fcontent)
                _encoders.encode_base64(part)
                part.add_header("Content-Disposition", "attachment", filename=fname)
                outer.attach(part)
            raw = _b64.urlsafe_b64encode(outer.as_bytes()).decode()
            service.users().messages().send(userId="me", body={"raw": raw}).execute()
            sent.append(em)
        except Exception as exc:
            errors.append({"email": em, "error": str(exc)})

    return {"sent": sent, "errors": errors, "count": len(sent)}


# ── Careers helper & CRUD ───────────────────────────────────────────────────

_DEFAULT_CAREERS = [
    "Ingeniería Mecatrónica",
    "Ingeniería en Sistemas",
    "Licenciatura en Sistemas",
    "Tecnicatura Universitaria en Ciencia de Datos",
    "Tecnicatura Universitaria en Desarrollo Web",
]


def seed_careers() -> None:
    """Seed default careers if the table is empty."""
    db = SessionLocal()
    try:
        if db.query(models.Career).count() == 0:
            for name in _DEFAULT_CAREERS:
                db.add(models.Career(name=name, is_active=True))
            db.commit()
            print(f"[INFO] Seeded {len(_DEFAULT_CAREERS)} default careers")
    finally:
        db.close()


@app.get("/careers", tags=["Careers"])
def list_careers(
    include_inactive: bool = False,
    db: Session = Depends(get_db),
):
    q = db.query(models.Career)
    if not include_inactive:
        q = q.filter(models.Career.is_active == True)
    careers = q.order_by(models.Career.name).all()
    teacher_cache: dict[int, str] = {}
    def teacher_name(tid):
        if not tid:
            return None
        if tid not in teacher_cache:
            t = db.query(models.Teacher).filter(models.Teacher.id == tid).first()
            teacher_cache[tid] = t.name if t else None
        return teacher_cache[tid]
    return [
        {
            "id": c.id,
            "name": c.name,
            "is_active": c.is_active,
            "created_at": c.created_at,
            "director_id": c.director_id,
            "secretario_id": c.secretario_id,
            "director_name": teacher_name(c.director_id),
            "secretario_name": teacher_name(c.secretario_id),
        }
        for c in careers
    ]


@app.post("/careers", response_model=schemas.CareerOut, status_code=201, tags=["Careers"])
def create_career(payload: schemas.CareerCreate, db: Session = Depends(get_db)):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="El nombre de la carrera no puede estar vacío.")
    existing = db.query(models.Career).filter(
        func.lower(models.Career.name) == func.lower(name)
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="Ya existe una carrera con ese nombre.")
    career = models.Career(name=name, is_active=payload.is_active)
    db.add(career)
    db.commit()
    db.refresh(career)
    return career


@app.patch("/careers/{career_id}", tags=["Careers"])
def update_career(career_id: int, payload: schemas.CareerUpdate, db: Session = Depends(get_db)):
    career = db.query(models.Career).filter(models.Career.id == career_id).first()
    if not career:
        raise HTTPException(status_code=404, detail="Carrera no encontrada.")
    try:
        update_data = payload.model_dump(exclude_unset=True)
    except AttributeError:
        update_data = payload.dict(exclude_unset=True)
    if "name" in update_data:
        name = update_data["name"].strip()
        if not name:
            raise HTTPException(status_code=422, detail="El nombre no puede estar vacío.")
        dup = db.query(models.Career).filter(
            func.lower(models.Career.name) == func.lower(name),
            models.Career.id != career_id,
        ).first()
        if dup:
            raise HTTPException(status_code=409, detail="Ya existe una carrera con ese nombre.")
        update_data["name"] = name
    for field, value in update_data.items():
        setattr(career, field, value)
    db.commit()
    db.refresh(career)
    # Return enriched dict with teacher names
    def tname(tid):
        if not tid: return None
        t = db.query(models.Teacher).filter(models.Teacher.id == tid).first()
        return t.name if t else None
    return {
        "id": career.id, "name": career.name, "is_active": career.is_active,
        "created_at": career.created_at,
        "director_id": career.director_id, "secretario_id": career.secretario_id,
        "director_name": tname(career.director_id), "secretario_name": tname(career.secretario_id),
    }


@app.delete("/careers/{career_id}", status_code=204, tags=["Careers"])
def delete_career(career_id: int, db: Session = Depends(get_db)):
    career = db.query(models.Career).filter(models.Career.id == career_id).first()
    if not career:
        raise HTTPException(status_code=404, detail="Carrera no encontrada.")
    db.delete(career)
    db.commit()
    return Response(status_code=204)


@app.on_event("startup")
def on_startup():
    print(f"[DEBUG] Loaded backend module: {__file__}")
    # Validate required environment variables
    required_env_vars = ["OPENAI_API_KEY"]
    missing_vars = [var for var in required_env_vars if not os.getenv(var)]
    if missing_vars:
        raise RuntimeError(f"Missing required environment variables: {', '.join(missing_vars)}. "
                          f"Please set them in .env or as environment variables.")
    
    # Initialize database
    init_db()
    sync_teachers_from_existing_proposals()
    seed_careers()
    seed_admin_user()
    seed_default_passwords()

    print("[OK] Backend startup validation passed")
    print("   - OpenAI API Key: configured")
    print("   - Database: initialized")
    print(f"   - Upload folder: {UPLOAD_FOLDER}")


@app.post("/upload", response_model=schemas.ProposalOut)
async def upload_proposal(
    file: UploadFile = File(...),
    uploader: str = Form(None),
    career: str = Form(None),
    subject: str = Form(None),
    db: Session = Depends(get_db),
    background: BackgroundTasks = None,
):
    try:
        dest_path = os.path.join(UPLOAD_FOLDER, file.filename)
        with open(dest_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    proposal = models.Proposal(
        title=subject or file.filename,  # Use subject as title, or filename as fallback
        filename=file.filename,
        original_filename=file.filename,
        source_type="upload",
        status="Importada",
        career=career,
        subject=subject,
    )
    db.add(proposal)
    db.commit()
    db.refresh(proposal)
    # launch extraction in background (creates embeddings and uploads to Pinecone)
    try:
        if background is not None:
            background.add_task(extract_agent.process_file, dest_path, proposal.id)
        else:
            # fallback synchronous
            extract_agent.process_file(dest_path, proposal.id)
    except Exception:
        # do not fail the upload if background task fails to schedule
        pass
    return proposal



@app.post("/search")
def semantic_search(q: str = Form(...), top_k: int = Form(5)):
    """Return top-k matching fragments from Pinecone for query `q`."""
    try:
        results = extract_agent.query_local(q, top_k=top_k)
        return {"matches": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/suggest")
def suggest_for_proposal(proposal_id: int = Form(...), prompt_context: str = Form(None)):
    """Generate a suggestion (text) for a proposal using nearby evidence and OpenAI.
    Returns suggestion text and used evidences.
    """
    try:
        # get top evidence for the proposal
        matches = extract_agent.query_local(f"proposal:{proposal_id}", top_k=5)
        evidence_texts = "\n\n".join([m.get("metadata", {}).get("text", "") for m in matches])
        system_prompt = (
            "Eres un asistente que ayuda a redactar la Fundamentación de una propuesta docente, "
            "usando la evidencia asociada. Devuelve un párrafo sugerido. "
            f"{TERMINOLOGY_POLICY_PROMPT}"
        )
        user_prompt = f"Evidencias:\n{evidence_texts}\n\nContexto adicional:{prompt_context or ''}\n\nGenera una sugerencia concisa para la Fundamentación."
        client = get_openai_client()
        resp = create_chat_completion_compatible(
            client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=300,
        )
        suggestion = normalize_terminology_text((resp.choices[0].message.content or "").strip())
        return {"suggestion": suggestion, "evidence_used": matches}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ai-generate")
def ai_generate(payload: AiPrompt):
    try:
        if not payload.prompt.strip():
            raise HTTPException(status_code=400, detail="Prompt is required")
        system_prompt = (
            "Eres un asistente que redacta contenido academico en espanol, claro y conciso. "
            f"{TERMINOLOGY_POLICY_PROMPT}"
        )
        client = get_openai_client()
        resp = create_chat_completion_compatible(
            client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": payload.prompt}
            ],
            max_tokens=500
        )
        content = normalize_terminology_text((resp.choices[0].message.content or "").strip())
        return {"status": "success", "content": content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ai-reformulate")
def ai_reformulate(payload: AiPrompt):
    try:
        if not payload.prompt.strip():
            raise HTTPException(status_code=400, detail="Prompt is required")
        system_prompt = (
            "Eres un asistente que reformula textos academicos manteniendo el significado. "
            "Devuelve UNICAMENTE el texto final, sin encabezados, sin preambulos y sin etiquetas como 'Sección:' o 'Texto original:'. "
            f"{TERMINOLOGY_POLICY_PROMPT}"
        )
        user_prompt = payload.prompt
        client = get_openai_client()
        resp = create_chat_completion_compatible(
            client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500
        )
        content = normalize_terminology_text((resp.choices[0].message.content or "").strip())
        return {"status": "success", "content": content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.get("/proposals/{proposal_id}", response_model=schemas.Proposal)
def get_proposal(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return build_proposal_response(db, proposal)


@app.get("/intelligent-controls", response_model=list[schemas.IntelligentControlOut])
def list_intelligent_controls(topic: str = "", db: Session = Depends(get_db)):
    query = db.query(models.IntelligentControl)
    if topic:
        query = query.filter(models.IntelligentControl.topic == normalize_intelligent_topic(topic))
    return query.order_by(
        models.IntelligentControl.topic.asc(),
        models.IntelligentControl.sort_order.asc(),
        models.IntelligentControl.id.asc(),
    ).all()


@app.get("/intelligent-controls/settings", response_model=schemas.IntelligentControlSettingsOut)
def get_intelligent_control_settings(db: Session = Depends(get_db)):
    settings = get_or_create_intelligent_settings(db)
    return serialize_intelligent_settings(settings)


@app.patch("/intelligent-controls/settings", response_model=schemas.IntelligentControlSettingsOut)
def update_intelligent_control_settings(payload: schemas.IntelligentControlSettingsUpdate, db: Session = Depends(get_db)):
    settings = get_or_create_intelligent_settings(db)
    defaults = default_intelligent_mode_settings()
    data = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)

    if "director_last_mode" in data and data["director_last_mode"] is not None:
        settings.director_last_mode = normalize_intelligent_mode(data["director_last_mode"], default="delfin")
    if "docente_mode" in data and data["docente_mode"] is not None:
        settings.docente_mode = normalize_intelligent_mode(data["docente_mode"], default="guepardo")

    for mode in ("guepardo", "delfin", "ballena"):
        mode_payload = data.get(mode)
        if not isinstance(mode_payload, dict):
            continue

        if "model" in mode_payload and mode_payload.get("model") is not None:
            model_value = str(mode_payload.get("model") or "").strip()
            if model_value:
                setattr(settings, f"{mode}_model", model_value)

        if "temperature" in mode_payload and mode_payload.get("temperature") is not None:
            setattr(
                settings,
                f"{mode}_temperature",
                sanitize_mode_temperature(mode_payload.get("temperature"), defaults[mode]["temperature"]),
            )

        if "max_tokens" in mode_payload and mode_payload.get("max_tokens") is not None:
            setattr(
                settings,
                f"{mode}_max_tokens",
                sanitize_mode_max_tokens(mode_payload.get("max_tokens"), defaults[mode]["max_tokens"]),
            )

    db.add(settings)
    db.commit()
    db.refresh(settings)
    return serialize_intelligent_settings(settings)


@app.post("/intelligent-controls", response_model=schemas.IntelligentControlOut)
def create_intelligent_control(payload: schemas.IntelligentControlCreate, db: Session = Depends(get_db)):
    main_topic = normalize_intelligent_topic(payload.topic)
    control = models.IntelligentControl(
        topic=main_topic,
        name=payload.name.strip(),
        instruction=payload.instruction.strip(),
        is_active=bool(payload.is_active) if payload.is_active is not None else True,
        sort_order=payload.sort_order,
        associated_topics=normalize_associated_topics(payload.associated_topics, main_topic),
    )
    db.add(control)
    db.commit()
    db.refresh(control)
    return control


@app.patch("/intelligent-controls/{control_id}", response_model=schemas.IntelligentControlOut)
def update_intelligent_control(control_id: int, payload: schemas.IntelligentControlUpdate, db: Session = Depends(get_db)):
    control = db.query(models.IntelligentControl).filter(models.IntelligentControl.id == control_id).first()
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")

    data = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    if "topic" in data:
        data["topic"] = normalize_intelligent_topic(data["topic"])
    if "name" in data and data["name"] is not None:
        data["name"] = str(data["name"]).strip()
    if "instruction" in data and data["instruction"] is not None:
        data["instruction"] = str(data["instruction"]).strip()
    if "associated_topics" in data:
        main_topic = data.get("topic", control.topic)
        data["associated_topics"] = normalize_associated_topics(data.get("associated_topics"), main_topic)

    for key, value in data.items():
        setattr(control, key, value)

    db.add(control)
    db.commit()
    db.refresh(control)
    return control


@app.delete("/intelligent-controls/{control_id}")
def delete_intelligent_control(control_id: int, db: Session = Depends(get_db)):
    control = db.query(models.IntelligentControl).filter(models.IntelligentControl.id == control_id).first()
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")
    db.query(models.ProposalIntelligentControlResult).filter(
        models.ProposalIntelligentControlResult.control_id == control_id
    ).delete()
    db.delete(control)
    db.commit()
    return {"status": "deleted", "id": control_id}


@app.get("/proposals/{proposal_id}/intelligent-controls/results", response_model=schemas.ProposalIntelligentControlsSummary)
def get_intelligent_control_results(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return build_intelligent_summary(db, proposal)


@app.post("/proposals/{proposal_id}/intelligent-controls/run", response_model=schemas.ProposalIntelligentControlsSummary)
def run_intelligent_controls(
    proposal_id: int,
    payload: schemas.IntelligentControlRunRequest = Body(default=None),
    db: Session = Depends(get_db),
):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    query = db.query(models.IntelligentControl).filter(models.IntelligentControl.is_active == True)
    requested_ids = []
    if payload and payload.control_ids:
        requested_ids = [int(value) for value in payload.control_ids if value is not None]
        if requested_ids:
            query = query.filter(models.IntelligentControl.id.in_(requested_ids))

    controls = query.order_by(
        models.IntelligentControl.topic.asc(),
        models.IntelligentControl.sort_order.asc(),
        models.IntelligentControl.id.asc(),
    ).all()

    selected_mode = normalize_intelligent_mode((payload.mode if payload else "delfin") or "delfin", default="delfin")

    if not controls:
        return build_intelligent_summary(db, proposal)

    for control in controls:
        topic_payload = build_topic_payload(proposal, control.topic)
        associated_context = build_associated_topic_payloads(proposal, getattr(control, "associated_topics", None))
        if not topic_payload.get("has_content"):
            llm_result = {
                "pass": False,
                "what_failed": f"El tópico '{control.topic}' no tiene contenido suficiente en la propuesta.",
                "why_failed": "No hay datos para aplicar el control inteligente.",
                "suggestion": "Completar este tópico en la propuesta y volver a ejecutar el control.",
                "proposed_text": "",
                "summary": "Control no ejecutable por falta de contenido",
                "raw_response": {"reason": "missing-topic-content"},
            }
        else:
            try:
                llm_result = evaluate_control_with_llm(
                    control,
                    proposal,
                    topic_payload,
                    mode=selected_mode,
                    associated_context=associated_context,
                )
            except Exception as exc:
                llm_result = {
                    "pass": False,
                    "what_failed": "No se pudo evaluar el control con el modelo.",
                    "why_failed": str(exc),
                    "suggestion": "Reintentar el control o revisar la configuración del modelo.",
                    "proposed_text": "",
                    "summary": "Error al ejecutar control inteligente",
                    "raw_response": {"error": str(exc)},
                }

        existing = db.query(models.ProposalIntelligentControlResult).filter(
            models.ProposalIntelligentControlResult.proposal_id == proposal.id,
            models.ProposalIntelligentControlResult.control_id == control.id,
        ).first()
        if not existing:
            existing = models.ProposalIntelligentControlResult(
                proposal_id=proposal.id,
                control_id=control.id,
            )
        existing.passed = bool(llm_result.get("pass"))
        existing.what_failed = llm_result.get("what_failed")
        existing.why_failed = llm_result.get("why_failed")
        existing.suggestion = llm_result.get("suggestion")
        existing.proposed_text = llm_result.get("proposed_text")
        existing.summary = llm_result.get("summary")
        existing.raw_response = llm_result.get("raw_response")
        existing.checked_at = datetime.utcnow()
        db.add(existing)
        db.commit()

    return build_intelligent_summary(db, proposal)


@app.patch("/proposals/{proposal_id}/intelligent-controls/results/{result_id}")
def update_intelligent_result(
    proposal_id: int,
    result_id: int,
    payload: schemas.ProposalIntelligentResultUpdate,
    db: Session = Depends(get_db),
):
    result = db.query(models.ProposalIntelligentControlResult).filter(
        models.ProposalIntelligentControlResult.id == result_id,
        models.ProposalIntelligentControlResult.proposal_id == proposal_id,
    ).first()
    if not result:
        raise HTTPException(status_code=404, detail="Resultado no encontrado")

    data = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    for key, value in data.items():
        setattr(result, key, value)
    result.checked_at = datetime.utcnow()
    db.add(result)
    db.commit()
    return {"status": "updated", "id": result.id}


# ── Lock / Unlock endpoints ──────────────────────────────────────────────────

class LockPayload(BaseModel):
    locked: bool

class BulkLockPayload(BaseModel):
    ids: list[int]
    locked: bool

@app.patch("/proposals/{proposal_id}/lock")
def set_proposal_lock(proposal_id: int, payload: LockPayload, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    proposal.editing_locked = payload.locked
    db.commit()
    db.refresh(proposal)
    return {"id": proposal.id, "editing_locked": proposal.editing_locked, "subject": proposal.subject}

@app.post("/proposals/bulk-lock")
def bulk_set_proposal_lock(payload: BulkLockPayload, db: Session = Depends(get_db)):
    updated = []
    for pid in payload.ids:
        proposal = db.query(models.Proposal).filter(models.Proposal.id == pid).first()
        if proposal:
            proposal.editing_locked = payload.locked
            updated.append(pid)
    db.commit()
    return {"updated": updated, "locked": payload.locked}

# ────────────────────────────────────────────────────────────────────────────

@app.patch("/proposals/{proposal_id}", response_model=schemas.Proposal)
def update_proposal(proposal_id: int, payload: schemas.ProposalUpdate, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    data = payload.model_dump(exclude_unset=True) if hasattr(payload, 'model_dump') else payload.dict(exclude_unset=True)

    # Si la propuesta está cerrada para edición y el payload no es solo editing_locked, rechazar
    non_lock_fields = {k for k in data if k != "editing_locked"}
    if proposal.editing_locked and non_lock_fields:
        raise HTTPException(status_code=403, detail="La propuesta está cerrada para edición. Solo un directivo puede habilitarla.")

    generic_items_raw = data.pop("generic_competencies_items", None)
    specific_items_raw = data.pop("specific_competencies_items", None)
    create_in_drive = bool(data.pop("create_in_drive", False))
    
    # Convert Pydantic models to dicts for JSON storage
    if 'learning_outcomes' in data and data['learning_outcomes']:
        data['learning_outcomes'] = [
            lo if isinstance(lo, dict) else (lo.model_dump() if hasattr(lo, 'model_dump') else lo.dict())
            for lo in data['learning_outcomes']
        ]
    if 'units' in data and data['units']:
        data['units'] = [
            u if isinstance(u, dict) else (u.model_dump() if hasattr(u, 'model_dump') else u.dict())
            for u in data['units']
        ]
    if 'practicals' in data and data['practicals']:
        data['practicals'] = [
            p if isinstance(p, dict) else (p.model_dump() if hasattr(p, 'model_dump') else p.dict())
            for p in data['practicals']
        ]
    if 'teaching_team' in data:
        data['teaching_team'] = [
            t if isinstance(t, dict) else (t.model_dump() if hasattr(t, 'model_dump') else t.dict())
            for t in (data['teaching_team'] or [])
        ]

    proposal_career = data.get("career") or proposal.career
    if generic_items_raw is not None:
        normalized = normalize_competency_items(generic_items_raw)
        data['generic_competencies'] = data.get('generic_competencies') or build_competencies_text(normalized)
        proposal_plan = data.get("study_plan") or proposal.study_plan
        ensure_competency_catalog(db, proposal_career, normalized, "generic", proposal_plan)
        replace_proposal_competencies(db, proposal.id, normalized, "generic")

    if specific_items_raw is not None:
        normalized = normalize_competency_items(specific_items_raw)
        data['specific_competencies'] = data.get('specific_competencies') or build_competencies_text(normalized)
        proposal_plan = data.get("study_plan") or proposal.study_plan
        ensure_competency_catalog(db, proposal_career, normalized, "specific", proposal_plan)
        replace_proposal_competencies(db, proposal.id, normalized, "specific")

    if 'teaching_team' in data:
        teacher_objs = []
        teacher_ids = []
        for entry in data['teaching_team'] or []:
            if not isinstance(entry, dict):
                continue
            teacher = upsert_teacher(db, entry)
            if teacher:  # Only process if teacher was created
                db.flush()
                ensure_teacher_career(db, teacher.id, proposal_career)
                teacher_objs.append(teacher)
                teacher_ids.append(teacher.id)
        if teacher_ids:
            replace_proposal_teachers(db, proposal.id, teacher_ids)
        data['teaching_team'] = build_teaching_team_payload(teacher_objs)
    
    for key, value in data.items():
        setattr(proposal, key, value)
    db.add(proposal)
    sync_subject_from_proposal(db, proposal)
    if create_in_drive and not proposal.gdoc_url:
        try_create_gdoc_for_proposal(db, proposal)
    db.commit()
    db.refresh(proposal)
    return build_proposal_response(db, proposal)


@app.delete("/proposals/{proposal_id}")
def delete_proposal(proposal_id: int, db: Session = Depends(get_db)):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    db.delete(proposal)
    db.commit()
    return {"status": "deleted", "id": proposal_id}


@app.get("/proposals", response_model=list[schemas.ProposalOut])
def list_proposals(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    proposals = db.query(models.Proposal).offset(skip).limit(limit).all()
    return [build_proposal_response(db, proposal) for proposal in proposals]


@app.post("/competencies/map-plans")
def map_competencies_to_plans(career: str = "", db: Session = Depends(get_db)):
    if not career:
        raise HTTPException(status_code=400, detail="Career is required")

    rows = db.query(models.ProposalCompetency, models.Proposal).join(models.Proposal).filter(
        models.Proposal.career == career
    ).all()

    entries = {}
    plans_by_code = {}
    for comp, proposal in rows:
        plan_name = (proposal.study_plan or "").strip()
        if not plan_name:
            continue
        code = (comp.code or "").strip()
        if not code:
            continue
        key = (comp.competency_type, code, plan_name)
        plans_by_code.setdefault((comp.competency_type, code), set()).add(plan_name)
        description = (comp.description or "").strip()
        if description and (key not in entries or len(description) > len(entries[key])):
            entries[key] = description

    created = 0
    updated = 0
    deleted = 0

    for (ctype, code, plan_name), description in entries.items():
        existing = db.query(models.CompetencyCatalog).filter(
            models.CompetencyCatalog.career == career,
            models.CompetencyCatalog.competency_type == ctype,
            models.CompetencyCatalog.code == code,
            models.CompetencyCatalog.plan_name == plan_name,
        ).first()
        if existing:
            if description and not existing.description:
                existing.description = description
                db.add(existing)
                updated += 1
            continue
        db.add(models.CompetencyCatalog(
            career=career,
            plan_name=plan_name,
            competency_type=ctype,
            code=code,
            description=description or "",
        ))
        created += 1

    unscoped_items = db.query(models.CompetencyCatalog).filter(
        models.CompetencyCatalog.career == career,
        models.CompetencyCatalog.plan_name.is_(None),
    ).all()
    for item in unscoped_items:
        code_key = (item.competency_type, (item.code or "").strip())
        plan_names = sorted(plans_by_code.get(code_key, set()))
        if len(plan_names) == 1:
            item.plan_name = plan_names[0]
            db.add(item)
            updated += 1
            continue
        if len(plan_names) > 1:
            for plan_name in plan_names:
                existing = db.query(models.CompetencyCatalog).filter(
                    models.CompetencyCatalog.career == career,
                    models.CompetencyCatalog.competency_type == item.competency_type,
                    models.CompetencyCatalog.code == item.code,
                    models.CompetencyCatalog.plan_name == plan_name,
                ).first()
                if existing:
                    continue
                description = entries.get((item.competency_type, item.code, plan_name)) or item.description or ""
                db.add(models.CompetencyCatalog(
                    career=career,
                    plan_name=plan_name,
                    competency_type=item.competency_type,
                    code=item.code,
                    description=description,
                ))
                created += 1
            db.delete(item)
            deleted += 1

    db.commit()

    catalog_items = db.query(models.CompetencyCatalog).filter(
        models.CompetencyCatalog.career == career,
    ).all()
    grouped = {}
    for item in catalog_items:
        key = (
            (item.plan_name or "").strip().lower(),
            (item.competency_type or "").strip().lower(),
            (item.code or "").strip().lower(),
        )
        grouped.setdefault(key, []).append(item)

    for key, items in grouped.items():
        if len(items) <= 1:
            continue
        def score(entry):
            description = (entry.description or "").strip()
            updated_at = entry.updated_at or entry.created_at
            timestamp = updated_at.timestamp() if updated_at else 0
            return (len(description), timestamp, entry.id)
        items_sorted = sorted(items, key=score, reverse=True)
        keep = items_sorted[0]
        for duplicate in items_sorted[1:]:
            db.delete(duplicate)
            deleted += 1
        db.add(keep)

    db.commit()
    return {
        "career": career,
        "created": created,
        "updated": updated,
        "deleted": deleted,
        "mapped": len(entries),
    }


@app.get("/competencies", response_model=list[schemas.CompetencyCatalogOut])
def list_competencies(
    career: str = "",
    competency_type: str = "",
    plan_name: str = "",
    db: Session = Depends(get_db),
):
    query = db.query(models.CompetencyCatalog)
    if career:
        query = query.filter(models.CompetencyCatalog.career == career)
    if competency_type:
        query = query.filter(models.CompetencyCatalog.competency_type == competency_type)
    if plan_name:
        query = query.filter(models.CompetencyCatalog.plan_name == plan_name)
    return query.order_by(models.CompetencyCatalog.code.asc()).all()


@app.post("/competencies", response_model=schemas.CompetencyCatalogOut)
def create_competency(payload: schemas.CompetencyCatalogCreate, db: Session = Depends(get_db)):
    if not payload.plan_name:
        raise HTTPException(status_code=400, detail="Competency plan is required")
    existing = db.query(models.CompetencyCatalog).filter(
        models.CompetencyCatalog.career == payload.career,
        models.CompetencyCatalog.competency_type == payload.competency_type,
        func.lower(models.CompetencyCatalog.code) == payload.code.strip().lower(),
        func.coalesce(models.CompetencyCatalog.plan_name, "") == (payload.plan_name or ""),
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="Competency code already exists")
    item = models.CompetencyCatalog(
        career=payload.career,
        plan_name=payload.plan_name,
        competency_type=payload.competency_type,
        code=payload.code,
        description=payload.description,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.patch("/competencies/{competency_id}", response_model=schemas.CompetencyCatalogOut)
def update_competency(competency_id: int, payload: schemas.CompetencyCatalogUpdate, db: Session = Depends(get_db)):
    item = db.query(models.CompetencyCatalog).filter(models.CompetencyCatalog.id == competency_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Competency not found")
    data = payload.model_dump(exclude_unset=True) if hasattr(payload, 'model_dump') else payload.dict(exclude_unset=True)
    target_career = data.get("career", item.career)
    target_type = data.get("competency_type", item.competency_type)
    target_code = data.get("code", item.code)
    target_plan = data.get("plan_name", item.plan_name)
    if target_code and target_career and target_type:
        existing = db.query(models.CompetencyCatalog).filter(
            models.CompetencyCatalog.id != competency_id,
            models.CompetencyCatalog.career == target_career,
            models.CompetencyCatalog.competency_type == target_type,
            func.lower(models.CompetencyCatalog.code) == str(target_code).strip().lower(),
            func.coalesce(models.CompetencyCatalog.plan_name, "") == (target_plan or ""),
        ).first()
        if existing:
            raise HTTPException(status_code=409, detail="Competency code already exists")
    for key, value in data.items():
        setattr(item, key, value)
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.get("/competencies/{competency_id}/usage")
def get_competency_usage(competency_id: int, db: Session = Depends(get_db)):
    item = db.query(models.CompetencyCatalog).filter(models.CompetencyCatalog.id == competency_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Competency not found")
    query = db.query(models.ProposalCompetency).join(models.Proposal).filter(
        models.Proposal.career == item.career,
        models.ProposalCompetency.competency_type == item.competency_type,
        models.ProposalCompetency.code == item.code,
    )
    if item.plan_name:
        query = query.filter(models.Proposal.study_plan == item.plan_name)
    rows = query.all()
    affected_ids = sorted({row.proposal_id for row in rows})
    proposals_info = []
    if affected_ids:
        proposals = db.query(models.Proposal).filter(models.Proposal.id.in_(affected_ids)).all()
        proposals_info = [
            {
                "id": proposal.id,
                "subject": proposal.subject,
                "career": proposal.career,
                "title": proposal.title,
            }
            for proposal in proposals
        ]
    return {
        "competency_id": competency_id,
        "affected_proposals": len(affected_ids),
        "affected_proposal_ids": affected_ids,
        "affected_proposals_info": proposals_info,
    }


@app.delete("/competencies/{competency_id}")
def delete_competency(competency_id: int, db: Session = Depends(get_db)):
    item = db.query(models.CompetencyCatalog).filter(models.CompetencyCatalog.id == competency_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Competency not found")
    query = db.query(models.ProposalCompetency).join(models.Proposal).filter(
        models.Proposal.career == item.career,
        models.ProposalCompetency.competency_type == item.competency_type,
        models.ProposalCompetency.code == item.code,
    )
    if item.plan_name:
        query = query.filter(models.Proposal.study_plan == item.plan_name)
    affected_rows = query.all()
    affected_ids = {row.proposal_id for row in affected_rows}
    if affected_ids:
        delete_query = db.query(models.ProposalCompetency).filter(
            models.ProposalCompetency.proposal_id.in_(affected_ids),
            models.ProposalCompetency.competency_type == item.competency_type,
            models.ProposalCompetency.code == item.code,
        )
        if item.plan_name:
            delete_query = delete_query.join(models.Proposal).filter(
                models.Proposal.study_plan == item.plan_name
            )
        delete_query.delete(synchronize_session=False)
    db.delete(item)
    for proposal_id in affected_ids:
        proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
        if not proposal:
            continue
        remaining_rows = db.query(models.ProposalCompetency).filter(
            models.ProposalCompetency.proposal_id == proposal_id,
            models.ProposalCompetency.competency_type == item.competency_type,
        ).all()
        remaining_items = [
            {
                "code": row.code,
                "description": row.description,
                "level": row.level,
            }
            for row in remaining_rows
        ]
        updated_text = build_competencies_text(remaining_items)
        if item.competency_type == "specific":
            proposal.specific_competencies = updated_text
        else:
            proposal.generic_competencies = updated_text
        db.add(proposal)
    db.commit()
    return {
        "status": "deleted",
        "id": competency_id,
        "affected_proposals": len(affected_ids),
        "affected_proposal_ids": sorted(affected_ids),
    }


@app.get("/teachers", response_model=list[schemas.TeacherOut])
def list_teachers(career: str = "", db: Session = Depends(get_db)):
    query = db.query(models.Teacher)
    if career:
        proposal_ids = db.query(models.Teacher.id).join(models.ProposalTeacher).join(models.Proposal).filter(
            models.Proposal.career == career
        ).distinct().all()
        career_ids = db.query(models.TeacherCareer.teacher_id).filter(
            models.TeacherCareer.career == career
        ).distinct().all()
        teacher_ids = {row[0] for row in proposal_ids} | {row[0] for row in career_ids}
        if not teacher_ids:
            return []
        query = query.filter(models.Teacher.id.in_(teacher_ids))
    teachers = query.order_by(models.Teacher.name.asc()).distinct().all()
    if not teachers:
        return []
    teacher_ids_list = [t.id for t in teachers]
    from sqlalchemy import func as sqlfunc
    proposal_counts = {
        row[0]: row[1]
        for row in db.query(models.ProposalTeacher.teacher_id, sqlfunc.count())
        .filter(models.ProposalTeacher.teacher_id.in_(teacher_ids_list))
        .group_by(models.ProposalTeacher.teacher_id)
        .all()
    }
    career_counts = {
        row[0]: row[1]
        for row in db.query(models.TeacherCareer.teacher_id, sqlfunc.count())
        .filter(models.TeacherCareer.teacher_id.in_(teacher_ids_list))
        .group_by(models.TeacherCareer.teacher_id)
        .all()
    }
    result = []
    for t in teachers:
        data = schemas.TeacherOut.model_validate(t)
        data.has_password = bool(t.password_hash)
        data.has_proposals = proposal_counts.get(t.id, 0) > 0
        data.has_career_links = career_counts.get(t.id, 0) > 0
        result.append(data)
    return result


@app.post("/teachers", response_model=schemas.TeacherOut)
def create_teacher(payload: schemas.TeacherCreate, db: Session = Depends(get_db)):
    data = payload.model_dump() if hasattr(payload, "model_dump") else payload.dict()
    career = (data.get("career") or "").strip() or None
    name = normalize_teacher_name(data.get("name"))
    email = (data.get("email") or "").strip().lower() or None
    data["dedication"] = normalize_teacher_dedication(data.get("dedication"))
    normalized_key = normalize_teacher_key(name)

    existing = None
    if email:
        existing = db.query(models.Teacher).filter(models.Teacher.email == email).first()
    if not existing and normalized_key:
        existing = db.query(models.Teacher).filter(models.Teacher.normalized_key == normalized_key).first()
    if not existing and name:
        existing = find_teacher_by_name_fuzzy(db, name)

    if existing:
        ensure_teacher_career(db, existing.id, career)
        db.commit()
        db.refresh(existing)
        return existing

    teacher = upsert_teacher(db, data)
    if not teacher:
        raise HTTPException(status_code=400, detail="El docente debe tener al menos un nombre o correo")
    db.flush()
    ensure_teacher_career(db, teacher.id, career)
    db.commit()
    db.refresh(teacher)
    return teacher


@app.patch("/teachers/{teacher_id}", response_model=schemas.TeacherOut)
def update_teacher(teacher_id: int, payload: schemas.TeacherUpdate, db: Session = Depends(get_db)):
    teacher = db.query(models.Teacher).filter(models.Teacher.id == teacher_id).first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")

    data = payload.model_dump(exclude_unset=True) if hasattr(payload, "model_dump") else payload.dict(exclude_unset=True)
    if not data:
        return teacher

    name = normalize_teacher_name(data.get("name") if "name" in data else teacher.name)
    email_raw = data.get("email") if "email" in data else teacher.email
    email = (email_raw or "").strip().lower() or None
    category = teacher.category
    if "category" in data:
        category = normalize_teacher_category(data.get("category"))
    dedication = teacher.dedication
    if "dedication" in data:
        dedication = normalize_teacher_dedication(data.get("dedication"))
    normalized_key = normalize_teacher_key(name)

    if email:
        existing_email = db.query(models.Teacher).filter(
            models.Teacher.email == email,
            models.Teacher.id != teacher_id,
        ).first()
        if existing_email:
            raise HTTPException(status_code=409, detail="Teacher already exists")

    if normalized_key:
        existing_name = db.query(models.Teacher).filter(
            models.Teacher.normalized_key == normalized_key,
            models.Teacher.id != teacher_id,
        ).first()
        if existing_name:
            raise HTTPException(status_code=409, detail="Teacher already exists")
    if name:
        fuzzy = find_teacher_by_name_fuzzy(db, name)
        if fuzzy and fuzzy.id != teacher_id:
            raise HTTPException(status_code=409, detail="Teacher already exists")

    teacher.name = name or teacher.name
    teacher.email = email
    teacher.category = category
    teacher.dedication = dedication or "Sin Informar"
    teacher.normalized_key = normalized_key or teacher.normalized_key
    # If no password set yet and email exists, seed initial password from email
    if not teacher.password_hash and (email or teacher.email):
        teacher.password_hash = hash_password(email or teacher.email)
        teacher.password_reset_at = datetime.now(timezone.utc)
    db.add(teacher)

    proposal_ids = db.query(models.ProposalTeacher.proposal_id).filter(
        models.ProposalTeacher.teacher_id == teacher_id
    ).distinct().all()
    for (proposal_id,) in proposal_ids:
        remaining_rows = db.query(models.ProposalTeacher).filter(
            models.ProposalTeacher.proposal_id == proposal_id
        ).all()
        remaining_ids = [row.teacher_id for row in remaining_rows]
        remaining_teachers = db.query(models.Teacher).filter(
            models.Teacher.id.in_(remaining_ids)
        ).all()
        proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
        if proposal:
            proposal.teaching_team = build_teaching_team_payload(remaining_teachers)
            db.add(proposal)

    db.commit()
    db.refresh(teacher)
    return teacher


@app.get("/teachers/{teacher_id}/usage")
def get_teacher_usage(teacher_id: int, career: str = "", db: Session = Depends(get_db)):
    teacher = db.query(models.Teacher).filter(models.Teacher.id == teacher_id).first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    query = db.query(models.Proposal).join(models.ProposalTeacher).filter(
        models.ProposalTeacher.teacher_id == teacher_id
    )
    if career:
        query = query.filter(models.Proposal.career == career)
    proposals = query.all()
    affected_ids = sorted({proposal.id for proposal in proposals})
    proposals_info = [
        {
            "id": proposal.id,
            "subject": proposal.subject,
            "career": proposal.career,
            "title": proposal.title,
            "quarter": proposal.quarter,
            "year_of_career": proposal.year_of_career,
            "academic_year": proposal.academic_year,
        }
        for proposal in proposals
    ]
    return {
        "teacher_id": teacher_id,
        "teacher_name": teacher.name,
        "affected_proposals": len(affected_ids),
        "affected_proposal_ids": affected_ids,
        "affected_proposals_info": proposals_info,
    }


@app.delete("/teachers/{teacher_id}")
def delete_teacher(teacher_id: int, career: str = "", db: Session = Depends(get_db)):
    teacher = db.query(models.Teacher).filter(models.Teacher.id == teacher_id).first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    query = db.query(models.ProposalTeacher).join(models.Proposal).filter(
        models.ProposalTeacher.teacher_id == teacher_id
    )
    if career:
        query = query.filter(models.Proposal.career == career)
    affected_rows = query.all()
    affected_ids = {row.proposal_id for row in affected_rows}
    if affected_ids:
        db.query(models.ProposalTeacher).filter(
            models.ProposalTeacher.teacher_id == teacher_id,
            models.ProposalTeacher.proposal_id.in_(affected_ids),
        ).delete(synchronize_session=False)

    if career:
        db.query(models.TeacherCareer).filter(
            models.TeacherCareer.teacher_id == teacher_id,
            models.TeacherCareer.career == career,
        ).delete(synchronize_session=False)

    proposals_info = []
    if affected_ids:
        proposals = db.query(models.Proposal).filter(models.Proposal.id.in_(affected_ids)).all()
        proposals_info = [
            {
                "id": proposal.id,
                "subject": proposal.subject,
                "career": proposal.career,
                "title": proposal.title,
            }
            for proposal in proposals
        ]
        for proposal in proposals:
            remaining_rows = db.query(models.ProposalTeacher).filter(
                models.ProposalTeacher.proposal_id == proposal.id
            ).all()
            if remaining_rows:
                remaining_ids = [row.teacher_id for row in remaining_rows]
                remaining_teachers = db.query(models.Teacher).filter(
                    models.Teacher.id.in_(remaining_ids)
                ).all()
                proposal.teaching_team = build_teaching_team_payload(remaining_teachers)
            else:
                proposal.teaching_team = []
            db.add(proposal)

    still_linked = db.query(models.ProposalTeacher).filter(
        models.ProposalTeacher.teacher_id == teacher_id
    ).first()
    still_career_link = db.query(models.TeacherCareer).filter(
        models.TeacherCareer.teacher_id == teacher_id
    ).first()
    if not still_linked and not still_career_link:
        db.delete(teacher)

    db.commit()
    return {
        "status": "deleted",
        "id": teacher_id,
        "affected_proposals": len(affected_ids),
        "affected_proposal_ids": sorted(affected_ids),
        "affected_proposals_info": proposals_info,
    }


@app.get("/proposals/{proposal_id}/docx")
def download_proposal_docx(
    proposal_id: int,
    db: Session = Depends(get_db),
    background: BackgroundTasks = None,
):
    return export_proposal(
        proposal_id=proposal_id,
        format="docx",
        db=db,
        background=background,
    )


@app.get("/proposals/{proposal_id}/export")
def export_proposal(
    proposal_id: int,
    format: str = "docx",
    db: Session = Depends(get_db),
    background: BackgroundTasks = None,
):
    proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    export_format = (format or "docx").strip().lower()
    allowed_formats = {"docx", "pdf", "json", "xml"}
    if export_format not in allowed_formats:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {export_format}. Formatos válidos: docx, pdf, json, xml",
        )

    basename = build_proposal_export_basename(proposal)
    payload = make_json_compatible(build_proposal_response(db, proposal))

    if export_format in {"docx", "pdf"}:
        if not os.path.exists(TEMPLATE_PATH):
            raise HTTPException(status_code=500, detail="Template Propuestas.docx not found")

        try:
            from .docx_export import generate_proposal_docx
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx no esta instalado")

        docx_path = generate_proposal_docx(proposal, TEMPLATE_PATH)
        output_dir = os.path.dirname(docx_path)
        if background is not None:
            background.add_task(shutil.rmtree, output_dir, ignore_errors=True)

        if export_format == "docx":
            filename = f"{basename}.docx"
            return FileResponse(
                docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=filename,
                headers=build_download_headers(filename),
            )

        renderer_used = "libreoffice"
        try:
            pdf_path = convert_docx_to_pdf_with_libreoffice(docx_path)
        except HTTPException:
            pdf_path, renderer_used = convert_docx_to_pdf(docx_path)
        filename = f"{basename}.pdf"
        pdf_headers = build_download_headers(filename)
        pdf_headers["X-PDF-Renderer"] = renderer_used
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers=pdf_headers,
        )

    if export_format == "json":
        filename = f"{basename}.json"
        body = json.dumps(payload, ensure_ascii=False, indent=2)
        return Response(
            content=body,
            media_type="application/json; charset=utf-8",
            headers=build_download_headers(filename),
        )

    filename = f"{basename}.xml"
    xml_bytes = build_proposal_xml_bytes(payload)
    return Response(
        content=xml_bytes,
        media_type="application/xml; charset=utf-8",
        headers=build_download_headers(filename),
    )


@app.post("/proposals")
def create_proposal(proposal: schemas.ProposalCreate, db: Session = Depends(get_db)):
    """Create a new proposal from form data (no file upload)."""
    try:
        subject_clean = re.sub(r'[<>:"/\\|?*]', '', proposal.subject or '').strip()
        safe_subject = subject_clean or "Sin asignatura"
        safe_year = proposal.academic_year or "Sin año"
        safe_quarter = proposal.quarter or "Sin cuatrimestre"
        filename = f"{safe_year} - {safe_quarter} - {safe_subject}.docx"

        # Convert Pydantic models to dicts for JSON storage
        learning_outcomes_dict = [lo.model_dump() if hasattr(lo, 'model_dump') else lo.dict() for lo in proposal.learning_outcomes] if proposal.learning_outcomes else []
        units_dict = [u.model_dump() if hasattr(u, 'model_dump') else u.dict() for u in proposal.units] if proposal.units else []
        practicals_dict = [p.model_dump() if hasattr(p, 'model_dump') else p.dict() for p in proposal.practicals] if proposal.practicals else []
        teaching_team_dict = [t.model_dump() if hasattr(t, 'model_dump') else t.dict() for t in proposal.teaching_team] if proposal.teaching_team else []
        teacher_objs = []
        teacher_ids = []
        for entry in teaching_team_dict:
            if not isinstance(entry, dict):
                continue
            teacher = upsert_teacher(db, entry)
            if teacher:  # Only process if teacher was created
                db.flush()
                ensure_teacher_career(db, teacher.id, proposal.career)
                teacher_objs.append(teacher)
                teacher_ids.append(teacher.id)
        teaching_team_payload = build_teaching_team_payload(teacher_objs)
        
        generic_items = normalize_competency_items(proposal.generic_competencies_items or [])
        specific_items = normalize_competency_items(proposal.specific_competencies_items or [])

        create_in_drive = bool(getattr(proposal, "create_in_drive", False))

        db_proposal = models.Proposal(
            title=proposal.title,
            filename=filename,
            career=proposal.career,
            subject=proposal.subject,
            study_plan=proposal.study_plan,
            academic_year=proposal.academic_year,
            year_of_career=proposal.year_of_career,
            quarter=proposal.quarter,
            character=proposal.character,
            regime=proposal.regime,
            theoretical_hours=proposal.theoretical_hours,
            practical_hours=proposal.practical_hours,
            total_hours=proposal.total_hours,
            weekly_hours=proposal.weekly_hours,
            minimum_content=proposal.minimum_content,
            generic_competencies=proposal.generic_competencies or build_competencies_text(generic_items),
            specific_competencies=proposal.specific_competencies or build_competencies_text(specific_items),
            fundamentals_part1=proposal.fundamentals_part1,
            fundamentals_part2=proposal.fundamentals_part2,
            learning_outcomes=learning_outcomes_dict,
            units=units_dict,
            practicals=practicals_dict,
            teaching_team=teaching_team_payload,
            methodology=proposal.methodology,
            evaluation=proposal.evaluation,
            bibliography=proposal.bibliography,
            observations=proposal.observations,
            original_filename="form_submission",
            source_type=proposal.source_type or ("gdoc" if proposal.gdoc_url else "manual"),
            gdoc_url=proposal.gdoc_url,
            status=proposal.status or "EnProceso"
        )
        db.add(db_proposal)
        db.flush()
        if teacher_ids:
            replace_proposal_teachers(db, db_proposal.id, teacher_ids)
        if generic_items:
            ensure_competency_catalog(db, proposal.career, generic_items, "generic", proposal.study_plan)
        if specific_items:
            ensure_competency_catalog(db, proposal.career, specific_items, "specific", proposal.study_plan)
        if generic_items:
            replace_proposal_competencies(db, db_proposal.id, generic_items, "generic")
        if specific_items:
            replace_proposal_competencies(db, db_proposal.id, specific_items, "specific")
        sync_subject_from_proposal(db, db_proposal)
        drive_creation_success = False
        drive_creation_error = None
        if create_in_drive and not db_proposal.gdoc_url:
            drive_creation_success, drive_creation_error = try_create_gdoc_for_proposal(db, db_proposal)
        
        db.commit()
        db.refresh(db_proposal)
        return {
            "id": db_proposal.id,
            "title": db_proposal.title,
            "career": db_proposal.career,
            "subject": db_proposal.subject,
            "study_plan": db_proposal.study_plan,
            "status": db_proposal.status,
            "created_at": db_proposal.created_at,
            "study_subject_id": db_proposal.study_subject_id,
            "gdoc_url": db_proposal.gdoc_url,
            "drive_creation_requested": create_in_drive,
            "drive_creation_success": drive_creation_success,
            "drive_creation_error": drive_creation_error,
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Error creating proposal: {str(e)}")


@app.post("/proposals/import-docx")
async def import_proposal_docx(file: UploadFile = File(...)):
    """
    Importa una propuesta desde un archivo DOCX.
    Extrae todos los datos y retorna JSON para previsualización.
    """
    try:
        # Guardar archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            # Extraer datos del DOCX (pasar filename original para parsing)
            extracted_data = import_proposal_from_docx(tmp_path, file.filename)
            return {
                "success": True,
                "data": extracted_data,
                "preview": {
                    "career": extracted_data.get('career', ''),
                    "subject": extracted_data.get('subject', ''),
                    "teachers": extracted_data.get('teachers', ''),
                    "year": extracted_data.get('year_of_career', ''),
                    "quarter": extracted_data.get('quarter', ''),
                    "total_hours": extracted_data.get('total_hours', ''),
                    "theoretical_hours": extracted_data.get('theoretical_hours', ''),
                    "practical_hours": extracted_data.get('practical_hours', ''),
                    "weekly_hours": extracted_data.get('weekly_hours', ''),
                    "regime": extracted_data.get('regime', ''),
                    "units_count": len(extracted_data.get('units', [])),
                    "practicals_count": len(extracted_data.get('practicals', [])),
                    "ra_count": len(extracted_data.get('learning_outcomes', [])),
                    "generic_comp_count": len(extracted_data.get('generic_competencies', [])),
                    "specific_comp_count": len(extracted_data.get('specific_competencies', [])),
                }
            }
        finally:
            # Limpiar archivo temporal
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error importing DOCX: {str(e)}")


@app.post("/proposals/{proposal_id}/sync-to-study-plan")
async def sync_proposal_to_study_plan(proposal_id: int, db: Session = Depends(get_db)):
    """
    Force synchronization of a proposal to the study plan.
    Useful for debugging why a proposal isn't linking to the plan.
    """
    db_proposal = db.query(models.Proposal).filter(models.Proposal.id == proposal_id).first()
    if not db_proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    
    # Log the current state
    log_message = f"Syncing proposal {proposal_id} to study plan:\n"
    log_message += f"  - career: {db_proposal.career}\n"
    log_message += f"  - subject: {db_proposal.subject}\n"
    log_message += f"  - study_plan: {db_proposal.study_plan}\n"
    log_message += f"  - year_of_career: {db_proposal.year_of_career}\n"
    log_message += f"  - quarter: {db_proposal.quarter}\n"
    print(f"[DEBUG] {log_message}")
    
    # Check prerequisites
    if not db_proposal.career:
        raise HTTPException(status_code=400, detail="Proposal has no career specified")
    if not db_proposal.subject:
        raise HTTPException(status_code=400, detail="Proposal has no subject specified")
    
    try:
        # Perform the sync
        sync_subject_from_proposal(db, db_proposal)
        db.commit()
        db.refresh(db_proposal)
        
        return {
            "success": True,
            "message": "Proposal synced to study plan",
            "study_subject_id": db_proposal.study_subject_id,
            "debug_info": log_message.replace("\n", " | ")
        }
    except Exception as e:
        db.rollback()
        print(f"[ERROR] Sync failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Sync failed: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# EVALUATIVE INSTRUMENTS
# ═══════════════════════════════════════════════════════════════════════════

@app.get("/instruments", response_model=list[schemas.EvaluativeInstrumentOut])
def list_instruments(
    career: str = None,
    subject: str = None,
    instrument_type: str = None,
    db: Session = Depends(get_db),
):
    q = db.query(models.EvaluativeInstrument)
    if career:
        q = q.filter(models.EvaluativeInstrument.career == career)
    if subject:
        q = q.filter(models.EvaluativeInstrument.subject == subject)
    if instrument_type:
        q = q.filter(models.EvaluativeInstrument.instrument_type == instrument_type)
    return q.order_by(models.EvaluativeInstrument.created_at.desc()).all()


@app.get("/instruments/summary")
def instruments_summary(career: str = None, db: Session = Depends(get_db)):
    """Returns list of subjects with counts per type."""
    q = db.query(models.EvaluativeInstrument)
    if career:
        q = q.filter(models.EvaluativeInstrument.career == career)
    instruments = q.all()
    summary: dict = {}
    for inst in instruments:
        key = (inst.career, inst.study_plan or "", inst.subject)
        if key not in summary:
            summary[key] = {
                "career": inst.career,
                "study_plan": inst.study_plan or "",
                "subject": inst.subject,
                "tp": 0,
                "parcial": 0,
                "final": 0,
            }
        t = inst.instrument_type
        if t == "TP":
            summary[key]["tp"] += 1
        elif t == "Parcial":
            summary[key]["parcial"] += 1
        elif t == "Final":
            summary[key]["final"] += 1
    # Enrich with folder URLs
    folders_q = db.query(models.EvaluativeInstrumentFolder)
    if career:
        folders_q = folders_q.filter(models.EvaluativeInstrumentFolder.career == career)
    for f in folders_q.all():
        key = (f.career, f.study_plan or "", f.subject)
        if key in summary:
            summary[key]["gdrive_folder_url"] = f.gdrive_folder_url or ""
    return list(summary.values())


@app.get("/instruments/folders")
def list_instrument_folders(career: str = None, db: Session = Depends(get_db)):
    """Returns all linked Drive folders for evaluative instruments, optionally filtered by career."""
    q = db.query(models.EvaluativeInstrumentFolder)
    if career:
        q = q.filter(models.EvaluativeInstrumentFolder.career == career)
    return [
        {"career": f.career, "study_plan": f.study_plan or "", "subject": f.subject, "gdrive_folder_url": f.gdrive_folder_url or ""}
        for f in q.all()
    ]


@app.get("/instruments/folder", response_model=schemas.EvaluativeInstrumentFolderOut | None)
def get_instrument_folder(
    career: str,
    subject: str,
    study_plan: str = None,
    db: Session = Depends(get_db),
):
    q = db.query(models.EvaluativeInstrumentFolder).filter(
        models.EvaluativeInstrumentFolder.career == career,
        models.EvaluativeInstrumentFolder.subject == subject,
    )
    if study_plan:
        q = q.filter(models.EvaluativeInstrumentFolder.study_plan == study_plan)
    return q.first()


@app.post("/instruments/folder/link", response_model=schemas.EvaluativeInstrumentFolderOut)
def link_instrument_folder(payload: dict = Body(...), db: Session = Depends(get_db)):
    career = (payload.get("career") or "").strip()
    subject = (payload.get("subject") or "").strip()
    study_plan = (payload.get("study_plan") or "").strip() or None
    folder_url = (payload.get("folder_url") or "").strip()
    if not career or not subject or not folder_url:
        raise HTTPException(status_code=400, detail="career, subject y folder_url son requeridos")
    folder_id = extract_drive_folder_id(folder_url)
    if not folder_id:
        raise HTTPException(status_code=400, detail="URL de carpeta de Drive inválida")
    existing = db.query(models.EvaluativeInstrumentFolder).filter(
        models.EvaluativeInstrumentFolder.career == career,
        models.EvaluativeInstrumentFolder.subject == subject,
    ).first()
    if existing:
        existing.gdrive_folder_url = folder_url
        existing.gdrive_folder_id = folder_id
        if study_plan:
            existing.study_plan = study_plan
        db.commit()
        db.refresh(existing)
        return existing
    new_folder = models.EvaluativeInstrumentFolder(
        career=career, study_plan=study_plan, subject=subject,
        gdrive_folder_url=folder_url, gdrive_folder_id=folder_id,
    )
    db.add(new_folder)
    db.commit()
    db.refresh(new_folder)
    return new_folder


@app.delete("/instruments/folder")
def unlink_instrument_folder(
    career: str,
    subject: str,
    db: Session = Depends(get_db),
):
    """Remove the Drive folder link for a subject (does not delete the folder from Drive)."""
    existing = db.query(models.EvaluativeInstrumentFolder).filter(
        models.EvaluativeInstrumentFolder.career == career,
        models.EvaluativeInstrumentFolder.subject == subject,
    ).first()
    if not existing:
        raise HTTPException(status_code=404, detail="No hay carpeta vinculada para esta asignatura")
    db.delete(existing)
    db.commit()
    return {"ok": True}


@app.post("/instruments/folder/create", response_model=schemas.EvaluativeInstrumentFolderOut)
def create_instrument_folder_in_drive(payload: dict = Body(...), db: Session = Depends(get_db)):
    career = (payload.get("career") or "").strip()
    subject = (payload.get("subject") or "").strip()
    study_plan = (payload.get("study_plan") or "").strip() or None
    if not career or not subject:
        raise HTTPException(status_code=400, detail="career y subject son requeridos")
    existing = db.query(models.EvaluativeInstrumentFolder).filter(
        models.EvaluativeInstrumentFolder.career == career,
        models.EvaluativeInstrumentFolder.subject == subject,
    ).first()
    if existing and existing.gdrive_folder_id:
        return existing
    drive_settings = db.query(models.DriveSettings).filter(models.DriveSettings.career == career).first()
    if not drive_settings or not drive_settings.root_folder_url:
        raise HTTPException(status_code=400, detail=f"No hay carpeta raíz de Drive configurada para la carrera '{career}'")
    parent_folder_id = extract_drive_folder_id(drive_settings.root_folder_url)
    if not parent_folder_id:
        raise HTTPException(status_code=400, detail="La URL de carpeta raíz de Drive es inválida")
    drive_service = get_google_drive_service()
    try:
        instr_folder_name = "Instrumentos Evaluativos"
        results = drive_service.files().list(
            q=f"name='{instr_folder_name}' and '{parent_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True,
        ).execute()
        instr_folders = results.get("files", [])
        if instr_folders:
            instr_folder_id = instr_folders[0]["id"]
        else:
            meta = {"name": instr_folder_name, "mimeType": "application/vnd.google-apps.folder", "parents": [parent_folder_id]}
            created = drive_service.files().create(body=meta, fields="id", supportsAllDrives=True).execute()
            instr_folder_id = created["id"]
        safe_subject = re.sub(r'[/\\:*?"<>|]', "_", subject)
        results2 = drive_service.files().list(
            q=f"name='{safe_subject}' and '{instr_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True,
        ).execute()
        subj_folders = results2.get("files", [])
        if subj_folders:
            subj_folder_id = subj_folders[0]["id"]
        else:
            meta2 = {"name": safe_subject, "mimeType": "application/vnd.google-apps.folder", "parents": [instr_folder_id]}
            created2 = drive_service.files().create(body=meta2, fields="id", supportsAllDrives=True).execute()
            subj_folder_id = created2["id"]
        gdrive_folder_url = f"https://drive.google.com/drive/folders/{subj_folder_id}"
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Error al crear carpeta en Drive: {exc}")
    if existing:
        existing.gdrive_folder_url = gdrive_folder_url
        existing.gdrive_folder_id = subj_folder_id
        db.commit()
        db.refresh(existing)
        return existing
    new_folder = models.EvaluativeInstrumentFolder(
        career=career, study_plan=study_plan, subject=subject,
        gdrive_folder_url=gdrive_folder_url, gdrive_folder_id=subj_folder_id,
    )
    db.add(new_folder)
    db.commit()
    db.refresh(new_folder)
    return new_folder


@app.post("/instruments/upload")
async def upload_instruments(
    files: list[UploadFile] = File(...),
    types: str = Form(...),
    titles: str = Form(None),
    career: str = Form(...),
    subject: str = Form(...),
    study_plan: str = Form(None),
    uploaded_by: str = Form(None),
    db: Session = Depends(get_db),
):
    """Upload one or more evaluative instrument files for a subject."""
    import uuid as _uuid
    types_list: list[str] = json.loads(types)
    titles_list: list = json.loads(titles) if titles else [None] * len(files)
    valid_types = {"TP", "Parcial", "Final"}
    if len(files) != len(types_list):
        raise HTTPException(status_code=400, detail="Los arrays 'files' y 'types' deben tener el mismo tamaño")
    for t in types_list:
        if t not in valid_types:
            raise HTTPException(status_code=400, detail=f"Tipo inválido '{t}'. Use: TP, Parcial, Final")

    # Resolve INSTRUMENTS_FOLDER relative to CWD
    instruments_dir = os.path.abspath(INSTRUMENTS_FOLDER)
    os.makedirs(instruments_dir, exist_ok=True)

    # Get linked drive folder if exists
    folder_record = db.query(models.EvaluativeInstrumentFolder).filter(
        models.EvaluativeInstrumentFolder.career == career,
        models.EvaluativeInstrumentFolder.subject == subject,
    ).first()

    results = []
    for i, (upload_file, itype) in enumerate(zip(files, types_list)):
        ext = os.path.splitext(upload_file.filename or "")[-1]
        stored_name = f"{_uuid.uuid4().hex}{ext}"
        dest_path = os.path.join(instruments_dir, stored_name)
        content = await upload_file.read()
        with open(dest_path, "wb") as fout:
            fout.write(content)

        gdrive_url = None
        gdrive_file_id = None
        # Auto-upload to Drive if folder is linked
        if folder_record and folder_record.gdrive_folder_id:
            try:
                from googleapiclient.http import MediaFileUpload as _MediaUpload
                _drive = get_google_drive_service()
                _media = _MediaUpload(dest_path, resumable=False)
                _meta = {"name": upload_file.filename, "parents": [folder_record.gdrive_folder_id]}
                _created = _drive.files().create(body=_meta, media_body=_media, fields="id,webViewLink", supportsAllDrives=True).execute()
                gdrive_url = _created.get("webViewLink") or f"https://drive.google.com/file/d/{_created.get('id')}/view"
                gdrive_file_id = _created.get("id")
            except Exception:
                pass  # local save is sufficient; Drive upload is best-effort

        title_val = (titles_list[i] if titles_list and i < len(titles_list) else None) or upload_file.filename
        inst = models.EvaluativeInstrument(
            career=career,
            study_plan=study_plan,
            subject=subject,
            instrument_type=itype,
            title=title_val,
            original_filename=upload_file.filename,
            stored_filename=stored_name,
            file_path=dest_path,
            file_size=len(content),
            mime_type=upload_file.content_type,
            gdrive_url=gdrive_url,
            gdrive_file_id=gdrive_file_id,
            uploaded_by=uploaded_by,
        )
        db.add(inst)
        results.append(inst)
    db.commit()
    for r in results:
        db.refresh(r)
    return [schemas.EvaluativeInstrumentOut.model_validate(r) for r in results]


@app.delete("/instruments/{instrument_id}")
def delete_instrument(instrument_id: int, db: Session = Depends(get_db)):
    inst = db.query(models.EvaluativeInstrument).filter(models.EvaluativeInstrument.id == instrument_id).first()
    if not inst:
        raise HTTPException(status_code=404, detail="Instrumento no encontrado")
    if inst.file_path and os.path.exists(inst.file_path):
        try:
            os.remove(inst.file_path)
        except Exception:
            pass
    db.delete(inst)
    db.commit()
    return {"deleted": instrument_id}


@app.get("/instruments/{instrument_id}/file")
def download_instrument_file(instrument_id: int, db: Session = Depends(get_db)):
    inst = db.query(models.EvaluativeInstrument).filter(models.EvaluativeInstrument.id == instrument_id).first()
    if not inst:
        raise HTTPException(status_code=404, detail="Instrumento no encontrado")
    if not inst.file_path or not os.path.exists(inst.file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado en el servidor")
    return FileResponse(
        inst.file_path,
        filename=inst.original_filename,
        media_type=inst.mime_type or "application/octet-stream",
    )
