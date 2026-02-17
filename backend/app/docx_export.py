from __future__ import annotations

import os
import re
import tempfile
from typing import Iterable, List, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")
HEADER_LABELS_TO_REMOVE = [
    "Cuatrimestre:",
    "Año de Carrera:",
    "Ciclo:",
    "Plan de Estudio:",
    "Asignatura:",
    "Carrera:",
]


def _iter_block_items(doc: Document) -> Iterable[Paragraph | Table]:
    for child in doc.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _replace_placeholders_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    """Reemplaza placeholders en un párrafo, preservando el formato."""
    if not mapping:
        return
    
    # Obtener el texto completo del párrafo
    full_text = paragraph.text
    
    # Buscar si hay algún placeholder que reemplazar
    has_replacement = False
    for key in mapping.keys():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            has_replacement = True
            break
    
    if not has_replacement:
        return
    
    # Hacer todos los reemplazos en el texto
    for key, value in mapping.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = full_text.replace(placeholder, value)
    
    # Limpiar todos los runs existentes
    for run in list(paragraph.runs):
        r = run._element
        r.getparent().remove(r)
    
    # Agregar el texto reemplazado como un nuevo run
    if full_text:
        paragraph.add_run(full_text)


def _replace_placeholders_in_table(table: Table, mapping: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, mapping)


def _replace_placeholders(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, mapping)
    for table in doc.tables:
        _replace_placeholders_in_table(table, mapping)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, mapping)
        for table in header.tables:
            _replace_placeholders_in_table(table, mapping)


def _remove_header_paragraphs_in_body(doc: Document) -> None:
    labels = [label.lower() for label in HEADER_LABELS_TO_REMOVE]
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if any(text.startswith(label.lower()) for label in labels):
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None


def _fill_first_empty_paragraph(paragraphs: List[Paragraph], value: str) -> None:
    for paragraph in paragraphs:
        if not paragraph.text.strip():
            paragraph.text = value
            return
    if paragraphs:
        # Agregar al último párrafo con salto de línea si no hay espacio vacío
        if paragraphs[-1].text:
            paragraphs[-1].text += "\n" + value
        else:
            paragraphs[-1].text = value


def _fill_after_label(paragraphs: List[Paragraph], label_keywords: List[str], value: str) -> bool:
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.lower()
        if all(keyword in text for keyword in label_keywords):
            for next_para in paragraphs[idx + 1:]:
                if not next_para.text.strip():
                    next_para.text = value
                    return True
            # Si no hay párrafo vacío después, agregar al actual con salto de línea
            if paragraph.text:
                paragraph.text += "\n" + value
            else:
                paragraph.text = value
            return True
    return False


def _fill_fundamentals(doc: Document, importance: str, profile: str) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                if _fill_after_label(paragraphs, ["importancia", "plan"], importance):
                    pass
                if _fill_after_label(paragraphs, ["perfil", "profesional"], profile):
                    pass


def _fill_learning_outcomes(doc: Document, outcomes: List[str]) -> None:
    if not outcomes:
        return
    ra_paragraphs: List[Paragraph] = []
    last_cell = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if re.search(r"RA\d+:", paragraph.text):
                        ra_paragraphs.append(paragraph)
                        last_cell = cell
    if not ra_paragraphs:
        return
    for idx, paragraph in enumerate(ra_paragraphs):
        if idx < len(outcomes):
            paragraph.text = outcomes[idx]
        else:
            paragraph.text = ""
    # Si hay más RAs que placeholders, agregar los adicionales a la celda
    if len(outcomes) > len(ra_paragraphs) and last_cell:
        for extra_outcome in outcomes[len(ra_paragraphs):]:
            last_cell.add_paragraph(extra_outcome)


def _clone_table_after(doc: Document, table: Table) -> Table:
    """Clona una tabla y la inserta después de la tabla original."""
    from copy import deepcopy
    new_tbl = deepcopy(table._element)
    table._element.addnext(new_tbl)
    return Table(new_tbl, doc)


def _clear_table_data_cells(table: Table) -> None:
    """Limpia el contenido de las celdas de datos (no los labels), preservando placeholders."""
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            # En la primera fila, limpiar solo la segunda celda en adelante (nombres de unidad/TP)
            if row_idx == 0 and cell_idx > 0:
                for paragraph in cell.paragraphs:
                    # No limpiar si contiene placeholders
                    if "{{" not in paragraph.text:
                        paragraph.text = ""
            # En filas posteriores, buscar y limpiar el contenido después de labels
            elif row_idx > 0:
                for paragraph in cell.paragraphs:
                    # NUNCA limpiar líneas que contengan placeholders
                    if "{{" in paragraph.text:
                        continue
                    
                    text_lower = paragraph.text.lower()
                    # Si es un label (termina en :), dividir y limpiar solo después del :
                    if ":" in paragraph.text:
                        # Preservar el label, limpiar el contenido después
                        label_match = re.match(r'([^:]+:)\s*.*', paragraph.text)
                        if label_match and any(kw in text_lower for kw in ["contenido", "bibliograf", "objetivo", "actividad", "material", "ámbito", "ambito"]):
                            paragraph.text = label_match.group(1)  # Solo el label con :
                        else:
                            # Si tiene : pero no es un label conocido, limpiar todo
                            paragraph.text = ""
                    # Si no tiene label, limpiar completamente
                    elif paragraph.text.strip():
                        paragraph.text = ""


def _find_unit_tables(doc: Document) -> List[Table]:
    tables = []
    for table in doc.tables:
        has_unidad = any("unidad" in cell.text.lower() for row in table.rows for cell in row.cells)
        has_contenidos = any("contenidos:" in cell.text.lower() for row in table.rows for cell in row.cells)
        if has_unidad and has_contenidos:
            tables.append(table)
    return tables


def _fill_units(doc: Document, units: List[dict]) -> None:
    tables = _find_unit_tables(doc)
    
    # Si faltan tablas, clonar la última
    if len(units) > len(tables) and tables:
        last_table = tables[-1]
        for _ in range(len(units) - len(tables)):
            new_table = _clone_table_after(doc, last_table)
            _clear_table_data_cells(new_table)
            tables.append(new_table)
            last_table = new_table
    
    for idx, table in enumerate(tables):
        if idx >= len(units):
            break
        unit = units[idx]
        unit_number = idx + 1
        name = unit.get("name") or ""
        content = unit.get("content") or ""
        bib_basic = unit.get("bibliography_basic") or ""
        bib_comp = unit.get("bibliography_complementary") or ""
        
        if table.rows:
            row = table.rows[0]
            if len(row.cells) > 0:
                # Reemplazar # con el número de unidad en la primera celda
                for paragraph in row.cells[0].paragraphs:
                    if "#" in paragraph.text:
                        paragraph.text = paragraph.text.replace("#", str(unit_number))
            if len(row.cells) > 1 and "unidad" in row.cells[0].text.lower():
                row.cells[1].text = name
        
        # Usar placeholders para reemplazar contenido
        mapping = {
            "content": content,
            "biblioBas": bib_basic,
            "biblioCom": bib_comp,
        }
        _replace_placeholders_in_table(table, mapping)


def _find_practical_tables(doc: Document) -> List[Table]:
    tables = []
    practical_placeholders = ("{{object}}", "{{activities}}", "{{maters}}", "{{ambito}}")
    for table in doc.tables:
        # Identificar por encabezado y por placeholders caracteristicos del TP.
        if not table.rows or not table.rows[0].cells:
            continue
        header_text = " ".join(cell.text.lower() for cell in table.rows[0].cells)
        has_practico_header = "practico" in header_text or "práctico" in header_text

        has_tp_placeholders = False
        if has_practico_header:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if any(ph in text for ph in practical_placeholders):
                        has_tp_placeholders = True
                        break
                if has_tp_placeholders:
                    break

        if has_practico_header and has_tp_placeholders:
            tables.append(table)
    return tables


def _fill_practicals(doc: Document, practicals: List[dict]) -> None:
    tables = _find_practical_tables(doc)
    
    # Si faltan tablas, clonar la última
    if len(practicals) > len(tables) and tables:
        last_table = tables[-1]
        for _ in range(len(practicals) - len(tables)):
            new_table = _clone_table_after(doc, last_table)
            _clear_table_data_cells(new_table)
            tables.append(new_table)
            last_table = new_table
    
    for idx, table in enumerate(tables):
        if idx >= len(practicals):
            break
        tp = practicals[idx]
        tp_number = idx + 1
        name = tp.get("name") or ""
        objective = tp.get("objective") or ""
        activities = tp.get("activities") or ""
        materials = tp.get("materials") or ""
        scope = tp.get("scope") or ""
        
        if table.rows:
            row = table.rows[0]
            if len(row.cells) > 0:
                # Reemplazar # con el número de TP en la primera celda
                for paragraph in row.cells[0].paragraphs:
                    if "#" in paragraph.text:
                        paragraph.text = paragraph.text.replace("#", str(tp_number))
            if len(row.cells) > 1 and ("practico" in row.cells[0].text.lower() or "práctico" in row.cells[0].text.lower()):
                row.cells[1].text = name
        
        # Usar placeholders para reemplazar contenido
        mapping = {
            "object": objective,
            "activities": activities,
            "maters": materials,
            "ambito": scope,  # Si es vacío, quedará vacío en el reemplazo
        }
        _replace_placeholders_in_table(table, mapping)


def _fill_section_after_heading(doc: Document, heading_text: str, value: str) -> None:
    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph) and heading_text in block.text.upper():
            for next_block in blocks[idx + 1:]:
                if isinstance(next_block, Table):
                    cell = next_block.cell(0, 0)
                    _fill_first_empty_paragraph(cell.paragraphs, value)
                    return


def _overwrite_table_after_heading(doc: Document, heading_text: str, value: str) -> None:
    if not value:
        return
    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph) and heading_text in block.text.upper():
            for next_block in blocks[idx + 1:]:
                if isinstance(next_block, Table):
                    cell = next_block.cell(0, 0)
                    # Limpiar todo el contenido existente y reemplazar por el valor.
                    for paragraph in cell.paragraphs:
                        paragraph.text = ""
                    if cell.paragraphs:
                        cell.paragraphs[0].text = value
                    else:
                        cell.text = value
                    return


def _split_bibliography(text: str) -> Tuple[str, str]:
    if not text:
        return "", ""
    lower = text.lower()
    if "complementaria" in lower:
        parts = re.split(r"complementaria\s*:|bibliografia\s+complementaria\s*:", text, flags=re.IGNORECASE)
        if len(parts) >= 2:
            basic = parts[0]
            comp = parts[1]
            basic = re.sub(r"bibliografia\s+basica\s*:|basica\s*:", "", basic, flags=re.IGNORECASE).strip()
            return basic.strip(), comp.strip()
    return text.strip(), ""


def generate_proposal_docx(proposal, template_path: str) -> str:
    doc = Document(template_path)

    # El template ya incluye el encabezado; eliminar duplicados del cuerpo.
    _remove_header_paragraphs_in_body(doc)

    def safe(value) -> str:
        if value is None:
            return ""
        return str(value)

    def format_year_of_career(value) -> str:
        text = safe(value).strip()
        if not text:
            return ""
        if text.endswith("º"):
            return text
        return f"{text}º"

    teaching_team = proposal.teaching_team or []
    doc1 = teaching_team[0] if len(teaching_team) > 0 else {}
    doc2 = teaching_team[1] if len(teaching_team) > 1 else {}
    doc3 = teaching_team[2] if len(teaching_team) > 2 else {}

    # Dividir bibliografía en básica y complementaria
    basic_bib, comp_bib = _split_bibliography(safe(proposal.bibliography))
    
    def normalize_quarter(value: str) -> str:
        text = safe(value).strip()
        if not text:
            return ""
        lower = text.lower()
        if "anual" in lower or lower.strip() == "a":
            return "A"
        if "1" in lower or "primer" in lower:
            return "1º"
        if "2" in lower or "segundo" in lower:
            return "2º"
        return text

    mapping = {
        "nombreCarrera": safe(proposal.career),
        "nombreAsignatura": safe(proposal.subject or proposal.title),
        "plan": safe(proposal.study_plan),
        "anio": format_year_of_career(proposal.year_of_career),
        "anioAcadem": safe(proposal.academic_year),
        "cuat": normalize_quarter(proposal.quarter),
        "caracter": safe(proposal.character),
        "regimen": safe(proposal.regime),
        "horasTeoria": safe(proposal.theoretical_hours),
        "horasPractica": safe(proposal.practical_hours),
        "cargaHoraria": safe(proposal.total_hours),
        "hsSe": safe(proposal.weekly_hours),
        "contenidoMinimo": safe(proposal.minimum_content),
        "compegen": safe(proposal.generic_competencies),
        "compeespe": safe(proposal.specific_competencies),
        "doc1": safe(doc1.get("name")),
        "doc2": safe(doc2.get("name")),
        "doc3": safe(doc3.get("name")),
        "catDoc1": safe(doc1.get("category")),
        "catDoc2": safe(doc2.get("category")),
        "catDoc3": safe(doc3.get("category")),
        "correoDoc1": safe(doc1.get("email")),
        "correoDoc2": safe(doc2.get("email")),
        "correoDoc3": safe(doc3.get("email")),
        "metodology": safe(proposal.methodology),
        "evaluation": safe(proposal.evaluation),
        "biblioBasAPA": basic_bib,
        "biblioComAPA": comp_bib,
        "observations": safe(proposal.observations),
    }

    _replace_placeholders(doc, mapping)

    _fill_fundamentals(doc, safe(proposal.fundamentals_part1), safe(proposal.fundamentals_part2))

    learning_outcomes = [
        f"RA {idx + 1}: {item.get('description') or ''}".strip()
        for idx, item in enumerate(proposal.learning_outcomes or [])
        if item.get("description")
    ]
    _fill_learning_outcomes(doc, learning_outcomes)

    _fill_units(doc, proposal.units or [])
    _fill_practicals(doc, proposal.practicals or [])
    _overwrite_table_after_heading(doc, "METODOLOGÍA", safe(proposal.methodology))

    output_dir = tempfile.mkdtemp(prefix="proposal_docx_")
    output_path = os.path.join(output_dir, f"Propuesta_{proposal.id}.docx")
    doc.save(output_path)
    return output_path
