"""
Módulo FINAL para importar datos desde archivos DOCX.
Estrategia: Buscar tablas y párrafos por contenido específico, no por índices.
"""
from __future__ import annotations

import re
from typing import List, Dict, Any
from docx import Document
from docx.table import Table


def extract_text_from_table_cell(cell) -> str:
    """Extrae todo el texto de una celda, preservando párrafos."""
    return '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])


def extract_table_for_content(doc: Document, *search_keywords) -> str:
    """
    Busca una tabla que contenga TODOS los keywords dados y extrae su contenido.
    Usado para encontrar tablas de "Contenidos Mínimos", "Fundamentos", etc.
    """
    for table in doc.tables:
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
        
        # Verificar que la tabla tenga TODOS los keywords
        if all(kw.lower() in table_text for kw in search_keywords):
            content_lines = []
            for row in table.rows[1:]:  # Saltar encabezado
                row_content = '\n'.join([extract_text_from_table_cell(cell) for cell in row.cells])
                if row_content.strip():
                    content_lines.append(row_content)
            return '\n---\n'.join(content_lines)
    
    return ''


def extract_programa_analitico(doc: Document) -> Dict[str, str]:
    """
    Extrae TODOS los campos de la tabla "Programa Analítico":
    Carácter, Régimen, Carga Horaria Tot, Hs Teóricas, Hs Prácticas, Hs Sem
    """
    program_data = {
        'character': '',
        'regime': '',
        'total_hours': '',
        'theoretical_hours': '',
        'practical_hours': '',
        'weekly_hours': '',
    }
    
    for table in doc.tables:
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        if ('régimen' in table_text or 'regimen' in table_text or 'carácter' in table_text):
            if len(table.rows) >= 2:
                header_row = table.rows[0]
                data_row = table.rows[1]
                
                # Mapear encabezados a índices
                for col_idx, header_cell in enumerate(header_row.cells):
                    header_lower = header_cell.text.lower().strip()
                    
                    if col_idx < len(data_row.cells):
                        value = data_row.cells[col_idx].text.strip()
                        
                        if 'carácter' in header_lower or 'caracter' in header_lower:
                            program_data['character'] = value
                        elif 'régimen' in header_lower or 'regimen' in header_lower:
                            program_data['regime'] = value
                        elif 'carga horaria tot' in header_lower or 'carga horaria total' in header_lower:
                            program_data['total_hours'] = value
                        elif 'hs teórica' in header_lower or 'hs teorica' in header_lower:
                            program_data['theoretical_hours'] = value
                        elif 'hs prácti' in header_lower or 'hs practi' in header_lower:
                            program_data['practical_hours'] = value
                        elif 'hs sem' in header_lower:
                            program_data['weekly_hours'] = value
    
    return program_data


def extract_equipo_docente(doc: Document) -> List[Dict[str, str]]:
    """Extrae tabla de Equipo Docente con profesores, categoría, correo."""
    teaching_team = []
    
    for table in doc.tables:
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        if ('equipo' in table_text and 'docente' in table_text) or \
           ('profesor' in table_text and 'categoría' in table_text and 'correo' in table_text):
            
            # Encontrar fila de encabezado
            header_row_idx = 0
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.lower() for cell in row.cells])
                if 'profesor' in row_text or 'categoría' in row_text or 'nombre' in row_text:
                    header_row_idx = row_idx
                    break
            
            # Mapear columnas desde encabezado
            header_row = table.rows[header_row_idx]
            name_idx, category_idx, email_idx = 0, 1, 2
            
            for col_idx, cell in enumerate(header_row.cells):
                cell_lower = cell.text.lower().strip()
                if 'profesor' in cell_lower or 'docente' in cell_lower or 'nombre' in cell_lower:
                    name_idx = col_idx
                elif 'categoría' in cell_lower or 'categoria' in cell_lower:
                    category_idx = col_idx
                elif 'correo' in cell_lower or 'email' in cell_lower:
                    email_idx = col_idx
            
            # Extraer filas de datos
            for row_idx in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[row_idx]
                
                if len(row.cells) > max(name_idx, category_idx, email_idx):
                    name = row.cells[name_idx].text.strip() if name_idx < len(row.cells) else ''
                    category = row.cells[category_idx].text.strip() if category_idx < len(row.cells) else ''
                    email = row.cells[email_idx].text.strip() if email_idx < len(row.cells) else ''
                    
                    if name and name not in ['', '{{doc1}}', '{{doc2}}', '{{doc3}}']:
                        teaching_team.append({'name': name, 'category': category, 'email': email})
    
    return teaching_team


def extract_units_from_docx(doc: Document) -> List[Dict[str, str]]:
    """Extrae unidades desde tablas con "unidad" y "contenidos"."""
    units = []
    unit_counter = 0
    
    for table in doc.tables:
        has_unidad = any('unidad' in cell.text.lower() for row in table.rows for cell in row.cells)
        has_contenidos = any('contenidos' in cell.text.lower() for row in table.rows for cell in row.cells)
        
        if has_unidad and has_contenidos:
            unit_counter += 1
            unit_data = {
                'number': str(unit_counter),
                'name': '',
                'content': '',
                'bibliography_basic': '',
                'bibliography_complementary': '',
            }
            
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    cell_lower = cell_text.lower()
                    
                    if 'unidad' in cell_lower:
                        match = re.search(r'unidad\s+n[°º]?\s*:?\s*(\d+)\s+(.*)', cell_text, re.IGNORECASE)
                        if match:
                            unit_data['name'] = match.group(2).strip()
                    
                    if 'contenidos:' in cell_lower:
                        unit_data['content'] = cell_text.replace('Contenidos:', '').replace('contenidos:', '').strip()
            
            if unit_data['name'] or unit_data['content']:
                units.append(unit_data)
    
    return units


def extract_practicals_from_docx(doc: Document) -> List[Dict[str, str]]:
    """Extrae prácticos desde tablas con "práctico" y "objetivo"."""
    practicals = []
    tp_counter = 0
    
    for table in doc.tables:
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
        
        if ('practico' in table_text or 'práctico' in table_text) and 'objetivo' in table_text:
            tp_counter += 1
            practical_data = {
                'number': str(tp_counter),
                'name': '',
                'objective': '',
                'activities': '',
            }
            
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    cell_lower = cell_text.lower()
                    
                    if 'practico' in cell_lower or 'práctico' in cell_lower:
                        match = re.search(r'pr[áa]ctico\s+n[°º]?\s*:?\s*(\d+)\s+(.*)', cell_text, re.IGNORECASE)
                        if match:
                            practical_data['name'] = match.group(2).strip()
                    
                    if 'objetivo' in cell_lower:
                        practical_data['objective'] = cell_text.replace('Objetivo', '').replace('objetivo', '').strip()
            
            if practical_data['name'] or practical_data['objective']:
                practicals.append(practical_data)
    
    return practicals


def extract_content_from_single_cell_table(table: Table) -> str:
    """Extrae todo el contenido de una tabla con 1 fila y 1 columna."""
    if len(table.rows) >= 1 and len(table.columns) >= 1:
        return extract_text_from_table_cell(table.rows[0].cells[0])
    return ''


def import_proposal_from_docx(file_path: str, filename: str = "") -> Dict[str, Any]:
    """
    Importa una propuesta desde DOCX.
    Estructura conocida:
    - Tabla 0: Programa Analítico (6 cols)
    - Tabla 1: Equipo Docente (3 cols)
    - Tabla 2+: Contenido (mayormente 1 celda)
    """
    doc = Document(file_path)
    
    # Extraer metadata del encabezado
    fields = {
        'career': '',
        'subject': '',
        'year_of_career': '',
        'quarter': '',
    }
    
    # Parsear filename
    if filename:
        year_match = re.search(r'(\d+)[°º]', filename)
        if year_match:
            fields['year_of_career'] = year_match.group(1)
        
        quarter_match = re.search(r'(\d+)[°º]\s*[-_]\s*(\d+)[°º]', filename)
        if quarter_match:
            fields['quarter'] = quarter_match.group(2)
        
        if ' - ' in filename:
            subject_part = filename.split(' - ', 1)[1]
            subject_part = subject_part.replace('.docx', '').replace('.DOCX', '').strip()
            if subject_part:
                fields['subject'] = subject_part
    
    # Buscar carrera en documento
    for para in doc.paragraphs[:100]:
        if 'carrera' in para.text.lower() and ':' in para.text:
            fields['career'] = para.text.split(':', 1)[1].strip()
            break
    
    # Extraer tablas principales
    programa_analitico = extract_programa_analitico(doc)
    teaching_team = extract_equipo_docente(doc)
    
    # Contenido de secciones - buscar en TODAS las tablas por contenido
    contenidos_minimos = ''
    importance = ''
    generic_comp = ''
    specific_comp = ''
    objectives = ''
    methodology = ''
    evaluation = ''
    bibliography = ''
    observations = ''
    
    # Estrategia: buscar tablas por contenido/keyword, pero siendo inteligente:
    # - Tabla después de Equipo Docente probablemente sea Contenidos Mínimos
    # - Si una tabla no tiene "unidad", "práctico", "caracter", etc., probablemente sea contenido
    
    for table_idx, table in enumerate(doc.tables):
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
        table_content = extract_content_from_single_cell_table(table) if len(table.rows) >= 1 else ''
        
        # Tabla 2 es típicamente Contenidos Mínimos (después de Programa Analítico y Equipo Docente)
        if table_idx == 2 and not contenidos_minimos and len(table_content) > 50:
            contenidos_minimos = table_content
        
        # Buscar cada sección por palabra clave
        if not contenidos_minimos and ('contenidos' in table_text and 'mínimos' in table_text):
            contenidos_minimos = table_content
        
        if not importance and ('importancia' in table_text or 'fundamentos' in table_text):
            importance = table_content
        
        if not generic_comp and 'competencias' in table_text and 'genéricas' in table_text:
            generic_comp = table_content
        
        if not specific_comp and 'competencias' in table_text and 'específicas' in table_text:
            specific_comp = table_content
        
        if not methodology and ('metodología' in table_text or 'metodologia' in table_text):
            methodology = table_content
        
        if not evaluation and ('evaluación' in table_text or 'evaluacion' in table_text):
            evaluation = table_content
        
        if not bibliography and ('bibliografía' in table_text or 'bibliografia' in table_text):
            bibliography = table_content
        
        if not observations and ('observaciones' in table_text or 'laboratorio' in table_text):
            observations = table_content
    
    # Extraer TODAS las unidades desde TODAS las tablas
    # Buscar cada celda que tenga "Unidad N°:" sin importar en qué tabla esté
    units = []
    seen_numbers = set()
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                # Buscar "Unidad N°:" en cualquier celda
                if 'unidad' in cell_text.lower():
                    match = re.search(r'unidad\s+n[°º]?\s*:?\s*(\d+)', cell_text, re.IGNORECASE)
                    if match:
                        unit_num = match.group(1)
                        if unit_num not in seen_numbers:  # Evitar duplicados
                            seen_numbers.add(unit_num)
                            unit_data = {'number': unit_num, 'name': '', 'content': '', 'bibliography_basic': '', 'bibliography_complementary': ''}
                            
                            # El nombre está en la siguiente celda de la misma fila
                            if cell_idx + 1 < len(row.cells):
                                unit_data['name'] = row.cells[cell_idx + 1].text.strip()
                            
                            # Buscar contenidos en la siguiente fila
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx + 1]
                                if len(next_row.cells) >= 1:
                                    content_text = extract_text_from_table_cell(next_row.cells[0])
                                    unit_data['content'] = content_text
                            
                            units.append(unit_data)
    
    # Extraer TODOS los prácticos desde TODAS las tablas
    # Buscar cada celda que tenga "Práctico Nº:" sin importar en qué tabla esté
    practicals = []
    seen_tp_numbers = set()
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                # Buscar "Práctico Nº:" en cualquier celda
                if ('práctico' in cell_text.lower() or 'practico' in cell_text.lower()):
                    match = re.search(r'pr[áa]ctico\s+n[°º]?\s*:?\s*(\d+)', cell_text, re.IGNORECASE)
                    if match:
                        tp_num = match.group(1)
                        if tp_num not in seen_tp_numbers:  # Evitar duplicados
                            seen_tp_numbers.add(tp_num)
                            practical_data = {'number': tp_num, 'name': '', 'objective': '', 'activities': ''}
                            
                            # El nombre está en la siguiente celda de la misma fila
                            if cell_idx + 1 < len(row.cells):
                                practical_data['name'] = row.cells[cell_idx + 1].text.strip()
                            
                            # Buscar objetivo en la siguiente fila
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx + 1]
                                if len(next_row.cells) >= 1:
                                    obj_text = extract_text_from_table_cell(next_row.cells[0])
                                    practical_data['objective'] = obj_text
                            
                            practicals.append(practical_data)
    
    # Buscar metodología, evaluación, bibliografía, observaciones en tablas finales (1 celda)
    # (Ya se buscaron en el loop anterior, pero podemos también buscar en paragrafos si no se encontraron)
    
    # Compilar respuesta
    return {
        'career': fields.get('career', ''),
        'subject': fields.get('subject', ''),
        'year_of_career': fields.get('year_of_career', ''),
        'quarter': fields.get('quarter', ''),
        'study_plan': '',
        
        'character': programa_analitico.get('character', ''),
        'regime': programa_analitico.get('regime', ''),
        'total_hours': programa_analitico.get('total_hours', ''),
        'theoretical_hours': programa_analitico.get('theoretical_hours', ''),
        'practical_hours': programa_analitico.get('practical_hours', ''),
        'weekly_hours': programa_analitico.get('weekly_hours', ''),
        
        'teachers': '; '.join([t['name'] + ' (' + t['category'] + ')' for t in teaching_team if t.get('name')]),
        'teaching_team': teaching_team,
        
        'minimum_content': contenidos_minimos,
        'importance': importance,
        'fundamentals': importance,
        'generic_competencies': generic_comp,
        'specific_competencies': specific_comp,
        'objectives': objectives,
        'learning_outcomes': [],
        
        'units': units,
        'practicals': practicals,
        
        'methodology': methodology,
        'evaluation': evaluation,
        'bibliography': bibliography,
        'observations': observations,
    }
