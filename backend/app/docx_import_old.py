"""
Módulo para importar datos desde archivos DOCX.
Extrae contenido de un DOCX generado por el sistema y lo convierte en JSON.
"""
from __future__ import annotations

import re
from typing import List, Dict, Any
from docx import Document
from docx.table import Table


def extract_text_from_table_cell(cell) -> str:
    """Extrae todo el texto de una celda, preservando párrafos."""
    return '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])


def extract_header_fields(doc: Document, filename: str = "") -> Dict[str, str]:
    """
    Extrae campos del encabezado del documento.
    Busca en cuatro lugares:
    1. Nombre del archivo (ej: "5°_2° - Proyecto de Ingeniería Mecatrónica")
    2. Párrafos iniciales (carrera, asignatura, plan, ciclo, año, cuatrimestre)
    3. Tabla de Programa Analítico (régimen, carga horaria, horas)
    4. Tabla de Equipo Docente (profesores, categorías, emails)
    """
    fields = {
        'career': '',
        'subject': '',
        'teachers': '',
        'year_of_career': '',
        'quarter': '',
        'hours': '',
        'regime': '',
        'teaching_team': [],
    }
    
    # 0. EXTRAER DEL NOMBRE DEL ARCHIVO
    # Formato típico: "5°_2° - Nombre de Asignatura.docx"
    if filename:
        # Buscar patrón como "5°" o "5°_2°" etc
        year_match = re.search(r'(\d+)[°º]', filename)
        if year_match:
            fields['year_of_career'] = year_match.group(1)
        
        # Buscar segundo número para cuatrimestre: "5°_2°"
        quarter_match = re.search(r'(\d+)[°º]\s*[-_]\s*(\d+)[°º]', filename)
        if quarter_match:
            fields['quarter'] = quarter_match.group(2)
        
        # Extraer asignatura: todo después del guion " - "
        if ' - ' in filename:
            subject_part = filename.split(' - ', 1)[1]
            subject_part = subject_part.replace('.docx', '').replace('.DOCX', '').strip()
            if subject_part:
                fields['subject'] = subject_part
    
    # 1. EXTRAER DEL ENCABEZADO (párrafos iniciales)
    # Limitar a los primeros 15 párrafos para evitar buscar en contenido
    for idx, para in enumerate(doc.paragraphs[:15]):
        if idx > 15:  # Safety limit
            break
        
        text = para.text.strip()
        if not text:
            continue
        
        text_lower = text.lower()
        
        # Carrera (exacto: debe empezar con "Carrera:")
        if text_lower.startswith('carrera:') and not fields['career']:
            fields['career'] = text.split(':', 1)[1].strip()
        
        # Asignatura/Nombre de Asignatura (no reemplazar si ya tenemos del filename)
        elif text_lower.startswith('asignatura:') and not fields['subject']:
            fields['subject'] = text.split(':', 1)[1].strip()
        
        # Plan de Estudio
        elif text_lower.startswith('plan') and ':' in text and not fields.get('study_plan'):
            fields['study_plan'] = text.split(':', 1)[1].strip()
        
        # Ciclo (puede decir "Ciclo:" o "Ciclo Lectivo:")
        elif 'ciclo' in text_lower and ':' in text and not fields.get('cycle'):
            fields['cycle'] = text.split(':', 1)[1].strip()
        
        # Año de Carrera - no reemplazar si ya tenemos del filename
        elif text_lower.startswith('año') and 'carrera' in text_lower and ':' in text and not fields['year_of_career']:
            fields['year_of_career'] = text.split(':', 1)[1].strip()
        
        # Cuatrimestre - no reemplazar si ya tenemos del filename
        elif 'cuatrimestre' in text_lower and ':' in text and not fields['quarter']:
            fields['quarter'] = text.split(':', 1)[1].strip()
    
    # 2. EXTRAER TABLA DE PROGRAMA ANALÍTICO (Carácter, Régimen, Carga Horaria)
    for table in doc.tables:
        # Buscar tabla que contiene "Programa Analítico de Asignatura" o similar
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        if 'régimen' in table_text or 'regimen' in table_text:
            # Esta es la tabla de características
            # Estructura: Row 0 = Headers, Row 1+ = Valores
            if len(table.rows) >= 2:
                header_row = table.rows[0]
                data_row = table.rows[1]
                
                # Encontrar índices de las columnas
                regime_idx = None
                hours_idx = None
                
                for cell_idx, cell in enumerate(header_row.cells):
                    cell_lower = cell.text.lower()
                    if 'régimen' in cell_lower or 'regimen' in cell_lower:
                        regime_idx = cell_idx
                    elif 'carga horaria' in cell_lower:
                        hours_idx = cell_idx
                
                # Extraer valores de la fila de datos
                if regime_idx is not None and regime_idx < len(data_row.cells):
                    regime_val = data_row.cells[regime_idx].text.strip()
                    if regime_val and regime_val not in ['{{regimen}}', '-']:
                        fields['regime'] = regime_val
                
                if hours_idx is not None and hours_idx < len(data_row.cells):
                    hours_val = data_row.cells[hours_idx].text.strip()
                    if hours_val and hours_val not in ['{{cargaHoraria}}', '-']:
                        fields['hours'] = hours_val
    
    # 3. EXTRAER TABLA DE EQUIPO DOCENTE
    for table in doc.tables:
        table_text = ' '.join([cell.text.lower() for row in table.rows for cell in row.cells])
        
        # Buscar tabla que contiene "Docente", "Profesor", "Equipo Docente"
        if ('equipo' in table_text and 'docente' in table_text) or \
           ('profesor' in table_text and 'categoría' in table_text and 'correo' in table_text):
            
            # Extraer docentes de las filas (saltando encabezado)
            header_row_idx = 0
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.lower() for cell in row.cells])
                if 'profesor' in row_text or 'docente' in row_text or 'categoría' in row_text:
                    header_row_idx = row_idx
                    break
            
            # Procesar filas de datos después del encabezado
            for row_idx in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[row_idx]
                
                # Necesitamos al menos nombre + categoría + email (3 celdas)
                if len(row.cells) >= 3:
                    name = row.cells[0].text.strip()
                    category = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                    email = row.cells[2].text.strip() if len(row.cells) > 2 else ''
                    
                    # Solo agregar si tiene al menos nombre válido
                    if name and name not in ['', '{{doc1}}', '{{doc2}}', '{{doc3}}']:
                        fields['teaching_team'].append({
                            'name': name,
                            'category': category,
                            'email': email,
                        })
    
    # Formatear docentes como string si es necesario (para compatibilidad)
    if fields['teaching_team']:
        fields['teachers'] = '; '.join([
            f"{t['name']} ({t['category']})" 
            for t in fields['teaching_team'] 
            if t.get('name')
        ])
    
    # 4. BÚSQUEDA ADICIONAL EN TABLAS Y PÁRRAFOS PARA CARRERA
    # Si no tenemos carrera aún, buscar en texto de "Importancia en Plan de estudio"
    if not fields['career']:
        # Buscar en todas las tablas y párrafos
        carrera_keywords = ['ingeniería mecatrónica', 'ingeniería en sistemas', 'ingeniería civil', 
                           'ingeniería industrial', 'ingeniería eléctrica', 'ingeniería electrónica',
                           'licenciatura', 'profesorado']
        
        # Buscar en párrafos
        for para in doc.paragraphs[:100]:  # Buscar en más párrafos
            text_lower = para.text.lower()
            for career_kw in carrera_keywords:
                if career_kw in text_lower:
                    # Intentar extraer directamente
                    for match in re.finditer(career_kw, text_lower):
                        start_idx = text_lower.rfind(' ', 0, match.start()) + 1
                        end_idx = text_lower.find(' ', match.end())
                        if end_idx == -1:
                            end_idx = len(text_lower)
                        extracted = para.text[start_idx:end_idx].strip()
                        if 8 <= len(extracted) <= 60 and any(c.isalpha() for c in extracted):
                            fields['career'] = extracted
                            break
            if fields['career']:
                break
        
        # Si aún no tenemos carrera, buscar en celdas de tablas
        if not fields['career']:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.lower()
                        for career_kw in carrera_keywords:
                            if career_kw in cell_text:
                                # Intentar extraer
                                for match in re.finditer(career_kw, cell_text):
                                    start_idx = cell_text.rfind(' ', 0, match.start()) + 1
                                    end_idx = cell_text.find(' ', match.end())
                                    if end_idx == -1:
                                        end_idx = len(cell_text)
                                    extracted = cell.text[start_idx:end_idx].strip()
                                    if 8 <= len(extracted) <= 60 and any(c.isalpha() for c in extracted):
                                        fields['career'] = extracted
                                        break
                if fields['career']:
                    break
    
    return fields


def extract_section_content(doc: Document, section_keyword: str, end_keyword: str = None) -> str:
    """Extrae contenido completo entre dos secciones."""
    content = []
    found = False
    
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        
        if section_keyword in text_lower:
            found = True
            # Agregar el contenido después de la etiqueta
            if ':' in para.text:
                content_after = para.text.split(':', 1)[1].strip()
                if content_after:
                    content.append(content_after)
        elif found:
            if end_keyword and end_keyword in text_lower:
                break
            if para.text.strip():
                content.append(para.text)
    
    return '\n'.join(content).strip()


def extract_units(doc: Document) -> List[Dict[str, str]]:
    """Extrae todas las unidades de las tablas."""
    units = []
    
    # Encontrar tablas de unidades
    for table in doc.tables:
        has_unidad = any("unidad" in cell.text.lower() for row in table.rows for cell in row.cells)
        has_contenidos = any("contenidos:" in cell.text.lower() for row in table.rows for cell in row.cells)
        
        if not (has_unidad and has_contenidos):
            continue
        
        unit = {
            'name': '',
            'content': '',
            'bibliography_basic': '',
            'bibliography_complementary': '',
        }
        
        # Primera fila contiene el número y nombre
        if table.rows:
            first_row = table.rows[0]
            if len(first_row.cells) > 1:
                # Extraer nombre de la unidad
                unit['name'] = first_row.cells[1].text.strip()
        
        # Buscar contenido, bibliografía básica y complementaria
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                
                if 'contenidos:' in cell_text.lower():
                    # Extraer contenido después del label
                    lines = cell_text.split('\n')
                    content_lines = []
                    found_label = False
                    for line in lines:
                        if 'contenidos:' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                content_lines.append(content_after)
                        elif found_label and line.strip() and 'bibliograf' not in line.lower():
                            content_lines.append(line.strip())
                        elif found_label and 'bibliograf' in line.lower():
                            break
                    unit['content'] = '\n'.join(content_lines).strip()
                
                elif 'bibliografía básica' in cell_text.lower():
                    # Extraer bibliografía básica
                    lines = cell_text.split('\n')
                    bib_lines = []
                    found_label = False
                    for line in lines:
                        if 'bibliografía básica' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                bib_lines.append(content_after)
                        elif found_label and line.strip() and 'complementaria' not in line.lower():
                            bib_lines.append(line.strip())
                        elif found_label and 'complementaria' in line.lower():
                            break
                    unit['bibliography_basic'] = '\n'.join(bib_lines).strip()
                
                elif 'bibliografía complementaria' in cell_text.lower():
                    # Extraer bibliografía complementaria
                    lines = cell_text.split('\n')
                    bib_lines = []
                    found_label = False
                    for line in lines:
                        if 'bibliografía complementaria' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                bib_lines.append(content_after)
                        elif found_label and line.strip():
                            bib_lines.append(line.strip())
                    unit['bibliography_complementary'] = '\n'.join(bib_lines).strip()
        
        if unit['name']:  # Solo agregar si tiene nombre
            units.append(unit)
    
    return units


def extract_practicals(doc: Document) -> List[Dict[str, str]]:
    """Extrae todos los trabajos prácticos de las tablas."""
    practicals = []
    
    # Encontrar tablas de TPs
    for table in doc.tables:
        has_practico = any(("practico" in cell.text.lower() or "práctico" in cell.text.lower())
                          for row in table.rows for cell in row.cells)
        has_objetivo = any("objetivo" in cell.text.lower() for row in table.rows for cell in row.cells)
        
        if not (has_practico and has_objetivo):
            continue
        
        practical = {
            'name': '',
            'objective': '',
            'activities': '',
            'materials': '',
            'scope': '',
        }
        
        # Primera fila contiene el número y nombre
        if table.rows:
            first_row = table.rows[0]
            if len(first_row.cells) > 1:
                practical['name'] = first_row.cells[1].text.strip()
        
        # Buscar objetivo, actividades, materiales, ámbito
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                
                if 'objetivo' in cell_text.lower():
                    lines = cell_text.split('\n')
                    obj_lines = []
                    found_label = False
                    for line in lines:
                        if 'objetivo' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                obj_lines.append(content_after)
                        elif found_label and line.strip() and 'actividad' not in line.lower():
                            obj_lines.append(line.strip())
                        elif found_label and 'actividad' in line.lower():
                            break
                    practical['objective'] = '\n'.join(obj_lines).strip()
                
                elif 'actividad' in cell_text.lower() and 'desarrollar' in cell_text.lower():
                    lines = cell_text.split('\n')
                    act_lines = []
                    found_label = False
                    for line in lines:
                        if 'actividad' in line.lower() and 'desarrollar' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                act_lines.append(content_after)
                        elif found_label and line.strip() and 'material' not in line.lower():
                            act_lines.append(line.strip())
                        elif found_label and 'material' in line.lower():
                            break
                    practical['activities'] = '\n'.join(act_lines).strip()
                
                elif 'material' in cell_text.lower():
                    lines = cell_text.split('\n')
                    mat_lines = []
                    found_label = False
                    for line in lines:
                        if 'material' in line.lower():
                            found_label = True
                            content_after = line.split(':', 1)[1].strip() if ':' in line else ''
                            if content_after:
                                mat_lines.append(content_after)
                        elif found_label and line.strip() and ('ámbito' in line.lower() or 'ambito' in line.lower()):
                            break
                        elif found_label and line.strip():
                            mat_lines.append(line.strip())
                    practical['materials'] = '\n'.join(mat_lines).strip()
                
                elif 'ámbito' in cell_text.lower() or 'ambito' in cell_text.lower():
                    lines = cell_text.split('\n')
                    scope_lines = []
                    found_label = False
                    for line in lines:
                        if ('ámbito' in line.lower() or 'ambito' in line.lower()) and ':' in line:
                            found_label = True
                            content_after = line.split(':', 1)[1].strip()
                            if content_after:
                                scope_lines.append(content_after)
                        elif found_label and line.strip():
                            scope_lines.append(line.strip())
                    practical['scope'] = '\n'.join(scope_lines).strip()
        
        if practical['name']:  # Solo agregar si tiene nombre
            practicals.append(practical)
    
    return practicals


def extract_learning_outcomes(doc: Document) -> List[str]:
    """Extrae los Resultados de Aprendizaje (RA)."""
    outcomes = []
    found_ra_section = False
    
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        
        if 'resultado' in text_lower and 'aprendizaje' in text_lower:
            found_ra_section = True
            continue
        
        if found_ra_section:
            if para.text.strip().startswith('RA ') or para.text.strip().startswith('ra '):
                outcomes.append(para.text.strip())
            elif any(keyword in text_lower for keyword in ['contenido', 'unidad', 'fundamental']):
                break
    
    return outcomes


def import_proposal_from_docx(file_path: str, filename: str = "") -> Dict[str, Any]:
    """
    Importa una propuesta completa desde un DOCX.
    Retorna un diccionario con todos los campos extraídos.
    
    Args:
        file_path: Ruta del archivo DOCX
        filename: Nombre original del archivo (para parsing de metadata). Si no se proporciona, se extrae de file_path.
    """
    doc = Document(file_path)
    
    # Usar el filename proporcionado, o extraer del path
    if not filename:
        import os
        filename = os.path.basename(file_path).replace('.docx', '').replace('.DOCX', '')
    
    # Extraer headers con la nueva lógica mejorada (incluyendo filename para parsing)
    header_fields = extract_header_fields(doc, filename)
    
    # Extraer teaching team si existe
    teaching_team = header_fields.pop('teaching_team', [])
    
    # Construir lista de docentes en el formato esperado
    teaching_team_list = []
    for idx, teacher in enumerate(teaching_team[:3]):  # Max 3 docentes
        teaching_team_list.append({
            'name': teacher.get('name', ''),
            'category': teacher.get('category', ''),
            'email': teacher.get('email', ''),
        })
    
    # Extraer todos los datos
    data = {
        # Encabezado (como se envía al API)
        'career': header_fields.get('career', ''),
        'subject': header_fields.get('subject', ''),
        'study_plan': header_fields.get('study_plan', ''),
        'cycle': header_fields.get('cycle', ''),
        'year_of_career': header_fields.get('year_of_career', ''),
        'quarter': header_fields.get('quarter', ''),
        'character': header_fields.get('character', ''),
        'regime': header_fields.get('regime', ''),
        'total_hours': header_fields.get('hours', ''),  # Carga Horaria
        'teaching_team': teaching_team_list,
        
        # Secciones generales
        'minimum_content': extract_section_content(doc, 'contenidos mínimos', 'importancia'),
        'importance': extract_section_content(doc, 'importancia', 'fundamental'),
        'fundamentals_part1': extract_section_content(doc, 'fundamental'),
        'fundamentals_part2': '',
        
        # Competencias
        'generic_competencies': extract_section_content(doc, 'competencias genérica', 'competencias específica'),
        'specific_competencies': extract_section_content(doc, 'competencias específica', 'resultado'),
        
        # Resultados de aprendizaje
        'learning_outcomes': extract_learning_outcomes(doc),
        
        # Unidades
        'units': extract_units(doc),
        
        # Trabajos prácticos
        'practicals': extract_practicals(doc),
        
        # Secciones finales
        'methodology': extract_section_content(doc, 'metodología', 'evaluación'),
        'evaluation': extract_section_content(doc, 'evaluación', 'bibliografía'),
        'bibliography': extract_section_content(doc, 'bibliografía', 'observaciones'),
        'observations': extract_section_content(doc, 'observaciones'),
    }
    
    # Limpiar valores vacíos en teaching_team
    data['teaching_team'] = [
        t for t in data['teaching_team'] 
        if t.get('name') and t['name'] not in ['', '-']
    ]
    
    return data
