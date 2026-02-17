#!/usr/bin/env python
"""Inspeccionador de tablas - Ver EXACTAMENTE qué hay en cada tabla"""
from docx import Document

file_path = r'c:\TesisMCD\backend\data\uploads\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(file_path)

print("=" * 100)
print("ANÁLISIS DETALLADO DE TODAS LAS TABLAS")
print("=" * 100)

for table_idx, table in enumerate(doc.tables):
    print(f"\n{'=' * 100}")
    print(f"TABLA {table_idx}: {len(table.rows)} filas x {len(table.columns)} columnas")
    print(f"{'=' * 100}")
    
    # Mostrar primeras 5 filas
    for row_idx, row in enumerate(table.rows[:5]):
        print(f"\n  Fila {row_idx}:")
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()[:150]  # Primeros 150 caracteres
            if cell_text:
                print(f"    [{col_idx}] {repr(cell_text)}")
            else:
                print(f"    [{col_idx}] [VACÍO]")
    
    if len(table.rows) > 5:
        print(f"\n  ... ({len(table.rows) - 5} filas más)")
    
    # Resumen de contenido
    all_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
    keywords_found = []
    for kw in ['unidad', 'contenido', 'practico', 'práctico', 'objetivo', 'metodología', 
               'evaluacion', 'evaluación', 'bibliografía', 'bibliografia', 'competencia', 
               'caracter', 'régimen', 'regimen', 'profesor', 'categoría', 'correo',
               'fundamentos', 'mínimos']:
        if kw in all_text:
            keywords_found.append(kw)
    
    if keywords_found:
        print(f"\n  Keywords: {', '.join(keywords_found)}")
    else:
        print(f"\n  Keywords: [NINGUNO]")

print("\n" + "=" * 100)
print("TOTAL DE TABLAS:", len(doc.tables))
print("=" * 100)
