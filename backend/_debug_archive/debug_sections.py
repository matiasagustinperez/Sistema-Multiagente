#!/usr/bin/env python3
"""Debug section finding"""

import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

from docx import Document
import re

DOCX_PATH = r"C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx"
doc = Document(DOCX_PATH)

print("=" * 80)
print("DEBUG: BUSQUEDA DE SECCIONES")
print("=" * 80)

section_keywords = {
    'CONTENIDOS MINIMOS': 0,
    'CONTENIDOS MÍNIMOS': 0,
    'FUNDAMENTOS': 1,
    'OBJETIVOS': 2,
    'CONTENIDOS DE LA ASIGNATURA': 3,
    'PROGRAMA DE TRABAJOS PRACTICOS': 4,
    'PROGRAMA DE TRABAJO PRACTICO': 4,
    'METODOLOGIA': 5,
    'METODOLOGÍA': 5,
    'EVALUACION': 6,
    'EVALUACIÓN': 6,
    'BIBLIOGRAFIA': 7,
    'BIBLIOGRAFÍA': 7,
    'OBSERVACIONES': 8,
}

print("\nBuscando en primeros 35 párrafos:")
print("-" * 80)

for idx, para in enumerate(doc.paragraphs[:35]):
    text = para.text.strip()
    text_upper = text.upper()
    
    print("Para %d: %s" % (idx, text[:70]))
    
    # Buscar coincidencia
    for keyword in section_keywords.keys():
        if text_upper.startswith(keyword):
            print("  --> ENCONTRADO: '%s'" % keyword)

print("\n\nIntentando ahora con función actualizada:")
print("-" * 80)

sections = {}
section_starts = {}
section_order = []

for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip().upper()
    
    for keyword, order in section_keywords.items():
        if text.startswith(keyword):
            section_name = para.text.strip().rstrip(':')
            section_starts[order] = (idx, section_name)
            section_order.append((order, idx, section_name))
            print("Seccion encontrada: ordem=%d, idx=%d, nombre=%s" % (order, idx, section_name[:60]))

print("\nResultado final:")
print("Secciones encontradas: %d" % len(section_order))
