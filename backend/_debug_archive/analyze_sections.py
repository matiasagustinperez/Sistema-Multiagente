#!/usr/bin/env python3
"""Analizar estructura de secciones en el DOCX real"""

from docx import Document
import re

DOCX_PATH = r"C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx"
doc = Document(DOCX_PATH)

print("=" * 80)
print("ANÁLISIS: ESTRUCTURA DE SECCIONES EN EL DOCX REAL")
print("=" * 80)

print("\n🔍 Primeros 50 párrafos (buscando secciones numeradas):")
print("-" * 80)

secciones_encontradas = []

for idx, para in enumerate(doc.paragraphs[:100]):
    text = para.text.strip()
    
    # Buscar patrones de secciones
    if text:
        # Mostrar línea
        preview = text[:100] if len(text) > 100 else text
        print(f"Para {idx:2d}: {preview}")
        
        # Buscar secciones numeradas
        match = re.match(r'^(\d+)\.\s+([A-Z][A-ZÁÉÍÓÚñáéíóú\s\-]+):?\s*$', text, re.IGNORECASE)
        if match:
            section_num = match.group(1)
            section_name = match.group(2)
            secciones_encontradas.append((idx, section_num, section_name))
            print(f"  ✓ SECCIÓN ENCONTRADA: {section_num}. {section_name}")

print("\n\n📋 RESUMEN DE SECCIONES ENCONTRADAS:")
print("-" * 80)
if secciones_encontradas:
    for idx, section_num, section_name in secciones_encontradas:
        print(f"Sección {section_num}: {section_name} (párrafo {idx})")
else:
    print("❌ No se encontraron secciones con patrón '1. NOMBRE:'")
    
    # Intentar buscar con otros patrones
    print("\nIntentando otros patrones:")
    for idx, para in enumerate(doc.paragraphs[:60]):
        text = para.text.strip()
        if re.search(r'^\d+\.\s+', text):
            print(f"Para {idx}: {text[:80]}")

print("\n\nTotal de párrafos: {len(doc.paragraphs)}")
print("Total de tablas: {len(doc.tables)}")
