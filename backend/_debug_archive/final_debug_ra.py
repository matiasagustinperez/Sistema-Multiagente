"""Máximo nivel de debugging - ver EXACTAMENTE qué recibe parse_practical_block"""
import sys
import os
sys.path.insert(0, os.getcwd())

from docx import Document
import re

def extract_text_from_table_cell(cell) -> str:
    return '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])

doc = Document(r"1º_1º - CBI - Álgebra I.docx")
table = doc.tables[7]

# Obtener TP1 (Fila 0 y 1)
row_idx = 0
row = table.rows[row_idx]

# Obtener bloque (Fila 1)
block_cells = table.rows[1].cells
block_text = '\n'.join([extract_text_from_table_cell(cell) for cell in block_cells]).strip()

print("="*80)
print("TEXTO PASADO A parse_practical_block PARA TP1")
print("="*80)

print(f"\nblock_text ({len(block_text)} chars):\n")
print(block_text[:1000])
print("\n... (truncado)")

# Ahora extraer objetivo exactamente
label_patterns = [
    ('objective', re.compile(r'(?i)objetivo(?:\s*\(.*?\))?\s*:?')),
]

print("\n" + "="*80)
print("EXTRACCIÓN DE OBJETIVO")
print("="*80)

matches = []
for label, pattern in label_patterns:
    match = pattern.search(block_text)
    if match:
        print(f"\nPatrón 'objective' encontrado en posición {match.start()}-{match.end()}")
        match_text = block_text[match.start():match.end()]
        print(f"  Texto del match: {match_text}")
        matches.append((match.start(), match.end(), label))

# Encontrar fin del segmento (buscar la siguiente etiqueta o fin)
label_patterns = [
    ('objective', re.compile(r'(?i)objetivo(?:\s*\(.*?\))?\s*:?')),
    ('activities', re.compile(r'(?i)actividades?\s+a\s+desarrollar(?:\s*\(.*?\))?\s*:?')),
]

# Buscar la siguiente
next_match = None
for label, pattern in label_patterns[1:]:  # Skip 'objective'
    m = pattern.search(block_text)
    if m:
        if next_match is None or m.start() < next_match[0]:
            next_match_pos = m.start()
            break

if next_match is None:
    next_match_pos = len(block_text)

# Extraer segmento objetivo
start, end, _ = matches[0]
objective_raw = block_text[end:next_match_pos].strip()

print(f"objective_raw extraído  ({len(objective_raw)} chars):\n")
print(objective_raw[:500])
print("\n... (truncado)" if len(objective_raw) > 500 else "")

# Buscar RAs
print("\n" + "="*80)
print("BÚSQUEDA DE RAs EN objective_raw")
print("="*80)

ra_codes = []
for match in re.finditer(r'RA\s*(\d+)', objective_raw, re.IGNORECASE):
    code = f"RA{match.group(1)}"
    if code not in ra_codes:
        ra_codes.append(code)
        print(f"\nEncontrado: {code}")

print(f"\nRAs finales: {ra_codes}")
