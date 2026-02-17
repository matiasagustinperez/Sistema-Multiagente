#!/usr/bin/env python
"""Debug para encontrar los contenidos mínimos"""
from docx import Document

file_path = r'c:\TesisMCD\backend\data\uploads\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(file_path)

print("Tabla 0: Programa Analítico")
print(f"  Tabla 1: Equipo Docente (siguiente a Programa)")
print(f"  Tabla 2: ¿Contenidos Mínimos?")

if len(doc.tables) > 2:
    table_2 = doc.tables[2]
    print(f"\nTabla 2 - {len(table_2.rows)} filas x {len(table_2.columns)} columnas:")
    print(f"  Content: {table_2.rows[0].cells[0].text[:150]}")
    print(f"  ¿Tiene 'contenidos'? {'contenidos' in table_2.rows[0].cells[0].text.lower()}")
    print(f"  ¿Tiene 'mínimos'? {'mínimos' in table_2.rows[0].cells[0].text.lower()}")

print("\nBuscando todas las tablas con 'contenidos':")
for idx, table in enumerate(doc.tables):
    table_text = ' '.join([cell.text for row in table.rows for cell in row.cells]).lower()
    if 'contenidos' in table_text:
        print(f"  Tabla {idx}: SÍ tiene 'contenidos'")
        if 'mínimos' in table_text:
            print(f"    → También tiene 'mínimos'")
        rows = len(table.rows)
        cols = len(table.columns)
        preview = table.rows[0].cells[0].text[:80]
        print(f"    Estructura: {rows} filas, {cols} cols")
        print(f"    Preview: {preview}")
