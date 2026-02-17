"""Verificar qué tablas existen"""
from docx import Document

doc = Document(r"1º_1º - CBI - Álgebra I.docx")

print(f"Total de tablas: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
    has_practico = 'practico' in table_text.lower()
    has_objetivo = 'objetivo' in table_text.lower()
    
    print(f"\nTabla {idx}:")
    print(f"  Tiene 'practico': {has_practico}")
    print(f"  Tiene 'objetivo': {has_objetivo}")
    print(f"  Primeras 100 chars: {table_text[:100]}")
