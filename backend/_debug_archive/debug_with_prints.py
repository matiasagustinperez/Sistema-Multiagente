"""Test extract_practicals_from_docx CON DEBUGGING INSERTADO"""
import sys
import os
sys.path.insert(0, os.getcwd())

from docx import Document
import re

doc_path = r"1º_1º - CBI - Álgebra I.docx"
doc = Document(doc_path)

# COPIAR LA FUNCIÓN COMPLETA Y AGREGAR PRINTS

def extract_text_from_table_cell(cell) -> str:
    """Extrae todo el texto de una celda, preservando párrafos."""
    return '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])

def extract_practicals_from_docx_DEBUG(doc):
    """Version CON DEBUGGING"""
    
    def parse_practical_block(text: str):
        if not text:
            return {'objective': '', 'activities': '', 'materials': '', 'scope': '', 'ra_codes': []}

        label_patterns = [
            ('objective', re.compile(r'(?i)objetivo(?:\s*\(.*?\))?\s*:?')),
            ('activities', re.compile(r'(?i)actividades?\s+a\s+desarrollar(?:\s*\(.*?\))?\s*:?')),
            ('materials', re.compile(r'(?i)materiales?\s*:?')),
            ('scope', re.compile(r'(?i)[áa]mbito(?:\s+de\s+pr[áa]ctica)?\s*:?')),
        ]

        def strip_trailing_labels(segment: str) -> str:
            if not segment:
                return ''
            earliest = None
            for _, pattern in label_patterns:
                match = pattern.search(segment)
                if match:
                    earliest = match.start() if earliest is None else min(earliest, match.start())
            if earliest is None:
                return segment.strip()
            return segment[:earliest].strip()

        matches = []
        for label, pattern in label_patterns:
            match = pattern.search(text)
            if match:
                matches.append((match.start(), match.end(), label))

        matches.sort(key=lambda x: x[0])
        segments = {}

        for idx, (start, end, label) in enumerate(matches):
            next_start = matches[idx + 1][0] if idx + 1 < len(matches) else len(text)
            segment = text[end:next_start].strip()
            segments[label] = strip_trailing_labels(segment)

        objective_raw = segments.get('objective', '')
        print(f"    [DEBUG] objective_raw extraído ({len(objective_raw)} chars)")
        
        ra_codes = []
        for match in re.finditer(r'RA\s*(\d+)', objective_raw, re.IGNORECASE):
            code = f"RA{match.group(1)}"
            if code not in ra_codes:
                ra_codes.append(code)
                print(f"      -> Encontrado RA: {code}")
        
        print(f"    [DEBUG] RAs encontrados: {ra_codes}")
        
        return {
            'objective': objective_raw,
            'activities': segments.get('activities', '').strip(),
            'materials': segments.get('materials', '').strip(),
            'scope': segments.get('scope', '').strip(),
            'ra_codes': ra_codes,
        }

    practicals = []
    header_pattern = re.compile(r'pr[áa]ctico\s*n[°º]?\s*:?\s*(\d+)\s*(.*)', re.IGNORECASE)

    for table in doc.tables:
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
        table_text_lower = table_text.lower()

        if 'practico' not in table_text_lower or 'objetivo' not in table_text_lower:
            continue

        print(f"\n✓ Tabla con 'práct' y 'objetivo' encontrada")

        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]
            header_match = None
            header_cell_idx = None

            for cell_idx, cell_text in enumerate(row_cells):
                match = header_pattern.search(cell_text)
                if match:
                    header_match = match
                    header_cell_idx = cell_idx
                    break

            if not header_match:
                continue

            tp_number = header_match.group(1).strip() if header_match.group(1) else ''
            tp_name = header_match.group(2).strip() if header_match.group(2) else ''
            
            print(f"\n  Práctico {tp_number} encontrado (Fila {row_idx})")

            practical = {'number': tp_number or str(len(practicals) + 1), 'name': tp_name}

            if row_idx + 1 < len(table.rows):
                print(f"    Extrayendo bloque de Fila {row_idx + 1}")
                block_cells = table.rows[row_idx + 1].cells
                block_text = '\n'.join([extract_text_from_table_cell(cell) for cell in block_cells]).strip()
                
                parsed = parse_practical_block(block_text)
                practical['objective'] = parsed.get('objective', '')
                practical['activities'] = parsed.get('activities', '')
                practical['materials'] = parsed.get('materials', '')
                practical['scope'] = parsed.get('scope', '')
                
                if parsed.get('ra_codes'):
                    practical['ra_codes'] = parsed.get('ra_codes')
                    print(f"    ✓ ra_codes ASIGNADOS: {practical['ra_codes']}")
                else:
                    print(f"    NO SE ASIGNARON ra_codes")

            practicals.append(practical)

    return practicals

# Ejecutar
print("="*80)
print("EJECUTANDO extract_practicals_from_docx_DEBUG")
print("="*80)

practicals = extract_practicals_from_docx_DEBUG(doc)

print("\n" + "="*80)
print("RESULTADO FINAL")
print("="*80)

for idx, p in enumerate(practicals, 1):
    ra_codes = p.get('ra_codes', 'NO PRESENTE')
    print(f"\nTP{idx}: {p.get('name')} → RAs: {ra_codes}")
