"""Analizar estructura completa del documento"""
from docx import Document

doc_path = r"1º_1º - CBI - Álgebra I.docx"
doc = Document(doc_path)

print("="*80)
print("ESTRUCTURA DEL DOCUMENTO")
print("="*80)

print(f"\nTotal de párrafos: {len(doc.paragraphs)}")
print(f"Total de tablas: {len(doc.tables)}\n")

print("TABLAS EN EL DOCUMENTO:")
print("-"*80)
for idx, table in enumerate(doc.tables):
    # Obtener primeras líneas del contenido de la tabla
    first_cells = ' '.join([cell.text[:50] for cell in table.rows[0].cells if cell.text]) if table.rows else ""
    table_text = '\n'.join([' '.join([cell.text for cell in row.cells]) for row in table.rows])
    has_objetivo = 'objetivo' in table_text.lower()
    has_actividades = 'actividades' in table_text.lower()
    
    texto_celda1 = table.rows[0].cells[0].text[:60] if table.rows else ""
    
    print(f"\nTabla {idx}: Filas={len(table.rows)}, Cols={len(table.columns) if table.rows else 0}")
    print(f"  Contenido (primeros 100 chars): {first_cells[:100]}")
    print(f"  Tiene 'Objetivo': {has_objetivo}, Tiene 'Actividades': {has_actividades}")

print("\n" + "="*80)
print("BÚSQUEDA DE 'PRÁCTICO' EN EL DOCUMENTO")
print("="*80)

# Buscar en paragrafos
for idx, para in enumerate(doc.paragraphs):
    if 'práctico' in para.text.lower() or 'tp' in para.text.lower():
        print(f"\nPárrafo {idx}: {para.text[:100]}")

# Buscar en tablas
for table_idx, table in enumerate(doc.tables):
    table_text = '\n'.join([' '.join([cell.text for cell in row.cells]) for row in table.rows])
    if 'práctico' in table_text.lower():
        # Contar cuántas veces aparece
        count = table_text.lower().count('práctico')
        print(f"\nTabla {table_idx}: Contiene 'Práctico' {count} veces")
        # Mostrar líneas con "Práctico"
        for line in table_text.split('\n'):
            if 'práctico' in line.lower():
                print(f"  → {line[:90]}")
