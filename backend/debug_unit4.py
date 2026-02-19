#!/usr/bin/env python
import sys
sys.path.insert(0, '.')
from app.docx_import import import_proposal_from_docx
from app.main import get_docx_from_gdoc_url
import tempfile
import os
from docx import Document

gdoc_url = 'https://docs.google.com/document/d/1_7uykNmL_3QM0f_WjyvxJgLCw5nl00ME-fMZoVFDdQs/edit?tab=t.0'

try:
    print('[DESCARGAR] Descargando documento desde Google Docs...')
    docx_bytes = get_docx_from_gdoc_url(gdoc_url)
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name
    
    print('[OK] Documento descargado')
    
    # Analisis de estructura
    doc = Document(tmp_path)
    print(f'[INFO] ESTRUCTURA DEL DOCUMENTO:')
    print(f'[INFO]   Parrafos totales: {len(doc.paragraphs)}')
    print(f'[INFO]   Tablas totales: {len(doc.tables)}')
    
    # Ver parrafos con "Unidad"
    print(f'[DEBUG] PARRAFOS CON "UNIDAD":')
    for i, para in enumerate(doc.paragraphs):
        if 'unidad' in para.text.lower():
            print(f'[DEBUG]   [{i}] {para.text[:100]}')
    
    # Ver estructura de tablas
    print(f'[DEBUG] ESTRUCTURA DE TABLAS:')
    for t_idx, table in enumerate(doc.tables):
        print(f'[DEBUG] Tabla {t_idx}: {len(table.rows)} filas x {len(table.columns)} columnas')
        print(f'[DEBUG]   Primera fila: {[cell.text[:50] for cell in table.rows[0].cells]}')
        if 'unidad' in str(table.rows[0].cells[0].text).lower() or 'contenido' in str([cell.text for cell in table.rows[0].cells]).lower():
            print(f'[DEBUG]   >>> Esta tabla tiene "Unidad" o "Contenido"')
            for r_idx, row in enumerate(table.rows[:5]):
                print(f'[DEBUG]      Fila {r_idx}: {[cell.text[:40] for cell in row.cells]}')
    
    print('[PROCESS] Extrayendo propuesta...')
    data = import_proposal_from_docx(tmp_path, 'propuesta_test.docx')
    
    print('='*80)
    print('UNIDADES EXTRAIDAS')
    print('='*80)
    
    for i, unit in enumerate(data.get('units', []), 1):
        print(f'\nUnidad {i}: {unit.get("name")}')
        contenidos = unit.get('contenidos', '')
        print(f'  Longitud contenidos: {len(contenidos)} caracteres')
        if len(contenidos) > 0:
            print(f'  [OK] Primeros 200 caracteres: {contenidos[:200]}')
        else:
            print(f'  [ERROR] VACIO - Sin contenidos extraidos')
        
        if i == 4:
            print(f'\n  === UNIDAD 4 COMPLETO ===')
            print(f'  {contenidos if contenidos else "VACIO"}')
            print(f'  === FIN UNIDAD 4 ===')
    
    os.unlink(tmp_path)
    
except Exception as e:
    print(f'Error: {e}')
    import traceback
    traceback.print_exc()
