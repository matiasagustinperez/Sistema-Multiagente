#!/usr/bin/env python3
"""Análisis detallado de los datos en el DOCX real"""
import sys
sys.path.insert(0, r'C:\TesisMCD\backend')

from docx import Document

docx_path = r'C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(docx_path)

print("="*100)
print("CONTENIDO DE TABLAS RELEVANTES")
print("="*100)

# Tabla 3: Importancia y Fundamentos
print("\nTABLA 3: FUNDAMENTOS")
print("-" * 100)
table = doc.tables[3]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        text = cell.text.strip()[:200]
        print(f"  Row {row_idx}, Cell {cell_idx}: {text}")

# Tabla 4: Competencias Genéricas
print("\nTABLA 4: COMPETENCIAS GENÉRICAS")
print("-" * 100)
table = doc.tables[4]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        text = cell.text.strip()[:200]
        print(f"  Row {row_idx}, Cell {cell_idx}: {text}")

# Tabla 5-6: Unidades
print("\nTABLA 5: UNIDADES (parte 1)")
print("-" * 100)
table = doc.tables[5]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        text = cell.text.strip()[:150]
        print(f"  Row {row_idx}, Cell {cell_idx}: {text}")

print("\nTABLA 6: UNIDADES (parte 2)")
print("-" * 100)
table = doc.tables[6]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        text = cell.text.strip()[:150]
        print(f"  Row {row_idx}, Cell {cell_idx}: {text}")

# Tabla 7: Trabajos Prácticos
print("\nTABLA 7: TRABAJOS PRÁCTICOS")
print("-" * 100)
table = doc.tables[7]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        text = cell.text.strip()[:150]
        print(f"  Row {row_idx}, Cell {cell_idx}: {text}")

# Buscar Resultados de Aprendizaje en párrafos
print("\n\nBÚSQUEDA EN PÁRRAFOS: RESULTADOS DE APRENDIZAJE")
print("-" * 100)
for idx, para in enumerate(doc.paragraphs):
    if "RA" in para.text and ":" in para.text:
        print(f"  Párrafo {idx}: {para.text.strip()[:150]}")
