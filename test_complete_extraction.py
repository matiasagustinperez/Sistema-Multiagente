#!/usr/bin/env python3
"""
Crear un DOCX de test COMPLETO con todas las secciones
"""
import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

from docx import Document
from docx.shared import Pt, RGBColor
import tempfile

def create_complete_test_docx():
    """Crea un DOCX con estructura completa para testing"""
    doc = Document()
    
    # ENCABEZADO
    p = doc.add_paragraph()
    p.add_run("Carrera: ").bold = True
    p.add_run("Ingeniería Mecatrónica")
    
    p = doc.add_paragraph()
    p.add_run("Asignatura: ").bold = True
    p.add_run("Proyecto de Ingeniería Mecatrónica")
    
    # TABLA PROGRAMA ANALÍTICO
    doc.add_heading("Programa Analítico", level=2)
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'
    header = table.rows[0].cells
    header[0].text = 'Carácter'
    header[1].text = 'Régimen'
    header[2].text = 'Carga Horaria Total'
    header[3].text = 'Horas Teóricas'
    header[4].text = 'Horas Prácticas'
    header[5].text = 'Horas Semanales'
    
    data_row = table.rows[1].cells
    data_row[0].text = 'OBLIGATORIA'
    data_row[1].text = 'CUATRIMESTRAL'
    data_row[2].text = '180'
    data_row[3].text = '60'
    data_row[4].text = '120'
    data_row[5].text = '12'
    
    # TABLA EQUIPO DOCENTE
    doc.add_heading('EQUIPO DOCENTE', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Profesor'
    hdr[1].text = 'Categoría'
    hdr[2].text = 'Email'
    
    row1 = table.rows[1].cells
    row1[0].text = 'SMITH, JOHN'
    row1[1].text = 'TITULAR'
    row1[2].text = 'john.smith@univ.edu'
    
    row2 = table.rows[2].cells
    row2[0].text = 'JONES, MARY'
    row2[1].text = 'ADJUNTO'
    row2[2].text = 'mary.jones@univ.edu'
    
    # CONTENIDOS MÍNIMOS
    doc.add_heading('CONTENIDOS MÍNIMOS', level=2)
    doc.add_paragraph('Conceptos fundamentales de ingeniería de proyectos. Metodologías de gestión. Herramientas de planificación.')
    
    # FUNDAMENTOS
    doc.add_heading('FUNDAMENTOS', level=2)
    doc.add_heading('Importancia en el Plan de estudio:', level=3)
    doc.add_paragraph('Proyecto de Ingeniería Mecatrónica es la asignatura integradora horizontal y vertical, en el último nivel de Ingeniería Mecatrónica, para lograr una adecuada integración y aplicación del aprendizaje de las y los estudiantes en la Carrera. Incluye contenido y actividades didácticas destinadas a que las y los estudiantes logren capacidades para la formulación, planificación, ejecución, desarrollo y control de proyectos.')
    
    doc.add_heading('Relación con el perfil profesional esperado:', level=3)
    doc.add_paragraph('Esta asignatura contribuye a que el futuro ingeniero/a pueda formular, planificar, gestionar y desarrollar un proyecto seleccionando y utilizando las metodologías y herramientas más adecuadas. Facilita el desarrollo de competencias en el trabajo en equipo y en forma individual.')
    
    # OBJETIVOS
    doc.add_heading('OBJETIVOS', level=2)
    
    doc.add_heading('Competencias genéricas a las que aporta la asignatura', level=3)
    doc.add_paragraph('- CGT1 - Identificar, formular y resolver problemas de ingeniería mecatrónica - Alto')
    doc.add_paragraph('- CGT2 - Concebir, diseñar y desarrollar proyectos de ingeniería mecatrónica - Alto')
    doc.add_paragraph('- CGT3 - Gestionar, planificar, ejecutar y controlar proyectos de ingeniería mecatrónica - Alto')
    
    doc.add_heading('Competencias específicas a las que aporta la asignatura', level=3)
    doc.add_paragraph('- CE1 - Analizar la funcionalidad y aplicabilidad de máquinas, equipos y sistemas - Alto')
    doc.add_paragraph('- CE2 - Diseñar, calcular e implementar soluciones tecnológicas - Alto')
    doc.add_paragraph('- CE3 - Proyectar, dirigir y controlar la construcción y operación - Alto')
    
    doc.add_heading('Resultados de aprendizaje de la asignatura', level=3)
    doc.add_paragraph('- RA1 - Establece la forma y realiza la evaluación previa del proyecto, mediante estudio de factibilidad técnica/operativa, para decidir la viabilidad.')
    doc.add_paragraph('- RA2 - Formula la planificación del proyecto mediante evaluación y selección de la metodología más adecuada.')
    doc.add_paragraph('- RA3 - Desarrolla el proyecto mediante la gestión de recursos, entregables y documentación.')
    doc.add_paragraph('- RA4 - Analiza críticamente los resultados obtenidos y redacta información técnica.')
    doc.add_paragraph('- RA5 - Expone eficazmente de forma oral los resultados obtenidos en el Proyecto.')
    
    # UNIDADES
    doc.add_heading('CONTENIDOS DE LA ASIGNATURA: UNIDADES', level=2)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    h = table.rows[0].cells
    h[0].text = 'Unidad 1'
    h[1].text = 'Introducción a Gestión de Proyectos'
    d = table.rows[1].cells
    d[0].text = 'Contenidos:'
    d[1].text = 'Historia, conceptos básicos, tipos de proyectos, ciclo de vida.'
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    h = table.rows[0].cells
    h[0].text = 'Unidad 2'
    h[1].text = 'Planificación y Control'
    d = table.rows[1].cells
    d[0].text = 'Contenidos:'
    d[1].text = 'Cronograma, recursos, riesgos, seguimiento y control.'
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    h = table.rows[0].cells
    h[0].text = 'Unidad 3'
    h[1].text = 'Ejecución y Cierre'
    d = table.rows[1].cells
    d[0].text = 'Contenidos:'
    d[1].text = 'Ejecución de tareas, comunicación, cierre del proyecto.'
    
    # TRABAJOS PRÁCTICOS
    doc.add_heading('PROGRAMA DE TRABAJOS PRÁCTICOS', level=2)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    h = table.rows[0].cells
    h[0].text = 'Práctico 1'
    h[1].text = 'Planificación de Proyecto Piloto'
    d = table.rows[1].cells
    d[0].text = 'Objetivo:'
    d[1].text = 'Crear plan inicial para proyecto piloto seleccionado.'
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    h = table.rows[0].cells
    h[0].text = 'Práctico 2'
    h[1].text = 'Ejecución y Seguimiento'
    d = table.rows[1].cells
    d[0].text = 'Objetivo:'
    d[1].text = 'Ejecutar proyecto y monitorear progreso vs plan.'
    
    # METODOLOGÍA
    doc.add_heading('METODOLOGÍA', level=2)
    doc.add_paragraph('Clases teóricas con ejercicios prácticos. Trabajo en equipo. Desarrollo de proyecto real.')
    
    # EVALUACIÓN
    doc.add_heading('EVALUACIÓN', level=2)
    doc.add_paragraph('Parciales, trabajos prácticos, proyecto final, presentación oral.')
    
    # Guardar
    temp_path = r"C:\TesisMCD\test_complete_docx.docx"
    doc.save(temp_path)
    return temp_path

# Crear DOCX
docx_path = create_complete_test_docx()
print(f"DOCX creado: {docx_path}")

# Ahora testear la importación
from app.docx_import import import_proposal_from_docx

try:
    data = import_proposal_from_docx(docx_path)
    
    print("\n" + "="*80)
    print("EXTRACCIÓN DE DATOS")
    print("="*80)
    
    print(f"\nCarrera: {data.get('career')}")
    print(f"Asignatura: {data.get('subject')}")
    
    print(f"\n>>> COMPETENCIAS GENÉRICAS ({len(data.get('generic_competencies', []))}):")
    for comp in data.get("generic_competencies", []):
        print(f"  {comp.get('code')}: {comp.get('description')[:60]}... [{comp.get('level')}]")
    
    print(f"\n>>> COMPETENCIAS ESPECÍFICAS ({len(data.get('specific_competencies', []))}):")
    for comp in data.get("specific_competencies", []):
        print(f"  {comp.get('code')}: {comp.get('description')[:60]}... [{comp.get('level')}]")
    
    print(f"\n>>> RESULTADOS DE APRENDIZAJE ({len(data.get('learning_outcomes', []))}):")
    for ra in data.get("learning_outcomes", []):
        print(f"  {ra.get('code')}: {ra.get('description')[:70]}...")
    
    print(f"\n>>> IMPORTANCIA (primeros 150 chars):")
    importance = data.get("importance", "")
    print(f"  {importance[:150]}...")
    
    print(f"\n>>> PERFIL PROFESIONAL (primeros 150 chars):")
    profile = data.get("professional_profile", "")
    print(f"  {profile[:150]}...")
    
    print(f"\n>>> UNIDADES: {len(data.get('units', []))}")
    print(f">>> TRABAJOS PRÁCTICOS: {len(data.get('practicals', []))}")
    
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
