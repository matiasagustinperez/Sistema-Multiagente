#!/usr/bin/env python3
"""Buscar campos de Carrera y Asignatura"""
import sys
sys.path.insert(0, r'C:\TesisMCD\backend')

from docx import Document

docx_path = r'C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(docx_path)

print("="*100)
print("BÚSQUEDA DE CARRERA Y ASIGNATURA")
print("="*100)

# Buscar en los primeros párrafos
print("\nPRIMEROS PÁRRAFOS (búsqueda de metadatos):")
for idx, para in enumerate(doc.paragraphs[:15]):
    text = para.text.strip()
    if text:
        print(f"  {idx}: {text[:120]}")

# Buscar en Tabla 0 (Programa Analítico)
print("\n\nTABLA 0 (Programa Analítico) - primera tabla:")
table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        text = cell.text.strip()
        if text and len(text) > 3:
            print(f"  [{text[:50]}]")

# Buscar en tabla 1 (Docentes)
print("\n\nTABLA 1 (Docentes):")
table = doc.tables[1]
for row_idx, row in enumerate(table.rows[:3]):
    cells_text = " | ".join([cell.text.strip()[:40] for cell in row.cells])
    print(f"  Row {row_idx}: {cells_text}")

# Buscar propiedades del documento
print("\n\nPROPIEDADES DEL DOCUMENTO:")
props = doc.core_properties
print(f"  Title: [{props.title}]")
print(f"  Subject: [{props.subject}]")
print(f"  Author: [{props.author}]")
print(f"  Comments: [{props.comments}]")

# Buscar archivos incrustados o metadatos
print("\n\nBÚSQUEDA POR FILENAME:")
# El nombre del archivo contiene información
import os
filename = os.path.basename(docx_path)
print(f"  Filename: {filename}")
# Parsearlo
parts = filename.replace('.docx', '').replace('°', '').replace('_', ' ').split(' - ')
print(f"  Parsed: {parts}")
