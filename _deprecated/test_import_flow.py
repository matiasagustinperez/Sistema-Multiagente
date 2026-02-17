#!/usr/bin/env python3
"""
Test end-to-end del flujo de importación de DOCX.
"""
import sys
sys.path.insert(0, r'C:\TesisMCD\backend')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.docx_import import import_proposal_from_docx
import json
import os

# Cambiar al directorio backend para las importaciones
os.chdir(r'C:\TesisMCD\backend')

def create_test_docx():
    """Crea un DOCX de prueba con la estructura que espera docx_import.py"""
    doc = Document()
    
    # === ENCABEZADO ===
    header = doc.add_paragraph()
    header.add_run("PROGRAMA ANALÍTICO\n").bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === TABLA DE DATOS BÁSICOS ===
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Llenar con datos reales
    data = [
        ('Carrera:', 'Ingeniería Mecatrónica'),
        ('Asignatura:', 'Proyecto Integrador'),
        ('Año de Carrera:', '4'),
        ('Cuatrimestre:', '1'),
        ('Plan de Estudio:', '2023'),
        ('Año Académico:', '2024'),
    ]
    
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # === PROGRAMA ANALÍTICO ===
    doc.add_heading('PROGRAMA ANALÍTICO', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    prog_data = [
        ('Carácter', 'Obligatoria'),
        ('Régimen', 'Cuatrimestral'),
        ('Carga Horaria Total', '96'),
        ('Hs. Teóricas', '32'),
        ('Hs. Prácticas', '64'),
        ('Hs. Semanales', '6'),
        ('Código', 'PM401'),
    ]
    
    for i, (label, value) in enumerate(prog_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # === EQUIPO DOCENTE ===
    doc.add_heading('Equipo Docente', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Nombre'
    table.rows[0].cells[1].text = 'Categoría'
    table.rows[0].cells[2].text = 'Email'
    table.rows[1].cells[0].text = 'Dr. Juan Pérez'
    table.rows[1].cells[1].text = 'Profesor Titular'
    table.rows[1].cells[2].text = 'jperez@edu.ar'
    table.rows[2].cells[0].text = 'Ing. María García'
    table.rows[2].cells[1].text = 'JTP'
    table.rows[2].cells[2].text = 'mgarcia@edu.ar'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === CONTENIDOS MÍNIMOS ===
    doc.add_heading('CONTENIDOS MÍNIMOS', level=2)
    doc.add_paragraph('Diseño mecánico. Simulación. Prototipado. Integración de sistemas. Manufactura asistida por computadora.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === FUNDAMENTOS ===
    doc.add_heading('FUNDAMENTOS', level=2)
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Importancia en el Plan de Estudio'
    table.rows[1].cells[0].text = 'Esta asignatura constituye el proyecto integrador de la carrera, consolidando todos los conocimientos adquiridos en las disciplinas previas. Permite al estudiante aplicar dichos conocimientos en un proyecto real de ingeniería mecatrónica.'
    
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Relación con el Perfil Profesional Esperado'
    table.rows[1].cells[0].text = 'Prepara al futuro ingeniero mecatrónico para enfrentar desafíos reales en diseño, manufactura y automatización de sistemas mecatrónicos. Desarrolla habilidades de liderazgo y trabajo en equipo.'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === OBJETIVOS ===
    doc.add_heading('OBJETIVOS', level=2)
    
    doc.add_heading('Competencias Genéricas a las que aporta la asignatura', level=3)
    doc.add_paragraph('- CGT1: Capacidad de análisis y síntesis - [Alto]')
    doc.add_paragraph('- CGT2: Capacidad de resolución de problemas complejos - [Alto]')
    doc.add_paragraph('- CGT3: Trabajo efectivo en equipo - [Medio]')
    
    doc.add_heading('Competencias Específicas a las que aporta la asignatura', level=3)
    doc.add_paragraph('- CE1: Diseño mecánico avanzado mediante CAD - [Alto]')
    doc.add_paragraph('- CE2: Simulación de sistemas mediante análisis FEA - [Alto]')
    doc.add_paragraph('- CE3: Control industrial y automatización - [Medio]')
    
    doc.add_heading('Resultados de Aprendizaje de la asignatura', level=3)
    doc.add_paragraph('- RA1: Establece la forma más conveniente de abordar un problema mecatrónico real considerando aspectos técnico-económicos.')
    doc.add_paragraph('- RA2: Diseña y desarrolla soluciones mecatrónicas integradas.')
    doc.add_paragraph('- RA3: Implementa prototipación rápida utilizando herramientas modernas.')
    doc.add_paragraph('- RA4: Valida diseños mediante simulación y ensayos experimentales.')
    doc.add_paragraph('- RA5: Comunica resultados de forma clara y profesional.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === UNIDADES ===
    doc.add_heading('CONTENIDOS DE LA ASIGNATURA', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Unidad'
    table.rows[0].cells[1].text = 'Tema'
    table.rows[0].cells[2].text = 'Contenidos'
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Especificación del proyecto'
    table.rows[1].cells[2].text = 'Análisis de requisitos, definición del alcance, restricciones técnicas y económicas.'
    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = 'Diseño detallado'
    table.rows[2].cells[2].text = 'Modelado 3D, análisis de cargas, selección de materiales y procesos de manufactura.'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === TRABAJOS PRÁCTICOS ===
    doc.add_heading('PROGRAMA DE TRABAJOS PRÁCTICOS', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Práctico Nº'
    table.rows[0].cells[1].text = 'Nombre'
    table.rows[0].cells[2].text = 'Objetivo'
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Diseño CAD Básico'
    table.rows[1].cells[2].text = 'Aprender a utilizar software CAD profesional para modelado 3D'
    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = 'Simulación FEA'
    table.rows[2].cells[2].text = 'Realizar análisis de elementos finitos en estructuras'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === METODOLOGÍA ===
    doc.add_heading('METODOLOGÍA', level=2)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = 'Clases expositivas con ejercitación práctica. Trabajo en equipo en proyectos integradores. Uso de herramientas CAD/CAM modernas.'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === EVALUACIÓN ===
    doc.add_heading('EVALUACIÓN', level=2)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = 'Informe de proyecto (40%), Presentación final (40%), Participación en prácticos (20%)'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === BIBLIOGRAFÍA ===
    doc.add_heading('BIBLIOGRAFÍA', level=2)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = 'Shigley, J. E. (2020). Mechanical Engineering Design. McGraw Hill. ISBN: 978-1259-11408-9'
    
    # Guardar
    docx_path = r'C:\TesisMCD\test_import_flow.docx'
    doc.save(docx_path)
    print(f"✓ DOCX creado: {docx_path}")
    return docx_path

def test_import():
    """Prueba la importación y extracción de datos."""
    # Usar el DOCX real del usuario en lugar de crear uno de prueba
    docx_path = r'C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
    filename = '5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
    print(f"\n✓ Usando DOCX real: {docx_path}\n")
    
    print("\n" + "="*80)
    print("EXTRAYENDO DATOS DEL DOCX...")
    print("="*80 + "\n")
    
    data = import_proposal_from_docx(docx_path, filename)
    
    # Mostrar resultados
    print(f"✓ Carrera: {data.get('career')}")
    print(f"✓ Asignatura: {data.get('subject')}")
    print(f"✓ Año: {data.get('year_of_career')}")
    print(f"✓ Cuatrimestre: {data.get('quarter')}")
    print()
    
    # Competencias Genéricas
    gen_comp = data.get('generic_competencies', [])
    print(f"✓ COMPETENCIAS GENÉRICAS ({len(gen_comp)}):")
    for comp in gen_comp:
        print(f"  - {comp.get('code', '?')}: {comp.get('description', '')[:60]}... [{comp.get('level', 'N/A')}]")
    print()
    
    # Competencias Específicas
    spec_comp = data.get('specific_competencies', [])
    print(f"✓ COMPETENCIAS ESPECÍFICAS ({len(spec_comp)}):")
    for comp in spec_comp:
        print(f"  - {comp.get('code', '?')}: {comp.get('description', '')[:60]}... [{comp.get('level', 'N/A')}]")
    print()
    
    # Resultados de Aprendizaje
    learning = data.get('learning_outcomes', [])
    print(f"✓ RESULTADOS DE APRENDIZAJE ({len(learning)}):")
    for ra in learning:
        print(f"  - {ra.get('code', '?')}: {ra.get('description', '')[:60]}...")
    print()
    
    # Fundamentos
    importance = data.get('importance', '')
    prof_profile = data.get('professional_profile', '')
    print(f"✓ IMPORTANCIA: {importance[:100]}... ({len(importance)} chars)")
    print(f"✓ PERFIL PROFESIONAL: {prof_profile[:100]}... ({len(prof_profile)} chars)")
    print()
    
    # Equipo docente
    team = data.get('teaching_team', [])
    print(f"✓ EQUIPO DOCENTE ({len(team)}):")
    for doc_info in team:
        print(f"  - {doc_info.get('name', '?')} ({doc_info.get('category', '?')})")
    print()
    
    # Unidades y TPs
    units = data.get('units', [])
    practicals = data.get('practicals', [])
    print(f"✓ UNIDADES: {len(units)}")
    print(f"✓ TRABAJOS PRÁCTICOS: {len(practicals)}")
    print()
    
    # Verificar que los datos clave están presentes
    checks = [
        ('Carrera', bool(data.get('career'))),
        ('Asignatura', bool(data.get('subject'))),
        ('Competencias Genéricas', len(gen_comp) > 0),
        ('Competencias Específicas', len(spec_comp) > 0),
        ('Resultados de Aprendizaje', len(learning) > 0),
        ('Importancia', bool(importance)),
        ('Perfil Profesional', bool(prof_profile)),
        ('Equipo Docente', len(team) > 0),
    ]
    
    print("="*80)
    print("VALIDACIÓN FINAL:")
    print("="*80)
    for check_name, result in checks:
        status = "✓" if result else "✗"
        print(f"{status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    print()
    if all_pass:
        print("✓ TODOS LOS TESTS PASARON")
    else:
        print("✗ ALGUNOS TESTS FALLARON")
    
    return data

if __name__ == '__main__':
    test_import()
