import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

from docx import Document

# Cargar el DOCX
docx_path = r"C:\TesisMCD\backend\data\uploads\1°_2° - Estructuras de Datos.docx"
doc = Document(docx_path)

print("="*70)
print(f"ANALIZANDO TABLAS DEL DOCX")
print("="*70)

print(f"\nTotal de tablas: {len(doc.tables)}\n")

# Buscar tablas con competencias
for table_idx, table in enumerate(doc.tables):
    # Ver si tiene "competencia" u otro indicador
    full_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
    
    if 'CG' in full_text or 'CGT' in full_text or 'CGS' in full_text:
        print(f"TABLA {table_idx}: Contiene competencias")
        print(f"  Dimensiones: {len(table.rows)} filas x {len(table.columns)} columnas")
        
        # Mostrar primeras 3 filas
        for row_idx, row in enumerate(table.rows[:3]):
            print(f"\n  FILA {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                content = cell.text[:100]
                print(f"    [{col_idx}]: {content}...")
        
        print("\n" + "-"*70 + "\n")

print("="*70)
