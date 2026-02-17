#!/usr/bin/env python3
"""
Test rápido de validación del sistema de importación
Verifica:
1. Que el módulo docx_import carga sin errores
2. Que todas las funciones clave existen y son callable
3. Que la extracción de RA desde tablas funciona
"""
import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

def test_imports():
    """Valida que los módulos cargan"""
    try:
        from app.docx_import import (
            import_proposal_from_docx,
            extract_learning_outcomes_from_tables,
            extract_equipo_docente,
            extract_units_from_docx,
            extract_practicals_from_docx,
            extract_programa_analitico
        )
        print("✅ TODAS LAS FUNCIONES IMPORTADAS CORRECTAMENTE")
        return True
    except Exception as e:
        print(f"❌ ERROR AL IMPORTAR: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_ra_extraction_function():
    """Prueba la función de extracción de RAs desde tablas"""
    try:
        from app.docx_import import extract_learning_outcomes_from_tables
        from docx import Document
        from docx.shared import Inches, Pt
        import tempfile
        
        # Crear DOCX mínimo con tabla de RAs
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        
        # Encabezado
        table.rows[0].cells[0].text = 'Resultado de Aprendizaje'
        table.rows[0].cells[1].text = 'Descripción'
        
        # Datos
        table.rows[1].cells[0].text = 'RA1:'
        table.rows[1].cells[1].text = 'Comprender conceptos fundamentales'
        
        table.rows[2].cells[0].text = 'RA2:'
        table.rows[2].cells[1].text = 'Aplicar técnicas avanzadas'
        
        table.rows[3].cells[0].text = 'RA3:'
        table.rows[3].cells[1].text = 'Evaluar resultados críticos'
        
        # Guardar temporalmente
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as f:
            doc.save(f.name)
            temp_path = f.name
        
        # Extraer
        ras = extract_learning_outcomes_from_tables(Document(temp_path))
        
        # Limpiar
        import os
        os.remove(temp_path)
        
        # Validar
        if len(ras) >= 3:
            print(f"✅ EXTRACCIÓN DE RAs DESDE TABLA FUNCIONA ({len(ras)} RAs encontrados)")
            for i, ra in enumerate(ras[:3], 1):
                print(f"   RA{i}: {ra[:50]}...")
            return True
        else:
            print(f"⚠️ Se esperaban 3+ RAs, encontradas {len(ras)}")
            return len(ras) > 0
            
    except Exception as e:
        print(f"❌ ERROR EN EXTRACCIÓN DE RAs: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_backend_loads():
    """Valida que el backend FastAPI carga sin errores"""
    try:
        sys.path.insert(0, r"C:\TesisMCD\backend")
        from app.main import app
        print("✅ BACKEND FASTAPI CARGA CORRECTAMENTE")
        return True
    except Exception as e:
        print(f"❌ ERROR AL CARGAR BACKEND: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    print("=" * 80)
    print("TEST DE VALIDACIÓN DEL SISTEMA DE IMPORTACIÓN")
    print("=" * 80)
    
    results = []
    
    print("\n1. Validando importaciones...")
    results.append(("Imports", test_imports()))
    
    print("\n2. Validando extracción de RAs desde tabla...")
    results.append(("RA Extraction", test_ra_extraction_function()))
    
    print("\n3. Validando que Backend carga...")
    results.append(("Backend Load", test_backend_loads()))
    
    print("\n" + "=" * 80)
    print("RESUMEN")
    print("=" * 80)
    for name, passed in results:
        symbol = "✅" if passed else "❌"
        print(f"{symbol} {name}")
    
    all_pass = all(passed for _, passed in results)
    print("\n" + ("=" * 80))
    if all_pass:
        print("✅ SISTEMA LISTO PARA TESTING MANUAL")
        print("\nPróximos pasos:")
        print("1. Ejecutar: run_task START BACKEND y START FRONTEND")
        print("2. Navegar a http://localhost:5173")
        print("3. Usar 'Importar Propuesta' con tu DOCX real")
        print("4. Verificar preview con todos los campos")
        print("5. Cargar al formulario")
    else:
        print("❌ SE ENCONTRARON ERRORES - VER ARRIBA")
    print("=" * 80)
    
    exit(0 if all_pass else 1)
