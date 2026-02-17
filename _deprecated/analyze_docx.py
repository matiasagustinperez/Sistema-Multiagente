#!/usr/bin/env python3
"""Analizar la estructura del DOCX real"""
import sys
sys.path.insert(0, r'C:\TesisMCD\backend')

from docx import Document

docx_path = r'C:\TesisMCD\backend\5°_2° - Proyecto de Ingeniería Mecatrónica.docx'
doc = Document(docx_path)

print("="*80)
print("ANÁLISIS DEL DOCUMENTO")
print("="*80)

# Mostrar párrafos
print(f"\nTOTAL DE PÁRRAFOS: {len(doc.paragraphs)}\n")
print("PRIMEROS 50 PÁRRAFOS:")
print("-" * 80)
for i, para in enumerate(doc.paragraphs[:50]):
    text = para.text.strip()
    if text:
        level = para.style.name
        print(f"{i:3d}. [{level:20s}] {text[:100]}")

# Mostrar tablas
print(f"\n\nTOTAL DE TABLAS: {len(doc.tables)}\n")
print("TABLAS:")
print("-" * 80)
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    # Obtener primer texto de la tabla
    first_text = ""
    if table.rows and table.rows[0].cells:
        first_text = table.rows[0].cells[0].text[:80]
    
    has_competencias = any("competencia" in cell.text.lower() for row in table.rows for cell in row.cells)
    has_carrera = any("carrera" in cell.text.lower() for row in table.rows for cell in row.cells)
    has_docentes = any("docente" in cell.text.lower() for row in table.rows for cell in row.cells)
    
    tags = []
    if has_competencias:
        tags.append("COMPETENCIAS")
    if has_carrera:
        tags.append("CARRERA")
    if has_docentes:
        tags.append("DOCENTES")
    
    tag_str = f" [{', '.join(tags)}]" if tags else ""
    
    print(f"\nTabla {i}: {rows}x{cols} {tag_str}")
    print(f"  Contenido: {first_text}")
    
    # Mostrar headers
    if table.rows:
        headers = [cell.text[:40] for cell in table.rows[0].cells]
        print(f"  Headers: {headers}")

print("\n" + "="*80)
print("BÚSQUEDA DE PALABRAS CLAVE")
print("="*80)

# Buscar secciones
keywords = ['Carrera', 'Asignatura', 'Competencias', 'Genérica', 'Específica', 
            'Resultados', 'Aprendizaje', 'Importancia', 'Perfil', 'Unidad', 
            'Práctico', 'Docente', 'Metodología', 'Evaluación']

for keyword in keywords:
    found = 0
    for para in doc.paragraphs:
        if keyword.lower() in para.text.lower():
            found += 1
    
    if found > 0:
        print(f"\n'{keyword}': {found} ocurrencias")
        # Mostrar primeras 3 ocurrencias
        count = 0
        for para in doc.paragraphs:
            if keyword.lower() in para.text.lower() and count < 3:
                text = para.text.strip()[:100]
                print(f"  - {text}")
                count += 1
