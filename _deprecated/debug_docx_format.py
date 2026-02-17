import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

from docx import Document
import os

# Cargar el DOCX
docx_path = r"C:\TesisMCD\backend\data\uploads\1°_2° - Estructuras de Datos.docx"
doc = Document(docx_path)

print("="*70)
print("ANALIZANDO ESTRUCTURA DEL DOCX")
print("="*70)

# Buscar sección de objetivos
for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'OBJETIVO' in text.upper() or 'COMPETENCIA' in text.upper():
        print(f"\nPárrafo {para_idx}: {text[:80]}")
        
        # Mostrar los siguientes párrafos que podrían tener competencias
        for i in range(para_idx + 1, min(para_idx + 10, len(doc.paragraphs))):
            next_para = doc.paragraphs[i].text.strip()
            if next_para and not any(word in next_para.upper() for word in ['CONTENIDO', 'UNIDAD', 'EVALUACIÓN', 'BIBLIOGRAFÍA']):
                print(f"\nPárrafo {i}:")
                print(f"CONTENIDO: {next_para[:200]}")
                print(f"LONGITUD: {len(next_para)} caracteres")
                
                # Si contiene competencias, mostrarlo completo
                if 'CG' in next_para or 'CE' in next_para or 'RA' in next_para:
                    print(f"\nFULL TEXT:")
                    print(repr(next_para[:400]))
                    print("...")
                    break

print("\n" + "="*70)
