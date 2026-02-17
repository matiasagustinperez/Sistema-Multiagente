#!/usr/bin/env python3
"""
Test completo de la mejora de import/export
Verifica que se extraigan correctamente:
- Units (todos, sin truncar)
- Practicals (todos, sin truncar)
- RAs desde tablas
- Horas (todas: total, teórica, práctica, semanal)
- Docentes como array individual
"""
import sys
sys.path.insert(0, r"C:\TesisMCD\backend")

from app.docx_import import import_proposal_from_docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

def create_test_docx_with_ras():
    """Crea un DOCX de prueba con RAs en tabla"""
    doc = Document()
    
    # Encabezados
    p = doc.add_paragraph()
    p.add_run("Carrera: ").bold = True
    p.add_run("Ingeniería en Sistemas")
    
    p = doc.add_paragraph()
    p.add_run("Asignatura: ").bold = True
    p.add_run("Bases de Datos II")
    
    # Tabla de Programa Analítico
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Carácter'
    hdr_cells[1].text = 'Régimen'
    hdr_cells[2].text = 'Carga Horaria Total'
    hdr_cells[3].text = 'Horas Teóricas'
    hdr_cells[4].text = 'Horas Prácticas'
    hdr_cells[5].text = 'Horas Semanales'
    
    data_cells = table.rows[1].cells
    data_cells[0].text = 'OBLIGATORIA'
    data_cells[1].text = 'CUATRIMESTRAL'
    data_cells[2].text = '120'
    data_cells[3].text = '60'
    data_cells[4].text = '60'
    data_cells[5].text = '8'
    
    # Tabla de Docentes
    doc.add_heading('EQUIPO DOCENTE', level=3)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Profesor'
    hdr[1].text = 'Categoría'
    hdr[2].text = 'Email'
    
    row1 = table.rows[1].cells
    row1[0].text = 'SMITH, JOHN'
    row1[1].text = 'TITULAR'
    row1[2].text = 'john.smith@univ.edu'
    
    row2 = table.rows[2].cells
    row2[0].text = 'JONES, MARY'
    row2[1].text = 'ADJUNTO'
    row2[2].text = 'mary.jones@univ.edu'
    
    # Secciones de contenido
    doc.add_heading('CONTENIDOS MÍNIMOS', level=2)
    doc.add_paragraph('Conceptos fundamentales de bases de datos. Modelos de datos. SQL.')
    
    doc.add_heading('FUNDAMENTOS', level=2)
    doc.add_paragraph('Importancia en el Plan: Esta materia es fundamental para ingeniería.')
    
    doc.add_heading('OBJETIVOS DE APRENDIZAJE', level=2)
    
    # TABLA CON RAs - Esta es lo nuevo que probamos
    doc.add_paragraph('Resultados de Aprendizaje:')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    ras_data = [
        ('RA1:', 'Comprender conceptos fundamentales de bases de datos'),
        ('RA2:', 'Diseñar esquemas relacionales normalizados'),
        ('RA3:', 'Escribir consultas SQL complejas')
    ]
    
    for idx, (ra_num, ra_desc) in enumerate(ras_data, 1):
        row = table.rows[idx].cells
        row[0].text = ra_num
        row[1].text = ra_desc
    
    # Unidades
    doc.add_heading('CONTENIDOS DE LA ASIGNATURA: UNIDADES', level=2)
    doc.add_paragraph('Unidad 1: Introducción a BD').bold = True
    doc.add_paragraph('Contenidos: Historia, conceptos básicos, tipos de BD.')
    
    doc.add_paragraph('Unidad 2: Modelo Relacional').bold = True
    doc.add_paragraph('Contenidos: Tablas, claves, relaciones, integridad.')
    
    doc.add_paragraph('Unidad 3: SQL Avanzado').bold = True
    doc.add_paragraph('Contenidos: Joins, subconsultas, agregaciones, vistas.')
    
    doc.add_paragraph('Unidad 4: Optimización').bold = True
    doc.add_paragraph('Contenidos: Índices, ejecución, benchmarking.')
    
    # Prácticos
    doc.add_heading('PROGRAMA DE TRABAJOS PRÁCTICOS', level=2)
    doc.add_paragraph('TP 1: Diseño de ER').bold = True
    doc.add_paragraph('Objetivo: Crear diagrama ER de caso de estudio.')
    
    doc.add_paragraph('TP 2: Normalización').bold = True
    doc.add_paragraph('Objetivo: Normalizar esquemas hasta 3FN.')
    
    doc.add_paragraph('TP 3: SQL DML').bold = True
    doc.add_paragraph('Objetivo: Realizar operaciones CRUD complejas.')
    
    # Guardar
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=r'C:\TesisMCD') as f:
        doc.save(f.name)
        return f.name

def test_import():
    """Test la importación y extracción"""
    print("=" * 80)
    print("PRUEBA COMPLETA DE EXTRACCIÓN MEJORADA")
    print("=" * 80)
    
    # Crear DOCX de prueba
    print("\n1. Generando DOCX de prueba con RAs en tabla...")
    test_file = create_test_docx_with_ras()
    print(f"   ✓ Creado: {test_file}")
    
    # Importar
    print("\n2. Extrayendo datos...")
    try:
        data = import_proposal_from_docx(test_file)
        print("   ✓ Extracción exitosa")
    except Exception as e:
        print(f"   ✗ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # Verificación
    print("\n" + "-" * 80)
    print("RESULTADOS DE EXTRACCIÓN")
    print("-" * 80)
    
    # 1. Horas
    print(f"\n✓ HORAS (debe haber 4 campos):")
    hours_ok = (
        data.get('total_hours'),
        data.get('theoretical_hours'),
        data.get('practical_hours'),
        data.get('weekly_hours')
    )
    print(f"  Total: {data.get('total_hours')} hs")
    print(f"  Teórica: {data.get('theoretical_hours')} hs")
    print(f"  Práctica: {data.get('practical_hours')} hs")
    print(f"  Semanal: {data.get('weekly_hours')} hs")
    if all(hours_ok):
        print("  ✅ TODOS LOS CAMPOS PRESENTES")
    else:
        print("  ❌ FALTA ALGÚN CAMPO")
    
    # 2. Docentes como array
    print(f"\n✓ DOCENTES (debe ser array de 2):")
    teaching_team = data.get('teaching_team', [])
    print(f"  Cantidad: {len(teaching_team)}")
    if isinstance(teaching_team, list):
        for doc in teaching_team:
            print(f"    - {doc.get('name')} ({doc.get('category')}) - {doc.get('email')}")
        print("  ✅ FORMATO CORRECTO (array de objetos)")
    else:
        print(f"  ❌ FORMATO INCORRECTO: {type(teaching_team)}")
    
    # 3. Units (debe haber 4)
    print(f"\n✓ UNIDADES EXTRAÍDAS:")
    units = data.get('units', [])
    print(f"  Cantidad: {len(units)}")
    for unit in units:
        print(f"    - Unidad {unit.get('number', '?')}: {unit.get('name', 'SIN NOMBRE')}")
    if len(units) == 4:
        print("  ✅ TODAS LAS UNIDADES (4/4)")
    else:
        print(f"  ⚠️ Encontradas {len(units)} de 4 esperadas")
    
    # 4. Practicals (debe haber 3)
    print(f"\n✓ TRABAJOS PRÁCTICOS EXTRAÍDOS:")
    practicals = data.get('practicals', [])
    print(f"  Cantidad: {len(practicals)}")
    for tp in practicals:
        print(f"    - TP {tp.get('number', '?')}: {tp.get('name', 'SIN NOMBRE')}")
    if len(practicals) >= 3:
        print(f"  ✅ PRÁCTICOS EXTRAÍDOS ({len(practicals)}/3+)")
    else:
        print(f"  ⚠️ Encontrados {len(practicals)} de 3+ esperados")
    
    # 5. RAs desde tabla
    print(f"\n✓ RESULTADOS DE APRENDIZAJE (desde TABLA):")
    ras = data.get('learning_outcomes', [])
    print(f"  Cantidad: {len(ras)}")
    for ra in ras[:5]:  # Mostrar primeros 5
        preview = ra[:50] + '...' if len(ra) > 50 else ra
        print(f"    - {preview}")
    if len(ras) >= 3:
        print(f"  ✅ RAs EXTRAÍDOS DESDE TABLA ({len(ras)}/3+)")
    else:
        print(f"  ⚠️ Encontrados {len(ras)} de 3+ esperados")
    
    # Resumen final
    print("\n" + "=" * 80)
    print("RESUMEN")
    print("=" * 80)
    
    checks = [
        ("Todas las horas (4 campos)", all(hours_ok)),
        ("Docentes como array", isinstance(teaching_team, list) and len(teaching_team) == 2),
        ("Unidades sin truncar (4 unidades)", len(units) == 4),
        ("Prácticos sin truncar (3+ TP)", len(practicals) >= 3),
        ("RAs extraídos desde tabla (3+ RA)", len(ras) >= 3),
    ]
    
    for check_name, result in checks:
        symbol = "✅" if result else "❌"
        print(f"{symbol} {check_name}")
    
    all_pass = all(result for _, result in checks)
    
    print("\n" + ("=" * 80))
    if all_pass:
        print("✅ TODAS LAS PRUEBAS PASARON - SISTEMA FUNCIONA CORRECTAMENTE")
    else:
        print("❌ ALGUNAS PRUEBAS FALLARON - VER ARRIBA")
    print("=" * 80)
    
    # Limpiar
    try:
        os.remove(test_file)
    except:
        pass
    
    return all_pass

if __name__ == '__main__':
    success = test_import()
    exit(0 if success else 1)
